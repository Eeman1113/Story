#!pip install langchain langchain-ollama langchain-community langchain-core python-dotenv pypdf python-docx
# --- Imports ---
import os
import docx
import traceback
import random
import time
from dotenv import load_dotenv

# --- Langchain Imports ---
from langchain_ollama import OllamaLLM
from langchain_core.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.output_parsers import StrOutputParser
from langchain.schema import Document # To handle document objects

# --- Configuration ---
load_dotenv()
# Consider using a more powerful model available via Ollama if Gemma struggles with complexity
# DEFAULT_MODEL = "mixtral:latest" # Example: Or llama3:70b if available and powerful enough
DEFAULT_MODEL = "gemma2:9b" # Use Gemma 2 9B as a starting point
OUTPUT_FOLDER = './generated_novels'
RESUME_FOLDER = './docs' # Separate folder for input resumes

# --- LLM Initialization ---
def create_llm(temperature=0.7, top_p=0.9):
    """Creates a configured OllamaLLM instance."""
    print(f"Initializing LLM (Model: {DEFAULT_MODEL}, Temp: {temperature}, Top-P: {top_p})")
    # Add more parameters if needed (e.g., top_k, repeat_penalty)
    # Check Ollama documentation for available parameters for your model
    return OllamaLLM(
        model=DEFAULT_MODEL,
        temperature=temperature,
        top_p=top_p, # Helps control randomness along with temperature
        # Add other parameters as needed, e.g.:
        # top_k=40,
        # repeat_penalty=1.1
    )

# --- Utility Functions ---
def sanitize_filename(name):
    """Removes or replaces characters invalid for filenames."""
    name = str(name) # Ensure it's a string
    name = name.replace(':', '-').replace('/', '-').replace('\\', '-').replace('?', '').replace('*', '')
    name = name.replace('"', '').replace('<', '').replace('>', '').replace('|', '')
    name = name.replace('\n', ' ').replace('\r', '') # Replace newlines
    name = "".join(c for c in name if c.isalnum() or c in (' ', '.', '_', '-')).strip()
    name = name.replace(' ', '_')
    # Truncate if too long (OS limits)
    max_len = 100
    if len(name) > max_len:
        name = name[:max_len]
    while name.endswith('.') or name.endswith('_') or name.endswith('-'):
        name = name[:-1]
    if not name:
        name = "Untitled_Novel_" + str(int(time.time())) # Fallback
    return name

# --- Resume Loading ---
def load_resume_text(file_name):
    """Loads text content from a PDF file located in RESUME_FOLDER."""
    file_path = os.path.join(RESUME_FOLDER, file_name)
    if not os.path.exists(file_path):
         raise FileNotFoundError(f"Resume file not found at: {file_path}")
    try:
        print(f"Loading PDF from: {file_path}")
        loader = PyPDFLoader(file_path)
        docs = loader.load() # Returns list of Document objects
        if not docs:
            print(f"Warning: PyPDFLoader didn't extract any documents/pages from {file_name}.")
            return ""
        print(f"Successfully loaded {len(docs)} page(s) from PDF.")
        # Combine text from all pages
        full_text = '\n\n'.join([doc.page_content for doc in docs if doc.page_content])
        if not full_text.strip():
             print("Warning: Resume content appears empty after extraction.")
             return ""
        return full_text
    except Exception as e:
         print(f"Error loading PDF {file_path}: {e}")
         traceback.print_exc()
         raise

# --- Core Generation Components ---

class CharacterGeneratorChain:
    CHARACTER_PROFILE_PROMPT = """
    Based on the provided resume text and story concept, create a detailed character profile for the main protagonist.
    Flesh out their personality, motivations, and backstory beyond just the professional details. Make them feel like a real, complex person suitable for a novel.
    Infer potential personality traits, hobbies, relationships, fears, and driving desires that align with the resume and the story's subject/genre.

    Resume Text:
    {resume_text}

    Story Subject: {subject}
    Story Genre: {genre}

    Generate the following character details:
    - Full Name: (If not explicitly mentioned, create a plausible one)
    - Age (Approximate):
    - Appearance Snippet: (Brief description, 2-3 sentences)
    - Personality: (Describe core traits, positive and negative, e.g., ambitious, meticulous, introverted, cynical, resilient)
    - Core Motivation/Goal: (What drives them fundamentally in life or at the start of the story?)
    - Major Flaw/Internal Conflict: (What holds them back or creates internal struggle?)
    - Greatest Fear: (What are they most afraid of?)
    - Key Relationships (Initial Sketch): (Mention 1-2 important potential relationships - family, friend, rival, mentor - briefly)
    - Backstory Summary (Relevant to Subject): (2-4 sentences connecting their past to the story's premise, potentially using resume details)
    - Quirks/Hobbies (Optional but Recommended): (1-2 details to make them unique)

    Format the output clearly, using headings for each section.
    """

    def __init__(self):
        self.llm = create_llm(temperature=0.7, top_p=0.9) # Slightly creative for character
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.CHARACTER_PROFILE_PROMPT),
            verbose=True,
            output_parser=StrOutputParser()
        )

    def run(self, resume_text, subject, genre):
        try:
            print("Generating Detailed Character Profile...")
            profile_text = self.chain.invoke({
                "resume_text": resume_text,
                "subject": subject,
                "genre": genre
            })
            print("Character Profile Generated.")
            # Basic parsing into a dictionary (can be made more robust)
            profile_dict = {}
            current_key = None
            for line in profile_text.strip().split('\n'):
                line = line.strip()
                if not line: continue
                if ':' in line:
                    key, value = line.split(':', 1)
                    current_key = key.strip().replace(' ', '_').lower() # Normalize key
                    profile_dict[current_key] = value.strip()
                elif current_key: # Append to previous key if line continues
                    profile_dict[current_key] += "\n" + line
            
            # Ensure essential keys exist, provide defaults if missing
            defaults = {
                'full_name': 'Unknown Protagonist',
                'age_(approximate)': 'Late 20s / Early 30s',
                'personality': 'To be discovered',
                'core_motivation/goal': 'Survival and finding meaning',
                'major_flaw/internal_conflict': 'Self-doubt',
                'greatest_fear': 'Failure',
                'key_relationships_(initial_sketch)': 'None initially specified',
                'backstory_summary_(relevant_to_subject)': 'Background is currently unclear.'
            }
            for key, default_val in defaults.items():
                if key not in profile_dict or not profile_dict[key]:
                    profile_dict[key] = default_val
                    
            profile_dict['current_mood'] = "Neutral" # Add dynamic state
            profile_dict['short_term_goal'] = "Assess the situation" # Initial goal
            
            return profile_dict

        except Exception as e:
            print(f"Error generating character profile: {e}")
            traceback.print_exc()
            return {"error": str(e), **defaults} # Return error state with defaults

class ThemeGeneratorChain:
    THEME_PROMPT = """
    Analyze the provided story concept, genre, and initial character profile.
    Identify and articulate 3-5 core thematic questions or statements that could underpin this novel.
    Themes should be resonant and offer potential for exploration throughout the narrative. Think about deeper meanings and universal human experiences related to the concept.

    Story Subject: {subject}
    Story Genre: {genre}
    Main Character Profile Snippet:
    - Core Motivation: {motivation}
    - Major Flaw: {flaw}
    - Backstory Hint: {backstory}

    Generate 3-5 distinct themes. Present them as concise questions or statements.
    Examples:
    - Can technology truly replace human connection?
    - The struggle between individual ambition and societal responsibility.
    - How does one maintain hope in a dystopian world?
    - What is the cost of uncovering forbidden knowledge?
    - Finding identity when stripped of familiar structures.

    Core Themes:
    1. [Theme 1]
    2. [Theme 2]
    3. [Theme 3]
    (Continue up to 5 if appropriate)
    """

    def __init__(self):
        self.llm = create_llm(temperature=0.6, top_p=0.8) # More focused for themes
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.THEME_PROMPT),
            verbose=True,
            output_parser=StrOutputParser()
        )

    def run(self, subject, genre, character_profile):
        try:
            print("Generating Core Themes...")
            themes_text = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "motivation": character_profile.get('core_motivation/goal', 'Unknown'),
                "flaw": character_profile.get('major_flaw/internal_conflict', 'Unknown'),
                "backstory": character_profile.get('backstory_summary_(relevant_to_subject)', 'Unknown')
            })
            themes = [t.strip() for t in themes_text.strip().split('\n') if t.strip()]
            # Filter out potential numbering/list markers
            themes = [t.split('.', 1)[-1].strip() if t.strip()[0].isdigit() else t for t in themes]
            print(f"Themes Generated: {themes}")
            return themes if themes else ["Theme generation failed, using fallback: The struggle for survival."]
        except Exception as e:
            print(f"Error generating themes: {e}")
            traceback.print_exc()
            return ["Error generating themes: " + str(e)]

class WorldBibleManager:
    """Manages key world-building elements."""
    WORLD_BIBLE_PROMPT = """
    Based on the story subject, genre, and relevant details from the character's resume (locations, skills), establish foundational elements for the novel's world ('World Bible').
    Focus on details crucial for consistency and atmosphere, especially for the Sci-Fi / post-UFO context. Use real places mentioned in the resume as anchors.

    Story Subject: {subject}
    Story Genre: {genre}
    Resume Snippet (Locations/Skills): {resume_context}
    Character Profile Hint: {character_context}

    Generate entries for the following World Bible categories:
    - Key Locations (Expand on resume places): Describe 2-3 specific locations mentioned or implied in the resume (e.g., {resume_place_1}, {resume_place_2}), adapting them to the post-UFO world. What are they like now? Their significance?
    - UFO Domination Context: Briefly explain the nature of the UFO presence or its aftermath. How does it impact daily life, society, technology?
    - Factions/Groups (If applicable): Are there distinct groups (human resistance, collaborators, alien types, governing bodies)? Briefly name and describe 1-2 if relevant to the subject.
    - Technology Level & Style: What's the general tech level? Is it advanced alien tech, decaying human tech, a mix? Describe the aesthetic (e.g., retro-futuristic, bio-mechanical, makeshift).
    - World Rules/Atmosphere: Note 1-2 key 'rules' of this world (e.g., curfews, resource scarcity, strange environmental effects) and the overall mood (e.g., oppressive, hopeful, mysterious, desolate).

    Format clearly with headings for each category. Keep descriptions concise but evocative.
    """

    def __init__(self, subject, genre, resume_text, character_profile):
        self.subject = subject
        self.genre = genre
        self.resume_text = resume_text # Keep full text for potential future lookups
        self.character_profile = character_profile
        self.world_data = {}
        self.llm = create_llm(temperature=0.65, top_p=0.85) # Moderately creative for world details
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.WORLD_BIBLE_PROMPT),
            verbose=True,
            output_parser=StrOutputParser()
        )

    def _extract_resume_context(self, text):
        # Basic extraction - can be improved with NLP/regex
        lines = text.split('\n')
        locations = []
        skills = []
        # Look for headings like "Locations", "Projects", "Skills", "Experience"
        # This is highly dependent on resume format
        # For Thalen.pdf specifically, look for addresses or project locations
        potential_locations = ["Pune", "Maharashtra", "India", "CEPT University", "Ahmedabad", "Gujarat"] # Add more based on resume inspection
        for loc in potential_locations:
            if loc.lower() in text.lower():
                locations.append(loc)
        
        # Simple skill check (very basic)
        if "architect" in text.lower(): skills.append("Architecture")
        if "design" in text.lower(): skills.append("Design")
        if "software" in text.lower(): skills.append("Relevant Software (unspecified)")

        loc_str = ", ".join(list(set(locations))) if locations else "Specific locations not clearly identified in resume"
        skill_str = ", ".join(list(set(skills))) if skills else "Specific skills not clearly identified in resume"
        
        # Placeholder for specific places needed in prompt (replace manually or improve extraction)
        resume_place_1 = locations[0] if locations else "a key city"
        resume_place_2 = locations[1] if len(locations)>1 else "another significant area"

        return f"Potential Locations: {loc_str}. Potential Skills: {skill_str}.", resume_place_1, resume_place_2


    def generate_initial_bible(self):
        print("Generating World Bible...")
        resume_context, place1, place2 = self._extract_resume_context(self.resume_text)
        char_context = f"Protagonist is {self.character_profile.get('age_(approximate)', 'an adult')} {self.character_profile.get('personality', '')} driven by {self.character_profile.get('core_motivation/goal', '')}."

        try:
            bible_text = self.chain.invoke({
                "subject": self.subject,
                "genre": self.genre,
                "resume_context": resume_context,
                "character_context": char_context,
                "resume_place_1": place1,
                "resume_place_2": place2
            })
            # Simple parsing (can be improved)
            self.world_data = {}
            current_key = None
            for line in bible_text.strip().split('\n'):
                line = line.strip()
                if not line: continue
                # Check if line looks like a heading (e.g., ends with ':', all caps, etc.) - heuristic
                if line.endswith(':') or line.isupper():
                    current_key = line.replace(':','').strip()
                    self.world_data[current_key] = ""
                elif current_key:
                    self.world_data[current_key] += line + "\n"
                else: # Orphan line, maybe add to a general notes section
                    if "General Notes" not in self.world_data: self.world_data["General Notes"] = ""
                    self.world_data["General Notes"] += line + "\n"

            # Clean up trailing newlines
            for key in self.world_data:
                 self.world_data[key] = self.world_data[key].strip()

            print("World Bible Generated.")
            # Add resume locations explicitly if not captured well
            if 'Key Locations' not in self.world_data: self.world_data['Key Locations'] = ""
            extracted_locs = list(set(re.findall(r'\b[A-Z][a-zA-Z]+(?: [A-Z][a-zA-Z]+)*\b', resume_context))) # Simple Noun Phrase extraction
            for loc in extracted_locs:
                if loc not in self.world_data['Key Locations']:
                    self.world_data['Key Locations'] += f"\n- {loc}: (Needs description in post-UFO context)"


            if not self.world_data:
                 print("Warning: World Bible generation yielded empty results. Using fallback.")
                 self.world_data = {
                    "Key Locations": f"Primarily based around {place1} and surrounding areas, adapted to the new reality.",
                    "UFO Domination Context": "Subtle but pervasive influence; society is rebuilding or adapting under observation/control.",
                    "Technology Level & Style": "Mix of struggling pre-UFO tech and scavenged/poorly understood alien elements.",
                    "World Rules/Atmosphere": "Sense of unease, resource management is key, underlying mystery about the 'visitors'."
                }

        except Exception as e:
            print(f"Error generating World Bible: {e}")
            traceback.print_exc()
            self.world_data = {"error": str(e)}

        return self.world_data

    def get(self, key, default="Not specified"):
        """Safely retrieve a world data entry."""
        return self.world_data.get(key, default)

    def get_relevant_context(self, *keys):
        """Get a concatenated string of specified world data entries."""
        context = ""
        for key in keys:
            if key in self.world_data:
                context += f"{key}:\n{self.world_data[key]}\n\n"
        return context.strip() if context else "No specific world context available."

# --- Plot Promise System (Enhanced) ---

class PromiseGeneratorChain:
    PROMISE_PROMPT = """
    Generate a set of 5-7 interconnected and compelling 'plot promises' for a novel based on these details. These promises will form the backbone of the story's major arcs and subplots.
    Ensure variety: include promises related to the main external conflict, character development/internal conflict, relationships, and mysteries/world exploration.
    Consider the established themes and world context.

    Novel Title: "{title}"
    Genre: {genre}
    Themes: {themes}
    Main Character Profile: {profile_summary}
    World Context Snippet: {world_context}
    Subject: {subject}

    A plot promise is a narrative thread introduced early, developed with rising stakes and turning points, and culminating in a satisfying payoff.

    For each promise, provide:
    - A clear description of the promise (e.g., "Character will uncover the truth behind the UFO's arrival")
    - An importance score (1-10, higher = more central plotline)
    - A brief note on the expected payoff/climax (e.g., "Confrontation with the lead alien entity and a choice that determines humanity's future")
    - Potential Connection: (Optional: Mention if it connects to another promise, e.g., "Connects to Promise #2")

    Generate 5-7 varied plot promises. Format each promise exactly as:

    PROMISE: [Description]
    IMPORTANCE: [1-10]/10
    PAYOFF: [Brief description of eventual resolution]
    CONNECTION: [Optional connection note]
    """
    # (Keep parse_promises method similar to original, maybe add CONNECTION parsing)
    def __init__(self):
        self.llm = create_llm(temperature=0.75, top_p=0.9) # High creativity for plot ideas
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMISE_PROMPT),
            verbose=True,
            output_parser=StrOutputParser()
        )

    def parse_promises(self, raw_text):
        """Parse the LLM response into structured promise objects"""
        promises = []
        current_promise = {}
        lines = raw_text.strip().split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                # Save previous promise if it exists and is valid
                if current_promise and 'description' in current_promise:
                     # Ensure minimum fields before saving
                     if 'importance' not in current_promise: current_promise['importance'] = 5
                     if 'payoff' not in current_promise: current_promise['payoff'] = "Reaches a satisfying conclusion."
                     current_promise['progress'] = 0
                     current_promise['complete'] = False
                     current_promise['last_progressed_scene'] = 0
                     current_promise['target_chapter'] = None # To be set by outline
                     current_promise['target_act'] = None # To be set by outline
                     promises.append(current_promise)
                current_promise = {} # Reset for next promise block
                continue

            if line.startswith("PROMISE:"):
                 # If starting a new promise block but previous wasn't saved (e.g. no blank line), save previous first
                if current_promise and 'description' in current_promise:
                    if 'importance' not in current_promise: current_promise['importance'] = 5
                    if 'payoff' not in current_promise: current_promise['payoff'] = "Reaches a satisfying conclusion."
                    current_promise['progress'] = 0
                    current_promise['complete'] = False
                    current_promise['last_progressed_scene'] = 0
                    current_promise['target_chapter'] = None
                    current_promise['target_act'] = None
                    promises.append(current_promise)

                current_promise = {'description': line[len("PROMISE:"):].strip()}
            elif line.startswith("IMPORTANCE:") and current_promise:
                try:
                    importance_text = line[len("IMPORTANCE:"):].strip()
                    importance = int(importance_text.split('/')[0])
                    current_promise['importance'] = min(max(importance, 1), 10)
                except ValueError:
                    current_promise['importance'] = 5 # Default
            elif line.startswith("PAYOFF:") and current_promise:
                current_promise['payoff'] = line[len("PAYOFF:"):].strip()
            elif line.startswith("CONNECTION:") and current_promise:
                 current_promise['connection'] = line[len("CONNECTION:"):].strip()
            # Handle cases where the line might be a continuation of a field (less likely with strict format but possible)
            elif current_promise and 'description' in current_promise and not any(line.startswith(kw) for kw in ["IMPORTANCE:", "PAYOFF:", "CONNECTION:"]):
                 # Append to the last added field (heuristic: assume description or payoff if long)
                 if 'payoff' in current_promise:
                     current_promise['payoff'] += "\n" + line
                 elif 'description' in current_promise:
                     current_promise['description'] += "\n" + line


        # Add the last promise if it exists and is valid
        if current_promise and 'description' in current_promise:
            if 'importance' not in current_promise: current_promise['importance'] = 5
            if 'payoff' not in current_promise: current_promise['payoff'] = "Reaches a satisfying conclusion."
            current_promise['progress'] = 0
            current_promise['complete'] = False
            current_promise['last_progressed_scene'] = 0
            current_promise['target_chapter'] = None
            current_promise['target_act'] = None
            promises.append(current_promise)

        print(f"Parsed {len(promises)} promises.")
        # Add index to each promise dict for easier reference
        for i, p in enumerate(promises):
            p['id'] = i

        return promises


    def run(self, title, genre, themes, profile_summary, world_context, subject):
        try:
            print("Generating Initial Plot Promises...")
            result = self.chain.invoke({
                "title": title,
                "genre": genre,
                "themes": "\n- ".join(themes),
                "profile_summary": profile_summary, # Pass a summarized version
                "world_context": world_context, # Pass relevant snippets
                "subject": subject
            })

            promises = self.parse_promises(result)

            if not promises:
                 print("Warning: Failed to parse promises. Using fallbacks.")
                 # Create more relevant fallbacks based on input
                 promises = [
                    {
                        'id': 0, 'description': f"The protagonist must navigate the challenges of daily life in post-UFO {world_context.get('Key Locations', 'the city')}",
                        'importance': 7, 'payoff': "Achieves a stable footing or understanding of the new world.", 'progress': 0, 'complete': False, 'last_progressed_scene': 0, 'target_chapter': None, 'target_act': None
                    },
                    {
                        'id': 1, 'description': f"Explore the central theme: '{themes[0] if themes else 'Survival'}' through a specific challenge.",
                        'importance': 8, 'payoff': "Protagonist undergoes significant internal change related to the theme.", 'progress': 0, 'complete': False, 'last_progressed_scene': 0, 'target_chapter': None, 'target_act': None
                    },
                    {
                        'id': 2, 'description': f"Uncover a local mystery related to the UFO presence or its effects on {world_context.get('Key Locations', 'the area')}.",
                        'importance': 6, 'payoff': "Reveals a piece of the larger puzzle about the visitors or the world's state.", 'progress': 0, 'complete': False, 'last_progressed_scene': 0, 'target_chapter': None, 'target_act': None
                    }
                ]

            return promises
        except Exception as e:
            print(f"An error occurred in PromiseGeneratorChain.run: {e}")
            traceback.print_exc()
            return []


class PromiseManager:
    """Manages plot promises, linking them to the outline and tracking progress."""
    def __init__(self, initial_promises):
        self.promises = {p['id']: p for p in initial_promises} # Store by ID
        self.completed_promises = {}
        self.story_events = [] # Tracks events related to promises
        self.current_scene_index = 0

    def assign_promise_to_outline(self, promise_id, act, chapter):
        if promise_id in self.promises:
            self.promises[promise_id]['target_act'] = act
            self.promises[promise_id]['target_chapter'] = chapter
            print(f"Assigned Promise {promise_id} to Act {act}, Chapter {chapter}")
            return True
        return False

    def mark_promise_progressed(self, promise_id, scene_index, event_description, progress_increment=1):
        if promise_id in self.promises:
            promise = self.promises[promise_id]
            promise['progress'] += progress_increment
            promise['last_progressed_scene'] = scene_index
            
            event = {
                'scene': scene_index,
                'promise_id': promise_id,
                'promise': promise['description'],
                'event': event_description,
                'progress_level': promise['progress']
            }
            self.story_events.append(event)
            print(f"Scene {scene_index}: Progress on Promise {promise_id} - {event_description} (New Level: {promise['progress']})")

            # Simple completion check (can be made more sophisticated, e.g., based on expected payoff)
            # Let's say progress >= 4 means complete for now
            if promise['progress'] >= 4:
                promise['complete'] = True
                self.completed_promises[promise_id] = self.promises.pop(promise_id)
                print(f"Promise {promise_id} completed and moved: '{promise['description']}'")
            return True
        print(f"Warning: Tried to progress non-existent promise ID {promise_id}")
        return False

    def suggest_promises_for_chapter(self, target_act, target_chapter, count=3):
        """Suggest promises relevant to the current chapter/act, prioritizing those assigned or needing progress."""
        active_promises = list(self.promises.values())
        if not active_promises:
            return []

        scored_promises = []
        for promise in active_promises:
            score = 0
            # High priority if assigned to this specific chapter/act
            if promise['target_act'] == target_act:
                score += 50
                if promise['target_chapter'] == target_chapter:
                    score += 100

            # Boost based on importance
            score += promise['importance'] * 5

            # Boost if not progressed recently
            recency_factor = self.current_scene_index - promise['last_progressed_scene']
            score += recency_factor # More scenes passed = higher score boost

            # Penalize if already progressed significantly (unless it's the target chapter for payoff)
            if promise['progress'] >= 2 and promise['target_chapter'] != target_chapter:
                 score -= 20 * promise['progress']
            
            # Ensure target promises for this chapter get top priority if not progressed yet
            if promise['target_chapter'] == target_chapter and promise['progress'] == 0:
                score += 200


            scored_promises.append((promise['id'], promise['description'], score))

        # Sort by score (highest first)
        scored_promises.sort(key=lambda x: x[2], reverse=True)

        # Return ID and description
        return [(pid, desc) for pid, desc, _ in scored_promises[:count]]

    def get_active_promises_summary(self):
        return [f"ID {pid}: {p['description']} (Importance: {p['importance']}, Progress: {p['progress']})" for pid, p in self.promises.items()]

    def get_promise_details(self, promise_id):
        return self.promises.get(promise_id)

    def increment_scene_counter(self):
         self.current_scene_index += 1

    def get_story_summary(self, recent_count=None):
         events_to_show = self.story_events[-recent_count:] if recent_count else self.story_events
         return [f"Scene {e['scene']}: {e['event']} (Promise: {e['promise'][:50]}...)" for e in events_to_show]


# --- Outline and Structure ---

class OutlineGeneratorChain:
    OUTLINE_PROMPT = """
    You are a master novelist outlining a compelling story. Create a structured outline for a novel based on the provided elements.
    Divide the story into logical Acts (e.g., 3 Acts: Setup, Confrontation, Resolution) or Parts.
    Within each Act/Part, define several Chapters. For each Chapter, specify:
    - A concise Chapter Title (evocative, hint at content)
    - A brief Goal/Purpose for the chapter (e.g., "Introduce the central conflict," "Protagonist makes a crucial decision," "Explore Location X," "Raise the stakes")
    - Key Plot Promise(s) to be advanced or introduced in this chapter (Refer to Promise IDs: e.g., "Advance Promise #1, Introduce Promise #4")
    - Key Character Arc Moment(s) (e.g., "Character faces their fear," "Relationship with Y develops," "Character questions their motivation")
    - Estimated Scene Count: (A rough estimate like 2-4 scenes)

    Ensure the outline builds narrative momentum, escalates stakes, explores the themes, and allows for character development consistent with their profile. Weave the plot promises naturally into the structure. Aim for approximately {num_chapters} total chapters across the Acts/Parts.

    Novel Title: "{title}"
    Genre: {genre}
    Themes: {themes}
    Main Character Profile: {profile_summary}
    World Context Snippet: {world_context}
    Available Plot Promises:
    {plot_promises_list}

    Generate the outline using clear headings for Acts/Parts and Chapters. Be detailed enough to guide the writing process.

    Example Act/Chapter Structure:

    ACT I: The Setup (~ Chapters 1-{end_act1})

    Chapter 1: [Title]
    Goal: Introduce protagonist and the 'normal' world (post-UFO context). Hint at underlying tension.
    Promises: Introduce Promise #0 (Daily Life).
    Character Moment: Showcase protagonist's core personality and current struggle.
    Scenes: 2-3

    Chapter 2: [Title]
    Goal: Inciting Incident - event that disrupts normalcy and kicks off the main plot.
    Promises: Introduce Promise #2 (Mystery). Advance Promise #0.
    Character Moment: Protagonist forced to react/make an initial choice.
    Scenes: 3-4
    ...

    ACT II: Confrontation (~ Chapters {start_act2}-{end_act2})
    ...

    ACT III: Resolution (~ Chapters {start_act3}-{num_chapters})
    ...

    --- BEGIN OUTLINE ---
    """

    def __init__(self, num_chapters=15):
        self.llm = create_llm(temperature=0.6, top_p=0.8) # More structured, less random
        self.num_chapters = num_chapters
        # Calculate approximate act boundaries for the prompt
        self.end_act1 = num_chapters // 3
        self.start_act2 = self.end_act1 + 1
        self.end_act2 = (num_chapters * 2) // 3
        self.start_act3 = self.end_act2 + 1

        self.prompt = PromptTemplate.from_template(self.OUTLINE_PROMPT)
        self.chain = LLMChain(llm=self.llm, prompt=self.prompt, verbose=True, output_parser=StrOutputParser())

    def parse_outline(self, outline_text):
        """Parses the structured outline text into a list of dictionaries."""
        outline = {"acts": []}
        current_act = None
        current_chapter = None

        lines = outline_text.strip().split('\n')
        chapter_counter = 0

        for line in lines:
            line = line.strip()
            if not line: continue

            # Detect Act heading (simple heuristic: starts with ACT or PART)
            if line.upper().startswith("ACT ") or line.upper().startswith("PART "):
                current_act = {"title": line, "chapters": []}
                outline["acts"].append(current_act)
                continue

            # Detect Chapter heading (simple heuristic: starts with Chapter)
            if line.lower().startswith("chapter "):
                chapter_counter += 1
                # Extract title if possible
                title_match = re.match(r"chapter \d+:\s*(.*)", line, re.IGNORECASE)
                chapter_title = title_match.group(1).strip() if title_match else f"Chapter {chapter_counter}: Untitled"

                current_chapter = {
                    "number": chapter_counter,
                    "title": chapter_title,
                    "goal": "",
                    "promises": [],
                    "character_moment": "",
                    "estimated_scenes": 2 # Default
                }
                if current_act:
                    current_act["chapters"].append(current_chapter)
                else:
                    # Handle chapters outside of acts (fallback)
                    if not outline["acts"]: # Create a default act if none found
                         current_act = {"title": "Part 1", "chapters": []}
                         outline["acts"].append(current_act)
                    current_act["chapters"].append(current_chapter) # Add to last known act
                continue

            # Parse chapter details if inside a chapter
            if current_chapter:
                if line.lower().startswith("goal:"):
                    current_chapter["goal"] = line[len("Goal:"):].strip()
                elif line.lower().startswith("promises:"):
                     # Extract promise IDs (handles "Promise #1", "ID 2", "Advance 3")
                    ids_text = line[len("Promises:"):].strip()
                    ids = re.findall(r'#?(\d+)', ids_text)
                    current_chapter["promises"] = [int(id) for id in ids]
                elif line.lower().startswith("character moment:") or line.lower().startswith("character arc moment"):
                    current_chapter["character_moment"] = line.split(':', 1)[-1].strip()
                elif line.lower().startswith("scenes:") or line.lower().startswith("estimated scenes:"):
                    try:
                        scenes_text = line.split(':', 1)[-1].strip()
                        # Handle ranges like "2-3" -> use upper bound? or average? Let's take first digit found.
                        match = re.search(r'\d+', scenes_text)
                        if match:
                            current_chapter["estimated_scenes"] = int(match.group(0))
                        else:
                             current_chapter["estimated_scenes"] = 2 # Fallback if no number found
                    except ValueError:
                        current_chapter["estimated_scenes"] = 2 # Fallback
                # Append continuation lines to the last parsed field (heuristic)
                elif not any(line.lower().startswith(kw) for kw in ["goal:", "promises:", "character", "scenes:"]) and current_chapter["goal"]: # Assume continuation of goal if set
                     current_chapter["goal"] += " " + line
                elif not any(line.lower().startswith(kw) for kw in ["goal:", "promises:", "character", "scenes:"]) and current_chapter["character_moment"]: # Or character moment
                    current_chapter["character_moment"] += " " + line


        print(f"Parsed outline with {len(outline['acts'])} Acts/Parts and {chapter_counter} total chapters.")
        if chapter_counter == 0:
            print("Warning: Outline parsing failed to identify chapters.")
            # Add a fallback structure?
            return {"acts": [{"title": "Act 1", "chapters": [{"number": 1, "title": "Fallback Chapter", "goal": "Start the story", "promises": [], "character_moment": "Introduce character", "estimated_scenes": 3}]}]}
        return outline


    def run(self, title, genre, themes, profile_summary, world_context, promises):
        print("Generating Novel Outline...")
        promises_list_str = ""
        for p_id, p_data in promises.items():
            promises_list_str += f"- Promise #{p_id}: {p_data['description']} (Importance: {p_data['importance']})\n"

        try:
            outline_text = self.chain.invoke({
                "title": title,
                "genre": genre,
                "themes": "\n- ".join(themes),
                "profile_summary": profile_summary,
                "world_context": world_context,
                "plot_promises_list": promises_list_str.strip(),
                "num_chapters": self.num_chapters,
                "end_act1": self.end_act1,
                "start_act2": self.start_act2,
                "end_act2": self.end_act2,
                "start_act3": self.start_act3
            })

            parsed_outline = self.parse_outline(outline_text)
            print("Outline Generated and Parsed.")
            return parsed_outline

        except Exception as e:
            print(f"Error generating outline: {e}")
            traceback.print_exc()
            return {"error": str(e)}

# --- Scene and Chapter Generation (Enhanced) ---

class SceneGeneratorChain:
    SCENE_WRITING_PROMPT = """
    You are {author}, writing a scene for the {genre} novel "{title}".
    This scene is part of **Chapter {chapter_number}: {chapter_title}**.
    The Goal for this chapter is: **{chapter_goal}**

    **Current Story Context:**
    - Protagonist: {character_name} ({character_description})
    - Current Mood/State: {character_mood}
    - Short-term Goal: {character_short_term_goal}
    - Relevant World Details: {world_context_snippet}
    - Core Themes to touch upon (subtly if possible): {themes}
    - Recent Key Events (Last 2-3 scenes):
    {recent_events_summary}

    **This Scene's Focus:**
    - Primary Plot Promise to Advance: **Promise #{promise_id}: {promise_description}**
    - Current Progress Level of this Promise: {promise_progress} / 4 (Higher = closer to payoff)
    - Expected Payoff for this Promise: {promise_payoff}
    - Key Character Moment for this Chapter: {chapter_character_moment}

    **Instructions for Writing:**
    1.  **Advance the Plot:** Ensure the scene significantly progresses the chosen plot promise ({promise_description}). Show, don't just tell, the progress.
    2.  **Character Consistency:** Write {character_name} authentically based on their profile, current mood, and goals. Show their internal thoughts, reactions, and decisions. Incorporate the chapter's character moment naturally if possible.
    3.  **Engaging Prose:** Write vivid, immersive prose in the style of {author}. Use sensory details, strong verbs, and varied sentence structure. Craft believable dialogue that reveals character and moves the plot.
    4.  **World Integration:** Weave in details from the world bible ({world_context_snippet}) naturally to enhance atmosphere and setting.
    5.  **Pacing and Tone:** Maintain a tone consistent with the genre ({genre}) and the current point in the story (consider the Act/Chapter number). Ensure the scene contributes to the chapter's overall goal ({chapter_goal}).
    6.  **Length:** Aim for a substantial scene, approximately **600-1000 words**.

    **Begin Scene:**
    ---
    """

    EVENT_SUMMARY_PROMPT = """
    Based on the scene you just wrote, provide two summaries:
    1.  **One-Sentence Event Summary:** Concisely state the most crucial plot development or character action that occurred in this scene related to the promise being advanced.
    2.  **Character State Change:** Briefly describe how the protagonist's mood, goal, or understanding might have shifted as a result of the scene's events. (e.g., "Became more determined," "Grew suspicious of X," "Felt a surge of hope," "Now aims to investigate Y").

    Scene Content Snippet (First/Last few lines):
    {scene_snippet}

    Plot Promise Advanced:
    {promise_description}

    Character Name:
    {character_name}

    ---
    1. Event Summary: [Your one-sentence summary]
    2. Character State Change: [Your description of the character's change]
    ---
    """

    def __init__(self, author_style):
        self.author_style = author_style
        # Use lower temp for summarizer, higher for creative writing
        self.scene_writer_llm = create_llm(temperature=0.75, top_p=0.95) # Creative writing
        self.event_summarizer_llm = create_llm(temperature=0.4, top_p=0.7) # Factual summary

        self.scene_writer_chain = LLMChain(
            llm=self.scene_writer_llm,
            prompt=PromptTemplate.from_template(self.SCENE_WRITING_PROMPT),
            verbose=True,
            output_parser=StrOutputParser()
        )
        self.event_summarizer_chain = LLMChain(
            llm=self.event_summarizer_llm,
            prompt=PromptTemplate.from_template(self.EVENT_SUMMARY_PROMPT),
            verbose=True,
            output_parser=StrOutputParser()
        )

    def write_scene(self, chapter_context, character_state, world_bible, themes,
                    recent_events, promise_to_advance):
        """Generates the scene content."""
        promise_details = promise_to_advance # It's already the dict
        
        # Prepare inputs
        profile_summary = f"Personality: {character_state.get('personality','N/A')}. Flaw: {character_state.get('major_flaw/internal_conflict','N/A')}."
        world_context_snippet = world_bible.get_relevant_context("Key Locations", "UFO Domination Context", "World Rules/Atmosphere")
        recent_events_summary = "\n".join(recent_events) if recent_events else "This is the first scene of the chapter."
        
        print(f"--- Writing Scene (Promise #{promise_details['id']}) ---")
        try:
            scene_content = self.scene_writer_chain.invoke({
                "author": self.author_style,
                "genre": chapter_context['genre'],
                "title": chapter_context['title'],
                "chapter_number": chapter_context['number'],
                "chapter_title": chapter_context['chapter_title'],
                "chapter_goal": chapter_context['goal'],
                "character_name": character_state.get('full_name', 'Protagonist'),
                "character_description": profile_summary,
                "character_mood": character_state.get('current_mood', 'Neutral'),
                "character_short_term_goal": character_state.get('short_term_goal', 'Unclear'),
                "world_context_snippet": world_context_snippet,
                "themes": ", ".join(themes),
                "recent_events_summary": recent_events_summary,
                "promise_id": promise_details['id'],
                "promise_description": promise_details['description'],
                "promise_progress": promise_details['progress'],
                "promise_payoff": promise_details['payoff'],
                "chapter_character_moment": chapter_context['character_moment']
            })
            
            scene_content = scene_content.strip()
            if not scene_content or len(scene_content) < 50: # Basic check for empty/failed generation
                print("Warning: Scene generation produced very short or empty content. Using fallback.")
                scene_content = f"[Scene generation failed for Promise #{promise_details['id']}. The narrative requires a scene here focusing on: {promise_details['description']} within the context of Chapter '{chapter_context['chapter_title']}'. Protagonist ({character_state.get('full_name')}) should react to recent events ({recent_events_summary.splitlines()[0] if recent_events else 'the start'}) and take a step towards {promise_details['payoff']}.]"

            return scene_content

        except Exception as e:
            print(f"Error writing scene: {e}")
            traceback.print_exc()
            return f"[Error occurred while writing scene for Promise #{promise_details['id']}: {e}]"

    def summarize_scene_impact(self, scene_content, promise_description, character_name):
        """Generates the event summary and character state change."""
        print("--- Summarizing Scene Impact ---")
        try:
             # Provide snippet for context
             snippet = scene_content[:200] + "...\n..." + scene_content[-200:]

             summary_text = self.event_summarizer_chain.invoke({
                "scene_snippet": snippet,
                "promise_description": promise_description,
                "character_name": character_name
            })

             # Parse the two parts
             event_summary = "Event summary generation failed."
             char_change = "Character state change unclear."
             lines = summary_text.strip().split('\n')
             for line in lines:
                  if line.lower().startswith("1. event summary:"):
                      event_summary = line.split(":", 1)[1].strip()
                  elif line.lower().startswith("2. character state change:"):
                      char_change = line.split(":", 1)[1].strip()

             # Basic fallback if parsing failed
             if event_summary.startswith("Event summary generation"):
                 event_summary = f"Progress made on '{promise_description[:50]}...'"
             if char_change.startswith("Character state change"):
                 # Attempt to infer from summary if possible (simple keyword check)
                 if any(kw in event_summary.lower() for kw in ["decided", "realized", "confronted", "learned"]):
                      char_change = f"Experienced a realization or made a decision related to the event."
                 else:
                      char_change = "Experienced the events of the scene."


             return event_summary, char_change

        except Exception as e:
            print(f"Error summarizing scene impact: {e}")
            traceback.print_exc()
            return f"Error summarizing event related to '{promise_description}'", "State change indeterminate due to error."


class ChapterGenerator:
     """Orchestrates the generation of scenes within a single chapter."""
     def __init__(self, chapter_data, novel_context, character_manager, world_bible, promise_manager, scene_generator):
          self.chapter_data = chapter_data # Outline info for this chapter
          self.novel_context = novel_context # Title, genre, author
          self.character_manager = character_manager
          self.world_bible = world_bible
          self.promise_manager = promise_manager
          self.scene_generator = scene_generator
          self.scenes = []
          self.chapter_summary = ""

     def generate(self):
          print(f"\n=== Generating Chapter {self.chapter_data['number']}: {self.chapter_data['title']} ===")
          print(f"Goal: {self.chapter_data['goal']}")
          print(f"Target Promises: {self.chapter_data['promises']}")
          print(f"Character Moment: {self.chapter_data['character_moment']}")
          
          num_scenes_to_generate = max(1, self.chapter_data.get('estimated_scenes', 2)) # Ensure at least 1 scene
          print(f"Estimated scenes: {num_scenes_to_generate}")

          # Get promise suggestions prioritized for this chapter
          suggested_promise_ids = [p[0] for p in self.promise_manager.suggest_promises_for_chapter(
              self.chapter_data.get('act', 1), # Need act number here if available
              self.chapter_data['number'],
              count=num_scenes_to_generate + 2 # Get a few extra options
          )]
          
          # Prioritize promises explicitly listed in the outline for this chapter
          target_promises = self.chapter_data.get('promises', [])
          promises_for_chapter = target_promises + [pid for pid in suggested_promise_ids if pid not in target_promises]
          
          if not promises_for_chapter:
               print("Warning: No specific promises identified for this chapter. Will attempt generic scene.")
               # We might need a fallback promise or mechanism here if none are suitable
               # For now, let's create a dummy placeholder task
               fallback_promise = {
                    'id': -1, 'description': f"General progression towards chapter goal: {self.chapter_data['goal']}",
                    'progress': 0, 'payoff': 'Move story forward'
               }
               promises_for_chapter.append(-1) # Use ID -1 for the fallback
               self.promise_manager.promises[-1] = fallback_promise # Temporarily add


          scene_count_in_chapter = 0
          while scene_count_in_chapter < num_scenes_to_generate and promises_for_chapter:
               promise_id_to_advance = promises_for_chapter.pop(0) # Take the highest priority one
               
               promise_details = self.promise_manager.get_promise_details(promise_id_to_advance)
               
               if not promise_details:
                    print(f"Warning: Could not find details for promise ID {promise_id_to_advance}. Skipping.")
                    continue
               
               # Check if promise is already complete - shouldn't happen with suggestion logic, but good failsafe
               if promise_details.get('complete', False):
                    print(f"Skipping Promise {promise_id_to_advance} as it's already complete.")
                    continue

               scene_count_in_chapter += 1
               current_scene_global_index = self.promise_manager.current_scene_index + 1 # Index for the scene about to be written

               # Prepare context for scene generator
               chapter_context = {
                    'number': self.chapter_data['number'],
                    'title': self.novel_context['title'],
                    'genre': self.novel_context['genre'],
                    'chapter_title': self.chapter_data['title'],
                    'goal': self.chapter_data['goal'],
                    'character_moment': self.chapter_data['character_moment']
               }
               character_state = self.character_manager.get_current_state()
               recent_events = self.promise_manager.get_story_summary(recent_count=3)

               # --- Generate Scene ---
               scene_content = self.scene_generator.write_scene(
                    chapter_context,
                    character_state,
                    self.world_bible,
                    self.novel_context['themes'],
                    recent_events,
                    promise_details
               )
               self.scenes.append({
                    "scene_number_global": current_scene_global_index,
                    "scene_number_chapter": scene_count_in_chapter,
                    "promise_advanced_id": promise_id_to_advance,
                    "promise_description": promise_details['description'],
                    "content": scene_content
               })
               
               # --- Summarize Scene Impact ---
               event_summary, char_change = self.scene_generator.summarize_scene_impact(
                    scene_content,
                    promise_details['description'],
                    character_state.get('full_name', 'Protagonist')
               )
               
               # --- Update Managers ---
               if promise_id_to_advance != -1: # Don't track progress for fallback promise
                   self.promise_manager.mark_promise_progressed(
                       promise_id_to_advance,
                       current_scene_global_index,
                       event_summary
                   )
               self.character_manager.update_state_from_summary(char_change)
               self.promise_manager.increment_scene_counter() # Increment global scene counter AFTER processing scene

               # Add a small delay to avoid overwhelming the LLM API if running locally fast
               time.sleep(1) # Adjust as needed

          # Generate chapter summary (optional, could use LLM)
          self.chapter_summary = f"This chapter focused on '{self.chapter_data['goal']}' and advanced promises related to {self.chapter_data['promises']}."
          print(f"=== Finished Chapter {self.chapter_data['number']} with {len(self.scenes)} scenes. ===")
          
          # Clean up fallback promise if added
          if -1 in self.promise_manager.promises:
               del self.promise_manager.promises[-1]

          return self.scenes, self.chapter_summary


# --- Character State Management ---
class CharacterManager:
    """Holds the detailed character profile and tracks dynamic state."""
    CHARACTER_UPDATE_PROMPT = """
    Given the character's current state and a summary of what just happened to them in a scene, update their CURRENT MOOD and SHORT-TERM GOAL.
    Be subtle and realistic. Moods shift, goals adapt based on events.

    Character Name: {name}
    Current Mood: {mood}
    Current Short-Term Goal: {goal}

    Scene Impact Summary: {impact_summary}

    Updated Mood: [Provide the new mood, e.g., Hopeful, Anxious, Determined, Resigned, Curious]
    Updated Short-Term Goal: [Provide the new immediate goal, e.g., Find shelter, Warn Person X, Investigate the strange noise, Get answers about Y]
    """
    def __init__(self, initial_profile):
        self.profile = initial_profile.copy() # Store the detailed dict
        self.llm = create_llm(temperature=0.5, top_p=0.8) # Focused update
        self.update_chain = LLMChain(
             llm=self.llm,
             prompt=PromptTemplate.from_template(self.CHARACTER_UPDATE_PROMPT),
             verbose=False, # Less noisy for state updates
             output_parser=StrOutputParser()
        )

    def get_profile_summary(self):
        """Returns a concise summary string for prompts."""
        return (f"Name: {self.profile.get('full_name', 'N/A')}. "
                f"Age: {self.profile.get('age_(approximate)', 'N/A')}. "
                f"Personality: {self.profile.get('personality', 'N/A')}. "
                f"Motivation: {self.profile.get('core_motivation/goal', 'N/A')}. "
                f"Flaw: {self.profile.get('major_flaw/internal_conflict', 'N/A')}.")

    def get_full_profile(self):
        return self.profile

    def get_current_state(self):
        """Returns the full profile including dynamic state."""
        return self.profile # Since mood/goal are added to the profile dict

    def update_state_from_summary(self, impact_summary):
        """Uses LLM to update mood and short-term goal based on scene impact."""
        print(f"Updating character state based on: {impact_summary}")
        try:
             update_text = self.update_chain.invoke({
                  "name": self.profile.get('full_name', 'Protagonist'),
                  "mood": self.profile.get('current_mood', 'Neutral'),
                  "goal": self.profile.get('short_term_goal', 'Assess'),
                  "impact_summary": impact_summary
             })

             new_mood = self.profile.get('current_mood', 'Neutral') # Default to old if parse fails
             new_goal = self.profile.get('short_term_goal', 'Assess')

             lines = update_text.strip().split('\n')
             for line in lines:
                 if line.lower().startswith("updated mood:"):
                     new_mood = line.split(":", 1)[1].strip()
                 elif line.lower().startswith("updated short-term goal:"):
                      new_goal = line.split(":", 1)[1].strip()
             
             # Only update if a plausible value was extracted
             if len(new_mood) > 2 and len(new_mood) < 30: # Basic sanity check
                 self.profile['current_mood'] = new_mood
             if len(new_goal) > 5 and len(new_goal) < 100:
                 self.profile['short_term_goal'] = new_goal

             print(f"State Updated - Mood: {self.profile['current_mood']}, Goal: {self.profile['short_term_goal']}")

        except Exception as e:
            print(f"Error updating character state: {e}")
            # Keep previous state on error

# --- Document Writing ---
class DocWriter:
    def __init__(self, output_folder=OUTPUT_FOLDER):
        self.output_folder = output_folder
        os.makedirs(self.output_folder, exist_ok=True)

    def _add_metadata_section(self, doc, title, data):
         doc.add_heading(title, level=1)
         if isinstance(data, dict):
             for key, value in data.items():
                  # Improve formatting for keys
                  heading = key.replace('_', ' ').replace('(approximate)', '').replace('(initial sketch)', '').replace('(relevant to subject)', '').title()
                  doc.add_paragraph(f"**{heading}:** {value}")
         elif isinstance(data, list):
              for item in data:
                   doc.add_paragraph(f"- {item}")
         else:
              doc.add_paragraph(str(data))
         doc.add_paragraph() # Add spacing

    def _add_outline_section(self, doc, outline):
         doc.add_heading("Novel Outline", level=1)
         if "acts" not in outline:
             doc.add_paragraph("Outline generation failed or was not structured correctly.")
             return

         for act in outline['acts']:
              doc.add_heading(act.get('title', 'Untitled Act'), level=2)
              if not act.get('chapters'):
                  doc.add_paragraph("No chapters defined for this act.")
                  continue
              for chapter in act['chapters']:
                  doc.add_heading(f"Chapter {chapter.get('number', '?')}: {chapter.get('title', 'Untitled')}", level=3)
                  doc.add_paragraph(f"**Goal:** {chapter.get('goal', 'N/A')}")
                  doc.add_paragraph(f"**Promises:** {chapter.get('promises', 'N/A')}")
                  doc.add_paragraph(f"**Character Moment:** {chapter.get('character_moment', 'N/A')}")
                  doc.add_paragraph(f"**Estimated Scenes:** {chapter.get('estimated_scenes', 'N/A')}")
                  doc.add_paragraph() # Spacing between chapters
         doc.add_paragraph()


    def write_novel(self, novel_title, novel_context, character_profile, themes, world_bible, outline, chapters_content):
        """Writes the complete novel and its metadata to a DOCX file."""
        doc = docx.Document()
        
        # --- Preamble ---
        doc.add_heading(novel_title, level=0)
        self._add_metadata_section(doc, "Genre", novel_context.get('genre', 'N/A'))
        self._add_metadata_section(doc, "Author Style Influence", novel_context.get('author', 'N/A'))
        self._add_metadata_section(doc, "Core Themes", themes)
        self._add_metadata_section(doc, "Character Profile", character_profile)
        self._add_metadata_section(doc, "World Bible Snippets", world_bible) # Show the generated bible data
        self._add_outline_section(doc, outline)

        # --- Novel Content ---
        doc.add_page_break()
        doc.add_heading("Novel Content", level=1)

        current_chapter_num = 0
        for chapter in chapters_content:
             # Add chapter heading only once per chapter
             if chapter['number'] != current_chapter_num:
                  current_chapter_num = chapter['number']
                  doc.add_heading(f"Chapter {chapter['number']}: {chapter['title']}", level=2)
                  # Add chapter summary if available (though we didn't generate one in detail yet)
                  # doc.add_paragraph(f"*{chapter.get('summary', '')}*")
                  # doc.add_paragraph()

             # Add scenes for this chapter
             for scene in chapter['scenes']:
                  # Optional: Add scene number or marker?
                  # doc.add_heading(f"Scene {scene['scene_number_chapter']} (Global: {scene['scene_number_global']})", level=3)
                  doc.add_paragraph(scene['content'])
                  doc.add_paragraph() # Add space between scenes
             doc.add_page_break() # Page break after each chapter


        # --- Save Document ---
        safe_basename = sanitize_filename(novel_title)
        filename = safe_basename + '.docx'
        output_path = os.path.join(self.output_folder, filename)

        retries = 3
        for i in range(retries):
            try:
                doc.save(output_path)
                print(f"--- Novel saved successfully as: {output_path} ---")
                return True # Success
            except PermissionError:
                print(f"\nError: Permission denied trying to save '{output_path}'.")
                print("Please check if the file is open or if you have write permissions.")
                if i < retries - 1:
                     alt_filename = f"{safe_basename}_({i+1}).docx"
                     output_path = os.path.join(self.output_folder, alt_filename)
                     print(f"Retrying with filename: {alt_filename}")
                else:
                     print("Giving up after multiple permission errors.")
                     return False # Failed
            except Exception as e:
                print(f"\nError saving document to {output_path}: {e}")
                traceback.print_exc()
                return False # Failed

        return False # Should not be reached unless loop logic fails

# --- Main Orchestration ---

def main():
    print("--- Initializing Advanced Novel Generation System ---")
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    os.makedirs(RESUME_FOLDER, exist_ok=True)
    print(f"Input Resume Folder: {RESUME_FOLDER}")
    print(f"Output Novel Folder: {OUTPUT_FOLDER}")

    # --- Configuration ---
    # !! Ensure 'Thalen.pdf' exists in the './docs' folder !!
    resume_filename = 'Thalen.pdf'
    # More complex subject leveraging resume details
    subject = ("A slice-of-life story about Thalen, a young architect trying to rebuild both structures and "
               "a sense of normalcy in Pune, Maharashtra, years after a disruptive UFO event reshaped the world. "
               "The story explores themes of adaptation, the meaning of 'home' in a changed landscape, and the tension "
               "between human creativity (architecture) and overwhelming, poorly understood alien influence. "
               "Use locations like Pune and references to architectural studies (like CEPT) to ground the story.")
    author_style = 'Ted Chiang' # Example: Known for thoughtful Sci-Fi
    genre = 'Slice of Life, Soft Sci-Fi, Architectural Fiction, Post-Contact'
    target_chapters = 20 # Aim for a longer novel

    print("\n--- Novel Parameters ---")
    print(f"Resume File: {resume_filename}")
    print(f"Subject: {subject}")
    print(f"Author Style: {author_style}")
    print(f"Genre: {genre}")
    print(f"Target Chapters: {target_chapters}")

    # --- Instantiate Core Components ---
    try:
        print("\n--- Step 1: Loading Resume ---")
        resume_text = load_resume_text(resume_filename)
        if not resume_text:
             print("ERROR: Resume text could not be loaded. Cannot proceed.")
             return

        print("\n--- Step 2: Generating Character Profile ---")
        char_generator = CharacterGeneratorChain()
        character_profile = char_generator.run(resume_text, subject, genre)
        if "error" in character_profile:
            print(f"ERROR: Character profile generation failed: {character_profile['error']}")
            return
        character_manager = CharacterManager(character_profile)
        print("Character Profile Snippet:")
        print(character_manager.get_profile_summary())

        print("\n--- Step 3: Generating Core Themes ---")
        theme_generator = ThemeGeneratorChain()
        themes = theme_generator.run(subject, genre, character_profile)
        print(f"Generated Themes: {themes}")

        # Temporary Title Chain (can be enhanced later)
        print("\n--- Step 4: Generating Title ---")
        title_llm = create_llm(temperature=0.8)
        title_prompt = PromptTemplate.from_template("Generate a compelling and unique novel title based on: Subject '{subject}', Genre '{genre}', Themes '{themes}'. Title:")
        title_chain_simple = LLMChain(llm=title_llm, prompt=title_prompt, output_parser=StrOutputParser())
        novel_title = title_chain_simple.invoke({"subject": subject, "genre": genre, "themes": ", ".join(themes)}).strip().strip('"')
        novel_title = novel_title if novel_title else f"Untitled_{genre.split(',')[0]}_Novel"
        print(f"Generated Title: '{novel_title}'")

        print("\n--- Step 5: Generating World Bible ---")
        world_bible_manager = WorldBibleManager(subject, genre, resume_text, character_profile)
        world_bible_data = world_bible_manager.generate_initial_bible()
        if "error" in world_bible_data:
             print(f"ERROR: World Bible generation failed: {world_bible_data['error']}")
             # Maybe allow proceeding with defaults? For now, exit.
             return
        print("World Bible Snippets Generated:")
        print(world_bible_manager.get_relevant_context("Key Locations", "UFO Domination Context"))

        print("\n--- Step 6: Generating Initial Plot Promises ---")
        promise_generator = PromiseGeneratorChain()
        profile_summary_for_prompt = character_manager.get_profile_summary()
        world_context_for_prompt = world_bible_manager.get_relevant_context("Key Locations", "UFO Domination Context", "World Rules/Atmosphere")
        initial_promises = promise_generator.run(novel_title, genre, themes, profile_summary_for_prompt, world_context_for_prompt, subject)
        if not initial_promises:
            print("ERROR: Failed to generate initial plot promises. Cannot proceed.")
            return
        promise_manager = PromiseManager(initial_promises)
        print("Initial Promises:")
        for p_summary in promise_manager.get_active_promises_summary():
            print(f"- {p_summary}")

        print("\n--- Step 7: Generating Novel Outline ---")
        outline_generator = OutlineGeneratorChain(num_chapters=target_chapters)
        outline = outline_generator.run(novel_title, genre, themes, profile_summary_for_prompt, world_context_for_prompt, promise_manager.promises)
        if "error" in outline or not outline.get("acts"):
             print(f"ERROR: Outline generation failed: {outline.get('error', 'Unknown structure error')}")
             return
        # Basic sanity check on outline promises
        all_outline_pids = set()
        for act in outline.get('acts',[]):
            for chap in act.get('chapters',[]):
                all_outline_pids.update(chap.get('promises',[]))
        print(f"Outline references {len(all_outline_pids)} unique promise IDs.")
        # Assign promises to outline structure in PromiseManager (optional step, but good for tracking)
        # This requires parsing the outline again or doing it during generation/parsing.
        # For simplicity now, we rely on the suggestion logic during chapter generation.

        # --- Instantiate Scene/Chapter Generators ---
        scene_generator = SceneGeneratorChain(author_style=author_style)
        doc_writer = DocWriter(output_folder=OUTPUT_FOLDER)

        # --- Novel Generation Loop ---
        print("\n--- Step 8: Generating Novel Content (Chapter by Chapter) ---")
        full_novel_content = [] # List to hold generated chapter data [{number, title, scenes:[{...}]}]
        novel_context_for_chapters = {
             'title': novel_title,
             'genre': genre,
             'author': author_style,
             'themes': themes
        }

        # Iterate through the generated outline
        if "acts" in outline:
            chapter_count = 0
            for act in outline['acts']:
                print(f"\n-- Starting {act.get('title', 'Act')} --")
                if "chapters" in act:
                    for chapter_data in act['chapters']:
                        chapter_count += 1
                        # Add act info to chapter data if needed
                        chapter_data['act'] = act.get('title', 'Unknown Act')
                        
                        # Create and run chapter generator
                        chapter_gen = ChapterGenerator(
                            chapter_data,
                            novel_context_for_chapters,
                            character_manager,
                            world_bible_manager,
                            promise_manager,
                            scene_generator
                        )
                        scenes, chapter_summary = chapter_gen.generate()
                        
                        # Store chapter results
                        full_novel_content.append({
                             "number": chapter_data['number'],
                             "title": chapter_data['title'],
                             "summary": chapter_summary, # Store summary if generated
                             "scenes": scenes
                        })
                else:
                    print(f"Warning: No chapters found in {act.get('title', 'this act')}")
        else:
             print("ERROR: Outline structure invalid, cannot generate chapters.")
             return

        print(f"\n--- Novel Generation Complete: {chapter_count} chapters generated with {promise_manager.current_scene_index} total scenes. ---")
        print(f"Completed Promises: {len(promise_manager.completed_promises)}")
        print(f"Remaining Active Promises: {len(promise_manager.promises)}")


        # --- Save to Document ---
        print("\n--- Step 9: Saving Novel to Document ---")
        doc_writer.write_novel(
            novel_title,
            novel_context_for_chapters,
            character_manager.get_full_profile(),
            themes,
            world_bible_data,
            outline,
            full_novel_content
        )

    except FileNotFoundError as e:
        print(f"\nCRITICAL ERROR: Required file not found.")
        print(e)
        print(f"Please ensure '{resume_filename}' exists in the '{RESUME_FOLDER}' directory.")
    except ImportError as e:
         print(f"\nCRITICAL ERROR: Missing required library.")
         print(e)
         print("Please ensure all dependencies are installed (pip install -r requirements.txt if available).")
         print("Key dependencies include: langchain, langchain-ollama, langchain-community, python-dotenv, pypdf, python-docx")
    except Exception as e:
        print(f"\n--- An Unexpected Error Occurred ---")
        print(f"Error Type: {type(e).__name__}")
        print(f"Error Details: {e}")
        print("\n--- Traceback ---")
        traceback.print_exc()
        print("-----------------")
        print("Novel generation process halted due to error.")

    finally:
        print("\n--- Novel Generation Process Finished ---")


if __name__ == "__main__":
    main()
