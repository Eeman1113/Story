# app_v2.py
from dotenv import load_dotenv
import os
import docx
import traceback # For printing detailed errors
import time # To avoid overwhelming the LLM API if needed
import re # For robust parsing

# --- Langchain Imports ---
from langchain_ollama import OllamaLLM
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import PyPDFLoader
# --- End Imports ---

# Load environment variables (optional, but good practice)
load_dotenv()

# --- Constants ---
# Model Selection: Choose the most capable model available via Ollama
# Examples: "gemma2:latest", "llama3:70b", "mistral:latest"
# Ensure the model is actually pulled and running in Ollama.
DEFAULT_MODEL = os.getenv("OLLAMA_MODEL", "gemma3:4b")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OUTPUT_FOLDER = './docs'
# Optional: Delay between LLM calls. Increase if hitting rate limits or instability.
LLM_CALL_DELAY_SECONDS = 0.5
# Placeholder for the resume file - MAKE SURE THIS FILE EXISTS IN OUTPUT_FOLDER
# Or adjust the path logic as needed.
DEFAULT_RESUME_FILENAME = 'divi_1.pdf' # Example filename

# --- LLM Initialization Function ---
def create_llm(temperature=0.7, top_p=0.9, top_k=40):
    """
    Create and return an OllamaLLM instance with specified parameters.
    Allows tuning for different generation tasks.
    """
    print(f"--- Connecting to Ollama at: {OLLAMA_BASE_URL} with Model: {DEFAULT_MODEL} ---")
    try:
        llm = OllamaLLM(
            model=DEFAULT_MODEL,
            temperature=temperature,
            top_p=top_p, # Controls nucleus sampling
            top_k=top_k, # Controls top-k sampling
            base_url=OLLAMA_BASE_URL,
            request_timeout=180.0, # Increased timeout for potentially longer generations
            # Add other Ollama parameters if needed (e.g., num_ctx, stop sequences)
            # num_predict=512, # Example: Limit max tokens per call if needed
        )
        # Optional: Simple invoke test to check connection early
        # llm.invoke("Test connection.")
        print(f"--- LLM Instance Created (Temp: {temperature}, Top P: {top_p}, Top K: {top_k}) ---")
        return llm
    except Exception as e:
        print(f"--- FATAL ERROR: Could not create OllamaLLM instance ---")
        print(f"Error: {e}")
        print("Ensure Ollama service is running, the model is pulled, and accessible at the specified base URL.")
        print(f"Attempted Base URL: {OLLAMA_BASE_URL}")
        print(f"Attempted Model: {DEFAULT_MODEL}")
        traceback.print_exc()
        raise # Reraise the exception to stop execution

# --- Chain Classes ---

class MainCharacterChain:
    # Enhanced prompt for deeper character understanding
    PROMPT = """
    Analyze the following resume text to create a rich, multi-faceted character profile suitable for a literary narrative. Go beyond surface-level skills.

    Infer and extrapolate:
    1.  **Name:** (If mentioned, otherwise suggest one based on context or leave blank)
    2.  **Core Identity & Demeanor:** A brief summary (2-3 sentences) of their apparent personality, job/role, and how they present themselves.
    3.  **Key Skills/Strengths:** List 3-5 notable skills or positive attributes relevant to potential plot points.
    4.  **Potential Motivations:** What might drive this character? What deep-seated desires or goals could be hinted at? (Infer 2-3)
    5.  **Potential Flaws/Weaknesses:** What vulnerabilities, biases, or negative traits might they possess? (Infer 2-3)
    6.  **Internal Conflicts:** What inner struggles or contradictions could define their character arc? (Infer 1-2)
    7.  **Secrets/Hidden Depths:** What might this character be hiding from others, or even themselves? (Infer 1-2)
    8.  **Narrative Potential:** Briefly suggest (1-2 sentences) how these traits could fuel a compelling story in the {genre} genre.

    Resume Text:
    {text}

    Detailed Character Profile:"""

    def __init__(self):
        # Slightly lower temperature for extraction, but allow some inference
        self.llm = create_llm(temperature=0.6, top_p=0.85)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def load_resume(self, file_name):
        """Loads text content from a PDF file located in OUTPUT_FOLDER."""
        file_path = os.path.join(OUTPUT_FOLDER, file_name)
        if not os.path.exists(file_path):
             raise FileNotFoundError(f"Resume file not found at: {file_path}")
        try:
            print(f"Loading PDF from: {file_path}")
            loader = PyPDFLoader(file_path)
            # load_and_split can sometimes be better, but load is fine for moderate PDFs
            docs = loader.load()
            if not docs:
                print(f"Warning: PyPDFLoader didn't extract any documents/pages from {file_name}.")
                return None
            print(f"Successfully loaded {len(docs)} page(s) from PDF.")
            full_text = '\n\n'.join([doc.page_content for doc in docs if doc.page_content])
            # Basic text cleaning (optional)
            full_text = re.sub(r'\s+', ' ', full_text).strip()
            return full_text
        except Exception as e:
             print(f"Error loading or processing PDF {file_path}: {e}")
             traceback.print_exc()
             raise

    def run(self, file_name, genre):
        """Loads resume and generates the detailed character profile."""
        try:
            resume_text = self.load_resume(file_name)
            if not resume_text or not resume_text.strip():
                 print("Could not load or resume content is empty.")
                 return "Error: Could not generate profile due to missing or empty resume content."

            print("Invoking MainCharacterChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({"text": resume_text, "genre": genre})
            profile = result.get('text', "Error: Profile generation failed.").strip()

            # More robust check for valid profile
            if not profile or profile.startswith("Error:") or len(profile) < 100: # Increased minimum length
                print(f"Warning: Generated profile seems invalid or too short: '{profile}'")
                return f"Error: Failed to generate a meaningful profile. LLM Output: {profile}"
            return profile

        except Exception as e:
            print(f"An error occurred in MainCharacterChain.run: {e}")
            traceback.print_exc()
            return f"Error generating profile: {e}"


class SettingChain:
    PROMPT = """
    You are a world-building assistant specializing in atmospheric settings for novels.
    Based on the novel's subject, genre, and main character profile, generate a concise description of the primary setting(s).

    Focus on:
    1.  **Key Locations:** Identify 2-4 significant places where the story unfolds.
    2.  **Time Period/Atmosphere:** Describe the general era, mood, and sensory feeling (e.g., oppressive, nostalgic, futuristic, decaying, magical).
    3.  **Sensory Details:** Suggest 3-5 recurring sensory elements (specific sights, sounds, smells, textures) that define the world's feel.
    4.  **Relation to Character/Plot:** Briefly explain (1-2 sentences) how the setting reflects or influences the main character and potential plot events.

    Novel Subject: {subject}
    Genre(s): {genre}
    Main Character Profile: {profile}

    Setting Description:"""

    def __init__(self):
        # Moderate temperature for creative but focused description
        self.llm = create_llm(temperature=0.7, top_p=0.9)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, profile):
        """Generates the setting description."""
        try:
            print("Invoking SettingChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "profile": profile
            })
            setting = result.get('text', "Error: Setting generation failed.").strip()
            if not setting or setting.startswith("Error:") or len(setting) < 50:
                print(f"Warning: Generated setting seems invalid or too short: '{setting}'")
                return f"Error: Failed to generate a meaningful setting description. LLM Output: {setting}"
            return setting
        except Exception as e:
            print(f"An error occurred in SettingChain.run: {e}")
            traceback.print_exc()
            return f"Error generating setting: {e}"


class ThemeChain:
    PROMPT = """
    You are a literary analyst identifying core themes.
    Based on the novel's subject, genre, character profile, and setting, identify 2-4 central themes that the story explores.
    Themes should be abstract concepts (e.g., "Loss and Memory," "Identity vs. Society," "The Nature of Reality," "Redemption," "Man vs. Nature").
    Provide a brief (1-sentence) explanation for each theme, linking it to the provided context.

    Novel Subject: {subject}
    Genre(s): {genre}
    Main Character Profile: {profile}
    Setting Description: {setting}

    Identified Themes:
    1.  [Theme 1]: [1-sentence explanation]
    2.  [Theme 2]: [1-sentence explanation]
    3.  [Theme 3 (Optional)]:[1-sentence explanation]
    4.  [Theme 4 (Optional)]:[1-sentence explanation]
    """

    def __init__(self):
        # Lower temperature for analytical task
        self.llm = create_llm(temperature=0.5, top_p=0.8)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse_themes(self, response):
        """Parses the numbered list of themes and explanations."""
        themes = {}
        # Regex to capture "X. [Theme Name]: [Explanation]"
        pattern = re.compile(r"^\s*\d+\.\s*\[?([^:\]]+)\]?:\s*(.*)", re.MULTILINE)
        matches = pattern.findall(response)
        if not matches:
             # Fallback: Try splitting by lines if regex fails
             lines = [line.strip() for line in response.strip().split('\n') if ':' in line]
             for line in lines:
                 parts = line.split(':', 1)
                 # Clean up theme name (remove potential number/bullet)
                 theme_name = re.sub(r"^\s*\d+\.\s*", "", parts[0]).strip()
                 explanation = parts[1].strip()
                 if theme_name and explanation:
                     themes[theme_name] = explanation
        else:
             for match in matches:
                 theme_name = match[0].strip()
                 explanation = match[1].strip()
                 if theme_name and explanation:
                     themes[theme_name] = explanation

        if not themes:
            print(f"Warning: Could not parse themes from response:\n---\n{response}\n---")
            return {"Error": "Theme parsing failed."}
        return themes # Returns dict { 'Theme Name': 'Explanation' }

    def run(self, subject, genre, profile, setting):
        """Generates and parses the core themes."""
        try:
            print("Invoking ThemeChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "profile": profile,
                "setting": setting
            })
            raw_themes = result.get('text', "").strip()
            return self.parse_themes(raw_themes)
        except Exception as e:
            print(f"An error occurred in ThemeChain.run: {e}")
            traceback.print_exc()
            return {"Error": f"Theme generation failed: {e}"}


class TitleChain:
    # Prompt now includes themes and setting for better context
    PROMPT = """
    You are a creative book title generator known for evocative and genre-appropriate titles.
    Generate ONE compelling novel title based on the provided details.
    The title must be highly consistent with the genre(s), author's style, subject, main character, setting, and core themes.
    It should be intriguing and memorable.

    Subject: {subject}
    Genre(s): {genre}
    Author's Style Inspiration: {author}
    Main Character Profile: {profile}
    Setting Description: {setting}
    Core Themes: {themes}

    Return ONLY the generated title itself, without any quotation marks, labels (like "Title:"), or explanatory text.

    Novel Title:"""

    def __init__(self):
        # Higher temperature for creative title generation
        self.llm = create_llm(temperature=0.85, top_p=0.95)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile, setting, themes_str):
        """Generates the novel title."""
        try:
            print("Invoking TitleChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile,
                "setting": setting,
                "themes": themes_str
            })
            title = result.get('text', "Untitled Novel").strip()
            # Clean the output string thoroughly
            title = re.sub(r'^(Title:|Novel Title:)\s*', '', title, flags=re.IGNORECASE)
            title = title.strip('"\'')

            if not title or len(title) < 3 or title == "Untitled Novel":
                 print(f"Warning: Generated title seems invalid: '{title}'. Using placeholder.")
                 # Generate a slightly more descriptive placeholder
                 genre_tag = genre.split('/')[0].split(',')[0].strip()
                 return f"Placeholder - {genre_tag} Story about {profile.split('.')[0]}" # Use first sentence of profile
            return title
        except Exception as e:
            print(f"An error occurred in TitleChain.run: {e}")
            traceback.print_exc()
            return "Error Generating Title"


class PlotChain:
    # Enhanced prompt incorporating setting and themes
    PROMPT = """
    You are a master storyteller and plot architect, crafting narratives in the style of {author}.
    Develop a detailed, multi-act plot outline (e.g., Act I, Act II, Act III or Beginning, Middle, End) for a novel.
    The plot must be engaging, coherent, thematically resonant, and build towards a satisfying climax and resolution appropriate for the {genre} genre.

    Integrate the following elements seamlessly:
    - Main Character Arc: Show how the character (profile below) changes or is challenged, driven by their motivations/flaws.
    - Core Themes: Weave the themes ({themes}) into the events and character journey.
    - Setting Influence: Show how the setting ({setting}) impacts the mood, obstacles, or events.
    - Compelling Attributes: Incorporate story elements like: {features}.
    - Supporting Characters: Introduce necessary allies, antagonists, or foils.
    - Conflict: Include both internal (character-based) and external (plot-based) conflicts.
    - Turning Points: Define key moments that shift the narrative direction.

    Novel Subject: {subject}
    Genre(s): {genre}
    Novel Title: "{title}"
    Main Character Profile: {profile}
    Setting Description: {setting}
    Core Themes: {themes}

    Detailed Plot Outline:"""

    HELPER_PROMPT = """
    Generate a comma-separated list of 5-7 diverse and compelling story attributes or narrative devices suitable for the specified genre(s) and author style. Avoid generic terms. Be specific and evocative.
    Genre(s): {genre}
    Author's Style Inspiration: {author}

    Examples for different genres/styles:
    Psychological/Ocean Vuong: Fragmented memories, Lyrical prose focus, Visceral sensory details, Intergenerational trauma echoes, Non-linear emotional arcs, Symbolism in mundane objects.
    Sci-Fi/Ted Chiang: Thought-provoking central conceit, Exploration of philosophical implications, Rigorous internal logic, Emotional core within intellectual framework, Sense of wonder or unease.
    Fantasy/N.K. Jemisin: Intricate world-building reveals, High-stakes societal conflict, Complex moral ambiguity, Unique magic systems with consequences, Characters challenging power structures.
    Thriller/Gillian Flynn: Unreliable narrator, Toxic relationships, Dark secrets unraveling, Sharp social commentary, Unexpected plot twists rooted in character psychology.

    List of Attributes (comma-separated):"""

    def __init__(self):
        # Balanced temperature for structured creativity
        self.llm = create_llm(temperature=0.75, top_p=0.9)
        # Main chain for plot generation
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )
        # Helper chain to generate dynamic features
        self.helper_chain = LLMChain(
            llm=self.llm, # Use the same LLM instance
            prompt=PromptTemplate.from_template(self.HELPER_PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile, title, setting, themes_str):
        """Generates the detailed novel plot outline."""
        try:
            # Generate dynamic features
            print(f"Generating plot features for genre '{genre}' and author style '{author}'...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            features_result = self.helper_chain.invoke({"genre": genre, "author": author})
            features = features_result.get('text', "Compelling conflict, Character depth, Unexpected twists").strip()
            print(f"Generated plot features: {features}")

            # Generate the main plot outline
            print(f"Generating main plot outline for title: {title}")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            plot_result = self.chain.invoke({
                "features": features,
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile,
                "title": title,
                "setting": setting,
                "themes": themes_str
            })
            plot = plot_result.get('text', "Error: Plot generation failed.").strip()

            # Check for valid plot (e.g., minimum length, structure)
            if not plot or plot.startswith("Error:") or len(plot) < 200: # Increased minimum length
                 print(f"Warning: Generated plot seems invalid or too short: '{plot}'")
                 return f"Error: Failed to generate a detailed plot. LLM Output: {plot}"
            return plot
        except Exception as e:
            print(f"An error occurred in PlotChain.run: {e}")
            traceback.print_exc()
            return f"Error generating plot: {e}"


class ChaptersChain:
    # Prompt updated to request consistency with themes/setting
    PROMPT = """
    You are a meticulous book editor outlining chapter structure in the style of {author}.
    Based on the detailed plot outline, generate a list of chapter titles AND brief, one-sentence descriptions capturing the core focus or turning point of each chapter.
    Aim for a realistic number of chapters (e.g., 15-35) that logically progress through the plot.
    Include a Prologue and/or Epilogue ONLY if appropriate for the story structure, genre, and author's style.

    Ensure strict consistency with:
    - Novel Title: "{title}"
    - Genre(s): {genre}
    - Author's Style: {author}
    - Main Character Arc: {profile}
    - Setting: {setting}
    - Core Themes: {themes}
    - Detailed Plot Outline (provided below)

    Use this EXACT format for each entry, with NO blank lines between entries:
    [Chapter Type] [Number (if not Prologue/Epilogue)]: [Concise, evocative one-sentence Description]

    Example Format:
    Prologue: Whispers of the past in the decaying manor.
    Chapter 1: A mundane routine shattered by a cryptic message.
    Chapter 2: The first venture into the forbidden zone, guided by fear.
    ...
    Chapter N: The confrontation forces an impossible choice, echoing the core theme of sacrifice.
    Epilogue: Lingering questions in the changed landscape.

    Detailed Plot Outline:
    <PLOT_START>
    {plot}
    <PLOT_END>

    Chapters List (Strict Format Adherence Required):"""

    def __init__(self):
        # Slightly lower temperature for structured output, higher K for variety
        self.llm = create_llm(temperature=0.65, top_p=0.9, top_k=50)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse(self, response):
        """Parses the LLM response string into an ordered dictionary of chapters."""
        chapters = {} # Using dict to store { 'Title': 'Description' }
        lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
        print(f"\nRaw Chapter Response Lines ({len(lines)}):")
        # for i, line in enumerate(lines): print(f"  Line {i}: {line}") # Uncomment for deep debug

        parsed_count = 0
        # Regex to capture variations like "Chapter 1: Desc", "Prologue: Desc", "Epilogue: Desc"
        # Allows optional space after number, handles Prologue/Epilogue case-insensitively
        pattern = re.compile(r"^(Prologue|Epilogue|Chapter\s*(\d+))\s*:\s*(.+)$", re.IGNORECASE | re.MULTILINE)

        # First pass using regex for well-formatted lines
        processed_lines = set()
        matches = pattern.findall(response) # Search the whole response block
        for match in matches:
            full_title = match[0].strip() # e.g., "Prologue", "Chapter 1"
            # Reconstruct title properly if it's a chapter number
            if match[1]: # If chapter number was captured
                full_title = f"Chapter {match[1]}"
            description = match[2].strip()
            if description:
                chapters[full_title] = description
                parsed_count += 1
                # Mark the raw line containing this match as processed (approximate)
                # This is tricky if descriptions contain newlines, but helps avoid double parsing
                raw_line_approx = f"{match[0]}:{description}"
                for i, line in enumerate(lines):
                    if raw_line_approx in line: # Check if the core part exists
                         processed_lines.add(i)
                         break # Mark only the first occurrence
            else:
                print(f"  Warning: Skipping match - Empty description: '{match[0]}'")


        # Second pass for lines potentially missed by regex (less strict)
        print(f"Attempting fallback parsing for remaining lines...")
        for i, line in enumerate(lines):
            if i in processed_lines: continue # Skip already processed lines

            if ':' in line:
                parts = line.split(':', 1)
                potential_title = parts[0].strip()
                potential_desc = parts[1].strip()

                # Basic check for plausible chapter titles
                is_prologue = "prologue" in potential_title.lower()
                is_epilogue = "epilogue" in potential_title.lower()
                is_chapter = "chapter" in potential_title.lower() and any(char.isdigit() for char in potential_title)

                # Avoid adding if already parsed via regex (handles slight format variations)
                if potential_title not in chapters and (is_prologue or is_epilogue or is_chapter):
                    if potential_desc:
                        print(f"  Fallback Parsed Line {i}: '{potential_title}' -> '{potential_desc}'") # Debug fallback
                        chapters[potential_title] = potential_desc
                        parsed_count += 1
                    else:
                        print(f"  Warning: Skipping line {i} (fallback) - Empty description: '{line}'")
                # else:
                #      print(f"  Info: Skipping line {i} (fallback) - Doesn't look like chapter title or already parsed: '{line}'")
            # else:
            #     print(f"  Info: Skipping line {i} (fallback) - No colon found: '{line}'")


        if not chapters:
             print(f"CRITICAL WARNING: Could not parse *any* chapters adhering to the expected format.")
             print(f"LLM Raw Response was:\n---\n{response}\n---")
             return {} # Return empty dict to signal failure

        print(f"Successfully parsed {parsed_count} chapters.")
        # Return sorted chapters directly from parse function
        return dict(sort_chapters(chapters)) # Return sorted dict {title: description}

    def run(self, subject, genre, author, profile, title, plot, setting, themes_str):
        """Generates and parses the chapter list."""
        try:
            print("Invoking ChaptersChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            response_result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile,
                "title": title,
                "plot": plot,
                "setting": setting,
                "themes": themes_str
            })
            raw_response = response_result.get('text', "")
            return self.parse(raw_response)
        except Exception as e:
            print(f"An error occurred in ChaptersChain.run: {e}")
            traceback.print_exc()
            return {} # Return empty dict on error


class EventChain:
    # Prompt now includes profile and themes for better context
    PROMPT = """
    You are a narrative strategist breaking down a chapter into key scenes or events, following the style of {author}.
    Based on the overall plot, the chapter's details, character profile, and themes, generate an ordered list of 3-7 key plot events or scenes that MUST occur within this chapter.
    Focus on actions, decisions, revelations, significant interactions, or internal character moments that advance the plot AND reflect the character's journey or the novel's themes.
    Ensure events flow logically from the chapter summary and overall plot.

    Overall Plot Summary:
    <PLOT_SUMMARY>
    {plot}
    </PLOT_SUMMARY>

    Main Character Profile:
    <CHARACTER>
    {profile}
    </CHARACTER>

    Core Themes: {themes}

    Current Chapter Title: "{chapter_title}"
    Current Chapter Summary: {chapter_summary}

    Return ONLY the numbered list of events, one event per line. Start numbering from 1. Be concise but evocative.

    Example (Psychological Fiction):
    1. Divi awakens to the oppressive silence, the weight of unspoken words heavy in the air.
    2. A fragmented memory surfaces: a childhood promise broken.
    3. She attempts to write, but the words feel alien, mirroring her internal disconnect.
    4. A sound from outside triggers a cascade of anxious thoughts, linking to the theme of vulnerability.
    5. The chapter ends with her staring at her reflection, questioning the face looking back.

    Numbered Event List for Chapter "{chapter_title}":
    """

    def __init__(self):
        # Moderate temperature for focused event generation
        self.llm = create_llm(temperature=0.7, top_p=0.9)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse_events(self, response):
        """Parses the numbered list response into a list of strings."""
        events = []
        # Regex to find lines starting with a number, optional punctuation, and then text
        pattern = re.compile(r"^\s*\d+[\.\)\s]+(.*)", re.MULTILINE)
        matches = pattern.findall(response)

        if matches:
            for event_text in matches:
                if event_text.strip():
                    events.append(event_text.strip())
        else:
            # Fallback: split by lines and try to clean up
            lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
            for line in lines:
                 # Attempt to remove leading numbers/bullets if regex failed
                 cleaned_line = re.sub(r"^\s*\d+[\.\)\s]*", "", line).strip()
                 if cleaned_line:
                     events.append(cleaned_line)

        if not events:
             print(f"  Warning: Could not parse any numbered events from response:\n---\n{response}\n---")
             # Fallback: return the raw response lines (better than nothing)
             raw_lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
             return raw_lines if raw_lines else ["Placeholder event due to parsing failure."]
        return events

    def run(self, plot, profile, themes_str, chapter_title, chapter_summary, author):
        """Generates and parses the event list for a single chapter."""
        try:
            print(f"Invoking EventChain for: {chapter_title}")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "plot": plot,
                "profile": profile,
                "themes": themes_str,
                "chapter_title": chapter_title,
                "chapter_summary": chapter_summary,
                "author": author # Pass author style for event tone
            })
            raw_events = result.get('text', "").strip()
            return self.parse_events(raw_events)
        except Exception as e:
            print(f"  Error generating events for '{chapter_title}': {e}")
            # Don't print traceback here, handled in main loop
            # Return placeholder on error
            return [f"Event generation error 1 for '{chapter_summary}'", f"Event generation error 2 for '{chapter_summary}'"]


# --- Enhanced Writer Chain ---
class WriterChain:
    # Significantly enhanced prompt focusing on depth, theme, setting, and "show, don't tell"
    PROMPT = """
    You are a master novelist embodying the distinct writing style of {author}. Your task is to write a segment of the novel "{title}", a {genre} work.

    **ESTABLISHED CONTEXT (DO NOT REPEAT):**

    <MAIN_CHARACTER_PROFILE>
    {profile}
    </MAIN_CHARACTER_PROFILE>

    <SETTING_DESCRIPTION>
    {setting}
    </SETTING_DESCRIPTION>

    <CORE_THEMES>
    {themes}
    </CORE_THEMES>

    <NOVEL_PLOT_SUMMARY>
    {plot}
    </NOVEL_PLOT_SUMMARY>

    <CURRENT_CHAPTER_INFO>
    Chapter Title: {chapter_name}
    Chapter Summary: {summary}
    </CURRENT_CHAPTER_INFO>

    <NARRATIVE_HISTORY>
    Key events already occurred (prior chapters/events):
    {previous_events}
    </NARRATIVE_HISTORY>

    <CHAPTER_PROGRESS_SO_FAR>
    Narrative written previously within THIS chapter ({chapter_name}):
    {previous_paragraphs}
    </CHAPTER_PROGRESS_SO_FAR>

    **YOUR CURRENT WRITING TASK:**

    Bring the following specific event to life. Focus *exclusively* on this moment.

    <CURRENT_EVENT_TO_WRITE>
    {current_event}
    </CURRENT_EVENT_TO_WRITE>

    **WRITING INSTRUCTIONS (CRITICAL):**

    1.  **Style & Tone:** Write 1-3 detailed paragraphs (approx. 150-400 words total for this event, adjust based on significance) strictly in the voice and style of {author}. Match their typical sentence structure, vocabulary, pacing, and tone.
    2.  **Show, Don't Tell:** Demonstrate emotions, character thoughts, and plot developments through actions, dialogue (if appropriate for the event and character), internal monologue, and sensory details, rather than stating them directly.
    3.  **Sensory Integration:** Weave in specific sensory details (sight, sound, smell, touch, taste) consistent with the established <SETTING_DESCRIPTION>. Make the scene immersive.
    4.  **Character Depth:** Reflect the <MAIN_CHARACTER_PROFILE> (motivations, flaws, internal conflicts) in their reactions, thoughts, and actions during this event.
    5.  **Thematic Resonance:** Subtly connect the events or the character's internal experience to the <CORE_THEMES> where appropriate, without being overt.
    6.  **Continuity:** Ensure a seamless transition from the <CHAPTER_PROGRESS_SO_FAR>. Do NOT repeat information already covered. Do NOT summarize the event itself.
    7.  **Output:** Generate ONLY the newly written narrative paragraphs for the <CURRENT_EVENT_TO_WRITE>. No explanations, labels, or summaries.

    New Narrative Paragraphs (Style: {author}):"""

    def __init__(self):
        # High temperature for creative prose, high top-p for diversity, moderate top-k
        self.llm = create_llm(temperature=0.9, top_p=0.95, top_k=60)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True # Set to False in production if too noisy
        )

    def run(self, genre, author, title, profile, plot, setting, themes_str, chapter_name,
            previous_events_history_str, chapter_summary, previous_paragraphs_str, current_event):
        """Generates narrative paragraphs for a specific event."""
        # Ensure placeholders are clear if context is empty
        if not previous_events_history_str: previous_events_history_str = "None (This is the beginning of the story)."
        if not previous_paragraphs_str: previous_paragraphs_str = "None (This is the beginning of the chapter)."

        try:
            print(f"Invoking WriterChain for event: {current_event[:80]}...") # Log truncated event
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "genre": genre,
                "author": author,
                "title": title,
                "profile": profile,
                "plot": plot,
                "setting": setting,
                "themes": themes_str,
                "chapter_name": chapter_name,
                "previous_events": previous_events_history_str,
                "summary": chapter_summary,
                "previous_paragraphs": previous_paragraphs_str,
                "current_event": current_event
            })
            generated_text = result.get('text', f"[ERROR: Writing failed for event: '{current_event}']").strip()

            # Basic validation of output
            if not generated_text or generated_text.startswith("[ERROR:") or len(generated_text) < 50: # Increased min length
                 print(f"  Warning: WriterChain generated potentially invalid/short output for event '{current_event}'. Output: '{generated_text}'")
                 return f"[Writer Error - Event: '{current_event}'. LLM Output: '{generated_text}']"

            # Optional: Add post-processing to clean up LLM artifacts (e.g., repeated phrases, stray quotes)
            # generated_text = re.sub(r'^"|"$', '', generated_text) # Example: remove leading/trailing quotes

            return generated_text
        except Exception as e:
            # Log the error but return a placeholder to avoid crashing the whole process
            print(f"  ERROR occurred in WriterChain.run for event '{current_event}': {e}")
            # Consider logging traceback only if a debug flag is set
            # traceback.print_exc()
            return f"[FATAL WRITER ERROR - Event: '{current_event}'. Check logs for details.]"


# --- NEW: Refinement Chain (Optional Post-Processing) ---
class RefinementChain:
    PROMPT = """
    You are an expert literary editor refining a draft chapter written in the style of {author}.
    The chapter is part of the novel "{title}" ({genre}).

    **CONTEXT:**
    <MAIN_CHARACTER_PROFILE>{profile}</MAIN_CHARACTER_PROFILE>
    <SETTING_DESCRIPTION>{setting}</SETTING_DESCRIPTION>
    <CORE_THEMES>{themes}</CORE_THEMES>
    <NOVEL_PLOT_SUMMARY>{plot}</NOVEL_PLOT_SUMMARY>
    <CHAPTER_INFO>Title: {chapter_name}, Summary: {summary}</CHAPTER_INFO>

    **DRAFT TEXT FOR REFINEMENT:**
    <DRAFT_START>
    {draft_text}
    <DRAFT_END>

    **YOUR TASK:**
    Review and refine the provided draft text. Focus on ONE OR TWO of the following aspects (choose based on perceived need or specify focus):
    1.  **Stylistic Consistency:** Enhance adherence to {author}'s unique voice, sentence structure, and vocabulary.
    2.  **Prose Flow & Pacing:** Improve readability, sentence variety, and transitions between paragraphs. Smooth out awkward phrasing.
    3.  **Sensory Detail Enhancement:** Weave in more vivid and relevant sensory details based on the setting.
    4.  **Dialogue Polish:** Make dialogue sound more natural and character-specific (if applicable).
    5.  **Thematic Resonance:** Subtly strengthen connections to the core themes ({themes}).
    6.  **Show, Don't Tell:** Convert any instances of telling into showing through action, internal thought, or description.
    7.  **Conciseness:** Trim unnecessary words or redundant phrases without losing meaning or style.

    **INSTRUCTIONS:**
    - Make subtle but impactful changes. Do NOT rewrite entire sections unless absolutely necessary.
    - Preserve the original plot points and events of the draft.
    - Output ONLY the refined text for the chapter. Do not include explanations, summaries, or comments about your changes.

    Refined Chapter Text (Style: {author}):
    """

    def __init__(self):
        # Moderate temperature for controlled editing, high top-p for nuanced choices
        self.llm = create_llm(temperature=0.6, top_p=0.95, top_k=50)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, author, title, genre, profile, setting, themes_str, plot, chapter_name, summary, draft_text):
        """Refines the draft text of a chapter."""
        if not draft_text or draft_text.startswith("[Content generation skipped") or draft_text.startswith("[Writer Error"):
            print(f"  Skipping refinement for '{chapter_name}' due to missing or errored draft content.")
            return draft_text # Return original if it's bad

        try:
            print(f"Invoking RefinementChain for chapter: {chapter_name}...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "author": author,
                "title": title,
                "genre": genre,
                "profile": profile,
                "setting": setting,
                "themes": themes_str,
                "plot": plot,
                "chapter_name": chapter_name,
                "summary": summary,
                "draft_text": draft_text
            })
            refined_text = result.get('text', f"[ERROR: Refinement failed for chapter: '{chapter_name}']").strip()

            # Basic validation
            if not refined_text or refined_text.startswith("[ERROR:") or len(refined_text) < len(draft_text) * 0.5: # Check if it didn't shrink too much
                 print(f"  Warning: RefinementChain generated potentially invalid output for '{chapter_name}'. Output: '{refined_text[:100]}...'")
                 return draft_text # Fallback to original draft if refinement seems broken

            return refined_text
        except Exception as e:
            print(f"  ERROR occurred in RefinementChain.run for chapter '{chapter_name}': {e}")
            # traceback.print_exc() # Optional: uncomment for debugging
            return draft_text # Fallback to original draft on error


# --- Book Writing Orchestration ---

def format_themes_string(themes_dict):
    """Formats the theme dictionary into a string for prompts."""
    if not themes_dict or "Error" in themes_dict:
        return "N/A"
    return "; ".join([f"{name}: {desc}" for name, desc in themes_dict.items()])

def sort_chapters(chapter_dict):
    """Sorts chapter dictionary keys (titles) into a logical order. Returns list of (title, description) tuples."""
    def chapter_sort_key(chapter_title):
        title_lower = chapter_title.lower().strip()
        if title_lower == "prologue": return (-1, 0)
        if title_lower == "epilogue": return (9999, 0)
        match = re.match(r"chapter\s*(\d+)", title_lower)
        if match:
            try:
                num = int(match.group(1))
                return (1, num)
            except (IndexError, ValueError):
                return (2, chapter_title) # Fallback sort by title if number parsing fails
        return (3, chapter_title) # Sort any other titles alphabetically after chapters

    # Sort the items (title, description pairs) based on the key (title)
    return sorted(chapter_dict.items(), key=lambda item: chapter_sort_key(item[0]))


def generate_events_for_all_chapters(plot, profile, themes_str, chapter_dict, author):
    """Generates event lists for all chapters using EventChain."""
    print("\n--- Generating Events for Each Chapter ---")
    event_generator = EventChain()
    event_dict = {} # { 'Chapter Title': ['Event 1', 'Event 2', ...] }

    # Use the pre-sorted chapter list from the ChaptersChain parsing
    sorted_chapter_items = chapter_dict.items() # chapter_dict should already be sorted

    for chapter_title, summary in sorted_chapter_items:
         print(f"Generating events for: {chapter_title}")
         events = event_generator.run(plot, profile, themes_str, chapter_title, summary, author)
         if events and not any("error" in evt.lower() for evt in events):
            event_dict[chapter_title] = events
            print(f"  Generated {len(events)} events for '{chapter_title}'.")
            # for i, ev in enumerate(events): print(f"    {i+1}. {ev[:100]}...") # Log truncated events
         else:
            print(f"  WARNING: No valid events generated or parsed for '{chapter_title}'. Chapter content might be basic or skipped.")
            event_dict[chapter_title] = [] # Store empty list to indicate failure/skip

    return event_dict


def write_book(genre, author, title, profile, plot, setting, themes_str, chapter_dict, event_dict, refine_chapters=False):
    """Orchestrates the writing of the full book content, event by event, with optional refinement."""
    print("\n--- Starting Detailed Book Writing Process ---")
    writer_chain = WriterChain()
    refiner_chain = RefinementChain() if refine_chapters else None # Instantiate refiner only if needed

    book_content = {} # Stores final text: {chapter_title: "Full chapter text..."}
    previous_events_history = [] # Running list of event descriptions written so far

    # Use the pre-sorted chapter list
    sorted_chapter_items = chapter_dict.items()
    total_chapters = len(sorted_chapter_items)

    for chap_idx, (chapter_title, chapter_summary) in enumerate(sorted_chapter_items):
        print(f"\n--- Writing Chapter {chap_idx+1}/{total_chapters}: {chapter_title} ---")
        print(f"   Summary: {chapter_summary}")

        chapter_events = event_dict.get(chapter_title, [])
        if not chapter_events:
            print(f"   WARNING: No events found for '{chapter_title}'. Skipping content generation.")
            book_content[chapter_title] = f"[Content generation skipped: No events defined for this chapter.]"
            continue

        # Accumulator for text within the *current* chapter
        chapter_paragraphs_accumulator = ""
        total_events = len(chapter_events)

        for event_idx, event_description in enumerate(chapter_events):
            # Limit history length passed to LLM to avoid excessive context window usage
            # Keep maybe the last 20-30 events? Or based on token count?
            MAX_HISTORY_EVENTS = 30
            limited_history = previous_events_history[-MAX_HISTORY_EVENTS:]
            history_str = '\n'.join(f"- {evt}" for evt in limited_history) if limited_history else "None (Start of Story)"
            progress_str = chapter_paragraphs_accumulator if chapter_paragraphs_accumulator else "None (Start of Chapter)"

            print(f"   Writing Event {event_idx+1}/{total_events}: {event_description[:100]}...") # Log truncated event

            # Call the writer chain for the current event
            new_paragraphs = writer_chain.run(
                genre=genre, author=author, title=title, profile=profile, plot=plot,
                setting=setting, themes_str=themes_str, chapter_name=chapter_title,
                previous_events_history_str=history_str, chapter_summary=chapter_summary,
                previous_paragraphs_str=progress_str, current_event=event_description
            )

            # Append the newly generated block, handling potential errors
            if new_paragraphs.startswith("[FATAL WRITER ERROR") or new_paragraphs.startswith("[Writer Error"):
                print(f"   ERROR detected writing event {event_idx+1}. Adding error message to content.")
            if chapter_paragraphs_accumulator and not chapter_paragraphs_accumulator.endswith('\n\n'):
                chapter_paragraphs_accumulator += "\n\n" # Ensure separation
            chapter_paragraphs_accumulator += new_paragraphs

            # Add the *description* of the event we just attempted to write to the global history
            # Even if writing failed, log the attempt.
            previous_events_history.append(f"[{chapter_title}] {event_description}")

        # --- Optional Refinement Step ---
        final_chapter_text = chapter_paragraphs_accumulator
        if refiner_chain:
            print(f"   Refining chapter: {chapter_title}...")
            final_chapter_text = refiner_chain.run(
                author=author, title=title, genre=genre, profile=profile, setting=setting,
                themes_str=themes_str, plot=plot, chapter_name=chapter_title,
                summary=chapter_summary, draft_text=chapter_paragraphs_accumulator
            )
            if final_chapter_text != chapter_paragraphs_accumulator:
                 print(f"   Refinement applied for '{chapter_title}'.")
            else:
                 print(f"   Refinement skipped or resulted in no changes for '{chapter_title}'.")


        # Store the fully assembled (and potentially refined) text for the chapter
        book_content[chapter_title] = final_chapter_text
        print(f"--- Finished Chapter: {chapter_title} ---")

    print("\n--- Book Writing Process Complete ---")
    return book_content

# --- Document Writing Class ---

class DocWriter:
    def __init__(self, output_folder=OUTPUT_FOLDER):
        self.output_folder = output_folder
        os.makedirs(self.output_folder, exist_ok=True)
        print(f"Ensured output directory exists: {os.path.abspath(self.output_folder)}")

    def _sanitize_filename(self, name):
        """Removes or replaces characters invalid for filenames more aggressively."""
        if not isinstance(name, str): name = str(name)
        # Remove invalid chars
        name = re.sub(r'[\\/*?:"<>|]', "", name)
        # Replace spaces and multiple dots/hyphens
        name = re.sub(r'\s+', '_', name)
        name = re.sub(r'[_.-]{2,}', '_', name)
        # Trim leading/trailing whitespace/underscores/dots/hyphens
        name = name.strip('._- ')
        # Limit length
        max_len = 180 # Slightly shorter for safety across filesystems
        if len(name) > max_len:
            # Try to keep the beginning and end
            name = name[:max_len//2] + "..." + name[-(max_len//2 - 3):]
        if not name:
            name = "Untitled_Novel"
        return name

    def write_doc(self, book_content, chapter_dict, title, genre, author, themes_dict, setting):
        """Writes the generated book content to a .docx file with enhanced front matter."""
        print("\n--- Assembling and Writing Document ---")
        doc = docx.Document()

        # --- Enhanced Front Matter ---
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Genre: {genre}")
        doc.add_paragraph(f"Inspired by the style of: {author}")
        doc.add_paragraph("\n") # Add spacing

        # Add Themes
        doc.add_heading("Core Themes", level=2)
        themes_str = format_themes_string(themes_dict)
        if themes_str == "N/A" or "Error" in themes_dict:
             doc.add_paragraph("Themes could not be generated or parsed.")
        else:
            for theme_name, theme_desc in themes_dict.items():
                doc.add_paragraph(f"- **{theme_name}:** {theme_desc}", style='List Bullet')
        doc.add_paragraph("\n")

        # Add Setting Summary
        doc.add_heading("Setting Summary", level=2)
        if setting.startswith("Error"):
             doc.add_paragraph("Setting description could not be generated.")
        else:
             # Add setting description, potentially splitting into paragraphs
             setting_paras = [p.strip() for p in setting.split('\n') if p.strip()]
             for para in setting_paras:
                doc.add_paragraph(para)
        doc.add_paragraph("\n")

        # Add Table of Contents (Chapter List)
        doc.add_heading("Chapters", level=1)
        sorted_chapter_items = chapter_dict.items() # Already sorted
        for chapter_title, description in sorted_chapter_items:
            doc.add_paragraph(f"**{chapter_title}:** {description}", style='List Bullet')
        doc.add_page_break() # Page break after ToC

        # --- Body ---
        print("Adding chapter content to document...")
        total_chapters = len(sorted_chapter_items)
        for i, (chapter_title, description) in enumerate(sorted_chapter_items):
            chapter_heading = chapter_title.strip()
            doc.add_heading(chapter_heading, level=1)
            print(f"  Adding Chapter {i+1}/{total_chapters}: {chapter_heading}")

            chapter_text = book_content.get(chapter_title, "[Error: Chapter content not found]")

            # Add chapter text, splitting by double newlines used as separators
            # Also handle single newlines within LLM-generated paragraphs gracefully
            paragraphs = chapter_text.split('\n\n')
            for para_block in paragraphs:
                if para_block.strip():
                    # Add paragraph, preserving single newlines within the block if desired
                    # For standard docx, usually treat each block as one paragraph
                    doc.add_paragraph(para_block.strip())

            # Add a page break after each chapter for clear separation
            if i < total_chapters - 1:
                doc.add_page_break()
            print(f"    Added content for '{chapter_title}'")

        # --- Filename Creation ---
        safe_basename = self._sanitize_filename(title)
        safe_author = self._sanitize_filename(author)
        safe_genre = self._sanitize_filename(genre.split('/')[0].split(',')[0])
        filename = f"{safe_basename}_by_{safe_author}_{safe_genre}.docx"
        output_path = os.path.join(self.output_folder, filename)

        # --- Saving Document ---
        try:
            print(f"\nAttempting to save document to: {output_path}")
            doc.save(output_path)
            print(f"--- Document saved successfully! ---")
            return output_path # Return path if successful
        except PermissionError:
            print(f"\nERROR: Permission denied trying to save '{output_path}'.")
            print("Check file permissions or if the file is open elsewhere.")
            print(f"Folder: {os.path.abspath(self.output_folder)}")
        except Exception as e:
            print(f"\nERROR saving document to {output_path}: {e}")
            traceback.print_exc()
        return None # Return None if saving failed

# --- Main Execution ---

def main():
    process_start_time = time.time()
    print("=============================================")
    print("=== ENHANCED NOVEL GENERATION SYSTEM V2 ===")
    print(f"=== Timestamp: {time.strftime('%Y-%m-%d %H:%M:%S')} ===")
    print(f"=== Using Model: {DEFAULT_MODEL} ===")
    print("=============================================")

    # --- Configuration ---
    # Ensure output directory exists
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    print(f"Output Folder: {os.path.abspath(OUTPUT_FOLDER)}")

    # --- Inputs ---
    # ** IMPORTANT: Replace with the actual path to your PDF resume file **
    # Ensure it's accessible relative to where you run the script, or provide an absolute path.
    # It's currently expected inside the OUTPUT_FOLDER.
    resume_filename = DEFAULT_RESUME_FILENAME
    subject = ("""
Divi is a quiet observer of life — articulate, sensitive, and prone to spiraling thoughts that no one around her seems to notice. By day, she moves through life calmly, performing routines, attending to obligations, and maintaining the appearance of normalcy. But by night, especially in the hours before dawn — the “silver hush” — her mind becomes a strange and unfiltered terrain. Here, she is not just herself; she is every version of herself she’s ever buried.
Each chapter takes place over the course of a different night, exploring a specific loop or thread that her mind follows in the liminal silence before morning. Her insomnia is more than sleeplessness — it's a ritual, a confrontation. As the moonlight turns darkness into silver, Divi finds herself reliving memories that never quite healed, conversations that were never finished, and fears that echo too loudly when the world is quiet.
These episodes manifest almost like dream sequences: visceral, fragmented, yet real. Sometimes, she’s speaking with her younger self. Sometimes, it’s an imagined version of someone she loved or feared. Sometimes it’s just silence — heavy, metallic, alive. Her room becomes a surreal psychological stage where emotions take form and time becomes fluid.
Over the course of the story, Divi begins to fear that she may be slipping into something deeper — not just insomnia, but a fracture between reality and introspection. She starts journaling these episodes in a desperate attempt to anchor herself. But the act of writing becomes its own mirror, forcing her to question if the self she’s trying to hold onto has already changed.
By the final night, as dawn approaches and the silver haze turns to golden glow, Divi must choose: does she return fully to the waking world, or does she surrender to the nocturnal clarity where everything feels realer? The story ends in ambiguity — peace, perhaps, or madness cloaked in lucidity.""")
    author_style = 'Ocean Vuong'
    genre = 'Psychological Literary Fiction / Nocturnal Realism'
    # Set to True to enable the experimental refinement pass (uses more LLM calls)
    ENABLE_REFINEMENT_PASS = False

    print("\n--- Input Parameters ---")
    print(f"Resume File: {resume_filename} (Expected in: {OUTPUT_FOLDER})")
    print(f"Subject: {subject[:150]}...") # Print truncated subject
    print(f"Author Style Inspiration: {author_style}")
    print(f"Genre(s): {genre}")
    print(f"Refinement Pass Enabled: {ENABLE_REFINEMENT_PASS}")
    print("------------------------")

    # --- Instantiate Chains & Writer ---
    try:
        print("\n--- Initializing Core Components ---")
        # Note: LLMs are created within each chain now
        main_character_chain = MainCharacterChain()
        setting_chain = SettingChain()
        theme_chain = ThemeChain()
        title_chain = TitleChain()
        plot_chain = PlotChain()
        chapters_chain = ChaptersChain()
        # EventChain, WriterChain, RefinementChain instantiated later
        doc_writer = DocWriter(output_folder=OUTPUT_FOLDER)
        print("--- Core Components Initialized Successfully ---")
    except Exception as e:
        print(f"\nFATAL ERROR during component initialization: {e}")
        # Error message likely printed during LLM creation attempt
        return # Stop execution

    # --- Generate Novel Components Sequentially ---
    generation_successful = True
    try:
        # 1. Generate Profile
        print("\n--- Step 1: Generating Main Character Profile ---")
        start = time.time()
        profile = main_character_chain.run(resume_filename, genre)
        print(f"Profile Generation Time: {time.time() - start:.2f}s")
        if profile.startswith("Error:"):
            print(f"PROFILE GENERATION FAILED: {profile}")
            generation_successful = False
        else:
            print(f"Profile Generated:\n---\n{profile}\n---")

        # 2. Generate Setting (Requires Profile)
        setting = "Error: Setting generation skipped due to profile error."
        if generation_successful:
            print("\n--- Step 2: Generating Setting Description ---")
            start = time.time()
            setting = setting_chain.run(subject, genre, profile)
            print(f"Setting Generation Time: {time.time() - start:.2f}s")
            if setting.startswith("Error:"):
                print(f"SETTING GENERATION FAILED: {setting}")
                generation_successful = False # Allow continuing, but flag issue
            else:
                print(f"Setting Generated:\n---\n{setting}\n---")

        # 3. Generate Themes (Requires Profile & Setting)
        themes_dict = {"Error": "Theme generation skipped due to prior errors."}
        themes_str = "N/A"
        if generation_successful: # Only proceed if profile/setting likely okay
            print("\n--- Step 3: Generating Core Themes ---")
            start = time.time()
            themes_dict = theme_chain.run(subject, genre, profile, setting)
            print(f"Theme Generation Time: {time.time() - start:.2f}s")
            if "Error" in themes_dict:
                print(f"THEME GENERATION/PARSING FAILED: {themes_dict['Error']}")
                # Don't stop, but note the failure
            else:
                themes_str = format_themes_string(themes_dict)
                print(f"Themes Generated:\n---\n{themes_str}\n---")

        # 4. Generate Title (Requires Profile, Setting, Themes)
        title = "Error - Title Generation Failed"
        if generation_successful: # Proceed even if themes failed, using "N/A"
            print("\n--- Step 4: Generating Novel Title ---")
            start = time.time()
            title = title_chain.run(subject, genre, author_style, profile, setting, themes_str)
            print(f"Title Generation Time: {time.time() - start:.2f}s")
            if title.startswith("Error") or "Placeholder" in title:
                 print(f"TITLE GENERATION FAILED/PLACEHOLDER: {title}")
                 # Allow continuing with placeholder/error title
            else:
                 print(f"Title Generated: '{title}'")

        # 5. Generate Plot (Requires Profile, Setting, Themes, Title)
        plot = "Error: Plot generation skipped due to prior errors."
        if generation_successful:
            print("\n--- Step 5: Generating Detailed Plot Outline ---")
            start = time.time()
            plot = plot_chain.run(subject, genre, author_style, profile, title, setting, themes_str)
            print(f"Plot Generation Time: {time.time() - start:.2f}s")
            if plot.startswith("Error"):
                print(f"PLOT GENERATION FAILED: {plot}")
                generation_successful = False # Cannot proceed without plot
            else:
                print(f"Plot Generated:\n---\n{plot}\n---")

        # 6. Generate Chapter List (Requires Plot, Profile, Setting, Themes, Title)
        chapter_dict = {}
        if generation_successful:
            print("\n--- Step 6: Generating Chapter List & Summaries ---")
            start = time.time()
            # chapter_dict should be { 'Chapter 1': 'Description...', ... } and sorted
            chapter_dict = chapters_chain.run(subject, genre, author_style, profile, title, plot, setting, themes_str)
            print(f"Chapter List Generation Time: {time.time() - start:.2f}s")
            if not chapter_dict:
                 print("ERROR: Failed to generate or parse chapters. Cannot write book content.")
                 generation_successful = False
            else:
                 print("Chapters Generated & Parsed:")
                 for chap_title, chap_desc in chapter_dict.items(): # Already sorted
                      print(f"  - {chap_title}: {chap_desc}")

        # 7. Generate Events for Each Chapter (Requires Plot, Profile, Themes, Chapters)
        event_dict = {}
        if generation_successful:
            print("\n--- Step 7: Generating Events for All Chapters ---")
            start = time.time()
            event_dict = generate_events_for_all_chapters(plot, profile, themes_str, chapter_dict, author_style)
            print(f"Event Generation Time: {time.time() - start:.2f}s")
            if not any(event_dict.values()): # Check if *any* events were generated
                print("CRITICAL WARNING: Failed to generate events for *any* chapter. Book content will likely be empty or skipped.")
                # Allow proceeding, but the book will be mostly empty chapters

        # --- Write the Book Content ---
        book_content = {}
        if generation_successful and chapter_dict: # Need chapters to write
            print("\n--- Step 8: Writing Full Book Content ---")
            start = time.time()
            book_content = write_book(
                genre, author_style, title, profile, plot, setting, themes_str,
                chapter_dict, event_dict,
                refine_chapters=ENABLE_REFINEMENT_PASS
            )
            print(f"Book Writing Time: {time.time() - start:.2f}s")
        elif not generation_successful:
             print("\n--- Skipping Book Writing due to errors in previous steps ---")
        else:
             print("\n--- Skipping Book Writing as no chapters were generated ---")


        # --- Save to Document ---
        if book_content: # Only save if some content was generated
            print("\n--- Step 9: Saving Document ---")
            start = time.time()
            saved_path = doc_writer.write_doc(book_content, chapter_dict, title, genre, author_style, themes_dict, setting)
            print(f"Document Saving Time: {time.time() - start:.2f}s")
            if saved_path:
                print(f"\nSuccess! Novel saved to: {saved_path}")
            else:
                print("\nFailed to save the document.")
        else:
            print("\n--- Skipping Document Saving as no book content was generated ---")


    except FileNotFoundError as e:
         print(f"\nCRITICAL FILE ERROR: {e}")
         print(f"Please ensure the required file ('{resume_filename}') exists in the correct directory.")
         print(f"Expected location: '{os.path.abspath(OUTPUT_FOLDER)}'")
         generation_successful = False
    except Exception as e:
        print(f"\n--- AN UNEXPECTED ERROR OCCURRED DURING NOVEL GENERATION ---")
        print(f"Error Type: {type(e).__name__}")
        print(f"Error Details: {e}")
        print("--- Traceback ---")
        traceback.print_exc()
        print("-----------------")
        generation_successful = False

    finally:
        end_time = time.time()
        total_time = end_time - process_start_time
        print("\n============================================")
        print(f"=== NOVEL GENERATION PROCESS {'FINISHED' if generation_successful else 'FINISHED WITH ERRORS'} ===")
        print(f"=== Total Execution Time: {total_time:.2f} seconds ===")
        print("============================================")


if __name__ == "__main__":
    # Ensure you have Ollama running with the specified DEFAULT_MODEL pulled.
    # Example: `ollama run gemma2` in your terminal.
    # Ensure the resume PDF (e.g., 'divi_1.pdf') is in the './docs_generated/' folder
    # or update DEFAULT_RESUME_FILENAME and OUTPUT_FOLDER constants.
    main()