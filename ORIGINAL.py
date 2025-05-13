# app.py
from dotenv import load_dotenv
import os
from langchain_openai import ChatOpenAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.document_loaders import PyPDFLoader
import docx

# Load environment variables
load_dotenv()

# Constants
OPENROUTER_API_BASE = "https://openrouter.ai/api/v1"
OPENROUTER_REFERRER = "https://github.com/langchain-example"
DEFAULT_MODEL = "meta-llama/llama-4-maverick:free"
# DEFAULT_MODEL = "google/gemini-2.0-pro-exp-02-05:free"

def create_llm():
    """Create and return a ChatOpenAI instance configured for OpenRouter"""
    api_key = os.getenv('OPENROUTER_API_KEY')
    if not api_key:
        raise ValueError("OPENROUTER_API_KEY not found in environment variables")
    
    return ChatOpenAI(
        temperature=0.7,
        model=DEFAULT_MODEL,
        api_key=api_key,
        base_url=OPENROUTER_API_BASE,
        default_headers={"HTTP-Referer": OPENROUTER_REFERRER}
    )

class MainCharacterChain:
    PROMPT = """
    You are provided with the resume of a person. 
    Describe the person's profile in a few sentences and include that person's name.

    Resume: {text}

    Profile:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def load_resume(self, file_name):
        folder = './docs'
        file_path = os.path.join(folder, file_name)
        loader = PyPDFLoader(file_path)
        docs = loader.load_and_split()
        return docs

    def run(self, file_name):
        docs = self.load_resume(file_name)
        resume = '\n\n'.join([doc.page_content for doc in docs])
        return self.chain.run(text=resume)

class TitleChain:
    PROMPT = """
    Your job is to generate the title for a novel about the following subject and main character. 
    Return a title and only a title!
    The title should be consistent with the genre of the novel.
    The title should be consistent with the style of the author.

    Subject: {subject}
    Genre: {genre}
    Author: {author}
    Main character's profile: {profile}

    Title:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile):
        return self.chain.run(
            subject=subject,
            genre=genre,
            author=author,
            profile=profile
        )

class PlotChain:
    PROMPT = """
    Your job is to generate the plot for a novel. Return a plot and only a plot!
    Describe the full plot of the story and don't hesitate to create new characters to make it compelling.
    You are provided the following subject, title and main character's profile.
    Make sure that the main character is at the center of the story 
    The plot should be consistent with the genre of the novel.
    The plot should be consistent with the style of the author.

    Consider the following attributes to write an exciting story:
    {features}

    Subject: {subject}
    Genre: {genre}
    Author: {author}
    Title: {title}
    Main character's profile: {profile}

    Plot:"""

    HELPER_PROMPT = """
    Generate a list of attributes that characterized an exciting story.
    
    Input: Generate attributes
    
    List of attributes:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile, title):
        # Create helper chain with input variable
        helper_chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate(
                template=self.HELPER_PROMPT,
                input_variables=["input"]
            )
        )
        # Run helper chain with input
        features = helper_chain.run(input="Generate attributes")

        return self.chain.run(
            features=features,
            subject=subject,
            genre=genre,
            author=author,
            profile=profile,
            title=title
        )

class ChaptersChain:
    PROMPT = """
    Your job is to generate a list of chapters. 
    ONLY the list and nothing more!
    You are provided with a title, a plot and a main character for a novel.
    Generate a list of chapters describing the plot of that novel.
    Make sure the chapters are consistent with the plot.
    The chapters should be consistent with the genre of the novel. 
    The chapters should be consistent with the style of the author. 

    Follow this template: 

    Prologue: [description of prologue]
    Chapter 1: [description of chapter 1]
    ...
    Epilogue: [description of epilogue]

    Make sure the chapter is followed by the character `:` and its description.
    
    Subject: {subject}
    Genre: {genre}
    Author: {author}
    Title: {title}
    Main character's profile: {profile}
    Plot: {plot}
    
    Chapters list:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse(self, response):
        chapter_list = response.strip().split('\n')
        chapter_list = [chapter for chapter in chapter_list if ':' in chapter]
        return dict([chapter.strip().split(':') for chapter in chapter_list])

    def run(self, subject, genre, author, profile, title, plot):
        response = self.chain.run(
            subject=subject,
            genre=genre,
            author=author,
            profile=profile,
            title=title,
            plot=plot
        )
        return self.parse(response)

class WriterChain:
    PROMPT = """
    You are a novel writer. The novel is described by a list of events. 
    You have already written the novel up to the last event. 
    Your job is to generate the paragraphs of the novel about the new event.
    Make sure the paragraphs are consistent with the plot of the chapter.
    The paragraphs should be consistent with the genre and author's style.

    Genre: {genre}
    Author: {author}
    Title: {title}
    Main character's profile: {profile}
    Novel's Plot: {plot}
    Previous events: {previous_events}
    Current Chapter summary: {summary}
    Previous paragraphs: {previous_paragraphs}
    New event to write about: {current_event}

    Paragraphs describing that event:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, genre, author, title, profile, plot, 
            previous_events, summary, previous_paragraphs, current_event):
        previous_events = '\n'.join(previous_events)
        return self.chain.run(
            genre=genre,
            author=author,
            title=title,
            profile=profile,
            plot=plot,
            previous_events=previous_events,
            summary=summary,
            previous_paragraphs=previous_paragraphs,
            current_event=current_event
        )

def write_book(genre, author, title, profile, plot, summaries_dict, event_dict):
    writer_chain = WriterChain()
    previous_events = []
    book = {}
    paragraphs = ''

    for chapter, event_list in event_dict.items():
        book[chapter] = []
        for event in event_list:
            paragraphs = writer_chain.run(
                genre=genre,
                author=author,
                title=title,
                profile=profile,
                plot=plot,
                previous_events=previous_events,
                summary=summaries_dict[chapter],
                previous_paragraphs=paragraphs,
                current_event=event
            )
            previous_events.append(event)
            book[chapter].append(paragraphs)
    return book

class DocWriter:
    def __init__(self):
        self.doc = docx.Document()

    def write_doc(self, book, chapter_dict, title):
        self.doc.add_heading(title, 0)
        for chapter, paragraphs_list in book.items():
            description = chapter_dict[chapter]
            chapter_name = f'{chapter.strip()}: {description.strip()}'
            self.doc.add_heading(chapter_name, 1)
            text = '\n\n'.join(paragraphs_list)
            self.doc.add_paragraph(text)
        
        # Create a valid filename from the title
        # Replace spaces with underscores and remove any invalid characters
        filename = "".join(x for x in title if x.isalnum() or x in (' ', '-', '_')).strip()
        filename = filename.replace(' ', '_') + '.docx'
        
        # Save with the book's title as filename
        self.doc.save(os.path.join('./docs', filename))

def main():
    subject = 'Machine War'
    author = 'Ernest Hemingway'
    genre = 'horror'

    main_character_chain = MainCharacterChain()
    profile = main_character_chain.run('Profile_Sak.pdf')
    doc_writer = DocWriter()

    # Generate book structure
    title_chain = TitleChain()
    plot_chain = PlotChain()
    chapters_chain = ChaptersChain()

    title = title_chain.run(subject, genre, author, profile)
    plot = plot_chain.run(subject, genre, author, profile, title)
    chapter_dict = chapters_chain.run(subject, genre, author, profile, title, plot)

    # Generate events and summaries
    summaries_dict = {}
    event_dict = {}
    
    for chapter in chapter_dict:
        # For brevity, we'll generate simple summaries and events
        summaries_dict[chapter] = f"Summary of {chapter}"
        event_dict[chapter] = [f"Event 1 of {chapter}", f"Event 2 of {chapter}"]

    # Write the book
    book = write_book(
        genre,
        author,
        title,
        profile,
        plot,
        summaries_dict,
        event_dict
    )

    # Save to document
    doc_writer.write_doc(book, chapter_dict, title)

if __name__ == "__main__":
    main()
