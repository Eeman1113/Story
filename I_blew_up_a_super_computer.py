import requests
import json
import time
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pypdf # Added for PDF processing

# --- Configuration ---
# Replace with your Ollama API endpoint and desired model
OLLAMA_BASE_URL = "http://localhost:11434/api/generate"
# Consider models like "llama3:latest", "gemma2:latest", "mistral:latest" or more specialized ones if available.
# Ensure the model is pulled in Ollama (e.g., `ollama pull llama3`)
OLLAMA_MODEL = "llama3:latest" 
# Longer timeout for potentially complex generation tasks
OLLAMA_TIMEOUT = 360 # 6 minutes, adjust as needed

# Output directory for the generated novel
OUTPUT_DIR = "generated_novel_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

class NovelGenerator:
    def __init__(self, resume_content, subject, author_style, genre):
        self.resume_content = resume_content
        self.subject = subject
        self.author_style = author_style
        self.genre = genre
        self.num_chapters = 0 # Will be determined by the AI

        # Core story elements - will be populated by generation methods
        self.characters = {}  # Detailed character objects/dictionaries
        self.world_details = {"name": "", "key_locations": [], "cultural_elements": [], "rules": [], "atmosphere": ""}
        self.themes_motifs = {"themes": [], "motifs": []}
        self.plot_outline = ""  # High-level plot (e.g., 3-act structure)
        self.novel_title = "Untitled Novel"

        # Detailed chapter-by-chapter plan
        self.chapter_plans = {} # Key: chapter_num, Value: dict with plan details

        # Generated content and continuity data
        self.generated_chapters_content = {} # Key: chapter_num, Value: full chapter text
        self.chapter_continuity_data = {} # Key: chapter_num, Value: dict with summary, char updates, timeline, emotional arc, flow_analysis

        print("NovelGenerator initialized.")
        print(f"  Subject: {self.subject[:100]}...")
        print(f"  Author Style: {self.author_style}")
        print(f"  Genre: {self.genre}")
        print(f"  Resume provided: {'Yes' if self.resume_content else 'No'}")
        print(f"  Number of chapters will be determined automatically.")


    def _ollama_generate(self, prompt, system_prompt="You are a helpful AI assistant.", temperature=0.7, top_p=0.9):
        """
        Helper function to make API calls to the Ollama server.
        """
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "system": system_prompt,
            "stream": False,
            "options": {
                "temperature": temperature,
                "top_p": top_p,
                # "num_ctx": 8192 # Example: Adjust context window if needed and supported by model like Llama3
            }
        }
        # print(f"\n--- Sending Prompt to LLM ({OLLAMA_MODEL}) ---\n{prompt[:300]}...\n---") # Debug: Show prompt start
        try:
            response = requests.post(OLLAMA_BASE_URL, json=payload, timeout=OLLAMA_TIMEOUT)
            response.raise_for_status()
            # print(f"--- LLM Response Received ---\n{response.json()['response'][:300]}...\n---") # Debug: Show response start
            return response.json()["response"].strip()
        except requests.exceptions.Timeout:
            print(f"ERROR: Ollama request timed out after {OLLAMA_TIMEOUT} seconds for prompt: {prompt[:100]}...")
            return f"[OLLAMA TIMEOUT ERROR for prompt: {prompt[:100]}...]"
        except requests.exceptions.RequestException as e:
            print(f"ERROR: Ollama request failed: {e} for prompt: {prompt[:100]}...")
            return f"[OLLAMA REQUEST ERROR: {e} for prompt: {prompt[:100]}...]"
        except json.JSONDecodeError as e:
            print(f"ERROR: Failed to decode JSON response from Ollama: {e}")
            print(f"Raw response text: {response.text}")
            return f"[OLLAMA JSON DECODE ERROR for prompt: {prompt[:100]}...]"

    def _parse_character_profiles(self, text_block):
        """
        Parses character profiles from a structured text block.
        Expects format like:
        CHARACTER NAME: [Name]
        ROLE: [Role]
        DESCRIPTION: [Description]
        MOTIVATION: [Motivation]
        INITIAL_ARC_SUMMARY: [Arc Summary]
        FLAWS: [Flaws]
        STRENGTHS: [Strengths]
        """
        characters = {}
        current_char_data = {}
        current_char_name = None

        for line in text_block.split('\n'):
            line = line.strip()
            if not line: # Handles blank lines between character blocks
                if current_char_name and current_char_data:
                    characters[current_char_name] = {
                        "name": current_char_name,
                        "role": current_char_data.get("ROLE", "N/A"),
                        "description": current_char_data.get("DESCRIPTION", "N/A"),
                        "motivation": current_char_data.get("MOTIVATION", "N/A"),
                        "arc_summary": current_char_data.get("INITIAL_ARC_SUMMARY", "N/A"),
                        "flaws": current_char_data.get("FLAWS", "N/A"),
                        "strengths": current_char_data.get("STRENGTHS", "N/A"),
                        "current_status": "alive", 
                        "current_location": "unknown",
                        "emotional_state": "neutral",
                        "knowledge": [], 
                        "relationships": {}, 
                        "first_appearance_chapter": 0,
                        "development_log": [] 
                    }
                    current_char_data = {}
                    current_char_name = None
                continue

            match = re.match(r"CHARACTER NAME:\s*(.*)", line, re.IGNORECASE)
            if match:
                if current_char_name and current_char_data: # Save previous character before starting new one
                     characters[current_char_name] = {
                        "name": current_char_name,
                        "role": current_char_data.get("ROLE", "N/A"),
                        "description": current_char_data.get("DESCRIPTION", "N/A"),
                        "motivation": current_char_data.get("MOTIVATION", "N/A"),
                        "arc_summary": current_char_data.get("INITIAL_ARC_SUMMARY", "N/A"),
                        "flaws": current_char_data.get("FLAWS", "N/A"),
                        "strengths": current_char_data.get("STRENGTHS", "N/A"),
                        "current_status": "alive", "current_location": "unknown", "emotional_state": "neutral",
                        "knowledge": [], "relationships": {}, "first_appearance_chapter": 0, "development_log": []
                    }
                current_char_name = match.group(1).strip()
                current_char_data = {} # Reset for the new character
                continue

            if current_char_name: # If we are currently parsing a character
                for key in ["ROLE", "DESCRIPTION", "MOTIVATION", "INITIAL_ARC_SUMMARY", "FLAWS", "STRENGTHS"]:
                    if line.upper().startswith(key + ":"):
                        current_char_data[key] = line[len(key)+1:].strip()
                        break
        
        if current_char_name and current_char_data: # Save the last character in the block
            characters[current_char_name] = {
                "name": current_char_name,
                "role": current_char_data.get("ROLE", "N/A"),
                "description": current_char_data.get("DESCRIPTION", "N/A"),
                "motivation": current_char_data.get("MOTIVATION", "N/A"),
                "arc_summary": current_char_data.get("INITIAL_ARC_SUMMARY", "N/A"),
                "flaws": current_char_data.get("FLAWS", "N/A"),
                "strengths": current_char_data.get("STRENGTHS", "N/A"),
                "current_status": "alive", "current_location": "unknown", "emotional_state": "neutral",
                "knowledge": [], "relationships": {}, "first_appearance_chapter": 0, "development_log": []
            }
        return characters

    # --- Phase 1: Foundation Methods ---
    def generate_foundational_elements(self):
        """
        Generates initial character profiles, world details, themes/motifs, high-level plot outline,
        and determines the number of chapters.
        """
        print("\n--- Generating Foundational Elements ---")

        # 1. Character Conception
        print("Step 1.1: Generating Character Profiles...")
        char_system_prompt = f"You are a master character creator for {self.genre} novels, inspired by {self.author_style}."
        char_prompt = f"""
        Based on the novel's subject, genre, and the provided resume snippet (if any), create detailed profiles for the main protagonist and 1-2 key supporting characters (e.g., antagonist, mentor, love interest).
        Novel Subject: {self.subject}
        Genre: {self.genre}
        Author Style Influence: {self.author_style}
        Resume Snippet (use to inspire the protagonist, inferring personality, potential skills, and background):
        ---
        {self.resume_content if self.resume_content else "No resume provided. Create protagonist based on subject and genre."}
        ---
        For each character, provide:
        CHARACTER NAME: [Suggest a fitting name]
        ROLE: [Protagonist, Antagonist, Key Supporting - specify type]
        DESCRIPTION: [Detailed appearance, key personality traits, mannerisms, background hints]
        MOTIVATION(S): [What drives them? What are their primary goals, conscious or subconscious?]
        INITIAL_ARC_SUMMARY: [How might they change or develop throughout the story? What is their potential journey?]
        FLAWS/WEAKNESSES: [What are their vulnerabilities, biases, or negative traits?]
        STRENGTHS/SKILLS: [What are their notable positive attributes or skills?]

        Ensure characters are relatable and have depth.
        """
        character_profiles_text = self._ollama_generate(char_prompt, char_system_prompt, temperature=0.75)
        if "[OLLAMA" in character_profiles_text: 
            print(f"ERROR generating character profiles: {character_profiles_text}")
        else:
            self.characters = self._parse_character_profiles(character_profiles_text)
            print(f"Generated {len(self.characters)} character profiles: {', '.join(self.characters.keys())}")

        # 2. Worldbuilding
        print("\nStep 1.2: Generating World Details...")
        world_system_prompt = f"You are a world-building expert for {self.genre} fiction, creating immersive settings like those by {self.author_style}."
        world_prompt = f"""
        Based on the novel's subject and genre, describe the primary world/setting.
        Novel Subject: {self.subject}
        Genre: {self.genre}
        Provide the following details:
        1.  WORLD NAME: [A unique and evocative name for the main setting/world/city]
        2.  KEY LOCATIONS: [List 3-5 significant recurring locations with brief descriptions. e.g., "The Sunken Library - an ancient repository of forbidden knowledge"]
        3.  CULTURAL ELEMENTS: [Describe 2-3 unique customs, societal norms, beliefs, or technologies. e.g., "Aether-tech - devices powered by ambient magical energy"]
        4.  ATMOSPHERE/TONE: [Describe the overall mood and feeling of the world (e.g., oppressive, wondrous, decaying, futuristic, magical)]
        5.  KEY RULES/LAWS (if applicable, e.g., for magic systems, societal structure): [List 1-3 fundamental rules that govern this world or its unique aspects]
        """
        world_details_text = self._ollama_generate(world_prompt, world_system_prompt, temperature=0.65)
        if "[OLLAMA" in world_details_text:
            print(f"ERROR generating world details: {world_details_text}")
        else:
            lines = world_details_text.split('\n')
            for line in lines:
                if "WORLD NAME:" in line.upper(): self.world_details["name"] = line.split(":", 1)[1].strip()
                elif "ATMOSPHERE/TONE:" in line.upper(): self.world_details["atmosphere"] = line.split(":", 1)[1].strip()

            key_loc_match = re.search(r"KEY LOCATIONS:\s*((?:-\s*.*|\d\.\s*.*(?:\n|$))+)", world_details_text, re.IGNORECASE | re.DOTALL)
            if key_loc_match:
                self.world_details["key_locations"] = [loc.strip("- ").strip().lstrip("123456789. ") for loc in key_loc_match.group(1).strip().split('\n') if loc.strip()]

            cultural_elem_match = re.search(r"CULTURAL ELEMENTS:\s*((?:-\s*.*|\d\.\s*.*(?:\n|$))+)", world_details_text, re.IGNORECASE | re.DOTALL)
            if cultural_elem_match:
                self.world_details["cultural_elements"] = [elem.strip("- ").strip().lstrip("123456789. ") for elem in cultural_elem_match.group(1).strip().split('\n') if elem.strip()]

            rules_match = re.search(r"KEY RULES/LAWS:\s*((?:-\s*.*|\d\.\s*.*(?:\n|$))+)", world_details_text, re.IGNORECASE | re.DOTALL)
            if rules_match:
                self.world_details["rules"] = [rule.strip("- ").strip().lstrip("123456789. ") for rule in rules_match.group(1).strip().split('\n') if rule.strip()]
            print(f"Generated World Details for '{self.world_details.get('name', 'Unnamed World')}'.")

        # 3. Themes and Motifs
        print("\nStep 1.3: Generating Themes and Motifs...")
        themes_system_prompt = f"You are a literary analyst identifying profound themes and recurring motifs for {self.genre} novels, in the vein of {self.author_style}."
        themes_prompt = f"""
        Based on the novel's subject, genre, and initial character concepts, identify:
        Novel Subject: {self.subject}
        Genre: {self.genre}
        Character Concepts: {json.dumps(self.characters, indent=2)}
        World Atmosphere: {self.world_details.get("atmosphere", "N/A")}

        1.  CORE THEMES (2-4): [List abstract concepts the story will explore, e.g., "Loss and Memory," "Identity vs. Society." Provide a brief (1-sentence) explanation for each, linking it to the context.]
        2.  RECURRING MOTIFS (3-5): [List concrete symbols, objects, phrases, or imagery. e.g., "A cracked pocket watch," "The phrase 'shadows remember'."]
        """
        themes_motifs_text = self._ollama_generate(themes_prompt, themes_system_prompt, temperature=0.6)
        if "[OLLAMA" in themes_motifs_text:
            print(f"ERROR generating themes/motifs: {themes_motifs_text}")
        else:
            themes_match = re.search(r"CORE THEMES:\s*((?:-\s*.*|\d\.\s*.*(?:\n|$))+)", themes_motifs_text, re.IGNORECASE | re.DOTALL)
            if themes_match:
                self.themes_motifs["themes"] = [theme.strip("- ").strip().lstrip("123456789. ") for theme in themes_match.group(1).strip().split('\n') if theme.strip()]

            motifs_match = re.search(r"RECURRING MOTIFS:\s*((?:-\s*.*|\d\.\s*.*(?:\n|$))+)", themes_motifs_text, re.IGNORECASE | re.DOTALL)
            if motifs_match:
                self.themes_motifs["motifs"] = [motif.strip("- ").strip().lstrip("123456789. ") for motif in motifs_match.group(1).strip().split('\n') if motif.strip()]
            print(f"Generated Themes: {self.themes_motifs['themes']}")
            print(f"Generated Motifs: {self.themes_motifs['motifs']}")

        # 4. High-Level Plot Outline & Determine Number of Chapters
        print("\nStep 1.4: Generating High-Level Plot Outline and Determining Chapter Count...")
        plot_system_prompt = f"You are a master storyteller, outlining engaging plots for {self.genre} novels in the style of {self.author_style}."
        plot_prompt = f"""
        Create a high-level plot outline for a novel based on the subject, characters, world, and themes.
        Novel Subject: {self.subject}
        Genre: {self.genre}
        Characters: {json.dumps(self.characters, indent=2)}
        World: {json.dumps(self.world_details, indent=2)}
        Themes: {json.dumps(self.themes_motifs["themes"])}

        The outline should follow a classic narrative structure (e.g., Three-Act Structure: Setup, Confrontation, Resolution).
        Describe:
        -   **Act I (Setup):** Introduction of protagonist, inciting incident, establishment of conflict and stakes.
        -   **Act II (Confrontation):** Rising action, character development through trials, introduction of key allies/antagonists, major turning points/complications.
        -   **Act III (Resolution):** Climax, falling action, resolution of main conflict, and thematic conclusion.
        
        For each Act, provide a 2-4 sentence summary of its key developments and objectives.
        This is a HIGH-LEVEL outline. Detailed chapter breakdowns will come next.
        Focus on a clear narrative trajectory and logical progression.

        IMPORTANT: Based on the scope and complexity of the plot you outline, conclude your response with a line specifying the suggested number of chapters for this novel. Use the format:
        SUGGESTED_CHAPTER_COUNT: [Number] (e.g., SUGGESTED_CHAPTER_COUNT: 20)
        This number should be reasonable for a novel of this nature, typically between 10 and 35 chapters.
        """
        self.plot_outline = self._ollama_generate(plot_prompt, plot_system_prompt, temperature=0.7)
        
        if "[OLLAMA" in self.plot_outline:
            print(f"ERROR generating plot outline: {self.plot_outline}")
            print("Defaulting number of chapters to 15 due to error.")
            self.num_chapters = 15
        else:
            print("Generated High-Level Plot Outline:")
            # Only print the plot part, not the chapter count line for cleaner display
            plot_display = re.sub(r"SUGGESTED_CHAPTER_COUNT:\s*\d+", "", self.plot_outline, flags=re.IGNORECASE).strip()
            print(plot_display)
            
            suggested_chapters_match = re.search(r"SUGGESTED_CHAPTER_COUNT:\s*(\d+)", self.plot_outline, re.IGNORECASE)
            if suggested_chapters_match:
                try:
                    self.num_chapters = int(suggested_chapters_match.group(1))
                    if not (10 <= self.num_chapters <= 50): # Reasonable range for a novel
                        print(f"Warning: LLM suggested {self.num_chapters} chapters, which is outside the typical 10-50 range. Clamping to 15.")
                        self.num_chapters = 15 
                    else:
                        print(f"LLM suggested {self.num_chapters} chapters for the novel.")
                except ValueError:
                    print("Warning: Could not parse suggested chapter count as integer. Defaulting to 15 chapters.")
                    self.num_chapters = 15 
            else:
                print("Warning: Could not find suggested chapter count in plot outline. Defaulting to 15 chapters.")
                self.num_chapters = 15
        
        if self.num_chapters < 3: # Ensure minimum for structure
            print(f"Warning: Number of chapters determined ({self.num_chapters}) is less than 3. Setting to 3.")
            self.num_chapters = 3
        
        return not ("[OLLAMA" in character_profiles_text or "[OLLAMA" in world_details_text or \
                    "[OLLAMA" in themes_motifs_text or "[OLLAMA" in self.plot_outline)


    # --- Phase 2: Detailed Planning Method ---
    def generate_detailed_chapter_plans(self):
        """
        Expands the high-level plot outline into detailed plans for each chapter.
        """
        print("\n--- Generating Detailed Chapter-by-Chapter Plans ---")
        if not self.plot_outline or not self.characters or not self.world_details or self.num_chapters == 0:
            print("ERROR: Cannot generate chapter plans without foundational elements (plot, characters, world, num_chapters).")
            return False

        system_prompt = f"You are a meticulous plot architect, detailing chapter structures for a {self.genre} novel in the style of {self.author_style}."
        character_summary_for_prompt = "\n".join([f"- {name} ({data.get('role', 'N/A')}): Motivations: {data.get('motivation', 'N/A')}. Arc: {data.get('arc_summary', 'N/A')}" for name, data in self.characters.items()])
        world_summary_for_prompt = f"World Name: {self.world_details.get('name', 'N/A')}\nKey Locations: {', '.join(self.world_details.get('key_locations',[]))}\nCultural Elements: {', '.join(self.world_details.get('cultural_elements',[]))}\nAtmosphere: {self.world_details.get('atmosphere', 'N/A')}"
        themes_for_prompt = f"Core Themes: {', '.join(self.themes_motifs.get('themes',[]))}\nRecurring Motifs: {', '.join(self.themes_motifs.get('motifs',[]))}"

        detailed_plan_prompt = f"""
        Novel Subject: {self.subject}
        High-Level Plot Outline:
        {self.plot_outline} 

        Character Summaries:
        {character_summary_for_prompt}

        World Summary:
        {world_summary_for_prompt}

        Themes & Motifs:
        {themes_for_prompt}

        Total Chapters to Plan: {self.num_chapters}

        Based on ALL the information above, create a VERY DETAILED plan for EACH chapter (1 through {self.num_chapters}).
        For EACH chapter, provide STRICTLY the following, in this order and clearly labeled:

        Chapter [Number] - [Evocative Title for this Chapter]:
        1.  CHAPTER GOAL: [A single paragraph (50-80 words) stating the primary narrative goal this chapter needs to achieve in the overall story and how it impacts the main character's arc or the central conflict.]
        2.  KEY SCENES (3-6 scenes): [Bulleted list. Each scene: "- Scene X: [Brief description of action/dialogue/internal monologue], Location: [Specific location from World Details or new minor one], Characters Involved: [List characters present & active], Key Revelation/Turning Point/Outcome: [What changes, is learned, or is achieved? How does it advance the plot or character?]]
        3.  CHARACTER DEVELOPMENT FOCUS: [For key characters appearing: How do their motivations, relationships, knowledge, or understanding change in THIS chapter? Be specific. e.g., "Jessica: Confronts her fear of X, strengthening her resolve but creating friction with Y. Learns Z about her past."]
        4.  PLOT ADVANCEMENT: [Specific ways the main plot and any subplots move forward. What new questions are raised or old ones answered? How does this chapter build on the previous and set up the next?]
        5.  TIMELINE & PACING: [e.g., "This chapter takes place over a few hours the next day.", "Pacing: Fast, with building tension.", "Spans one week, slower reflective pace initially, then accelerates."]
        6.  EMOTIONAL TONE (End of Chapter): [e.g., "Hopeful but wary," "Tense and suspenseful," "Melancholy and reflective," "Ominous and foreboding."]
        7.  CONNECTION TO NEXT CHAPTER (Setup/Hook): [Explicitly state 1-2 elements, questions, cliffhangers, or character decisions that directly lead into the next chapter's planned events or themes.]

        This detailed plan must ensure logical progression, character consistency, integration of themes/motifs, and effective pacing.
        A character's status (location, knowledge, emotional state) at the end of one chapter MUST be the starting point for the next.
        Ensure the plans for later chapters logically follow from the resolutions and developments of earlier ones.
        The first chapter should introduce the protagonist and the inciting incident. The final chapters should bring the main plot threads to a climax and resolution.
        Do not include the "SUGGESTED_CHAPTER_COUNT" line in this response.
        """
        print(f"Generating detailed plan text for {self.num_chapters} chapters (this may take a while)...")
        full_detailed_plan_text = self._ollama_generate(detailed_plan_prompt, system_prompt, temperature=0.65)

        if "[OLLAMA" in full_detailed_plan_text:
            print(f"ERROR generating detailed chapter plans: {full_detailed_plan_text}")
            return False

        current_chapter_num = 0
        # Split by "Chapter X -" pattern, ensuring it's at the start of a line (^)
        chapter_blocks = re.split(r"(?=^Chapter\s*\d+\s*[:-])", full_detailed_plan_text, flags=re.MULTILINE | re.IGNORECASE)
        
        for block in chapter_blocks:
            block = block.strip()
            if not block: continue

            title_match = re.match(r"Chapter\s*(\d+)\s*[:*-]?\s*(.*?)(?=\n\s*1\.\s*CHAPTER GOAL:|\n|$)", block, re.IGNORECASE) # Ensure title is followed by goal
            if title_match:
                current_chapter_num = int(title_match.group(1))
                chapter_title = title_match.group(2).strip()
                if not chapter_title: chapter_title = f"Chapter {current_chapter_num} (Untitled)"
                
                # Only add if chapter number is within expected range
                if current_chapter_num > self.num_chapters:
                    print(f"Warning: Parsed plan for Chapter {current_chapter_num}, which exceeds expected {self.num_chapters} chapters. Skipping.")
                    continue

                self.chapter_plans[current_chapter_num] = {"number": current_chapter_num, "title": chapter_title}
                plan_details = self.chapter_plans[current_chapter_num]

                # More robust parsing for each section, looking for the numbered headers
                goal_match = re.search(r"1\.\s*CHAPTER GOAL:\s*(.*?)(?=\n\s*(?:2\.\s*KEY SCENES:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["goal"] = goal_match.group(1).strip() if goal_match else "N/A"

                scenes_match = re.search(r"2\.\s*KEY SCENES:\s*(.*?)(?=\n\s*(?:3\.\s*CHARACTER DEVELOPMENT FOCUS:|#|$))", block, re.IGNORECASE | re.DOTALL)
                if scenes_match:
                    scenes_text = scenes_match.group(1).strip()
                    # Improved scene parsing: handles multiline scene descriptions better
                    plan_details["scenes"] = [s.strip() for s in re.split(r'\n\s*(?=-\s*Scene|\*\s*Scene|\d\.\s*Scene)', scenes_text) if s.strip()]
                    # Clean up individual scene lines if they still have list markers
                    plan_details["scenes"] = [re.sub(r'^(-\s*Scene\s*\d*:?\s*|-\s*|\*\s*Scene\s*\d*:?\s*|\*\s*|\d\.\s*Scene\s*\d*:?\s*|\d\.\s*)', '', s).strip() for s in plan_details["scenes"]]

                else: plan_details["scenes"] = []
                
                char_dev_match = re.search(r"3\.\s*CHARACTER DEVELOPMENT FOCUS:\s*(.*?)(?=\n\s*(?:4\.\s*PLOT ADVANCEMENT:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["character_development"] = char_dev_match.group(1).strip() if char_dev_match else "N/A"

                plot_adv_match = re.search(r"4\.\s*PLOT ADVANCEMENT:\s*(.*?)(?=\n\s*(?:5\.\s*TIMELINE & PACING:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["plot_advancement"] = plot_adv_match.group(1).strip() if plot_adv_match else "N/A"

                timeline_match = re.search(r"5\.\s*TIMELINE & PACING:\s*(.*?)(?=\n\s*(?:6\.\s*EMOTIONAL TONE \(End of Chapter\):|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["timeline_pacing"] = timeline_match.group(1).strip() if timeline_match else "N/A"

                emotion_match = re.search(r"6\.\s*EMOTIONAL TONE \(End of Chapter\):\s*(.*?)(?=\n\s*(?:7\.\s*CONNECTION TO NEXT CHAPTER:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["emotional_tone_end"] = emotion_match.group(1).strip() if emotion_match else "N/A"

                connection_match = re.search(r"7\.\s*CONNECTION TO NEXT CHAPTER:\s*(.*?)(?=$|\n\s*#)", block, re.IGNORECASE | re.DOTALL) # Use # as a potential end marker too
                plan_details["connection_to_next"] = connection_match.group(1).strip() if connection_match else "N/A"
        
        if not self.chapter_plans or len(self.chapter_plans) == 0:
            print("ERROR: No chapter plans were successfully parsed. The LLM output might not conform to the expected structure.")
            print("--- LLM Output Start (first 2000 chars) ---")
            print(full_detailed_plan_text[:2000]) 
            print("--- LLM Output End ---")
            return False

        print(f"Successfully parsed detailed plans for {len(self.chapter_plans)} chapters out of {self.num_chapters} expected.")
        if len(self.chapter_plans) != self.num_chapters:
             print(f"WARNING: Mismatch in parsed plans ({len(self.chapter_plans)}) and expected chapters ({self.num_chapters}). Review LLM output for plan structure.")
        return True

    # --- Phase 3: Prose Generation Loop ---
    def _get_continuity_context_for_chapter(self, chapter_num):
        """
        Gathers all relevant context from previous chapters and plans for generating the current chapter.
        """
        context = f"Overall Novel Subject: {self.subject}\n"
        context += f"Author Style: {self.author_style}, Genre: {self.genre}\n"
        context += f"High-Level Plot Outline:\n{self.plot_outline}\n\n" # Already cleaned from SUGGESTED_CHAPTER_COUNT
        context += f"World Details: {json.dumps(self.world_details)}\n"
        context += f"Themes & Motifs: {json.dumps(self.themes_motifs)}\n\n"

        context += "Character Profiles & Current Status (as of start of this chapter):\n"
        for name, data in self.characters.items():
            context += f"- {name} ({data.get('role','N/A')}):\n"
            context += f"  Description: {data.get('description','N/A')}\n"
            context += f"  Motivations: {data.get('motivation','N/A')}. Initial Arc: {data.get('arc_summary','N/A')}\n"
            context += f"  Current Status: {data.get('current_status','unknown')}, Location: {data.get('current_location','unknown')}, Emotion: {data.get('emotional_state','unknown')}\n"
            context += f"  Known Facts: {', '.join(data.get('knowledge',[]))}\n"
            if data.get('development_log'):
                relevant_logs = [log for log in data['development_log'] if log['chapter'] < chapter_num]
                if relevant_logs:
                    last_dev = relevant_logs[-1]
                    context += f"  Last Noted Development (Ch {last_dev['chapter']}): {last_dev.get('summary', 'N/A')}\n"
        context += "\n"

        if chapter_num > 1:
            prev_chap_num = chapter_num - 1
            prev_continuity = self.chapter_continuity_data.get(prev_chap_num, {})
            prev_plan = self.chapter_plans.get(prev_chap_num, {})
            context += f"Summary of Previous Chapter ({prev_chap_num} - '{prev_plan.get('title', 'Untitled')}'):\n"
            context += f"{prev_continuity.get('summary', 'N/A')}\n"
            context += f"Ended with Emotional Tone: {prev_continuity.get('emotional_tone_end_achieved_in_summary', 'N/A')}\n" # Use the one from continuity
            context += f"Timeline at end of Ch {prev_chap_num}: {prev_continuity.get('timeline_end', 'N/A')}\n"
            context += f"Hook for current chapter (from prev chapter's plan): {prev_plan.get('connection_to_next', 'N/A')}\n\n"
        
        return context

    def _generate_chapter_opener(self, chapter_num, current_chapter_plan):
        """Generates the opening paragraph(s) for the current chapter."""
        chapter_title_line = f"Chapter {chapter_num} - {current_chapter_plan.get('title', 'Untitled')}"
        # The chapter_prose in generate_novel_content will start with this title line.
        # The actual opener text will be generated and appended.
        
        if chapter_num == 1:
            # For chapter 1, the "opener" is just the beginning of the story, so we don't need a special transition prompt.
            # The first scene's generation will handle the actual opening narrative.
            # We return just the title part, which will be prefixed to the first scene's prose.
            return f"{chapter_title_line}\n\n" 

        prev_chap_num = chapter_num - 1
        prev_continuity = self.chapter_continuity_data.get(prev_chap_num, {})
        prev_plan = self.chapter_plans.get(prev_chap_num, {})
        
        system_prompt = f"You are a novelist in the style of {self.author_style}, skilled at crafting chapter openings that immediately re-orient the reader and smoothly transition from the previous chapter's ending."
        prompt = f"""
        You are writing the VERY FIRST paragraph(s) for Chapter {chapter_num}, titled "{current_chapter_plan.get('title', 'Untitled')}".
        This opening must seamlessly connect to the end of Chapter {prev_chap_num} and establish the immediate context for Chapter {chapter_num}.

        Context from End of Chapter {prev_chap_num} ('{prev_plan.get('title', 'Untitled')}'):
        - Summary of Ch {prev_chap_num}: {prev_continuity.get('summary', 'Previously...')[-1000:]}
        - Actual Ending Hook/Transition Text from Ch {prev_chap_num}: "{prev_continuity.get('ending_hook_text', 'The previous chapter ended.')}"
        - Emotional Tone at End of Ch {prev_chap_num}: {prev_continuity.get('emotional_tone_end_achieved_in_summary', 'Neutral')}
        - Timeline at End of Ch {prev_chap_num}: {prev_continuity.get('timeline_end', 'Unknown')}
        - Key Character States at End of Ch {prev_chap_num}: {json.dumps({name: {'status': data['current_status'], 'location': data['current_location'], 'emotion': data['emotional_state']} for name, data in self.characters.items()}, indent=2)}

        Plan for Current Chapter ({chapter_num} - "{current_chapter_plan.get('title', 'Untitled')}"):
        - Goal: {current_chapter_plan.get('goal', 'The story progresses.')}
        - First Planned Scene Hint: {current_chapter_plan.get('scenes', ['A new scene begins.'])[0] if current_chapter_plan.get('scenes') else 'A new scene begins.'}
        - Timeline & Pacing: {current_chapter_plan.get('timeline_pacing', 'As expected')}

        Write 1-2 compelling opening paragraphs (approx 100-200 words) for Chapter {chapter_num}. These paragraphs should:
        1. Directly acknowledge or subtly resolve the immediate hook/question left by Chapter {prev_chap_num}'s actual ending hook.
        2. Establish the setting, characters present, and the time elapsed since the previous chapter (if any significant passage, use the Timeline Indicators from this chapter's plan).
        3. Set the initial tone for Chapter {chapter_num}, which might be a continuation or a shift from the previous chapter's end.
        4. Orient the reader quickly without extensive exposition. Focus on "showing" the new situation.
        5. Make the transition feel natural and engaging, drawing the reader into the new chapter.
        6. This is the *opening text* of Chapter {chapter_num}. Do NOT repeat the chapter title.

        Opening paragraph(s) for Chapter {chapter_num}:
        """
        opener_text = self._ollama_generate(prompt, system_prompt, temperature=0.68)
        # Prepend the title line to the generated opener text
        return f"{chapter_title_line}\n\n{opener_text}\n\n"


    def _generate_scene_prose(self, chapter_num, scene_index, scene_description, current_chapter_plan, continuity_context, previous_scene_prose=""):
        """Generates prose for a single scene within a chapter."""
        system_prompt = f"You are a celebrated novelist in the style of {self.author_style}, writing a {self.genre} novel. Your prose is vivid, emotionally resonant, and drives the plot forward. You excel at 'showing, not telling' and making fantastical elements relatable."
        
        motif_to_weave = "N/A"
        if self.themes_motifs.get("motifs") and len(self.themes_motifs["motifs"]) > 0 :
            motif_to_weave = self.themes_motifs["motifs"][(chapter_num + scene_index -1) % len(self.themes_motifs["motifs"])]

        prompt = f"""
        {continuity_context}

        Current Chapter Plan ({chapter_num} - "{current_chapter_plan.get('title', 'Untitled')}"):
        - Chapter Goal: {current_chapter_plan.get('goal', 'N/A')}
        - Full Scene Breakdown for this chapter: {json.dumps(current_chapter_plan.get('scenes',[]))}
        - Character Development Focus for this chapter: {current_chapter_plan.get('character_development', 'N/A')}
        - Plot Advancement for this chapter: {current_chapter_plan.get('plot_advancement', 'N/A')}
        - Planned Emotional Tone (End of Chapter): {current_chapter_plan.get('emotional_tone_end', 'N/A')}
        - Timeline & Pacing for this chapter: {current_chapter_plan.get('timeline_pacing', 'N/A')}
        
        Prose written SO FAR in THIS CHAPTER (before this scene):
        ---
        {previous_scene_prose[-2000:] if previous_scene_prose else "This is the first scene of the chapter (after the chapter title and opener)."} 
        ---

        YOUR TASK: Write the narrative prose for THE FOLLOWING SPECIFIC SCENE ONLY.
        Scene {scene_index + 1} Description (from chapter plan): "{scene_description}"

        INSTRUCTIONS FOR THIS SCENE:
        1.  **Narrative Focus:** Write approximately 300-700 words (adjust based on scene importance) bringing THIS SCENE to life.
        2.  **Style & Tone:** Adhere to {self.author_style}'s style. Maintain the chapter's intended pacing and build towards its emotional goal.
        3.  **"Show, Don't Tell":** Demonstrate emotions, thoughts, and plot points through actions, dialogue, sensory details, and internal monologues. Make character reactions and decisions clear through their behavior.
        4.  **Relatability & Emotional Depth:** Even in fantastical situations, ground character experiences in relatable emotions. Describe what things *feel* like.
        5.  **Sensory Details:** Weave in vivid sensory details (sight, sound, smell, touch, taste) consistent with the world and scene location.
        6.  **Character Consistency:** Ensure characters act and speak consistently with their established profiles, motivations, current emotional state, and knowledge. Reflect their planned development for this chapter.
        7.  **Dialogue:** If dialogue is part of the scene, make it natural, character-specific, and purposeful (revealing character, advancing plot, or building tension).
        8.  **Motif Integration:** Subtly weave in the recurring motif: '{motif_to_weave}' if it fits naturally within this scene's events. Do not force it.
        9.  **Continuity:** Ensure this scene flows logically from any previous prose in this chapter. Do NOT repeat information.
        10. **Output:** Generate ONLY the newly written narrative paragraphs for THIS SCENE. Do not add scene numbers or headings.

        Begin Scene {scene_index + 1} prose now:
        """
        scene_prose = self._ollama_generate(prompt, system_prompt, temperature=0.72, top_p=0.92)
        return scene_prose

    def _analyze_inter_chapter_flow(self, previous_chapter_num, current_chapter_num, current_chapter_opening_text):
        """
        Analyzes the narrative flow between the end of the previous chapter and the start of the current chapter.
        """
        print(f"  Analyzing flow from Chapter {previous_chapter_num} to Chapter {current_chapter_num}...")
        if previous_chapter_num not in self.chapter_continuity_data or \
           previous_chapter_num not in self.chapter_plans or \
           current_chapter_num not in self.chapter_plans:
            print("    Skipping flow analysis: Missing data for previous or current chapter.")
            return "Flow analysis skipped due to missing data."

        prev_continuity = self.chapter_continuity_data[previous_chapter_num]
        prev_plan = self.chapter_plans[previous_chapter_num]
        current_plan = self.chapter_plans[current_chapter_num]

        prev_summary = prev_continuity.get("summary", "N/A")
        prev_hook = prev_continuity.get("ending_hook_text", "N/A") # This is the actual generated hook
        prev_title = prev_plan.get("title", "Untitled Previous Chapter")
        
        current_title = current_plan.get("title", "Untitled Current Chapter")
        current_goal = current_plan.get("goal", "N/A")

        # Extract just the opening paragraphs from the current chapter's text
        # current_chapter_opening_text already contains title + opener. We need just the opener part.
        opening_paragraphs_only = "\n".join(current_chapter_opening_text.split('\n\n')[1:]).strip() # Skip title line, join paragraphs

        system_prompt = "You are an expert literary editor specializing in narrative coherence and flow between chapters."
        prompt = f"""
        Analyze the transition and flow from the end of Chapter {previous_chapter_num} ("{prev_title}") to the beginning of Chapter {current_chapter_num} ("{current_title}").

        CONTEXT:
        End of Chapter {previous_chapter_num} ("{prev_title}"):
        - Summary: {prev_summary[-1000:]}
        - Actual Ending Hook/Transition Text: "{prev_hook}"
        - Key Character States at End (from continuity): {json.dumps({name: {'status': data['current_status'], 'location': data['current_location'], 'emotion': data['emotional_state']} for name, data in self.characters.items() if data['first_appearance_chapter'] <= previous_chapter_num and data['first_appearance_chapter'] > 0}, indent=2)}
        - Planned Connection from Ch {previous_chapter_num} to Ch {current_chapter_num}: "{prev_plan.get('connection_to_next', 'N/A')}"


        Beginning of Chapter {current_chapter_num} ("{current_title}"):
        - Chapter Goal: {current_goal}
        - Actual Opening Paragraph(s): "{opening_paragraphs_only}"

        EVALUATION TASK:
        Provide a brief (1-2 paragraph) analysis covering these points:
        1.  **Hook Resolution:** How well does the opening of Chapter {current_chapter_num} address or follow up on the ending hook/transition of Chapter {previous_chapter_num}?
        2.  **Plot Continuity:** Is there a logical and clear progression of plot events or situation from the end of the previous chapter to the start of the current one?
        3.  **Character Consistency:** Do the characters' states (emotional, physical, location, knowledge) seem consistent and logically follow from the previous chapter's end into the current chapter's opening?
        4.  **Tone & Pacing:** Is the transition in emotional tone and pacing smooth and appropriate, or jarring?
        5.  **Overall Coherence:** How effective is the overall flow between these two chapters? Does it feel natural and engaging for the reader?
        
        Be concise and constructive.
        Flow Analysis:
        """
        flow_analysis_text = self._ollama_generate(prompt, system_prompt, temperature=0.5)
        print(f"    Flow Analysis Result: {flow_analysis_text[:200]}...") # Print a snippet

        # Store the analysis
        if current_chapter_num not in self.chapter_continuity_data:
            self.chapter_continuity_data[current_chapter_num] = {}
        self.chapter_continuity_data[current_chapter_num]["flow_analysis_from_previous"] = flow_analysis_text
        
        return flow_analysis_text


    def _update_chapter_continuity_data(self, chapter_num, full_chapter_content, is_final_pass_for_chapter=False):
        """
        Analyzes generated chapter content to update continuity data (summary, character states, timeline, emotion).
        This is CRITICAL for informing the next chapter's generation.
        """
        print(f"Updating continuity data for Chapter {chapter_num} ({'final pass' if is_final_pass_for_chapter else 'interim pass'})...")
        entry = self.chapter_continuity_data.get(chapter_num, {}) 

        summary_system_prompt = "You are a literary analyst. Your task is to summarize chapter content accurately and concisely for continuity purposes."
        summary_prompt = f"""
        Create a concise yet detailed summary of the following chapter content (Chapter {chapter_num}).
        Focus on all key plot events, character actions and significant development, setting details, revelations, emotional shifts, and how it ends.
        This summary will be crucial context for writing the NEXT chapter.
        CHAPTER CONTENT (Chapter {chapter_num} - "{self.chapter_plans.get(chapter_num, {}).get('title', 'Untitled')}"):
        ---
        {full_chapter_content}
        ---
        Detailed Summary of Chapter {chapter_num}:
        """
        entry["summary"] = self._ollama_generate(summary_prompt, summary_system_prompt, temperature=0.5)

        if is_final_pass_for_chapter: # Character updates only on final pass
            active_chars_in_chapter = []
            current_chapter_plan = self.chapter_plans.get(chapter_num, {})
            if current_chapter_plan:
                for scene_desc in current_chapter_plan.get("scenes", []):
                    char_involved_match = re.search(r"Characters Involved:\s*(.*?)(?:\.\s*Key Revelation|$)", scene_desc, re.IGNORECASE)
                    if char_involved_match:
                        names_str = char_involved_match.group(1)
                        potential_names = re.split(r'[,\s]+and\s+|\s*,\s*|[,\s]+with\s+', names_str)
                        for char_name_candidate in potential_names:
                            clean_name = char_name_candidate.strip().rstrip('.').strip()
                            if clean_name and clean_name in self.characters and clean_name not in active_chars_in_chapter:
                                active_chars_in_chapter.append(clean_name)
                char_dev_focus = current_chapter_plan.get("character_development", "")
                for char_name in self.characters.keys():
                    if re.search(r'\b' + re.escape(char_name) + r'\b', char_dev_focus, re.IGNORECASE) and char_name not in active_chars_in_chapter:
                         active_chars_in_chapter.append(char_name)
            
            if not active_chars_in_chapter: active_chars_in_chapter = list(self.characters.keys())

            char_update_system_prompt = "You are a narrative continuity expert. Update character states based on chapter events."
            char_update_prompt = f"""
            Based on the FULL content of Chapter {chapter_num} below, update the status for EACH listed character.
            Chapter {chapter_num} ("{self.chapter_plans.get(chapter_num, {}).get('title', 'Untitled')}") Content:
            ---
            {full_chapter_content}
            ---
            For EACH of these characters who appeared or were central: {', '.join(active_chars_in_chapter) if active_chars_in_chapter else "Summarize general impact if no specific characters listed."}
            Provide the following updates. If a character did not appear or had no significant change for a field, state "No change" or "Did not appear."
            
            CHARACTER NAME: [Character's Name]
            -   STATUS CHANGE: [e.g., "Remains alive," "Injured (twisted ankle)," "Captured," "Learned X," "Decided Y."]
            -   LOCATION AT END OF CHAPTER: [Specific location.]
            -   EMOTIONAL STATE AT END OF CHAPTER: [e.g., "Hopeful," "Grieving," "Determined," "Suspicious."]
            -   KEY DEVELOPMENT/ACTION: [Significant actions, learnings, or internal changes.]
            -   RELATIONSHIP CHANGES: [e.g., "Strained with Z over X," "New alliance with W."]
            -   NEW KNOWLEDGE/SECRETS ACQUIRED: [New critical info, clues, secrets.]
            
            Format clearly for each character.
            """
            character_updates_text = self._ollama_generate(char_update_prompt, char_update_system_prompt, temperature=0.55)
            entry["character_updates_text"] = character_updates_text
            
            current_char_name_update = None
            parsed_updates_for_log = {} 
            for line in character_updates_text.split('\n'):
                line = line.strip()
                name_match = re.match(r"CHARACTER NAME:\s*(.*)", line, re.IGNORECASE)
                if name_match:
                    if current_char_name_update and parsed_updates_for_log and current_char_name_update in self.characters:
                        self.characters[current_char_name_update]["development_log"].append(
                            {"chapter": chapter_num, "summary": "Updates from chapter events", **parsed_updates_for_log}
                        )
                    current_char_name_update = name_match.group(1).strip()
                    parsed_updates_for_log = {} 
                    if current_char_name_update not in self.characters:
                        current_char_name_update = None 
                    elif self.characters[current_char_name_update].get("first_appearance_chapter", 0) == 0:
                         self.characters[current_char_name_update]["first_appearance_chapter"] = chapter_num
                    continue

                if current_char_name_update and current_char_name_update in self.characters:
                    char_obj = self.characters[current_char_name_update]
                    def get_value(text, key_phrase):
                        if text.upper().startswith(key_phrase.upper()):
                            val = text.split(":",1)[1].strip()
                            if val.lower() in ["no change", "did not appear", "n/a", "none", "no significant change."]:
                                return None 
                            return val
                        return "NO_MATCH" 

                    status_val = get_value(line, "- STATUS CHANGE:")
                    if status_val != "NO_MATCH": 
                        if status_val is not None: char_obj["current_status"] = status_val
                        parsed_updates_for_log["status"] = status_val if status_val is not None else "No change"
                        continue
                    loc_val = get_value(line, "- LOCATION AT END OF CHAPTER:")
                    if loc_val != "NO_MATCH":
                        if loc_val is not None: char_obj["current_location"] = loc_val
                        parsed_updates_for_log["location"] = loc_val if loc_val is not None else "No change"
                        continue
                    emo_val = get_value(line, "- EMOTIONAL STATE AT END OF CHAPTER:")
                    if emo_val != "NO_MATCH":
                        if emo_val is not None: char_obj["emotional_state"] = emo_val
                        parsed_updates_for_log["emotion"] = emo_val if emo_val is not None else "No change"
                        continue
                    dev_val = get_value(line, "- KEY DEVELOPMENT/ACTION:")
                    if dev_val != "NO_MATCH":
                        parsed_updates_for_log["development"] = dev_val if dev_val is not None else "No change"
                        continue
                    rel_val = get_value(line, "- RELATIONSHIP CHANGES:")
                    if rel_val != "NO_MATCH":
                        parsed_updates_for_log["relationships_changed"] = rel_val if rel_val is not None else "No change"
                        continue
                    know_val = get_value(line, "- NEW KNOWLEDGE/SECRETS ACQUIRED:")
                    if know_val != "NO_MATCH":
                        if know_val is not None and know_val not in char_obj["knowledge"]: 
                            char_obj["knowledge"].append(know_val)
                        parsed_updates_for_log["new_knowledge"] = know_val if know_val is not None else "No change"
                        continue
            if current_char_name_update and parsed_updates_for_log and current_char_name_update in self.characters: 
                self.characters[current_char_name_update]["development_log"].append(
                    {"chapter": chapter_num, "summary": "Updates from chapter events", **parsed_updates_for_log}
                )

        timeline_system_prompt = "You are a temporal analyst for narratives. Extract timeline information precisely."
        timeline_prompt = f"""
        Analyze Chapter {chapter_num}'s content for timeline information:
        Chapter {chapter_num} ("{self.chapter_plans.get(chapter_num, {}).get('title', 'Untitled')}") Content:
        ---
        {full_chapter_content}
        ---
        Determine:
        1.  APPROXIMATE TIME ELAPSED DURING THIS CHAPTER: [e.g., "Several hours," "One full day," "A week."]
        2.  TIME OF DAY/DATE AT THE END OF THIS CHAPTER: [e.g., "Evening of the third day," "Midnight," "Following morning."]
        3.  ANY SPECIFIC TIME MARKERS MENTIONED: [e.g., "After dawn," "Two moons passed," "Clock struck three."]
        Reply in format:
        ELAPSED: [answer]
        END_TIME: [answer]
        MARKERS: [answer]
        """
        timeline_text = self._ollama_generate(timeline_prompt, timeline_system_prompt, temperature=0.4)
        entry["timeline_elapsed"] = re.search(r"ELAPSED:\s*(.*)", timeline_text, re.IGNORECASE).group(1).strip() if re.search(r"ELAPSED:\s*(.*)", timeline_text, re.IGNORECASE) else "N/A"
        entry["timeline_end"] = re.search(r"END_TIME:\s*(.*)", timeline_text, re.IGNORECASE).group(1).strip() if re.search(r"END_TIME:\s*(.*)", timeline_text, re.IGNORECASE) else "N/A"
        entry["timeline_markers"] = re.search(r"MARKERS:\s*(.*)", timeline_text, re.IGNORECASE).group(1).strip() if re.search(r"MARKERS:\s*(.*)", timeline_text, re.IGNORECASE) else "N/A"

        entry["emotional_tone_end_achieved_in_summary"] = entry["summary"][-300:] 
        # Initialize ending_hook_text if not present, especially for the last chapter
        if "ending_hook_text" not in entry:
            entry["ending_hook_text"] = "N/A (Last chapter or hook not generated)"


        self.chapter_continuity_data[chapter_num] = entry

    def _generate_chapter_transition_hook(self, chapter_num, current_chapter_content, current_chapter_continuity):
        """Generates the transition hook/paragraph(s) for the end of the current chapter."""
        if chapter_num >= self.num_chapters: return "" 

        next_chap_num = chapter_num + 1
        next_chapter_plan = self.chapter_plans.get(next_chap_num, {})
        if not next_chapter_plan:
            print(f"Warning: No plan found for Chapter {next_chap_num} to generate hook from Chapter {chapter_num}.")
            return "\n\n(The story continues...)"

        system_prompt = f"You are a master storyteller in the style of {self.author_style}, crafting suspenseful and engaging chapter endings that seamlessly lead into the next."
        prompt = f"""
        You are writing the VERY LAST paragraph(s) for Chapter {chapter_num} ("{self.chapter_plans.get(chapter_num,{}).get('title','Untitled')}").
        This transition must create anticipation for Chapter {next_chap_num} ("{next_chapter_plan.get('title','Untitled')}").

        End of Chapter {chapter_num} Context:
        - Current Chapter Summary (focus on ending events): {current_chapter_continuity.get('summary', 'The chapter concluded.')[-1000:]}
        - Emotional Tone at End (inferred from summary): {current_chapter_continuity.get('emotional_tone_end_achieved_in_summary', 'Neutral')}
        - Last ~500 characters of Chapter {chapter_num} (before this hook): "{current_chapter_content[-500:]}"
        - Key Character States at End of Ch {chapter_num}: {json.dumps({name: {'status': data['current_status'], 'location': data['current_location'], 'emotion': data['emotional_state']} for name, data in self.characters.items()}, indent=2)}

        Plan for NEXT Chapter ({next_chap_num} - "{next_chapter_plan.get('title', 'Untitled')}"):
        - Next Chapter Goal: {next_chapter_plan.get('goal', 'The story continues.')}
        - Next Chapter Likely Opening Scene/Focus: {next_chapter_plan.get('scenes', ['A new challenge arises.'])[0] if next_chapter_plan.get('scenes') else 'A new challenge arises.'}
        - Next Chapter Character Development Focus: {next_chapter_plan.get('character_development', 'Further growth.')}
        - Planned connection from current Ch {chapter_num} to next Ch {next_chap_num}: "{self.chapter_plans.get(chapter_num,{}).get('connection_to_next','N/A')}"

        Write 1-2 compelling paragraphs (approx 75-150 words) that will be the *final text* of Chapter {chapter_num}. This hook should:
        1.  Provide a sense of immediate closure for Chapter {chapter_num}'s main events but leave the reader wanting more.
        2.  Directly foreshadow, question, or set up the conflict, theme, or situation of Chapter {next_chap_num} based on its plan AND the planned connection.
        3.  Maintain or slightly shift the established emotional tone to build suspense, curiosity, or dread for what's next.
        4.  Avoid clichés. Be original and impactful. Ensure it feels like a natural continuation of the narrative, not an abrupt summary.
        5.  This is the *final text* of Chapter {chapter_num}.

        Final transition paragraph(s) for Chapter {chapter_num}:
        """
        hook_text = self._ollama_generate(prompt, system_prompt, temperature=0.75)
        # Store the generated hook in the current chapter's continuity data
        if chapter_num in self.chapter_continuity_data:
            self.chapter_continuity_data[chapter_num]["ending_hook_text"] = hook_text.strip()
        else: 
            self.chapter_continuity_data[chapter_num] = {"ending_hook_text": hook_text.strip()}
            
        return f"\n\n{hook_text}"


    def generate_novel_content(self):
        """
        Main loop to generate content for all chapters, applying coherence measures.
        """
        print("\n--- Generating Full Novel Content (Chapter by Chapter) ---")
        if not self.chapter_plans or self.num_chapters == 0:
            print("ERROR: Cannot generate novel content without detailed chapter plans or chapter count.")
            return False

        for i in range(1, self.num_chapters + 1):
            print(f"\n--- Generating Chapter {i} of {self.num_chapters} ---")
            current_chapter_plan = self.chapter_plans.get(i)
            if not current_chapter_plan:
                print(f"ERROR: No plan found for Chapter {i}. Skipping.")
                self.generated_chapters_content[i] = f"[ERROR: No plan found for Chapter {i}]"
                self.chapter_continuity_data[i] = {"summary": "Error: No plan.", "character_updates_text": "", "timeline_end": "Unknown", "emotional_tone_end_achieved_in_summary": "Error", "ending_hook_text": "", "flow_analysis_from_previous": "N/A"}
                continue

            continuity_context = self._get_continuity_context_for_chapter(i)
            
            # Generate chapter opener (includes title line)
            chapter_opener_text_with_title = self._generate_chapter_opener(i, current_chapter_plan)
            
            # Initialize chapter_prose with the opener 
            chapter_prose = chapter_opener_text_with_title 
            
            # Analyze flow from previous chapter to this chapter's opening
            if i > 1:
                self._analyze_inter_chapter_flow(i - 1, i, chapter_opener_text_with_title) 

            scenes = current_chapter_plan.get("scenes", [])
            if not scenes:
                print(f"Warning: No scenes defined in plan for Chapter {i}. Chapter might be short or only opener/hook.")
                if chapter_prose.strip() == chapter_opener_text_with_title.strip(): 
                     chapter_prose += "\n\n[This chapter's plan had no specific scenes. The narrative continues based on the chapter goal.]\n\n"
            else:
                accumulated_scene_prose_for_chapter = chapter_opener_text_with_title 
                for scene_idx, scene_desc in enumerate(scenes):
                    print(f"  Generating Scene {scene_idx + 1} of {len(scenes)} for Chapter {i}: {scene_desc[:80]}...")
                    scene_specific_prose = self._generate_scene_prose(i, scene_idx, scene_desc, current_chapter_plan, continuity_context, accumulated_scene_prose_for_chapter)
                    if "[OLLAMA" in scene_specific_prose: 
                         print(f"    ERROR generating scene {scene_idx+1}: {scene_specific_prose}")
                         scene_specific_prose = f"\n\n[Error generating scene: {scene_desc[:50]}...]\n\n"
                    
                    chapter_prose += scene_specific_prose + "\n\n" 
                    accumulated_scene_prose_for_chapter += scene_specific_prose + "\n\n" 
                    time.sleep(0.2) 

            # Interim continuity update (based on content BEFORE the hook)
            self._update_chapter_continuity_data(i, chapter_prose.strip(), is_final_pass_for_chapter=False)

            if i < self.num_chapters:
                print(f"  Generating transition hook for Chapter {i}...")
                hook_text = self._generate_chapter_transition_hook(i, chapter_prose.strip(), self.chapter_continuity_data[i])
                chapter_prose += hook_text
            
            self.generated_chapters_content[i] = chapter_prose.strip()
            print(f"  Chapter {i} ('{current_chapter_plan.get('title', 'Untitled')}') content generated (approx length: {len(chapter_prose)} chars).")

            # FINAL continuity update for the chapter (with opener, scenes, and hook included)
            self._update_chapter_continuity_data(i, self.generated_chapters_content[i], is_final_pass_for_chapter=True)

            if i < self.num_chapters:
                print("Pausing briefly before next chapter...")
                time.sleep(0.5) 
        
        return True
        
    # --- NEW: Transition Checking Phase ---
    def _check_and_improve_transition(self, prev_chapter_num, current_chapter_num):
        """Checks transition from prev to current chapter and improves if needed."""
        print(f"\n--- Checking transition from Chapter {prev_chapter_num} to {current_chapter_num} ---")
        
        prev_chapter_content = self.generated_chapters_content.get(prev_chapter_num)
        current_chapter_content = self.generated_chapters_content.get(current_chapter_num)

        if not prev_chapter_content or not current_chapter_content:
            print("  Skipping transition check: Missing content for one or both chapters.")
            return

        system_prompt = """You are a professional editor specializing in narrative flow and chapter transitions."""
        prompt = f"""Analyze the transition between the end of the previous chapter and the beginning of the current chapter.

        END OF PREVIOUS CHAPTER ({prev_chapter_num}):
        ---
        {prev_chapter_content[-1000:]} 
        ---

        BEGINNING OF CURRENT CHAPTER ({current_chapter_num}):
        ---
        {current_chapter_content[:1000]} 
        ---

        If the transition is already smooth and logical, respond ONLY with the exact text:
        TRANSITION: SMOOTH

        Otherwise, provide an improved beginning for the current chapter (first 1-3 paragraphs, approx 100-250 words) that:
        1. Creates a smoother, more logical connection with the previous chapter's ending hook/state.
        2. Avoids repeating information already established.
        3. Maintains character and plot consistency.
        4. Progresses the timeline naturally.
        5. Matches the established author style ({self.author_style}).

        Start your response with the exact text "TRANSITION: REVISED" followed by the revised beginning paragraphs. Do NOT include the chapter title in the revised text.
        """
        transition_check_result = self._ollama_generate(prompt, system_prompt, temperature=0.6)

        if "[OLLAMA" in transition_check_result:
            print(f"  Error during transition check: {transition_check_result}")
        elif "TRANSITION: REVISED" in transition_check_result:
            try:
                # Extract the revised beginning (text after "TRANSITION: REVISED")
                revised_beginning = transition_check_result.split("TRANSITION: REVISED", 1)[1].strip()
                
                if revised_beginning:
                    print(f"  Transition needs improvement. Applying revised opening to Chapter {current_chapter_num}.")
                    # Find the original chapter title line
                    original_lines = current_chapter_content.split('\n', 1)
                    original_title_line = original_lines[0]
                    
                    # Find the end of the original opening paragraphs (heuristic: look for the third double newline)
                    original_body = original_lines[1] if len(original_lines) > 1 else ""
                    parts = original_body.split('\n\n', 3) # Split into max 4 parts (title, opener, rest)
                    original_rest_of_chapter = parts[3] if len(parts) > 3 else (parts[1] if len(parts) == 2 else "") # Handle short chapters

                    # Reconstruct the chapter with the revised opening
                    self.generated_chapters_content[current_chapter_num] = f"{original_title_line}\n\n{revised_beginning}\n\n{original_rest_of_chapter}".strip()
                    print(f"  Chapter {current_chapter_num} opening revised successfully.")
                    # Optional: Could re-run continuity update here for the revised chapter
                    # self._update_chapter_continuity_data(current_chapter_num, self.generated_chapters_content[current_chapter_num], is_final_pass_for_chapter=True)

                else:
                    print("  Transition check indicated revision needed, but no revised text was provided by LLM.")
            except Exception as e:
                print(f"  Error applying revised transition for Chapter {current_chapter_num}: {e}")
        elif "TRANSITION: SMOOTH" in transition_check_result:
            print("  Transition is smooth. No changes needed.")
        else:
            print("  Transition check response was unclear. No changes applied.")


    def _perform_final_transition_checks(self):
        """Loops through all chapters to check and improve transitions."""
        print("\n--- Performing Final Pass: Checking Chapter Transitions ---")
        if len(self.generated_chapters_content) < 2:
            print("  Skipping transition checks (less than 2 chapters generated).")
            return

        for i in range(2, self.num_chapters + 1): # Start from chapter 2
            if i in self.generated_chapters_content and (i - 1) in self.generated_chapters_content:
                self._check_and_improve_transition(i - 1, i)
                time.sleep(0.5) # Small delay
            else:
                print(f"  Skipping transition check for Chapter {i} (missing previous or current chapter content).")
        print("--- Finished Final Transition Checks ---")


    # --- Phase 4: Compilation & Output Methods ---
    def generate_novel_title(self):
        """Generates a compelling title for the novel."""
        print("\n--- Generating Novel Title ---")
        if not self.subject or not self.characters or not self.themes_motifs:
            self.novel_title = f"A {self.genre.replace('/', ' ')} Story" # Sanitize genre for title
            print(f"Warning: Insufficient data for title generation. Using placeholder: {self.novel_title}")
            return

        system_prompt = f"You are a creative book title generator, expert in {self.genre} and the style of {self.author_style}."
        prompt = f"""
        Generate ONE compelling and marketable novel title based on the following details.
        The title must be highly consistent with the genre, author's style, subject, main character(s), world, and core themes.
        It should be intriguing, memorable, and not overly generic.

        Novel Subject: {self.subject}
        Genre: {self.genre}
        Author's Style Inspiration: {self.author_style}
        Main Character(s): {', '.join([f'{name} ({data.get("role")})' for name, data in self.characters.items()])}
        World Name/Atmosphere: {self.world_details.get('name', 'N/A')} / {self.world_details.get('atmosphere', 'N/A')}
        Core Themes: {', '.join(self.themes_motifs.get('themes', []))}
        Recurring Motifs: {', '.join(self.themes_motifs.get('motifs', []))}

        Return ONLY the generated title itself, without any quotation marks, labels (like "Title:"), or explanatory text.
        Novel Title:
        """
        title_text = self._ollama_generate(prompt, system_prompt, temperature=0.8)
        if "[OLLAMA" in title_text or not title_text.strip():
            print(f"ERROR generating title: {title_text}. Using placeholder.")
            main_char_name = list(self.characters.keys())[0] if self.characters else 'Adventure'
            self.novel_title = f"A {self.genre.replace('/', ' ')} Tale of {main_char_name}"
        else:
            # Clean title: remove potential prefixes and quotes
            title_text = re.sub(r'^(title|novel title):?\s*', '', title_text, flags=re.IGNORECASE).strip()
            self.novel_title = title_text.strip('"\'')
        print(f"Generated Novel Title: {self.novel_title}")


    def compile_and_save_novel(self):
        """Compiles the generated content into a .docx file and saves it."""
        print("\n--- Compiling and Saving Novel ---")
        if not self.generated_chapters_content:
            print("ERROR: No chapter content generated. Cannot save novel.")
            return

        self.generate_novel_title() 

        doc = Document()
        # Define styles (can be customized further)
        # Ensure these styles exist or handle potential errors if they don't
        try:
            title_style = doc.styles['Title']
            title_style.font.name = 'Garamond' 
            title_style.font.size = Pt(28)
        except KeyError:
            print("Warning: 'Title' style not found. Using default.")
            title_style = 'Title' # Fallback to string name
        
        try:
            heading1_style = doc.styles['Heading 1']
            heading1_style.font.name = 'Garamond'
            heading1_style.font.size = Pt(18)
        except KeyError:
            print("Warning: 'Heading 1' style not found. Using default.")
            heading1_style = 'Heading 1' # Fallback

        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Garamond'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.space_after = Pt(0) # No extra space after paragraphs
            normal_style.paragraph_format.first_line_indent = Pt(24) # Indent first line of paragraph
        except KeyError:
            print("Warning: 'Normal' style not found. Using default.")
            normal_style = 'Normal' # Fallback


        doc.add_paragraph(self.novel_title, style=title_style).alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_line = doc.add_paragraph(f"Inspired by the style of {self.author_style}", style=normal_style)
        author_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if hasattr(author_line.paragraph_format, 'first_line_indent'): # Check if style object was retrieved
            author_line.paragraph_format.first_line_indent = None 

        genre_line = doc.add_paragraph(f"Genre: {self.genre}", style=normal_style)
        genre_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if hasattr(genre_line.paragraph_format, 'first_line_indent'):
            genre_line.paragraph_format.first_line_indent = None 
        doc.add_page_break()

        for i in range(1, self.num_chapters + 1):
            chapter_content = self.generated_chapters_content.get(i, f"[ERROR: Content for Chapter {i} not found]")
            
            paragraphs = chapter_content.split('\n\n')
            
            if paragraphs:
                ch_title_line = paragraphs[0].strip()
                # Attempt to extract just the title part for the heading if format is "Chapter X - Title"
                title_match = re.match(r"Chapter\s*\d+\s*[:*-]?\s*(.*)", ch_title_line, re.IGNORECASE)
                heading_text = title_match.group(1).strip() if title_match and title_match.group(1).strip() else ch_title_line
                
                ch_heading = doc.add_heading(heading_text, level=1) # Use extracted title or full line
                ch_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                if hasattr(ch_heading.paragraph_format, 'space_before'): 
                    ch_heading.paragraph_format.space_before = Pt(12)
                    ch_heading.paragraph_format.space_after = Pt(6)

                for para_block in paragraphs[1:]: 
                    if para_block.strip():
                        p = doc.add_paragraph(para_block.strip(), style=normal_style)
            else: 
                 doc.add_paragraph(chapter_content, style=normal_style) # Add raw content if split fails

            if i < self.num_chapters:
                doc.add_page_break()

        safe_title = re.sub(r'[^\w\s-]', '', self.novel_title).strip().replace(' ', '_')
        safe_genre = self.genre.replace('/','-').replace(' ','')
        filename = f"{safe_title[:50]}_Novel_{safe_genre}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)

        try:
            doc.save(filepath)
            print(f"Novel successfully saved to: {filepath}")
        except Exception as e:
            print(f"ERROR saving .docx file: {e}")

        metadata = {
            "title": self.novel_title,
            "subject": self.subject,
            "author_style": self.author_style,
            "genre": self.genre,
            "num_chapters_determined": self.num_chapters,
            "ollama_model_used": OLLAMA_MODEL,
            "generation_timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "characters_final": self.characters,
            "world_details": self.world_details,
            "themes_motifs": self.themes_motifs,
            "plot_outline": self.plot_outline,
            "chapter_plans": self.chapter_plans,
            "chapter_continuity_data": self.chapter_continuity_data, # Now includes flow_analysis
        }
        meta_filename = f"{safe_title[:50]}_Novel_METADATA.json"
        meta_filepath = os.path.join(OUTPUT_DIR, meta_filename)
        try:
            with open(meta_filepath, "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
            print(f"Metadata saved to: {meta_filepath}")
        except Exception as e:
            print(f"ERROR saving metadata JSON: {e}")


    def orchestrate_generation(self):
        """
        Main public method to run the entire novel generation pipeline.
        """
        start_time = time.time()
        print("--- Starting Novel Generation Pipeline ---")

        if not self.generate_foundational_elements():
            print("Halting: Foundational element generation failed.")
            return

        if not self.generate_detailed_chapter_plans():
            print("Halting: Detailed chapter plan generation failed.")
            return
        
        if not self.generate_novel_content():
            print("Halting: Novel content generation failed.")
            return
            
        # Perform final transition checks after all chapters are generated
        self._perform_final_transition_checks()

        # Compile and save the potentially revised chapters
        self.compile_and_save_novel()

        end_time = time.time()
        total_time_minutes = (end_time - start_time) / 60
        print(f"--- Novel Generation Pipeline Finished ---")
        print(f"Total time taken: {total_time_minutes:.2f} minutes.")


def get_user_input_multiline(prompt_message):
    print(prompt_message + " (Type 'ENDINPUT' on a new line when done, or just press Enter if input is short):")
    lines = []
    first_line = input()
    if not first_line.strip() and not lines:
        return first_line 
    
    if first_line.strip().upper() == 'ENDINPUT':
        return "\n".join(lines)
    
    lines.append(first_line)

    if len(first_line) > 70 or "ENDINPUT" not in first_line.strip().upper():
        while True:
            try:
                line = input()
                if line.strip().upper() == 'ENDINPUT':
                    break
                lines.append(line)
            except EOFError: 
                print("INFO: EOF reached while reading multiline input.")
                break
    return "\n".join(lines)

def load_resume_text(file_path):
    """Loads text from a file, attempting PDF extraction if it's a .pdf file."""
    if not file_path:
        return ""
    
    resume_text = ""
    try:
        if file_path.lower().endswith(".pdf"):
            try:
                with open(file_path, 'rb') as f: 
                    reader = pypdf.PdfReader(f)
                    if reader.is_encrypted:
                         print(f"Warning: PDF file '{file_path}' is encrypted. Cannot extract text.")
                         return "" # Return empty if encrypted
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text: # Ensure text was extracted
                             resume_text += page_text + "\n"
                if resume_text.strip():
                    print(f"Successfully extracted text from PDF: {file_path}")
                else:
                    print(f"Warning: No text could be extracted from PDF: {file_path}. It might be an image-based PDF, scanned, or corrupted.")
            except pypdf.errors.PdfReadError as pe: # Catch specific pypdf errors
                 print(f"Warning: PyPDF could not read PDF file '{file_path}': {pe}. It might be encrypted or corrupted. Proceeding without resume.")
                 resume_text = ""
            except Exception as e:
                print(f"Warning: Error processing PDF file '{file_path}': {e}. Proceeding without resume content from this file.")
                resume_text = "" 
        else: 
            encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
            loaded_successfully = False
            for enc in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        resume_text = f.read()
                    print(f"Resume loaded from text file: {file_path} (using {enc})")
                    loaded_successfully = True
                    break 
                except UnicodeDecodeError:
                    continue 
                except Exception as e_text: 
                    print(f"Warning: Error loading text file '{file_path}' with {enc}: {e_text}.")
                    break 
            if not loaded_successfully and not resume_text: 
                 print(f"Warning: Could not load resume from text file '{file_path}' after trying multiple encodings. Proceeding without it.")

    except FileNotFoundError:
        print(f"Warning: Resume file not found at '{file_path}'. Proceeding without resume.")
    except Exception as e_general: 
        print(f"Warning: An unexpected error occurred while trying to load '{file_path}': {e_general}. Proceeding without resume.")
    
    return resume_text.strip()


if __name__ == "__main__":
    print("Welcome to the AI Novel Generator!")
    print("Please ensure your Ollama server is running and the model is available.")
    print(f"Using Ollama Model: {OLLAMA_MODEL} at {OLLAMA_BASE_URL}")
    print("----------------------------------------------------")

    resume_file_path_input = input("Enter path to resume file (text or PDF) (or press Enter to skip): ").strip()
    resume_text_content = load_resume_text(resume_file_path_input)
    
    novel_subject_input = get_user_input_multiline("Enter the novel's subject/premise")
    author_style_input_str = input("Enter the desired author style (e.g., 'Stephen King', 'Jane Austen'): ").strip()
    genre_input_str = input("Enter the genre(s) (e.g., 'Sci-Fi/Thriller', 'Historical Romance'): ").strip()
    
    generator = NovelGenerator(
        resume_content=resume_text_content,
        subject=novel_subject_input,
        author_style=author_style_input_str,
        genre=genre_input_str
    )
    generator.orchestrate_generation()
    print("----------------------------------------------------")