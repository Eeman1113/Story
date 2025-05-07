# app.py
from dotenv import load_dotenv
import os
import docx
import traceback # For printing detailed errors
import time # To avoid overwhelming the LLM API if needed

# --- Langchain Imports ---
from langchain_ollama import OllamaLLM
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import PyPDFLoader
# --- End Imports ---

# Load environment variables (optional, but good practice)
load_dotenv()

# --- Constants ---
DEFAULT_MODEL = "gemma3:4b" # Using a potentially more capable model if available
# DEFAULT_MODEL = "llama3:latest" # Example alternative
OUTPUT_FOLDER = './docs'
# Optional: Add a small delay between LLM calls if hitting rate limits or for stability
LLM_CALL_DELAY_SECONDS = 0.5

# --- LLM Initialization Function ---
def create_llm(temperature=0.7):
    """Create and return an OllamaLLM instance with specified temperature."""
    # Define the base URL - uses environment variable or default
    ollama_base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    print(f"--- Connecting to Ollama at: {ollama_base_url} with Model: {DEFAULT_MODEL} ---")
    try:
        llm = OllamaLLM(
            model=DEFAULT_MODEL,
            temperature=temperature, # Allow variable temperature
            base_url=ollama_base_url,
            # Consider adding timeout if requests hang (e.g., request_timeout=120.0)
            # Add other Ollama parameters if needed (num_ctx, etc.)
        )
        # Optional: Test connection (though OllamaLLM might do lazy init)
        # llm.invoke("Test connection.")
        print("--- LLM Instance Created ---")
        return llm
    except Exception as e:
        print(f"--- FATAL ERROR: Could not create OllamaLLM instance ---")
        print(f"Error: {e}")
        print("Ensure Ollama service is running and accessible at the specified base URL.")
        print(f"Attempted Base URL: {ollama_base_url}")
        print(f"Attempted Model: {DEFAULT_MODEL}")
        traceback.print_exc()
        raise # Reraise the exception to stop execution

# --- Chain Classes ---

class MainCharacterChain:
    # Keep prompt simple for profile extraction
    PROMPT = """
    Analyze the following resume text. Extract the person's name (if mentioned) and create a concise profile (3-4 sentences) highlighting key skills, experiences, and potential personality traits relevant for a fictional narrative. Focus on elements that could drive a story.

    Resume Text:
    {text}

    Character Profile:"""

    def __init__(self):
        # Use slightly lower temperature for factual extraction
        self.llm = create_llm(temperature=0.5)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True # Enable verbose output for debugging
        )

    def load_resume(self, file_name):
        """Loads text content from a PDF file."""
        file_path = os.path.join(OUTPUT_FOLDER, file_name)
        if not os.path.exists(file_path):
             raise FileNotFoundError(f"Resume file not found at: {file_path}")
        try:
            print(f"Loading PDF from: {file_path}")
            loader = PyPDFLoader(file_path)
            docs = loader.load() # Returns list of Document objects (one per page)
            if not docs:
                print(f"Warning: PyPDFLoader didn't extract any documents/pages from {file_name}.")
                return None
            print(f"Successfully loaded {len(docs)} page(s) from PDF.")
            # Combine text from all pages into a single string
            full_text = '\n\n'.join([doc.page_content for doc in docs if doc.page_content])
            return full_text
        except Exception as e:
             print(f"Error loading or processing PDF {file_path}: {e}")
             traceback.print_exc()
             raise

    def run(self, file_name):
        """Loads resume and generates the profile description."""
        try:
            resume_text = self.load_resume(file_name)
            if not resume_text or not resume_text.strip():
                 print("Could not load or resume content is empty.")
                 return "Error: Could not generate profile due to missing or empty resume content."

            time.sleep(LLM_CALL_DELAY_SECONDS) # Optional delay
            result = self.chain.invoke({"text": resume_text})
            profile = result.get('text', "Error: Profile generation failed.").strip()

            if not profile or profile.startswith("Error:") or len(profile) < 20: # Basic check for valid profile
                print(f"Warning: Generated profile seems invalid or too short: '{profile}'")
                # Optional: Retry logic could be added here
                return f"Error: Failed to generate a meaningful profile. LLM Output: {profile}"
            return profile

        except Exception as e:
            print(f"An error occurred in MainCharacterChain.run: {e}")
            traceback.print_exc()
            return f"Error generating profile: {e}"


class TitleChain:
    PROMPT = """
    You are a creative book title generator.
    Generate ONE compelling and evocative novel title based on the provided details.
    The title must be highly consistent with the specified genre(s), author's style, subject matter, and main character.
    It should be intriguing and suitable for the target audience of the genre.

    Subject: {subject}
    Genre(s): {genre}
    Author's Style Inspiration: {author}
    Main Character Profile: {profile}

    Return ONLY the generated title itself, without any quotation marks, labels (like "Title:"), or explanatory text.

    Novel Title:"""

    def __init__(self):
        self.llm = create_llm(temperature=0.8) # Slightly higher temperature for creativity
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile):
        """Generates the novel title."""
        try:
            time.sleep(LLM_CALL_DELAY_SECONDS) # Optional delay
            result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile
            })
            title = result.get('text', "Untitled Novel").strip()
            # Clean the output string thoroughly
            if title.startswith('"') and title.endswith('"'): title = title[1:-1]
            if title.startswith("'") and title.endswith("'"): title = title[1:-1]
            if title.startswith('Title:'): title = title.replace('Title:', '').strip()

            if not title or len(title) < 3 or title == "Untitled Novel": # Check for invalid titles
                 print(f"Warning: Generated title seems invalid: '{title}'. Using placeholder.")
                 return f"Placeholder Title - {genre.split(',')[0].strip()} Story" # Generate a simple placeholder
            return title
        except Exception as e:
            print(f"An error occurred in TitleChain.run: {e}")
            traceback.print_exc()
            return "Error Generating Title"


class PlotChain:
    # Enhanced prompt for more detailed and structured plot
    PROMPT = """
    You are a master storyteller and plot architect.
    Develop a detailed, multi-act plot outline for a novel based on the provided elements.
    The plot must be engaging, coherent, and build towards a satisfying climax and resolution.
    Introduce necessary supporting characters, conflicts (internal and external), and key turning points.
    The main character must be central, and their journey should reflect the themes implied by the subject and genre.
    Ensure the plot aligns perfectly with the novel's genre(s), author's style, title, and subject.
    Incorporate compelling story attributes like: {features}. Structure the output logically (e.g., Act I, Act II, Act III).

    Novel Subject: {subject}
    Genre(s): {genre}
    Author's Style Inspiration: {author}
    Novel Title: "{title}"
    Main Character Profile: {profile}

    Detailed Plot Outline:"""

    HELPER_PROMPT = """
    Generate a comma-separated list of 5-7 diverse and compelling attributes that make a story plot exciting, unique, and engaging for the specified genre(s). Avoid generic terms.
    Genre(s): {genre}

    Examples for different genres:
    Sci-Fi: Mind-bending concepts, Dystopian societies, AI ethics conflict, Alien encounters, Technological marvels/failures.
    Fantasy: Elaborate magic system, Ancient prophecies, Mythical creatures, Epic quests, Political intrigue in kingdoms.
    Thriller: Unreliable narrator, Ticking clock scenario, High-stakes chase sequences, Shocking betrayals, Psychological manipulation.
    Romance: Forbidden love, Slow burn tension, Grand romantic gestures, Misunderstandings and resolutions, External obstacles to love.
    Slice of Life: Subtle character growth, Realistic interpersonal dynamics, Mundane turned meaningful, Exploration of routine, Quiet epiphanies.

    List of Attributes (comma-separated):"""

    def __init__(self):
        self.llm = create_llm(temperature=0.75) # Balanced temperature for structured creativity
        # Main chain for plot generation
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )
        # Helper chain to generate dynamic features based on genre
        self.helper_chain = LLMChain(
            llm=self.llm, # Can use the same LLM instance
            prompt=PromptTemplate.from_template(self.HELPER_PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile, title):
        """Generates the novel plot outline."""
        try:
            # Generate dynamic features using the helper chain, tailored to genre
            print(f"Generating plot features for genre: {genre}")
            time.sleep(LLM_CALL_DELAY_SECONDS) # Optional delay
            features_result = self.helper_chain.invoke({"genre": genre})
            features = features_result.get('text', "Compelling conflict, Character depth, Unexpected twists").strip()
            print(f"Generated plot features: {features}")

            # Generate the main plot outline
            print(f"Generating main plot outline for title: {title}")
            time.sleep(LLM_CALL_DELAY_SECONDS) # Optional delay
            plot_result = self.chain.invoke({
                "features": features,
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile,
                "title": title
            })
            plot = plot_result.get('text', "Error: Plot generation failed.").strip()

            if not plot or plot.startswith("Error:") or len(plot) < 100: # Check for valid plot
                 print(f"Warning: Generated plot seems invalid or too short: '{plot}'")
                 return f"Error: Failed to generate a detailed plot. LLM Output: {plot}"
            return plot
        except Exception as e:
            print(f"An error occurred in PlotChain.run: {e}")
            traceback.print_exc()
            return f"Error generating plot: {e}"


class ChaptersChain:
    # Prompt adjusted for clarity and robustness
    PROMPT = """
    You are a book editor outlining chapter structure.
    Based on the detailed plot outline provided, generate a list of chapter titles AND brief, one-sentence descriptions for a novel.
    Aim for a reasonable number of chapters (e.g., 15-30, depending on plot complexity), ensuring logical progression through the plot.
    Include a Prologue and/or Epilogue ONLY if appropriate for the story structure and genre.

    Ensure strict consistency with the novel's title, genre(s), author's style, main character, and the provided plot.
    The descriptions should capture the core focus or turning point of each chapter.

    Use this EXACT format for each entry, with NO blank lines between entries:
    [Chapter Type] [Number (if not Prologue/Epilogue)]: [One-sentence Description]

    Example Format:
    Prologue: Introduction to the world and the inciting incident's foreshadowing.
    Chapter 1: The protagonist faces a mundane challenge revealing their core desire.
    Chapter 2: An unexpected event disrupts the protagonist's routine.
    ...
    Chapter N: The climax unfolds, testing the protagonist to their limits.
    Epilogue: Resolution of the main conflict and a glimpse into the future.

    Novel Title: "{title}"
    Genre(s): {genre}
    Author's Style Inspiration: {author}
    Main Character Profile: {profile}
    Detailed Plot Outline:
    <PLOT_START>
    {plot}
    <PLOT_END>

    Chapters List (Strict Format Adherence Required):"""

    def __init__(self):
        self.llm = create_llm(temperature=0.6) # Lower temperature for structured output
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse(self, response):
        """Parses the LLM response string into an ordered dictionary of chapters."""
        chapters = {} # Using dict to store { 'Title': 'Description' }
        lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
        print(f"\nRaw Chapter Response Lines ({len(lines)}):")
        # for i, line in enumerate(lines): print(f"  Line {i}: {line}") # Uncomment for deep debug

        parsed_count = 0
        for i, line in enumerate(lines):
            # More robust parsing: requires a colon and checks structure
            if ':' in line:
                parts = line.split(':', 1)
                potential_title = parts[0].strip()
                potential_desc = parts[1].strip()

                # Check if title looks like Prologue, Epilogue, or Chapter X
                is_prologue = potential_title.lower() == "prologue"
                is_epilogue = potential_title.lower() == "epilogue"
                is_chapter = potential_title.lower().startswith("chapter ") and \
                             potential_title.split(" ")[-1].isdigit()

                if is_prologue or is_epilogue or is_chapter:
                    # Ensure description is not empty
                    if potential_desc:
                        chapters[potential_title] = potential_desc
                        parsed_count += 1
                        # print(f"    Parsed Line {i}: '{potential_title}' -> '{potential_desc}'") # Debug
                    else:
                        print(f"  Warning: Skipping line {i} - Empty description: '{line}'")
                else:
                     print(f"  Warning: Skipping line {i} - Non-standard title format: '{line}'")
            else:
                print(f"  Warning: Skipping line {i} - No colon found: '{line}'")

        if not chapters:
             print(f"CRITICAL WARNING: Could not parse *any* chapters adhering to the expected format.")
             print(f"LLM Raw Response was:\n---\n{response}\n---")
             # Consider adding a fallback mechanism or raising an error here
             return {} # Return empty dict to signal failure

        print(f"Successfully parsed {parsed_count} chapters.")
        return chapters # Return as dict {title: description}

    def run(self, subject, genre, author, profile, title, plot):
        """Generates and parses the chapter list."""
        try:
            time.sleep(LLM_CALL_DELAY_SECONDS) # Optional delay
            response_result = self.chain.invoke({
                "subject": subject, # Keep subject for context if useful
                "genre": genre,
                "author": author,
                "profile": profile,
                "title": title,
                "plot": plot
            })
            raw_response = response_result.get('text', "")
            return self.parse(raw_response)
        except Exception as e:
            print(f"An error occurred in ChaptersChain.run: {e}")
            traceback.print_exc()
            return {} # Return empty dict on error


# --- NEW: Event Generation Chain ---
class EventChain:
    PROMPT = """
    You are a narrative strategist breaking down a chapter into key scenes or events.
    Based on the overall plot, the specific chapter's title, and its summary, generate a concise, ordered list of 3-6 key plot events or scenes that MUST occur within this chapter to advance the story logically.
    Focus on actions, decisions, revelations, or significant interactions.
    Ensure the events directly relate to the chapter summary and the overall plot trajectory.

    Overall Plot Summary:
    <PLOT_SUMMARY>
    {plot}
    </PLOT_SUMMARY>

    Current Chapter Title: "{chapter_title}"
    Current Chapter Summary: {chapter_summary}

    Return ONLY the numbered list of events, one event per line. Start numbering from 1. Be concise.

    Example:
    1. Protagonist arrives at the mysterious location mentioned in the previous chapter.
    2. They discover a hidden clue or object.
    3. An unexpected encounter with an antagonist or ally occurs.
    4. The chapter ends with a crucial decision or a cliffhanger related to the discovery.

    Numbered Event List for Chapter "{chapter_title}":
    """

    def __init__(self):
        self.llm = create_llm(temperature=0.65) # Moderate temperature for focused event generation
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse_events(self, response):
        """Parses the numbered list response into a list of strings."""
        events = []
        lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
        for line in lines:
            # Expect lines like "1. Event description" or "1) Event description" etc.
            if line and (line[0].isdigit() or (len(line)>1 and line[1].isdigit())):
                 # Find the first space after the number/punctuation
                 first_space_index = -1
                 for i, char in enumerate(line):
                    if char.isspace():
                        first_space_index = i
                        break
                 if first_space_index != -1:
                    event_text = line[first_space_index:].strip()
                    if event_text: # Ensure we captured something
                         events.append(event_text)
                 else: # Handle cases like "1.Event"
                     parts = line.split('.', 1)
                     if len(parts) > 1 and parts[1].strip():
                         events.append(parts[1].strip())


        if not events:
             print(f"  Warning: Could not parse any numbered events from response:\n---\n{response}\n---")
             # Fallback: return the raw response lines as events (better than nothing)
             return lines if lines else ["Placeholder event due to parsing failure."]
        return events

    def run(self, plot, chapter_title, chapter_summary):
        """Generates and parses the event list for a single chapter."""
        try:
            time.sleep(LLM_CALL_DELAY_SECONDS) # Optional delay
            result = self.chain.invoke({
                "plot": plot,
                "chapter_title": chapter_title,
                "chapter_summary": chapter_summary
            })
            raw_events = result.get('text', "").strip()
            return self.parse_events(raw_events)
        except Exception as e:
            print(f"  Error generating events for '{chapter_title}': {e}")
            traceback.print_exc()
            # Return placeholder on error
            return [f"Event generation error 1 for '{chapter_summary}'", f"Event generation error 2 for '{chapter_summary}'"]

# --- Enhanced Writer Chain ---
class WriterChain:
    # Significantly more detailed prompt with clear context separation
    PROMPT = """
    You are a master novelist embodying the distinct writing style of {author}.
    You are writing a novel titled "{title}", a {genre} work.

    CONTEXT:
    <MAIN_CHARACTER>
    {profile}
    </MAIN_CHARACTER>

    <NOVEL_PLOT_SUMMARY>
    {plot}
    </NOVEL_PLOT_SUMMARY>

    <CURRENT_CHAPTER_INFO>
    Chapter Title: {chapter_name}
    Chapter Summary: {summary}
    </CURRENT_CHAPTER_INFO>

    <NARRATIVE_HISTORY>
    Key events that have already occurred in the story (prior chapters):
    {previous_events}
    </NARRATIVE_HISTORY>

    <CHAPTER_PROGRESS_SO_FAR>
    Paragraphs already written for THIS chapter ({chapter_name}):
    {previous_paragraphs}
    </CHAPTER_PROGRESS_SO_FAR>

    YOUR CURRENT TASK:
    Write the next part of the story, focusing *exclusively* on the following event. Bring this specific event to life with vivid description, character thoughts/actions, dialogue (if appropriate), and sensory details, adhering strictly to the established context.

    <CURRENT_EVENT_TO_WRITE>
    {current_event}
    </CURRENT_EVENT_TO_WRITE>

    INSTRUCTIONS:
    1. Write one or more detailed paragraphs (approx. 100-300 words total for this event, adjust based on significance) describing ONLY the CURRENT EVENT.
    2. Maintain 100% consistency with the plot, character profile, {genre} genre conventions, and the specific stylistic nuances of {author} (sentence structure, vocabulary, tone, pacing).
    3. Ensure seamless transition from the 'CHAPTER_PROGRESS_SO_FAR'.
    4. Do NOT repeat information already covered in previous paragraphs or events.
    5. Do NOT add explanations, summaries, or any text outside the narrative itself.
    6. Output ONLY the newly written paragraphs for the current event.

    New Narrative Paragraphs (Style: {author}):"""

    def __init__(self):
        # Higher temperature might be needed for creative prose generation
        self.llm = create_llm(temperature=0.85)
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, genre, author, title, profile, plot, chapter_name,
            previous_events_history_str, chapter_summary, previous_paragraphs_str, current_event):
        """Generates narrative paragraphs for a specific event."""
        # Ensure placeholders are clear if context is empty
        if not previous_events_history_str: previous_events_history_str = "None (This is the beginning of the story)."
        if not previous_paragraphs_str: previous_paragraphs_str = "None (This is the beginning of the chapter)."

        try:
            time.sleep(LLM_CALL_DELAY_SECONDS) # Optional delay
            result = self.chain.invoke({
                "genre": genre,
                "author": author,
                "title": title,
                "profile": profile,
                "plot": plot, # Pass the full plot outline
                "chapter_name": chapter_name,
                "previous_events": previous_events_history_str, # Pass cumulative history string
                "summary": chapter_summary, # Pass current chapter summary
                "previous_paragraphs": previous_paragraphs_str, # Pass text written *in this chapter so far*
                "current_event": current_event # The specific event to write now
            })
            generated_text = result.get('text', f"[ERROR: Writing failed for event: '{current_event}']").strip()

            # Basic validation of output
            if not generated_text or generated_text.startswith("[ERROR:") or len(generated_text) < 30:
                 print(f"  Warning: WriterChain generated potentially invalid/short output for event '{current_event}'. Output: '{generated_text}'")
                 # Return a placeholder indicating the issue
                 return f"[Writer Error - Event: '{current_event}'. LLM Output: '{generated_text}']"

            # Optional: Add post-processing to clean up LLM artifacts if necessary
            # e.g., remove leading/trailing quotes sometimes added by models

            return generated_text
        except Exception as e:
            print(f"  An error occurred in WriterChain.run for event '{current_event}': {e}")
            traceback.print_exc()
            return f"[FATAL WRITER ERROR - Event: '{current_event}'. See logs.]"

# --- Book Writing Orchestration ---

def generate_events_for_all_chapters(plot, chapter_dict):
    """Generates event lists for all chapters using EventChain."""
    print("\n--- Generating Events for Each Chapter ---")
    event_generator = EventChain()
    event_dict = {} # { 'Chapter Title': ['Event 1', 'Event 2', ...] }

    # Sort chapters for predictable processing order
    sorted_chapters = sort_chapters(chapter_dict)

    for chapter_title, summary in sorted_chapters:
         print(f"Generating events for: {chapter_title}")
         events = event_generator.run(plot, chapter_title, summary)
         if events: # Only add if events were successfully generated/parsed
            event_dict[chapter_title] = events
            print(f"  Generated {len(events)} events for '{chapter_title}'.")
            # for i, ev in enumerate(events): print(f"    {i+1}. {ev}") # Uncomment for detail
         else:
            print(f"  WARNING: No events generated or parsed for '{chapter_title}'. Chapter content may be skipped.")
            event_dict[chapter_title] = [] # Store empty list to indicate failure/skip

    return event_dict


def sort_chapters(chapter_dict):
    """Sorts chapter dictionary keys (titles) into a logical order."""
    def chapter_sort_key(chapter_title):
        title = chapter_title.lower()
        if title == "prologue": return (-1, 0) # Sort Prologue first
        if title == "epilogue": return (9999, 0) # Sort Epilogue last
        if title.startswith("chapter "):
            try:
                num = int(title.split(" ")[1])
                return (1, num) # Sort chapters by number
            except (IndexError, ValueError):
                return (2, title) # Fallback sorting for malformed chapter titles
        return (3, title) # Sort any other titles alphabetically after chapters

    # Sort the items (title, description pairs) based on the key (title)
    return sorted(chapter_dict.items(), key=lambda item: chapter_sort_key(item[0]))


def write_book(genre, author, title, profile, plot, chapter_dict, event_dict):
    """Orchestrates the writing of the full book content, event by event."""
    print("\n--- Starting Detailed Book Writing Process ---")
    writer_chain = WriterChain() # Instantiate the writer

    book_content = {} # Stores final text: {chapter_title: "Full chapter text..."}
    previous_events_history = [] # Running list of event descriptions written so far

    sorted_chapters = sort_chapters(chapter_dict) # Use the sorted list of (title, description) pairs

    total_chapters = len(sorted_chapters)
    for chap_idx, (chapter_title, chapter_summary) in enumerate(sorted_chapters):
        print(f"\n--- Writing Chapter {chap_idx+1}/{total_chapters}: {chapter_title} ---")
        print(f"   Summary: {chapter_summary}")

        chapter_events = event_dict.get(chapter_title, [])
        if not chapter_events:
            print(f"   WARNING: No events found for '{chapter_title}'. Skipping content generation.")
            book_content[chapter_title] = f"[Content generation skipped: No events defined for this chapter.]"
            continue

        # Accumulator for text within the *current* chapter
        chapter_paragraphs_accumulator = ""
        total_events = len(chapter_events)

        for event_idx, event_description in enumerate(chapter_events):
            print(f"   Writing Event {event_idx+1}/{total_events}: {event_description}")

            # Prepare context strings
            history_str = '\n'.join(f"- {evt}" for evt in previous_events_history) if previous_events_history else "None (Start of Story)"
            progress_str = chapter_paragraphs_accumulator if chapter_paragraphs_accumulator else "None (Start of Chapter)"

            # Call the writer chain for the current event
            new_paragraphs = writer_chain.run(
                genre=genre,
                author=author,
                title=title,
                profile=profile,
                plot=plot,
                chapter_name=chapter_title,
                previous_events_history_str=history_str,
                chapter_summary=chapter_summary,
                previous_paragraphs_str=progress_str,
                current_event=event_description
            )

            # Append the newly generated block to the accumulator for this chapter
            if chapter_paragraphs_accumulator: # Add separation if not the first block
                chapter_paragraphs_accumulator += "\n\n" # Double newline between event blocks
            chapter_paragraphs_accumulator += new_paragraphs

            # Add the *description* of the event we just wrote to the global history
            # This history tells the LLM what *topics* have been covered.
            previous_events_history.append(f"[{chapter_title}] {event_description}")

        # Store the fully assembled text for the chapter
        book_content[chapter_title] = chapter_paragraphs_accumulator
        print(f"--- Finished writing Chapter: {chapter_title} ---")

    print("\n--- Book Writing Process Complete ---")
    return book_content

# --- Document Writing Class ---

class DocWriter:
    def __init__(self, output_folder=OUTPUT_FOLDER):
        self.output_folder = output_folder
        # Ensure output directory exists
        os.makedirs(self.output_folder, exist_ok=True)

    def _sanitize_filename(self, name):
        """Removes or replaces characters invalid for filenames."""
        # Remove problematic characters
        name = name.replace(':', '-').replace('/', '-').replace('\\', '-').replace('?', '').replace('*', '')
        name = name.replace('"', '').replace('<', '').replace('>', '').replace('|', '')
        # Replace spaces (optional, underscores often preferred)
        name = name.replace(' ', '_')
        # Trim leading/trailing whitespace/underscores
        name = name.strip('._ ')
        # Limit length (e.g., 200 chars) to avoid filesystem issues
        max_len = 200
        name = name[:max_len]
        # Ensure filename is not empty after sanitization
        if not name:
            name = "Untitled_Novel"
        return name

    def write_doc(self, book_content, chapter_dict, title, genre, author):
        """Writes the generated book content to a .docx file."""
        print("\n--- Assembling and Writing Document ---")
        doc = docx.Document() # Create a new document for each book

        # --- Front Matter (Optional but nice) ---
        doc.add_heading(title, level=0) # Book title
        doc.add_paragraph(f"Genre: {genre}")
        doc.add_paragraph(f"Inspired by the style of: {author}")
        doc.add_paragraph("\n") # Add some spacing

        # --- Body ---
        sorted_chapters = sort_chapters(chapter_dict) # Get sorted (title, description) pairs

        total_chapters = len(sorted_chapters)
        for i, (chapter_title, description) in enumerate(sorted_chapters):
            # Add chapter heading
            # Use title and description for context, or just title
            # chapter_heading = f'{chapter_title.strip()}: {description.strip()}' # Option 1: Title + Desc
            chapter_heading = chapter_title.strip() # Option 2: Just Title
            doc.add_heading(chapter_heading, level=1) # Chapter heading (Level 1)
            print(f"  Adding Chapter {i+1}/{total_chapters} to Doc: {chapter_heading}")

            # Get the full generated text for this chapter
            chapter_text = book_content.get(chapter_title, "[Error: Chapter content not found]")

            # Add the chapter text as one or more paragraphs
            # Split text into paragraphs based on double newlines used as separators
            paragraphs = chapter_text.split('\n\n')
            for para in paragraphs:
                if para.strip(): # Avoid adding empty paragraphs
                    doc.add_paragraph(para.strip())

            # Add a page break after each chapter (optional, creates clear separation)
            if i < total_chapters - 1: # Don't add after the last chapter
                doc.add_page_break()

            print(f"    Added content for '{chapter_title}'")


        # --- Filename Creation ---
        safe_basename = self._sanitize_filename(title)
        # Add genre/author to filename for uniqueness if desired
        # filename = f"{safe_basename}_{self._sanitize_filename(genre.split(',')[0])}_{self._sanitize_filename(author)}.docx"
        filename = f"{safe_basename}.docx"
        output_path = os.path.join(self.output_folder, filename)

        # --- Saving Document ---
        try:
            print(f"\nAttempting to save document to: {output_path}")
            doc.save(output_path)
            print(f"--- Document saved successfully! ---")
        except PermissionError:
            print(f"\nERROR: Permission denied trying to save '{output_path}'.")
            print("Please check if the file is open in another application or if you have write permissions for the folder:")
            print(f"Folder: {os.path.abspath(self.output_folder)}")
        except Exception as e:
            print(f"\nERROR saving document to {output_path}: {e}")
            traceback.print_exc()

# --- Main Execution ---

def main():

    start_time = time.time()
    print("=============================================")
    print("=== COMPLEX NOVEL GENERATION SYSTEM START ===")
    print(f"=== Timestamp: {time.strftime('%Y-%m-%d %H:%M:%S')} ===")
    print("=============================================")

    # --- Configuration ---
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    print(f"Output Folder: {os.path.abspath(OUTPUT_FOLDER)}")

    # --- Inputs ---
    # Ensure the resume file exists in the OUTPUT_FOLDER
    # resume_filename = 'Your_Resume.pdf' # <--- IMPORTANT: REPLACE WITH YOUR PDF FILENAME
    resume_filename = 'divi_1.pdf' # Using the filename from the original code
    # Detailed Subject - crucial for guiding the story
    subject = ("""
Divi is a quiet observer of life — articulate, sensitive, and prone to spiraling thoughts that no one around her seems to notice. By day, she moves through life calmly, performing routines, attending to obligations, and maintaining the appearance of normalcy. But by night, especially in the hours before dawn — the “silver hush” — her mind becomes a strange and unfiltered terrain. Here, she is not just herself; she is every version of herself she’s ever buried.
Each chapter takes place over the course of a different night, exploring a specific loop or thread that her mind follows in the liminal silence before morning. Her insomnia is more than sleeplessness — it's a ritual, a confrontation. As the moonlight turns darkness into silver, Divi finds herself reliving memories that never quite healed, conversations that were never finished, and fears that echo too loudly when the world is quiet.
These episodes manifest almost like dream sequences: visceral, fragmented, yet real. Sometimes, she’s speaking with her younger self. Sometimes, it’s an imagined version of someone she loved or feared. Sometimes it’s just silence — heavy, metallic, alive. Her room becomes a surreal psychological stage where emotions take form and time becomes fluid.
Over the course of the story, Divi begins to fear that she may be slipping into something deeper — not just insomnia, but a fracture between reality and introspection. She starts journaling these episodes in a desperate attempt to anchor herself. But the act of writing becomes its own mirror, forcing her to question if the self she’s trying to hold onto has already changed.
By the final night, as dawn approaches and the silver haze turns to golden glow, Divi must choose: does she return fully to the waking world, or does she surrender to the nocturnal clarity where everything feels realer? The story ends in ambiguity — peace, perhaps, or madness cloaked in lucidity.""")
    author_style = 'Ocean Vuong' # Example: Known for poetic, lyrical prose
    # author_style = 'Ted Chiang' # Example: Known for thoughtful, character-driven Sci-Fi
    # author_style = 'N. K. Jemisin' # Example: Known for intricate world-building and deep characters
    # author_style = 'Becky Chambers' # Example: Known for hopeful, character-focused Sci-Fi
    genre = 'Psychological Literary Fiction / Nocturnal Realism'# Be specific

    print("\n--- Input Parameters ---")
    print(f"Resume File: {resume_filename}")
    print(f"Subject: {subject}")
    print(f"Author Style Inspiration: {author_style}")
    print(f"Genre(s): {genre}")
    print("------------------------")

    # --- Instantiate Chains & Writer ---
    # Wrapped in try-except to catch LLM initialization errors early
    try:
        print("\n--- Initializing Core Components ---")
        main_character_chain = MainCharacterChain()
        title_chain = TitleChain()
        plot_chain = PlotChain()
        chapters_chain = ChaptersChain()
        # EventChain and WriterChain are instantiated within the orchestration functions
        doc_writer = DocWriter(output_folder=OUTPUT_FOLDER)
        print("--- Core Components Initialized Successfully ---")
    except Exception as e:
        print(f"\nFATAL ERROR during initialization. Cannot proceed.")
        # Error message already printed in create_llm or here
        return # Stop execution

    # --- Generate Novel Components Sequentially ---
    try:
        # 1. Generate Profile
        print("\n--- Step 1: Generating Main Character Profile ---")
        profile = main_character_chain.run(resume_filename)
        if profile.startswith("Error:"):
            print(f"PROFILE GENERATION FAILED: {profile}")
            print("Cannot proceed without a character profile.")
            return
        print(f"Profile Generated:\n---\n{profile}\n---")

        # 2. Generate Title
        print("\n--- Step 2: Generating Novel Title ---")
        title = title_chain.run(subject, genre, author_style, profile)
        if title.startswith("Error") or "Placeholder Title" in title:
             print(f"TITLE GENERATION FAILED/PLACEHOLDER USED: {title}")
             # Decide if you want to proceed with a placeholder title or stop
             # return # Uncomment to stop if title generation fails critically
        print(f"Title Generated: '{title}'")

        # 3. Generate Plot
        print("\n--- Step 3: Generating Detailed Plot Outline ---")
        plot = plot_chain.run(subject, genre, author_style, profile, title)
        if plot.startswith("Error"):
            print(f"PLOT GENERATION FAILED: {plot}")
            print("Cannot proceed without a plot.")
            return
        print(f"Plot Generated:\n---\n{plot}\n---")

        # 4. Generate Chapter List
        print("\n--- Step 4: Generating Chapter List & Summaries ---")
        # chapter_dict should be { 'Chapter 1': 'Description...', ... }
        chapter_dict = chapters_chain.run(subject, genre, author_style, profile, title, plot)
        if not chapter_dict:
             print("ERROR: Failed to generate or parse chapters. Cannot write book content.")
             return # Stop if chapters couldn't be generated/parsed
        print("Chapters Generated:")
        # Print sorted chapters for clarity
        for chap_title, chap_desc in sort_chapters(chapter_dict):
             print(f"  - {chap_title}: {chap_desc}")

        # 5. Generate Events for Each Chapter
        # This is a new, crucial step before writing
        event_dict = generate_events_for_all_chapters(plot, chapter_dict)
        if not any(event_dict.values()): # Check if *any* events were generated for *any* chapter
            print("CRITICAL ERROR: Failed to generate events for any chapter. Cannot write book content.")
            return

        # --- Write the Book Content ---
        # This function now orchestrates event generation and paragraph writing
        book_content = write_book(
            genre,
            author_style,
            title,
            profile,
            plot,
            chapter_dict, # Pass the descriptions
            event_dict    # Pass the events
        )

        # --- Save to Document ---
        doc_writer.write_doc(book_content, chapter_dict, title, genre, author_style)

    except FileNotFoundError as e:
         print(f"\nCRITICAL FILE ERROR: {e}")
         print(f"Please ensure the required file ('{resume_filename}') exists in the '{os.path.abspath(OUTPUT_FOLDER)}' directory.")
    except Exception as e:
        print(f"\n--- AN UNEXPECTED ERROR OCCURRED DURING NOVEL GENERATION ---")
        print(f"Error Type: {type(e).__name__}")
        print(f"Error Details: {e}")
        print("--- Traceback ---")
        traceback.print_exc()
        print("-----------------")

    finally:
        end_time = time.time()
        total_time = end_time - start_time
        print("\n============================================")
        print(f"=== NOVEL GENERATION PROCESS FINISHED ===")
        print(f"=== Total Execution Time: {total_time:.2f} seconds ===")
        print("============================================")


if __name__ == "__main__":
    # It's recommended to run this from a terminal where you can see the verbose output.
    main()
