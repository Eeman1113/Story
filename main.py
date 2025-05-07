# app.py
from dotenv import load_dotenv
import os
import docx
import traceback # For printing detailed errors

# --- Langchain Imports - UPDATED ---
# Use the dedicated package for Ollama
from langchain_ollama import OllamaLLM
# Core Langchain components
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
# Community loader for PDF
from langchain_community.document_loaders import PyPDFLoader
# --- End Imports ---

# Load environment variables (optional, but good practice)
load_dotenv()

# --- Constants ---
DEFAULT_MODEL = "gemma3:4b" # Or specify a variant like "gemma3:latest", "gemma3:9b" etc.
OUTPUT_FOLDER = './docs' # Define output folder consistently

# --- LLM Initialization Function - UPDATED ---
def create_llm():
    """Create and return an OllamaLLM instance"""
    # Use OllamaLLM from the new package
    return OllamaLLM(
        model=DEFAULT_MODEL,
        temperature=0.7,
        # Add base_url if your Ollama runs elsewhere:
        # base_url="http://localhost:11434"
    )
# def create_llm():
#     """Create and return an OllamaLLM instance"""
#     # Define the base URL - uses environment variable or default
#     ollama_base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
#     print(f"--- Attempting to connect to Ollama at: {ollama_base_url} ---") # Added print
#     return OllamaLLM(
#         model=DEFAULT_MODEL,
#         temperature=0.7,
#         base_url=ollama_base_url # Explicitly pass the base URL
#     )
# --- End LLM Initialization ---

# --- Chain Classes ---

class MainCharacterChain:
    PROMPT = """
    You are provided with the resume text of a person.
    Describe the person's profile concisely in 2-3 sentences. Include the person's name if explicitly mentioned in the text.

    Resume Text:
    {text}

    Profile Description:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def load_resume(self, file_name):
        """Loads text content from a PDF file."""
        file_path = os.path.join(OUTPUT_FOLDER, file_name)
        if not os.path.exists(file_path):
             raise FileNotFoundError(f"Resume file not found at: {file_path}")
        try:
            print(f"Loading PDF from: {file_path}")
            loader = PyPDFLoader(file_path)
            # Use load() which returns a list of Document objects, one per page
            docs = loader.load()
            if not docs:
                print(f"Warning: PyPDFLoader didn't extract any documents/pages from {file_name}.")
                return None # Indicate no content found
            print(f"Successfully loaded {len(docs)} page(s) from PDF.")
            return docs
        except Exception as e:
             print(f"Error loading PDF {file_path}: {e}")
             traceback.print_exc() # Print full traceback for PDF errors
             raise # Reraise the exception

    def run(self, file_name):
        """Loads resume and generates the profile description."""
        try:
            docs = self.load_resume(file_name)
            if not docs:
                 print("Could not load resume content.")
                 return "Error: Could not generate profile due to missing resume content."

            # Combine text from all pages
            resume_text = '\n\n'.join([doc.page_content for doc in docs if doc.page_content])

            if not resume_text.strip():
                 print("Warning: Resume content appears empty after extraction.")
                 return "Error: Could not generate profile because the resume content is empty."

            # LangChain components now often prefer invoke over run
            # Use invoke which returns a dictionary
            result = self.chain.invoke({"text": resume_text})
            # Extract the actual response text (key might be 'text' or specific to the chain)
            # Check the verbose output or debug to confirm the key
            return result.get('text', "Error: Profile generation failed.").strip()

        except Exception as e:
            print(f"An error occurred in MainCharacterChain.run: {e}")
            traceback.print_exc()
            return f"Error generating profile: {e}"


class TitleChain:
    PROMPT = """
    Generate a compelling novel title based on the provided details.
    Return ONLY the title itself, without any quotation marks, labels, or extra text.
    The title must be consistent with the specified genre and author's style.

    Subject: {subject}
    Genre: {genre}
    Author's Style: {author}
    Main Character Profile: {profile}

    Novel Title:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile):
        """Generates the novel title."""
        try:
            result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile
            })
            # Clean the output string thoroughly
            title = result.get('text', "Untitled").strip()
            # Remove potential surrounding quotes
            if title.startswith('"') and title.endswith('"'):
                title = title[1:-1]
            if title.startswith("'") and title.endswith("'"):
                title = title[1:-1]
            return title
        except Exception as e:
            print(f"An error occurred in TitleChain.run: {e}")
            traceback.print_exc()
            return "Error Generating Title"


class PlotChain:
    PROMPT = """
    Generate a detailed plot outline for a novel. Return ONLY the plot description.
    Create a compelling narrative, introducing necessary supporting characters or conflicts.
    The main character must be central to the plot.
    Ensure the plot aligns with the novel's genre, author's style, title, and subject.
    Incorporate some of the following exciting story attributes: {features}.

    Subject: {subject}
    Genre: {genre}
    Author's Style: {author}
    Title: "{title}"
    Main Character Profile: {profile}

    Plot Outline:"""

    HELPER_PROMPT = """
    Generate a comma-separated list of 5-7 key attributes that make a story plot exciting and engaging.
    Examples: High stakes, Moral ambiguity, Unpredictable twists, Compelling antagonist, Emotional depth, Unique world-building, Fast pacing.

    List of attributes:"""

    def __init__(self):
        self.llm = create_llm()
        # Main chain for plot generation
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )
        # Helper chain to generate dynamic features
        self.helper_chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.HELPER_PROMPT),
            verbose=True # Make helper verbose to see feature generation
        )

    def run(self, subject, genre, author, profile, title):
        """Generates the novel plot."""
        try:
            # Generate exciting features using the helper chain
            # Note: Helper prompt doesn't have input variables, so pass dummy dict
            features_result = self.helper_chain.invoke({})
            features = features_result.get('text', "High stakes, Twists, Emotional depth").strip()
            print(f"Generated features for plot: {features}")

            # Generate the main plot
            plot_result = self.chain.invoke({
                "features": features,
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile,
                "title": title
            })
            return plot_result.get('text', "Error: Plot generation failed.").strip()
        except Exception as e:
            print(f"An error occurred in PlotChain.run: {e}")
            traceback.print_exc()
            return f"Error generating plot: {e}"


class ChaptersChain:
    PROMPT = """
    Generate a list of chapter titles and brief descriptions for a novel.
    Return ONLY the list, adhering strictly to the specified format. No extra text or explanations.
    The chapters should logically follow the provided plot outline.
    Ensure consistency with the novel's genre, author's style, title, and characters.

    Use this EXACT format for each entry (including Prologue/Epilogue if applicable):
    [Chapter Type] [Number (for chapters)]: [One-sentence description]

    Example Format:
    Prologue: Introduction to the setting and foreshadowing of conflict.
    Chapter 1: The main character receives a mysterious message.
    Chapter 2: An unexpected encounter changes everything.
    ...
    Chapter N: The final confrontation takes place.
    Epilogue: Aftermath and lingering questions.

    Do NOT add blank lines between entries.

    Novel Title: "{title}"
    Genre: {genre}
    Author's Style: {author}
    Main Character Profile: {profile}
    Plot Outline: {plot}

    Chapters List:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse(self, response):
        """Parses the LLM response string into a dictionary of chapters."""
        chapters = {}
        lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
        print(f"\nRaw Chapter Response Lines ({len(lines)}):") # Debug output
        # for i, line in enumerate(lines): print(f"  Line {i}: {line}") # Uncomment for detailed debug

        for line in lines:
            if ':' in line:
                parts = line.split(':', 1)
                chapter_title = parts[0].strip()
                chapter_desc = parts[1].strip()
                # Basic validation for expected chapter titles
                if chapter_title.startswith("Chapter") or chapter_title in ["Prologue", "Epilogue"]:
                    chapters[chapter_title] = chapter_desc
                    # print(f"    Parsed: '{chapter_title}' -> '{chapter_desc}'") # Debug
                else:
                    print(f"  Warning: Skipping line - unexpected chapter title format: '{line}'")
            else:
                print(f"  Warning: Skipping line - no colon found: '{line}'")

        if not chapters:
             print(f"Critical Warning: Could not parse *any* chapters from the response.")
             print(f"LLM Raw Response was:\n---\n{response}\n---")
        return chapters

    def run(self, subject, genre, author, profile, title, plot):
        """Generates and parses the chapter list."""
        try:
            response_result = self.chain.invoke({
                "subject": subject, # Subject might still provide context
                "genre": genre,
                "author": author,
                "profile": profile,
                "title": title,
                "plot": plot
            })
            raw_response = response_result.get('text', "")
            return self.parse(raw_response)
        except Exception as e:
            print(f"An error occurred in ChaptersChain.run: {e}")
            traceback.print_exc()
            return {} # Return empty dict on error


class WriterChain:
    PROMPT = """
    You are a skilled novelist writing in the distinctive style of {author}.
    The novel, titled "{title}", is a {genre} story.
    Main Character: {profile}
    Novel's Plot Summary: {plot}

    You are currently writing a section within the chapter "{chapter_name}".
    Chapter Summary: {summary}

    Events written so far in the novel:
    <PREVIOUS_EVENTS>
    {previous_events}
    </PREVIOUS_EVENTS>

    Paragraphs already written for THIS chapter ({chapter_name}):
    <PREVIOUS_PARAGRAPHS>
    {previous_paragraphs}
    </PREVIOUS_PARAGRAPHS>

    Your specific task now is to write the next part of the story, focusing ONLY on this event:
    <CURRENT_EVENT>
    {current_event}
    </CURRENT_EVENT>

    Generate one or more paragraphs vividly describing this event.
    Maintain absolute consistency with the established plot, characters, {genre} genre conventions, and the stylistic nuances of {author}.
    Output ONLY the newly written paragraphs for the current event. Do not repeat previous content or add explanations.

    New Paragraphs:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, genre, author, title, profile, plot, chapter_name,
            previous_events, summary, previous_paragraphs, current_event):
        """Generates paragraphs for a specific event."""
        try:
            # Format list of previous events into a string
            previous_events_str = '\n'.join(f"- {evt}" for evt in previous_events) if previous_events else "None"
            # Handle potentially empty previous paragraphs for the chapter start
            previous_paragraphs_str = previous_paragraphs if previous_paragraphs else "(Start of Chapter)"

            result = self.chain.invoke({
                "genre": genre,
                "author": author,
                "title": title,
                "profile": profile,
                "plot": plot,
                "chapter_name": chapter_name,
                "previous_events": previous_events_str,
                "summary": summary,
                "previous_paragraphs": previous_paragraphs_str,
                "current_event": current_event
            })
            return result.get('text', f"Error writing event: {current_event}").strip()
        except Exception as e:
            print(f"An error occurred in WriterChain.run for event '{current_event}': {e}")
            traceback.print_exc()
            return f"[Error occurred while writing event: {current_event}]"

# --- Book Writing Orchestration ---

def generate_events_for_chapter(chapter_title, summary, llm):
    """Uses LLM to generate a list of key events for a chapter based on its summary."""
    # Simple event generation chain (could be made more sophisticated)
    event_prompt_template = """
    Based on the chapter title and summary below, generate a short, ordered list of 3-5 key plot events that should occur within this chapter.
    Return ONLY the numbered list, one event per line. Example:
    1. Character discovers clue.
    2. Character faces obstacle.
    3. Character makes a decision.

    Chapter Title: {chapter_title}
    Chapter Summary: {summary}

    Numbered Event List:
    """
    event_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(event_prompt_template),
        verbose=True
    )
    try:
        result = event_chain.invoke({"chapter_title": chapter_title, "summary": summary})
        raw_events = result.get('text', "").strip()
        # Basic parsing of numbered list
        events = [line.strip().split('.', 1)[-1].strip()
                  for line in raw_events.split('\n')
                  if line.strip() and line.strip()[0].isdigit()]
        if not events:
            print(f"  Warning: Could not parse events for '{chapter_title}'. Using placeholders.")
            # Fallback placeholder events
            return [f"Key moment 1 related to '{summary}'", f"Key moment 2 related to '{summary}'"]
        print(f"  Generated {len(events)} events for '{chapter_title}'.")
        return events
    except Exception as e:
        print(f"  Error generating events for '{chapter_title}': {e}")
        traceback.print_exc()
        return [f"Event generation error 1 for '{summary}'", f"Event generation error 2 for '{summary}'"]


def write_book(genre, author, title, profile, plot, chapter_dict):
    """Generates the full book content chapter by chapter, event by event."""
    writer_chain = WriterChain() # Handles writing paragraphs for one event
    llm_for_events = create_llm() # Separate LLM instance for event generation if needed, or reuse

    # Generate events for each chapter dynamically using LLM
    print("\n--- Generating Events for Each Chapter ---")
    event_dict = {}
    for chapter_title, summary in chapter_dict.items():
         print(f"Generating events for: {chapter_title}")
         event_dict[chapter_title] = generate_events_for_chapter(chapter_title, summary, llm_for_events)
         # print(f"  Events: {event_dict[chapter_title]}") # Debug: Show generated events


    print("\n--- Starting Book Writing Process (Paragraph Generation) ---")
    book_content = {} # Stores the generated text paragraphs for each chapter {chapter_title: [paragraph_set_1, paragraph_set_2,...]}
    previous_events_history = [] # Keep track of all events written so far across chapters

    # Ensure chapters are processed in a somewhat logical order (e.g., sort by Chapter number)
    # Basic sort key function
    def chapter_sort_key(item):
        title = item[0]
        if title == "Prologue": return -1
        if title == "Epilogue": return 9999 # Large number for epilogue
        if title.startswith("Chapter "):
            try: return int(title.split(" ")[1])
            except: return 999 # Fallback for parsing error
        return 998 # Other unexpected format

    sorted_chapters = sorted(chapter_dict.items(), key=chapter_sort_key)

    # Iterate through sorted chapters
    for chapter_title, chapter_summary in sorted_chapters:
        print(f"Writing Chapter: {chapter_title} - {chapter_summary}")
        book_content[chapter_title] = [] # Initialize list for this chapter's paragraphs
        chapter_paragraphs_accumulator = "" # Accumulates paragraphs *within* the current chapter for context
        event_list = event_dict.get(chapter_title, []) # Get events for this chapter

        if not event_list:
             print(f"  Warning: No events found for {chapter_title}. Skipping content generation.")
             book_content[chapter_title].append(f"[Content generation skipped: No events defined for '{chapter_summary}']")
             continue

        for i, event in enumerate(event_list):
            print(f"  Writing event {i+1}/{len(event_list)}: {event}")
            # Call the writer chain for the current event
            new_paragraphs = writer_chain.run(
                genre=genre,
                author=author,
                title=title,
                profile=profile,
                plot=plot,
                chapter_name=chapter_title, # Pass current chapter name
                previous_events=previous_events_history, # Pass history of *all* previous events
                summary=chapter_summary, # Pass summary for the *current* chapter
                previous_paragraphs=chapter_paragraphs_accumulator, # Pass text written *so far in this chapter*
                current_event=event # The specific event to write about now
            )
            # Append the newly generated block of paragraphs to the chapter's content list
            book_content[chapter_title].append(new_paragraphs)
            # Update the accumulator for the *next* event in this chapter
            if chapter_paragraphs_accumulator: # Add separator if not the first block
                chapter_paragraphs_accumulator += "\n\n"
            chapter_paragraphs_accumulator += new_paragraphs
            # Add the current event to the overall history *after* writing about it
            previous_events_history.append(f"[{chapter_title}] {event}")

        print(f"  Finished writing content for {chapter_title}")

    print("--- Book Writing Complete ---")
    return book_content

# --- Document Writing Class ---

class DocWriter:
    def __init__(self, output_folder=OUTPUT_FOLDER):
        self.doc = docx.Document()
        self.output_folder = output_folder
        # Ensure output directory exists
        os.makedirs(self.output_folder, exist_ok=True)

    def _sanitize_filename(self, name):
        """Removes or replaces characters invalid for filenames."""
        # Remove characters known to be problematic
        name = name.replace(':', '-').replace('/', '-').replace('\\', '-').replace('?', '').replace('*', '')
        name = name.replace('"', '').replace('<', '').replace('>', '').replace('|', '')
        # Replace spaces with underscores
        name = name.replace(' ', '_')
        # Limit length if necessary (optional)
        # max_len = 100
        # name = name[:max_len]
        return name.strip()

    def write_doc(self, book_content, chapter_dict, title):
        """Writes the generated book content to a .docx file."""
        print("\n--- Writing Document ---")
        self.doc.add_heading(title, 0) # Book title as main heading

        # Use the same sorting logic as in write_book for consistency
        def chapter_sort_key(item):
            chap_title = item[0]
            if chap_title == "Prologue": return -1
            if chap_title == "Epilogue": return 9999
            if chap_title.startswith("Chapter "):
                try: return int(chap_title.split(" ")[1])
                except: return 999
            return 998

        sorted_chapters = sorted(chapter_dict.items(), key=chapter_sort_key)

        # Iterate through sorted chapters to add them to the document
        for chapter_title, description in sorted_chapters:
            # Combine chapter title and its description for the heading
            chapter_heading = f'{chapter_title.strip()}: {description.strip()}'
            self.doc.add_heading(chapter_heading, level=1) # Chapter heading (Level 1)
            print(f"  Adding Chapter to Doc: {chapter_heading}")

            # Get the list of generated paragraph blocks for this chapter
            paragraphs_list = book_content.get(chapter_title, [])

            if not paragraphs_list:
                self.doc.add_paragraph("[Content generation skipped or failed for this chapter]")
                print(f"    (Adding placeholder for empty chapter)")
                continue

            # Join the paragraph blocks (each block corresponds to one event's output)
            # Use double newline for separation between event blocks
            full_chapter_text = '\n\n'.join(paragraphs_list).strip()

            # Add the joined text as paragraphs to the document
            # (Alternatively, add each block as a separate paragraph if needed)
            self.doc.add_paragraph(full_chapter_text)
            print(f"    (Added {len(paragraphs_list)} content blocks)")


        # --- Filename Creation ---
        # Sanitize the title to create a safe filename
        safe_basename = self._sanitize_filename(title)
        if not safe_basename: # Handle empty title case
            safe_basename = "Untitled_Novel"
        filename = safe_basename + '.docx'
        output_path = os.path.join(self.output_folder, filename)

        # --- Saving Document ---
        try:
            self.doc.save(output_path)
            print(f"--- Document saved successfully as: {output_path} ---")
        except PermissionError:
            print(f"\nError: Permission denied trying to save '{output_path}'.")
            print("Please check if the file is open or if you have write permissions for the folder.")
        except Exception as e:
            print(f"\nError saving document to {output_path}: {e}")
            traceback.print_exc()

# --- Main Execution ---

def main():
    # --- Configuration ---
    # Ensure the output folder exists
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    print(f"Using output folder: {OUTPUT_FOLDER}")

    # --- Inputs ---
    resume_filename = 'Thalen.pdf' # MAKE SURE THIS FILE EXISTS in ./docs
    subject = "Story of a young architect living in a post ufo dominated world. (USE REAL PLACES FROM THE RESUME AS LOCATIONS IN THE STORY AND MAKE IT A Slice OF LIFE STORY)" 
    author_style = 'Ana Huang' # Example author style
    genre = '''SiFi, Fiction, Life, Work, Struggles''' # Example genre

    print("\n--- Starting Novel Generation ---")
    print(f"Resume File: {resume_filename}")
    print(f"Subject: {subject}")
    print(f"Author Style: {author_style}")
    print(f"Genre: {genre}")

    # --- Instantiate Chains & Writer ---
    try:
        main_character_chain = MainCharacterChain()
        title_chain = TitleChain()
        plot_chain = PlotChain()
        chapters_chain = ChaptersChain()
        doc_writer = DocWriter(output_folder=OUTPUT_FOLDER)
        print("Chains initialized successfully.")
    except Exception as e:
        print(f"Fatal Error: Failed to initialize LangChain components: {e}")
        traceback.print_exc()
        return # Cannot proceed if chains fail

    # --- Generate Novel Components Sequentially ---
    try:
        print("\nStep 1: Generating Main Character Profile...")
        profile = main_character_chain.run(resume_filename)
        if profile.startswith("Error:"):
            print(f"Profile Generation Failed: {profile}")
            print("Cannot proceed without a character profile.")
            return
        print(f"Profile Generated:\n---\n{profile}\n---")

        print("\nStep 2: Generating Title...")
        title = title_chain.run(subject, genre, author_style, profile)
        if title == "Error Generating Title":
             print("Title Generation Failed. Using placeholder.")
             title = "Placeholder Title - Generation Failed"
        print(f"Title Generated: '{title}'")

        print("\nStep 3: Generating Plot...")
        plot = plot_chain.run(subject, genre, author_style, profile, title)
        if plot.startswith("Error generating plot:"):
            print(f"Plot Generation Failed: {plot}")
            print("Cannot proceed without a plot.")
            return
        print(f"Plot Generated:\n---\n{plot}\n---")

        print("\nStep 4: Generating Chapters...")
        # chapter_dict should be { 'Chapter 1': 'Description...', ... }
        chapter_dict = chapters_chain.run(subject, genre, author_style, profile, title, plot)
        if not chapter_dict:
             print("ERROR: Failed to generate or parse chapters. Cannot write book.")
             return # Stop if chapters couldn't be generated/parsed
        print("Chapters Generated:")
        for name, desc in chapter_dict.items():
             print(f"  - {name}: {desc}")

        # --- Write the Book Content ---
        # This function now orchestrates event generation and paragraph writing
        book_content = write_book(
            genre,
            author_style,
            title,
            profile,
            plot,
            chapter_dict # Pass the generated chapter dictionary
        )

        # --- Save to Document ---
        doc_writer.write_doc(book_content, chapter_dict, title)

    except FileNotFoundError as e:
         print(f"\nCRITICAL ERROR: {e}")
         print(f"Please ensure the required file exists in the '{OUTPUT_FOLDER}' directory.")
    except Exception as e:
        print(f"\nAn unexpected error occurred during the main generation process: {e}")
        traceback.print_exc() # Print detailed traceback for unexpected errors

    print("\n--- Novel Generation Process Finished ---")


if __name__ == "__main__":
    main()