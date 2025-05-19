from dotenv import load_dotenv
import os
# Langchain imports
from langchain_community.chat_models import ChatOllama # Changed from ChatGroq
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import PyPDFLoader # Common location for PyPDFLoader
import docx

# Load environment variables (if any are still needed, e.g., for other services)
load_dotenv()

# Constants
DEFAULT_MODEL = "qwen3" # IMPORTANT: Replace "llama3" with the actual model name you have pulled in Ollama.
                         # For example, if you pulled "llama3:70b", use "llama3:70b".

def create_llm():
    """Create and return a ChatOllama instance"""
    # Ensure your Ollama service is running.
    # By default, ChatOllama tries to connect to http://localhost:11434
    print(f"Attempting to use Ollama model: {DEFAULT_MODEL}")
    return ChatOllama(
        model=DEFAULT_MODEL,
        temperature=0.7,
        # You can add base_url if your Ollama is not on localhost:11434
        # base_url="http://your_ollama_host:port" 
    )

class MainCharacterChain:
    PROMPT = """
    You are provided with the resume of a person. 
    Describe the person's profile in a few sentences and include that person's name.

    Resume: {text}

    Profile:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def load_resume(self, file_name):
        folder = './docs'
        file_path = os.path.join(folder, file_name)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}. Please ensure it's in a 'docs' subdirectory.")
        loader = PyPDFLoader(file_path)
        docs = loader.load_and_split()
        return docs

    def run(self, file_name):
        docs = self.load_resume(file_name)
        resume_text = '\n\n'.join([doc.page_content for doc in docs])
        return self.chain.run(text=resume_text)

class TitleChain:
    PROMPT = """
    Your job is to generate the title for a novel about the following subject and main character. 
    Return a title and only a title!
    The title should be consistent with the genre of the novel.
    The title should be consistent with the style of the author.

    Subject: {subject}
    Genre: {genre}
    Author: {author}
    Main character's profile: {profile}

    Title:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile):
        return self.chain.run(
            subject=subject,
            genre=genre,
            author=author,
            profile=profile
        )

class PlotChain:
    PROMPT = """
    Your job is to generate the plot for a novel. Return a plot and only a plot!
    Describe the full plot of the story and don't hesitate to create new characters to make it compelling.
    You are provided the following subject, title and main character's profile.
    Make sure that the main character is at the center of the story.
    The plot should be consistent with the genre of the novel.
    The plot should be consistent with the style of the author.

    Consider the following attributes to write an exciting story:
    {features}

    Subject: {subject}
    Genre: {genre}
    Author: {author}
    Title: {title}
    Main character's profile: {profile}

    Plot:"""

    HELPER_PROMPT = """
    Generate a list of 5-7 key attributes or elements that characterize an exciting and compelling story plot.
    Examples: Strong protagonist, clear conflict, rising tension, surprising twists, satisfying resolution, thematic depth, vivid world-building.
    
    Input: Generate attributes for an exciting story plot.
    
    List of attributes:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )
        self.helper_chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate(
                template=self.HELPER_PROMPT,
                input_variables=["input"]
            ),
            verbose=True 
        )

    def run(self, subject, genre, author, profile, title):
        features = self.helper_chain.run(input="Generate attributes for an exciting story plot.")
        return self.chain.run(
            features=features,
            subject=subject,
            genre=genre,
            author=author,
            profile=profile,
            title=title
        )

class ChaptersChain:
    PROMPT = """
    Your job is to generate a list of chapter titles and concise one-sentence summaries for a novel. 
    ONLY the list and nothing more!
    You are provided with a subject, genre, author style, title, main character's profile, and the overall plot for a novel.
    Generate a list of chapters (including a Prologue and Epilogue if appropriate for the genre) that logically break down the plot.
    Each chapter entry should describe a key stage or segment of the story.
    Make sure the chapters are consistent with the plot, genre, and author's style.

    Follow this template strictly:
    Prologue: [One-sentence summary of prologue]
    Chapter 1: [One-sentence summary of chapter 1]
    Chapter 2: [One-sentence summary of chapter 2]
    ...
    Chapter N: [One-sentence summary of chapter N]
    Epilogue: [One-sentence summary of epilogue]

    Ensure each chapter title (e.g., "Prologue", "Chapter 1") is followed by a colon and then its one-sentence summary.
    
    Subject: {subject}
    Genre: {genre}
    Author: {author}
    Title: {title}
    Main character's profile: {profile}
    Plot: {plot}
    
    Chapters list:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def parse(self, response):
        chapter_list = response.strip().split('\n')
        parsed_chapters = {}
        for chapter_entry in chapter_list:
            if ':' in chapter_entry:
                parts = chapter_entry.split(':', 1) 
                if len(parts) == 2:
                    chapter_title = parts[0].strip()
                    chapter_description = parts[1].strip()
                    if chapter_title and chapter_description:
                         parsed_chapters[chapter_title] = chapter_description
        return parsed_chapters

    def run(self, subject, genre, author, profile, title, plot):
        response = self.chain.run(
            subject=subject,
            genre=genre,
            author=author,
            profile=profile,
            title=title,
            plot=plot
        )
        return self.parse(response)

class ChapterFlowCheckerChain:
    PROMPT = """
    You are a meticulous literary editor responsible for ensuring narrative coherence and smooth progression in a novel.
    You have been provided with the overall plot, main character profile, novel title, genre, author's style, and a list of chapter titles with their respective summaries.

    Your primary task is to critically analyze the sequence of these chapters. Evaluate if they collectively form a continuous, logical, and engaging narrative flow. Consider the following:
    1.  **Continuity**: Does each chapter naturally follow from the preceding one? Are there any abrupt jumps, unexplained gaps, or inconsistencies in the timeline or character development between chapters?
    2.  **Pacing and Progression**: Does the story advance at an appropriate pace through the chapters? Does each chapter contribute meaningfully to the overall plot, or do some feel redundant or out of place?
    3.  **Character Arc**: Is the main character's journey and development consistently portrayed and advanced across the chapters?
    4.  **Thematic Cohesion**: Do the chapters work together to explore the novel's central themes effectively?
    5.  **Engagement**: Would a reader find the progression from one chapter to the next logical and compelling, or would they feel disoriented or that chapters seem like standalone pieces?

    Provide a concise overall assessment of the chapter flow (e.g., "Excellent flow," "Generally good flow with minor concerns," "Significant flow issues noted").
    If you identify any issues, please be specific:
    - Pinpoint the chapter(s) or transition(s) that are problematic.
    - Clearly explain *why* it feels disjointed, confusing, or detrimental to the narrative flow.
    - If possible, offer a brief suggestion on what might improve the flow (e.g., "Consider adding a bridging scene before Chapter X," or "The motivation for Y in Chapter Z needs to be clearer based on Chapter Y-1.").

    Novel Details:
    Title: {title}
    Genre: {genre}
    Author's Style: {author_style}
    Main Character Profile: {profile}
    Overall Plot: {plot}

    Chapter Summaries:
    {formatted_chapter_summaries}

    Editor's Assessment of Chapter Flow:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def _format_chapters_for_prompt(self, chapter_dict):
        """Converts the chapter dictionary to a formatted string for the prompt."""
        return "\n".join([f"- {title}: {desc}" for title, desc in chapter_dict.items()])

    def run(self, plot, profile, title, genre, author_style, chapter_dict):
        formatted_summaries = self._format_chapters_for_prompt(chapter_dict)
        return self.chain.run(
            plot=plot,
            profile=profile,
            title=title,
            genre=genre,
            author_style=author_style,
            formatted_chapter_summaries=formatted_summaries
        )

class WriterChain:
    PROMPT = """
    You are a master novelist, tasked with writing a segment of a novel based on a specific event.
    You must adhere strictly to the provided genre, author's style, and the established context of the novel, including the main character, overall plot, current chapter's summary, previous events in the story, and the immediately preceding paragraphs you've already written for this chapter.

    Your writing should seamlessly continue from the 'Previous paragraphs'. The 'New event to write about' is the immediate focus for the text you generate.
    Ensure your generated paragraphs:
    - Directly and vividly describe the 'New event to write about'.
    - Maintain consistency with the 'Novel's Plot' and the 'Current Chapter summary'.
    - Reflect the 'Main character's profile' in their actions, thoughts, and dialogue (if any).
    - Are stylistically aligned with the specified 'Genre' and 'Author's style'.
    - Logically follow from the 'Previous events' and 'Previous paragraphs', ensuring smooth narrative flow.
    - Avoid repetition and advance the story.

    Context:
    Genre: {genre}
    Author's Style: {author}
    Novel Title: {title}
    Main Character's Profile: {profile}
    Novel's Overall Plot: {plot}
    
    Current Chapter Details:
    Current Chapter Summary: {summary}
    Events Already Narrated in the Novel: {previous_events}
    Previously Written Paragraphs in This Chapter:
    ---
    {previous_paragraphs}
    ---

    Your Current Task:
    New event to write about for this segment: {current_event}

    Begin writing the paragraphs for this new event now:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, genre, author, title, profile, plot, 
            previous_events, summary, previous_paragraphs, current_event):
        previous_events_text = '\n'.join([f"- {event}" for event in previous_events]) if previous_events else "This is the beginning of the story."
        previous_paragraphs_text = previous_paragraphs if previous_paragraphs else "This is the beginning of the chapter."
        
        return self.chain.run(
            genre=genre,
            author=author,
            title=title,
            profile=profile,
            plot=plot,
            previous_events=previous_events_text,
            summary=summary,
            previous_paragraphs=previous_paragraphs_text,
            current_event=current_event
        )

def write_book(genre, author, title, profile, plot, summaries_dict, event_dict):
    writer_chain = WriterChain()
    all_previous_events_narrated = [] # Tracks all events written across all chapters
    book = {} # Stores {"Chapter Title": [paragraph_string_for_event_1, paragraph_string_for_event_2]}
    
    for chapter_title, event_list_for_chapter in event_dict.items():
        book[chapter_title] = []
        current_chapter_paragraphs_text = '' # Accumulates paragraphs for the current chapter
        print(f"\n--- Writing Chapter: {chapter_title} ---")
        
        current_chapter_summary = summaries_dict.get(chapter_title, f"Summary for {chapter_title} not found.")

        for i, event_description in enumerate(event_list_for_chapter):
            print(f"  - Writing event ({i+1}/{len(event_list_for_chapter)}): {event_description}")
            
            paragraphs_for_event = writer_chain.run(
                genre=genre,
                author=author,
                title=title,
                profile=profile,
                plot=plot,
                previous_events=all_previous_events_narrated, # Pass all previously narrated events in the book
                summary=current_chapter_summary,
                previous_paragraphs=current_chapter_paragraphs_text, # Pass text of current chapter so far
                current_event=event_description
            )
            book[chapter_title].append(paragraphs_for_event)
            # Add a clear separator if concatenating; often, each event's text is a distinct block.
            current_chapter_paragraphs_text += ("\n\n" if current_chapter_paragraphs_text else "") + paragraphs_for_event
            all_previous_events_narrated.append(f"In {chapter_title}: {event_description}")
        print(f"--- Finished Chapter: {chapter_title} ---")
    return book

class DocWriter:
    def __init__(self):
        self.doc = docx.Document()

    def write_doc(self, book, chapter_dict, title, output_folder='./docs'):
        os.makedirs(output_folder, exist_ok=True)
        self.doc.add_heading(title, 0)
        for chapter_key_from_book, paragraphs_list_for_chapter in book.items():
            description = chapter_dict.get(chapter_key_from_book.strip(), f"Description for {chapter_key_from_book} not found.")
            chapter_name_display = f'{chapter_key_from_book.strip()}: {description.strip()}'
            self.doc.add_heading(chapter_name_display, 1)
            
            full_chapter_text = '\n\n'.join(paragraphs_list_for_chapter)
            self.doc.add_paragraph(full_chapter_text)
        
        filename = "".join(x for x in title if x.isalnum() or x in (' ', '-', '_')).strip().replace(' ', '_')
        filename = (filename if filename else "Untitled_Novel") + '.docx'
        
        output_path = os.path.join(output_folder, filename)
        self.doc.save(output_path)
        print(f"Book saved as: {output_path}")

# --- Placeholder for a more advanced Event Generation Chain ---
class EventGeneratorChain:
    PROMPT = """
    You are an expert story structure analyst. Given a chapter title, a concise summary of that chapter, the overall novel plot, and the main character's profile, your task is to break down the chapter summary into a sequence of 2 to 4 distinct, actionable key events that must occur within this chapter.

    These events should:
    1. Logically progress the narrative described in the chapter summary.
    2. Involve or significantly impact the main character.
    3. Build upon each other to create a mini-arc within the chapter.
    4. Be distinct enough to be narrated as separate segments.

    Provide ONLY the list of events, each on a new line. Do not add any other text, preamble, or explanation.

    Novel Plot: {plot}
    Main Character Profile: {profile}
    Chapter Title: {chapter_title}
    Chapter Summary: {chapter_summary}

    Key events for this chapter:
    """
    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, plot, profile, chapter_title, chapter_summary):
        response = self.chain.run(
            plot=plot,
            profile=profile,
            chapter_title=chapter_title,
            chapter_summary=chapter_summary
        )
        # Parse the response: split by newline and filter out empty strings
        events = [event.strip() for event in response.strip().split('\n') if event.strip()]
        return events if events else [f"Key development based on summary: {chapter_summary[:100]}..."] # Fallback

def main():
    # --- Configuration ---
    subject = 'The Last Signal from Andromeda' 
    author_style = 'Isaac Asimov with a touch of modern pacing and detailed scientific explanations'
    genre = 'Hard Sci-Fi Mystery with elements of cosmic horror'
    resume_filename = 'kenji_gamer_resume.pdf' # Ensure this file exists in ./docs/
    # --- End Configuration ---

    os.makedirs('./docs', exist_ok=True)
    print("Starting novel generation process...")

    try:
        main_character_chain = MainCharacterChain()
        print(f"Loading profile from '{resume_filename}'...")
        profile = main_character_chain.run(resume_filename)
        print(f"Profile generated: {profile}\n")

        doc_writer = DocWriter()
        title_chain = TitleChain()
        plot_chain = PlotChain()
        chapters_chain = ChaptersChain()
        chapter_flow_checker_chain = ChapterFlowCheckerChain() # Instantiate new chain
        event_generator_chain = EventGeneratorChain() # Instantiate event generator

        print("Generating title...")
        title = title_chain.run(subject, genre, author_style, profile).strip().replace('"', '') # Clean title
        print(f"Title generated: {title}\n")

        print("Generating plot...")
        plot = plot_chain.run(subject, genre, author_style, profile, title)
        print(f"Plot generated: {plot}\n")

        print("Generating chapters...")
        chapter_dict = chapters_chain.run(subject, genre, author_style, profile, title, plot)
        if not chapter_dict:
            print("Chapter generation failed or returned empty. Exiting.")
            return
        print("Chapters generated:")
        for ch_title, ch_desc in chapter_dict.items():
            print(f"  {ch_title}: {ch_desc}")
        print("\n")

        # --- Check Chapter Flow ---
        print("Checking chapter flow and coherence...")
        flow_assessment = chapter_flow_checker_chain.run(
            plot=plot,
            profile=profile,
            title=title,
            genre=genre,
            author_style=author_style,
            chapter_dict=chapter_dict
        )
        print("--- Chapter Flow Assessment ---")
        print(flow_assessment)
        print("--- End of Assessment ---\n")
        # At this point, you could add logic to halt or modify if the assessment is poor.

        # --- Generate Events for each chapter using LLM ---
        summaries_dict = {} # This will be the same as chapter_dict for summaries
        event_dict = {}     # This will store {chapter_title: [event1, event2, ...]}
        
        print("Generating events for each chapter using LLM...")
        for chapter_title_key, chapter_summary in chapter_dict.items():
            summaries_dict[chapter_title_key] = chapter_summary 
            print(f"  Generating events for: {chapter_title_key} - {chapter_summary[:50]}...")
            events_for_chapter = event_generator_chain.run(
                plot=plot,
                profile=profile,
                chapter_title=chapter_title_key,
                chapter_summary=chapter_summary
            )
            event_dict[chapter_title_key] = events_for_chapter
            print(f"    Events: {events_for_chapter}")
        print("Event generation complete.\n")
        # --- End Event Generation ---

        print("Starting to write the book content...")
        book = write_book(
            genre,
            author_style,
            title,
            profile,
            plot,
            summaries_dict, 
            event_dict
        )
        print("Book content generation complete.\n")

        print("Saving book to .docx file...")
        doc_writer.write_doc(book, chapter_dict, title)
        print("Novel generation process finished successfully!")

    except FileNotFoundError as e:
        print(f"Error: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
