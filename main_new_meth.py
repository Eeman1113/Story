from dotenv import load_dotenv
import os
import docx
import traceback  # For printing detailed errors

# --- Langchain Imports - UPDATED ---
# Use the dedicated package for Ollama
from langchain_ollama import OllamaLLM
# Core Langchain components
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
# Community loader for PDF
from langchain_community.document_loaders import PyPDFLoader
# --- End Imports ---

# Load environment variables (optional, but good practice)
load_dotenv()

# --- Constants ---
DEFAULT_MODEL = "gemma3:4b"  # Or specify a variant like "gemma3:latest", "gemma3:9b" etc.
OUTPUT_FOLDER = './docs'  # Define output folder consistently

# --- LLM Initialization Function ---
def create_llm():
    """Create and return an OllamaLLM instance"""
    return OllamaLLM(
        model=DEFAULT_MODEL,
        temperature=0.7,
    )

# --- Original Character & Title Chains (These remain useful) ---

class MainCharacterChain:
    PROMPT = """
    You are provided with the resume text of a person.
    Describe the person's profile concisely in 2-3 sentences. Include the person's name if explicitly mentioned in the text.

    Resume Text:
    {text}

    Profile Description:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def load_resume(self, file_name):
        """Loads text content from a PDF file."""
        file_path = os.path.join(OUTPUT_FOLDER, file_name)
        if not os.path.exists(file_path):
             raise FileNotFoundError(f"Resume file not found at: {file_path}")
        try:
            print(f"Loading PDF from: {file_path}")
            loader = PyPDFLoader(file_path)
            # Use load() which returns a list of Document objects, one per page
            docs = loader.load()
            if not docs:
                print(f"Warning: PyPDFLoader didn't extract any documents/pages from {file_name}.")
                return None  # Indicate no content found
            print(f"Successfully loaded {len(docs)} page(s) from PDF.")
            return docs
        except Exception as e:
             print(f"Error loading PDF {file_path}: {e}")
             traceback.print_exc()  # Print full traceback for PDF errors
             raise  # Reraise the exception

    def run(self, file_name):
        """Loads resume and generates the profile description."""
        try:
            docs = self.load_resume(file_name)
            if not docs:
                 print("Could not load resume content.")
                 return "Error: Could not generate profile due to missing resume content."

            # Combine text from all pages
            resume_text = '\n\n'.join([doc.page_content for doc in docs if doc.page_content])

            if not resume_text.strip():
                 print("Warning: Resume content appears empty after extraction.")
                 return "Error: Could not generate profile because the resume content is empty."

            # Use invoke which returns a dictionary
            result = self.chain.invoke({"text": resume_text})
            return result.get('text', "Error: Profile generation failed.").strip()

        except Exception as e:
            print(f"An error occurred in MainCharacterChain.run: {e}")
            traceback.print_exc()
            return f"Error generating profile: {e}"


class TitleChain:
    PROMPT = """
    Generate a compelling novel title based on the provided details.
    Return ONLY the title itself, without any quotation marks, labels, or extra text.
    The title must be consistent with the specified genre and author's style.

    Subject: {subject}
    Genre: {genre}
    Author's Style: {author}
    Main Character Profile: {profile}

    Novel Title:"""

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )

    def run(self, subject, genre, author, profile):
        """Generates the novel title."""
        try:
            result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile
            })
            # Clean the output string thoroughly
            title = result.get('text', "Untitled").strip()
            # Remove potential surrounding quotes
            if title.startswith('"') and title.endswith('"'):
                title = title[1:-1]
            if title.startswith("'") and title.endswith("'"):
                title = title[1:-1]
            return title
        except Exception as e:
            print(f"An error occurred in TitleChain.run: {e}")
            traceback.print_exc()
            return "Error Generating Title"


# --- New Plot Promise System Classes ---

class PromiseGeneratorChain:
    PROMPT = """
    Generate a set of compelling 'plot promises' for a novel based on these details:
    
    Subject: {subject}
    Genre: {genre}
    Author's Style: {author}
    Title: "{title}"
    Main Character Profile: {profile}
    
    A plot promise is a narrative thread that gets introduced, developed, and eventually paid off.
    As Brandon Sanderson describes: "A plot promise is a promise of something that will happen later 
    in the story. It sets expectations early, then builds tension through obstacles, twists, and 
    turning points—culminating in a powerful, satisfying climax."
    
    Examples:
    1. "Character will learn a magic spell that gives them super-strength" (importance: 8/10)
    2. "The mysterious artifact will be discovered and its powers revealed" (importance: 7/10)
    3. "Two characters will develop a relationship despite initial conflict" (importance: 5/10)
    
    For each promise, provide:
    - A clear description of the promise
    - An importance score (1-10, higher = more central to the story)
    - A brief note on the expected payoff (what's the climactic moment?)
    
    Generate 5-7 varied plot promises that would create an engaging narrative together.
    Format each promise exactly as:
    
    PROMISE: [Description]
    IMPORTANCE: [1-10]/10
    PAYOFF: [Brief description of eventual resolution]
    """

    def __init__(self):
        self.llm = create_llm()
        self.chain = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMPT),
            verbose=True
        )
    
    def parse_promises(self, raw_text):
        """Parse the LLM response into structured promise objects"""
        promises = []
        current_promise = {}
        
        lines = raw_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith("PROMISE:"):
                # Save previous promise if it exists
                if current_promise and 'description' in current_promise:
                    promises.append(current_promise)
                # Start new promise
                current_promise = {'description': line[len("PROMISE:"):].strip()}
            elif line.startswith("IMPORTANCE:"):
                if current_promise:
                    try:
                        # Handle both "8/10" and just "8" formats
                        importance_text = line[len("IMPORTANCE:"):].strip()
                        if '/' in importance_text:
                            importance = int(importance_text.split('/')[0])
                        else:
                            importance = int(importance_text)
                        current_promise['importance'] = min(max(importance, 1), 10)  # Ensure 1-10 range
                    except ValueError:
                        current_promise['importance'] = 5  # Default if parsing fails
            elif line.startswith("PAYOFF:"):
                if current_promise:
                    current_promise['payoff'] = line[len("PAYOFF:"):].strip()
        
        # Add the last promise if it exists
        if current_promise and 'description' in current_promise:
            promises.append(current_promise)
            
        # Ensure all promises have necessary fields
        for promise in promises:
            if 'importance' not in promise:
                promise['importance'] = 5
            if 'payoff' not in promise:
                promise['payoff'] = "Reaches a satisfying conclusion"
            # Initialize state
            promise['progress'] = 0
            promise['complete'] = False
            promise['last_progressed'] = 0  # Scene counter
            
        return promises

    def run(self, subject, genre, author, profile, title):
        """Generates the initial set of plot promises."""
        try:
            result = self.chain.invoke({
                "subject": subject,
                "genre": genre,
                "author": author,
                "profile": profile,
                "title": title
            })
            
            raw_promises = result.get('text', "").strip()
            promises = self.parse_promises(raw_promises)
            
            if not promises:
                # Fallback if parsing fails
                print("Warning: Failed to parse promises. Using fallbacks.")
                promises = [
                    {
                        'description': f"The main character will face a major challenge related to {subject}",
                        'importance': 8,
                        'payoff': "Character overcomes the challenge",
                        'progress': 0,
                        'complete': False,
                        'last_progressed': 0
                    },
                    {
                        'description': f"A mystery related to {genre} will be introduced and solved",
                        'importance': 7,
                        'payoff': "Mystery is solved in a surprising way",
                        'progress': 0,
                        'complete': False,
                        'last_progressed': 0
                    }
                ]
                
            return promises
        except Exception as e:
            print(f"An error occurred in PromiseGeneratorChain.run: {e}")
            traceback.print_exc()
            return []  # Return empty list on error


class PromiseManager:
    """Manages the set of plot promises and their progression"""
    
    def __init__(self, initial_promises):
        self.promises = initial_promises.copy()  # Make a copy to avoid modifying original
        self.scene_counter = 0
        self.completed_promises = []
        self.story_events = []  # Track major events that have occurred
        
    def add_promise(self, description, importance=5, payoff="Reaches a satisfying conclusion"):
        """Add a new promise to the active set"""
        new_promise = {
            'description': description,
            'importance': importance,
            'payoff': payoff,
            'progress': 0,
            'complete': False,
            'last_progressed': self.scene_counter
        }
        self.promises.append(new_promise)
        print(f"Added new promise: '{description}' (Importance: {importance})")
        return len(self.promises) - 1  # Return index of new promise
        
    def mark_promise_progressed(self, index, event_description):
        """Mark a promise as having progressed and record the event"""
        if 0 <= index < len(self.promises):
            promise = self.promises[index]
            promise['progress'] += 1
            promise['last_progressed'] = self.scene_counter
            
            # Record this event
            self.story_events.append({
                'scene': self.scene_counter,
                'promise_index': index,
                'promise': promise['description'],
                'event': event_description,
                'progress_level': promise['progress']
            })
            
            # Check if this was the payoff (simplified logic)
            if promise['progress'] >= 3:  # Assuming 3+ steps means complete
                promise['complete'] = True
                self.completed_promises.append(promise.copy())  # Store a copy
                self.promises.pop(index)
                print(f"Promise completed and removed: '{promise['description']}'")
            
            return True
        return False
    
    def suggest_next_promises(self, count=3):
        """Suggest the next most appropriate promises to progress"""
        if not self.promises:
            return []
            
        # Score each promise based on importance and recency
        scored_promises = []
        for i, promise in enumerate(self.promises):
            # Higher score = more urgently needs progression
            # Factor in importance and how long since last progressed
            recency_factor = self.scene_counter - promise['last_progressed']
            score = promise['importance'] + (recency_factor * 0.5)
            scored_promises.append((i, promise, score))
        
        # Sort by score (highest first)
        scored_promises.sort(key=lambda x: x[2], reverse=True)
        
        # Return the indices and descriptions of top suggestions
        return [(idx, promise['description']) for idx, promise, _ in scored_promises[:count]]
        
    def increment_scene(self):
        """Increment the scene counter after a scene is written"""
        self.scene_counter += 1
        
    def get_active_promises(self):
        """Return all active promises"""
        return self.promises
        
    def get_story_summary(self, recent_count=None):
        """Get a summary of recent story events"""
        if recent_count is None:
            events = self.story_events
        else:
            events = self.story_events[-recent_count:]
            
        summary = []
        for event in events:
            summary.append(f"Scene {event['scene']}: {event['event']} ({event['promise']})")
        
        return summary


class SceneGeneratorChain:
    PROMISE_SELECTION_PROMPT = """
    You are an expert narrative designer choosing which plot thread to develop next.
    
    Current story status:
    - Title: "{title}"
    - Genre: {genre}
    - Main Character: {profile}
    - Current scene number: {scene_number}
    
    Recent story events:
    {recent_events}
    
    The system suggests these plot promises as potential focus points for the next scene:
    {suggested_promises}
    
    Based on narrative flow, pacing, and context from recent events, which ONE promise should be developed next?
    Consider what would create the most compelling next scene and maintain good story rhythm.
    
    Return ONLY the index number of your selected promise, nothing more.
    """
    
    SCENE_WRITING_PROMPT = """
    You are writing the next scene in a {genre} novel titled "{title}" in the style of {author}.
    
    Main Character: {profile}
    
    Recent story events:
    {recent_events}
    
    You are now developing this plot promise:
    "{selected_promise}"
    
    Current progress level: {progress_level} (higher numbers = closer to payoff)
    
    Guidelines for this scene:
    - The scene should clearly advance the selected plot promise
    - Maintain consistent characterization and world-building
    - Consider the natural progression from recent events
    - Write vivid, engaging prose in {author}'s style
    - If this is a climactic moment (progress level 3+), make it especially impactful
    
    Write a complete scene (500-800 words) that develops this plot promise.
    """
    
    EVENT_SUMMARY_PROMPT = """
    Summarize the key events that occurred in this scene in one concise sentence:
    
    Scene content:
    {scene_content}
    
    Plot promise being developed:
    {promise_description}
    
    One-sentence summary of what happened:
    """

    def __init__(self):
        self.llm = create_llm()
        self.promise_selector = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.PROMISE_SELECTION_PROMPT),
            verbose=True
        )
        self.scene_writer = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.SCENE_WRITING_PROMPT),
            verbose=True
        )
        self.event_summarizer = LLMChain(
            llm=self.llm,
            prompt=PromptTemplate.from_template(self.EVENT_SUMMARY_PROMPT),
            verbose=True
        )
    
    def select_promise(self, title, genre, profile, scene_number, recent_events, suggested_promises):
        """Use LLM to select which promise to progress next"""
        # Format the suggested promises for the prompt
        promise_list = "\n".join([f"{idx}: {desc}" for idx, desc in suggested_promises])
        recent_events_text = "\n".join(recent_events) if recent_events else "No previous events yet."
        
        try:
            result = self.promise_selector.invoke({
                "title": title,
                "genre": genre,
                "profile": profile,
                "scene_number": scene_number,
                "recent_events": recent_events_text,
                "suggested_promises": promise_list
            })
            
            # Extract the selected index from the response
            selection_text = result.get('text', "").strip()
            
            # Try to find a number in the response
            for word in selection_text.split():
                word = word.strip('.,;:()"\'')  # Remove potential punctuation
                if word.isdigit():
                    selected_index = int(word)
                    # Verify it's a valid index
                    for idx, _ in suggested_promises:
                        if idx == selected_index:
                            return selected_index
            
            # If we couldn't parse or find a valid index, use the first suggestion
            if suggested_promises:
                return suggested_promises[0][0]
            return None
            
        except Exception as e:
            print(f"Error selecting promise: {e}")
            # Fallback to first promise if there is one
            if suggested_promises:
                return suggested_promises[0][0]
            return None
    
    def write_scene(self, title, genre, author, profile, recent_events, 
                    selected_promise, progress_level):
        """Generate the actual scene content"""
        recent_events_text = "\n".join(recent_events) if recent_events else "This is the beginning of the story."
        
        try:
            result = self.scene_writer.invoke({
                "title": title,
                "genre": genre,
                "author": author,
                "profile": profile,
                "recent_events": recent_events_text,
                "selected_promise": selected_promise,
                "progress_level": progress_level
            })
            
            scene_content = result.get('text', "").strip()
            if not scene_content:
                scene_content = f"[Error generating scene for promise: {selected_promise}]"
            
            return scene_content
        
        except Exception as e:
            print(f"Error writing scene: {e}")
            return f"[Error occurred while writing scene: {e}]"
    
    def summarize_event(self, scene_content, promise_description):
        """Create a concise summary of what happened in this scene"""
        try:
            # Truncate scene content if it's very long to avoid token limits
            max_chars = 2000  # Reasonable limit for context
            truncated_content = scene_content[:max_chars]
            if len(scene_content) > max_chars:
                truncated_content += "... [content truncated]"
                
            result = self.event_summarizer.invoke({
                "scene_content": truncated_content,
                "promise_description": promise_description
            })
            
            summary = result.get('text', "").strip()
            if not summary:
                # Create a basic fallback summary
                summary = f"A scene related to '{promise_description}' occurred."
            
            return summary
            
        except Exception as e:
            print(f"Error summarizing event: {e}")
            return f"Scene related to '{promise_description}'"


class DocWriter:
    def __init__(self, output_folder=OUTPUT_FOLDER):
        self.doc = docx.Document()
        self.output_folder = output_folder
        # Ensure output directory exists
        os.makedirs(self.output_folder, exist_ok=True)

    def _sanitize_filename(self, name):
        """Removes or replaces characters invalid for filenames."""
        # Remove characters known to be problematic
        name = name.replace(':', '-').replace('/', '-').replace('\\', '-').replace('?', '').replace('*', '')
        name = name.replace('"', '').replace('<', '').replace('>', '').replace('|', '')
        # Replace spaces with underscores
        name = name.replace(' ', '_')
        return name.strip()

    def write_promise_book(self, book_content, title, profile, genre, story_summary):
        """Writes a book created with the Plot Promise system to a document"""
        self.doc.add_heading(title, 0)  # Book title as main heading
        
        # Add book metadata
        self.doc.add_heading("About This Book", level=1)
        self.doc.add_paragraph(f"Genre: {genre}")
        self.doc.add_paragraph(f"Main Character: {profile}")
        
        # Add story summary for reference
        self.doc.add_heading("Story Events Summary", level=1)
        for event in story_summary:
            self.doc.add_paragraph(event)
        
        # Add content
        self.doc.add_heading("Story Content", level=1)
        for scene in book_content:
            self.doc.add_heading(scene["title"], level=2)
            self.doc.add_paragraph(scene["content"])
        
        # Save document
        safe_basename = self._sanitize_filename(title)
        if not safe_basename:
            safe_basename = "Untitled_Novel"
        filename = safe_basename + '.docx'
        output_path = os.path.join(self.output_folder, filename)
        
        try:
            self.doc.save(output_path)
            print(f"--- Document saved successfully as: {output_path} ---")
        except PermissionError:
            print(f"\nError: Permission denied trying to save '{output_path}'.")
            print("Please check if the file is open or if you have write permissions for the folder.")
        except Exception as e:
            print(f"\nError saving document to {output_path}: {e}")
            traceback.print_exc()

# --- Book Writing Orchestration ---

def write_book_with_promises(genre, author, title, profile, initial_promises, total_scenes=15):
    """Generates a book using the Plot Promise approach"""
    # Initialize the promise manager with initial promises
    promise_manager = PromiseManager(initial_promises)
    scene_generator = SceneGeneratorChain()
    
    book_content = []  # Will store {"title": scene_title, "content": scene_content}
    
    print("\n--- Starting Book Generation with Plot Promise System ---")
    print(f"Initial promises: {len(initial_promises)}")
    
    # Generate a series of scenes
    for scene_num in range(1, total_scenes + 1):
        print(f"\nGenerating Scene {scene_num}/{total_scenes}")
        
        # Get suggestions for which promises to progress next
        suggested_promises = promise_manager.suggest_next_promises(count=3)
        if not suggested_promises:
            print("No active promises remaining. Creating a new promise to continue the story...")
            # Create a new simple promise to continue (in a real system, this would be more sophisticated)
            new_idx = promise_manager.add_promise(
                f"A new challenge emerges for the main character in scene {scene_num}", 
                importance=7
            )
            suggested_promises = [(new_idx, promise_manager.promises[new_idx]['description'])]
            
        # Get recent events for context
        recent_events = promise_manager.get_story_summary(recent_count=5)
        
        # Let the AI select which promise to progress
        selected_idx = scene_generator.select_promise(
            title, genre, profile, scene_num, recent_events, suggested_promises
        )
        
        if selected_idx is None:
            print("Error: Could not select a promise to progress. Using default selection.")
            # Use fallback selection
            selected_idx = suggested_promises[0][0]
            
        # Find the selected promise
        selected_promise = None
        for i, promise in enumerate(promise_manager.promises):
            if i == selected_idx:
                selected_promise = promise
                break
                
        if not selected_promise:
            print(f"Error: Could not find promise with index {selected_idx}. Ending scene generation.")
            break
            
        print(f"Selected promise to progress: {selected_promise['description']}")
        
        # Generate the scene content
        scene_content = scene_generator.write_scene(
            title, genre, author, profile, recent_events,
            selected_promise['description'], selected_promise['progress'] + 1
        )
        
        # Create a scene title based on the promise
        scene_title = f"Scene {scene_num}: Developing '{selected_promise['description']}'"
        
        # Store the scene
        book_content.append({
            "title": scene_title,
            "content": scene_content
        })
        
        # Summarize what happened in this scene
        event_summary = scene_generator.summarize_event(
            scene_content, selected_promise['description']
        )
        
        # Update the promise manager
        promise_manager.mark_promise_progressed(selected_idx, event_summary)
        promise_manager.increment_scene()
        
        print(f"Scene {scene_num} complete. Progress made on promise: '{selected_promise['description']}'")
        
    print(f"--- Book Generation Complete: {len(book_content)} scenes generated ---")
    return book_content, promise_manager.get_story_summary()

# --- Main Execution ---

def main():
    # --- Configuration ---
    # Ensure the output folder exists
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    print(f"Using output folder: {OUTPUT_FOLDER}")

    # --- Inputs ---
    resume_filename = 'Thalen.pdf'  # MAKE SURE THIS FILE EXISTS in ./docs
    subject = "Story of a young architect living in a post ufo dominated world. (USE REAL PLACES FROM THE RESUME AS LOCATIONS IN THE STORY AND MAKE IT A Slice OF LIFE STORY)" 
    author_style = 'Ana Huang'  # Example author style
    genre = 'SiFi, Fiction, Life, Work, Struggles'  # Example genre

    print("\n--- Starting Novel Generation with Plot Promise System ---")
    print(f"Resume File: {resume_filename}")
    print(f"Subject: {subject}")
    print(f"Author Style: {author_style}")
    print(f"Genre: {genre}")

    # --- Instantiate Chains & Writer ---
    try:
        main_character_chain = MainCharacterChain()
        title_chain = TitleChain()
        promise_generator = PromiseGeneratorChain()
        doc_writer = DocWriter(output_folder=OUTPUT_FOLDER)
        print("Chains initialized successfully.")
    except Exception as e:
        print(f"Fatal Error: Failed to initialize LangChain components: {e}")
        traceback.print_exc()
        return  # Cannot proceed if chains fail

    # --- Generate Novel Components Sequentially ---
    try:
        print("\nStep 1: Generating Main Character Profile...")
        profile = main_character_chain.run(resume_filename)
        if profile.startswith("Error:"):
            print(f"Profile Generation Failed: {profile}")
            print("Cannot proceed without a character profile.")
            return
        print(f"Profile Generated:\n---\n{profile}\n---")

        print("\nStep 2: Generating Title...")
        title = title_chain.run(subject, genre, author_style, profile)
        if title == "Error Generating Title":
            print("Title Generation Failed. Using placeholder.")
            title = "Placeholder Title - Generation Failed"
        print(f"Title Generated: '{title}'")

        print("\nStep 3: Generating Initial Plot Promises...")
        initial_promises = promise_generator.run(subject, genre, author_style, profile, title)
        if not initial_promises:
            print("Failed to generate initial plot promises. Cannot proceed.")
            return
        print(f"Generated {len(initial_promises)} Initial Plot Promises:")
        for i, promise in enumerate(initial_promises):
            print(f"  {i+1}. {promise['description']} (Importance: {promise['importance']})")

        # --- Write the Book Content using Plot Promise System ---
        print("\nStep 4: Writing Book Content with Plot Promise System...")
        book_content, story_summary = write_book_with_promises(
            genre,
            author_style,
            title,
            profile,
            initial_promises,
            total_scenes=15  # Adjust based on desired book length
        )

        # --- Save to Document ---
        print("\nStep 5: Saving Book to Document...")
        doc_writer.write_promise_book(book_content, title, profile, genre, story_summary)

    except FileNotFoundError as e:
        print(f"\nCRITICAL ERROR: {e}")
        print(f"Please ensure the required file exists in the '{OUTPUT_FOLDER}' directory.")
    except Exception as e:
        print(f"\nAn unexpected error occurred during the main generation process: {e}")
        traceback.print_exc()  # Print detailed traceback for unexpected errors

    print("\n--- Novel Generation Process Finished ---")


if __name__ == "__main__":
    main()

