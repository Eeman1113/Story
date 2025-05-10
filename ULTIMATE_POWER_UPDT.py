import requests
import json
import time
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pypdf # Added for PDF processing

# --- Configuration ---
# Replace with your Ollama API endpoint and desired model
OLLAMA_BASE_URL = "http://localhost:11434/api/generate"
# Consider models like "llama3:latest", "gemma2:latest", "mistral:latest" or more specialized ones if available.
# Ensure the model is pulled in Ollama (e.g., `ollama pull llama3`)
OLLAMA_MODEL = "gemma3:12b" # As per the user's last log, or they can change it
# Longer timeout for potentially complex generation tasks
OLLAMA_TIMEOUT = 360 # 6 minutes, adjust as needed

# Output directory for the generated novel
OUTPUT_DIR = "generated_novel_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

class NovelGenerator:
    def __init__(self, resume_content, subject, author_style, genre):
        self.resume_content = resume_content
        self.subject = subject
        self.author_style = author_style
        self.genre = genre
        self.num_chapters = 0 # Will be determined by the AI

        # Core story elements - will be populated by generation methods
        self.characters = {}  # Detailed character objects/dictionaries
        self.world_details = {"name": "", "key_locations": [], "cultural_elements": [], "rules": [], "atmosphere": ""}
        self.themes_motifs = {"themes": [], "motifs": []}
        self.plot_outline = ""  # High-level plot (e.g., 3-act structure)
        self.novel_title = "Untitled Novel"

        # Detailed chapter-by-chapter plan
        self.chapter_plans = {} # Key: chapter_num, Value: dict with plan details

        # Generated content and continuity data
        self.generated_chapters_content = {} # Key: chapter_num, Value: full chapter text
        self.chapter_continuity_data = {} # Key: chapter_num, Value: dict with summary, char updates, timeline, emotional arc, flow_analysis

        print("NovelGenerator initialized.")
        print(f"  Subject: {self.subject[:100]}...")
        print(f"  Author Style: {self.author_style}")
        print(f"  Genre: {self.genre}")
        print(f"  Resume provided: {'Yes' if self.resume_content else 'No'}")
        print(f"  Number of chapters will be determined automatically.")


    def _ollama_generate(self, prompt, system_prompt="You are a helpful AI assistant.", temperature=0.7, top_p=0.9):
        """
        Helper function to make API calls to the Ollama server.
        """
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "system": system_prompt,
            "stream": False,
            "options": {
                "temperature": temperature,
                "top_p": top_p,
                # "num_ctx": 8192 # Example: Adjust context window if needed and supported by model like Llama3
            }
        }
        # print(f"\n--- Sending Prompt to LLM ({OLLAMA_MODEL}) ---\n{prompt[:300]}...\n---") # Debug: Show prompt start
        try:
            response = requests.post(OLLAMA_BASE_URL, json=payload, timeout=OLLAMA_TIMEOUT)
            response.raise_for_status()
            # print(f"--- LLM Response Received ---\n{response.json()['response'][:300]}...\n---") # Debug: Show response start
            return response.json()["response"].strip()
        except requests.exceptions.Timeout:
            print(f"ERROR: Ollama request timed out after {OLLAMA_TIMEOUT} seconds for prompt: {prompt[:100]}...")
            return f"[OLLAMA TIMEOUT ERROR for prompt: {prompt[:100]}...]"
        except requests.exceptions.RequestException as e:
            print(f"ERROR: Ollama request failed: {e} for prompt: {prompt[:100]}...")
            return f"[OLLAMA REQUEST ERROR: {e} for prompt: {prompt[:100]}...]"
        except json.JSONDecodeError as e:
            print(f"ERROR: Failed to decode JSON response from Ollama: {e}")
            print(f"Raw response text: {response.text}") # It's response.text, not response.text()
            return f"[OLLAMA JSON DECODE ERROR for prompt: {prompt[:100]}...]"

    def _parse_character_profiles(self, text_block):
        """
        Parses character profiles from a structured text block.
        Expects format like:
        CHARACTER NAME: [Name]
        ROLE: [Role]
        DESCRIPTION: [Description]
        MOTIVATION: [Motivation]
        INITIAL_ARC_SUMMARY: [Arc Summary]
        FLAWS: [Flaws]
        STRENGTHS: [Strengths]
        """
        characters = {}
        current_char_data = {}
        current_char_name = None

        for line in text_block.split('\n'):
            line = line.strip()
            if not line: # Handles blank lines between character blocks
                if current_char_name and current_char_data:
                    characters[current_char_name] = {
                        "name": current_char_name,
                        "role": current_char_data.get("ROLE", "N/A"),
                        "description": current_char_data.get("DESCRIPTION", "N/A"),
                        "motivation": current_char_data.get("MOTIVATION", "N/A").replace("MOTIVATION(S):","").strip(), # Clean if key repeated
                        "arc_summary": current_char_data.get("INITIAL_ARC_SUMMARY", "N/A"),
                        "flaws": current_char_data.get("FLAWS", "N/A").replace("FLAWS/WEAKNESSES:","").strip(),
                        "strengths": current_char_data.get("STRENGTHS", "N/A").replace("STRENGTHS/SKILLS:","").strip(),
                        "current_status": "alive",
                        "current_location": "unknown",
                        "emotional_state": "neutral",
                        "knowledge": [],
                        "relationships": {},
                        "first_appearance_chapter": 0,
                        "development_log": []
                    }
                    current_char_data = {}
                    current_char_name = None
                continue

            match = re.match(r"CHARACTER NAME:\s*(.*)", line, re.IGNORECASE)
            if match:
                if current_char_name and current_char_data: # Save previous character before starting new one
                       characters[current_char_name] = {
                        "name": current_char_name,
                        "role": current_char_data.get("ROLE", "N/A"),
                        "description": current_char_data.get("DESCRIPTION", "N/A"),
                        "motivation": current_char_data.get("MOTIVATION", "N/A").replace("MOTIVATION(S):","").strip(),
                        "arc_summary": current_char_data.get("INITIAL_ARC_SUMMARY", "N/A"),
                        "flaws": current_char_data.get("FLAWS", "N/A").replace("FLAWS/WEAKNESSES:","").strip(),
                        "strengths": current_char_data.get("STRENGTHS", "N/A").replace("STRENGTHS/SKILLS:","").strip(),
                        "current_status": "alive", "current_location": "unknown", "emotional_state": "neutral",
                        "knowledge": [], "relationships": {}, "first_appearance_chapter": 0, "development_log": []
                    }
                current_char_name = match.group(1).strip()
                current_char_data = {} # Reset for the new character
                continue

            if current_char_name: # If we are currently parsing a character
                # Adjust keys to match potential variations in LLM output if necessary
                # For example, 'MOTIVATION(S)' vs 'MOTIVATION'
                keys_to_check = {
                    "ROLE": "ROLE",
                    "DESCRIPTION": "DESCRIPTION",
                    "MOTIVATION(S)": "MOTIVATION", # Store as "MOTIVATION"
                    "MOTIVATION": "MOTIVATION",
                    "INITIAL_ARC_SUMMARY": "INITIAL_ARC_SUMMARY",
                    "FLAWS/WEAKNESSES": "FLAWS", # Store as "FLAWS"
                    "FLAWS": "FLAWS",
                    "STRENGTHS/SKILLS": "STRENGTHS", # Store as "STRENGTHS"
                    "STRENGTHS": "STRENGTHS"
                }
                for key_llm, key_internal in keys_to_check.items():
                    if line.upper().startswith(key_llm + ":"):
                        current_char_data[key_internal] = line[len(key_llm)+1:].strip()
                        break
                    # Handle cases where LLM might not repeat the key but continues on next line (multiline description)
                    elif current_char_data and not any(line.upper().startswith(k + ":") for k in keys_to_check):
                        # This logic is tricky; for simplicity, we assume single-line values for now after the key
                        # Or append to the last known key if it's a multiline field like DESCRIPTION
                        if "DESCRIPTION" in current_char_data and not any(line.upper().startswith(k + ":") for k in keys_to_check if k != "DESCRIPTION"):
                             current_char_data["DESCRIPTION"] += "\n" + line
                             break


        if current_char_name and current_char_data: # Save the last character in the block
            characters[current_char_name] = {
                "name": current_char_name,
                "role": current_char_data.get("ROLE", "N/A"),
                "description": current_char_data.get("DESCRIPTION", "N/A"),
                "motivation": current_char_data.get("MOTIVATION", "N/A").replace("MOTIVATION(S):","").strip(),
                "arc_summary": current_char_data.get("INITIAL_ARC_SUMMARY", "N/A"),
                "flaws": current_char_data.get("FLAWS", "N/A").replace("FLAWS/WEAKNESSES:","").strip(),
                "strengths": current_char_data.get("STRENGTHS", "N/A").replace("STRENGTHS/SKILLS:","").strip(),
                "current_status": "alive", "current_location": "unknown", "emotional_state": "neutral",
                "knowledge": [], "relationships": {}, "first_appearance_chapter": 0, "development_log": []
            }
        return characters

    # --- Phase 1: Foundation Methods ---
    def generate_foundational_elements(self):
        """
        Generates initial character profiles, world details, themes/motifs, high-level plot outline,
        and determines the number of chapters.
        """
        print("\n--- Generating Foundational Elements ---")

        # 1. Character Conception
        print("Step 1.1: Generating Character Profiles...")
        char_system_prompt = f"You are a master character creator for {self.genre} novels, inspired by {self.author_style}."
        char_prompt = f"""
Based on the novel's subject, genre, and the provided resume snippet (if any), create detailed profiles for the main protagonist and 1-2 key supporting characters (e.g., antagonist, mentor, love interest).
Novel Subject: {self.subject}
Genre: {self.genre}
Author Style Influence: {self.author_style}
Resume Snippet (use to inspire the protagonist, inferring personality, potential skills, and background):
---
{self.resume_content if self.resume_content else "No resume provided. Create protagonist based on subject and genre."}
---
IMPORTANT: For EACH character, provide EXACTLY the following fields, starting each character block with 'CHARACTER NAME:'.
Your entire response must consist of one or more character profiles in this strict format. Do NOT add any introductory sentences, summaries outside of these fields, or any other details like worldbuilding or plot.

Format for each character:
CHARACTER NAME: [Suggest a fitting name]
ROLE: [Protagonist, Antagonist, Key Supporting - specify type]
DESCRIPTION: [Detailed appearance, key personality traits, mannerisms, background hints]
MOTIVATION(S): [What drives them? What are their primary goals, conscious or subconscious?]
INITIAL_ARC_SUMMARY: [How might they change or develop throughout the story? What is their potential journey?]
FLAWS/WEAKNESSES: [What are their vulnerabilities, biases, or negative traits?]
STRENGTHS/SKILLS: [What are their notable positive attributes or skills?]

Example of ONE character profile:
CHARACTER NAME: Jane Doe
ROLE: Protagonist
DESCRIPTION: Tall, with fiery red hair and a determined gaze. Often bites her lip when thinking. Former city guard.
MOTIVATION(S): To find her missing brother and expose the corruption in the council.
INITIAL_ARC_SUMMARY: Starts cynical and isolated, learns to trust others and becomes a leader.
FLAWS/WEAKNESSES: Impulsive, mistrustful of authority.
STRENGTHS/SKILLS: Skilled sword fighter, keen observer.

(Provide profiles for the main protagonist and 1-2 key supporting characters adhering strictly to this format.)
"""
        character_profiles_text = self._ollama_generate(char_prompt, char_system_prompt, temperature=0.75)
        if "[OLLAMA" in character_profiles_text:
            print(f"ERROR generating character profiles: {character_profiles_text}")
        else:
            self.characters = self._parse_character_profiles(character_profiles_text)
            if not self.characters:
                print("Warning: LLM output for character profiles was not in the expected format or was empty. Raw output (first 1000 chars):")
                print(character_profiles_text[:1000] + "..." if len(character_profiles_text) > 1000 else character_profiles_text)
            print(f"Generated {len(self.characters)} character profiles: {', '.join(self.characters.keys())}")

        # 2. Worldbuilding
        print("\nStep 1.2: Generating World Details...")
        world_system_prompt = f"You are a world-building expert for {self.genre} fiction, creating immersive settings like those by {self.author_style}."
        world_prompt = f"""
Based on the novel's subject and genre, describe the primary world/setting.
Novel Subject: {self.subject}
Genre: {self.genre}
IMPORTANT: Provide ONLY the following details for the world/setting. Use the exact numbered headers as specified below.
Do not include plot summaries, character descriptions, or any information not requested in this step.

1.  WORLD NAME: [A unique and evocative name for the main setting/world/city]
2.  KEY LOCATIONS: [List 3-5 significant recurring locations with brief descriptions. Each item should be on a new line, like: "- The Sunken Library - an ancient repository of forbidden knowledge"]
3.  CULTURAL ELEMENTS: [Describe 2-3 unique customs, societal norms, beliefs, or technologies. Each item should be on a new line, like: "- Aether-tech - devices powered by ambient magical energy"]
4.  ATMOSPHERE/TONE: [Describe the overall mood and feeling of the world (e.g., oppressive, wondrous, decaying, futuristic, magical)]
5.  KEY RULES/LAWS (if applicable, e.g., for magic systems, societal structure): [List 1-3 fundamental rules that govern this world or its unique aspects. Each item should be on a new line.]
"""
        world_details_text = self._ollama_generate(world_prompt, world_system_prompt, temperature=0.65)
        if "[OLLAMA" in world_details_text:
            print(f"ERROR generating world details: {world_details_text}")
        else:
            lines = world_details_text.split('\n')
            for line in lines:
                if re.match(r"1\.\s*WORLD NAME:", line, re.IGNORECASE) or line.upper().startswith("WORLD NAME:"):
                    name_val = re.split(r":", line, 1)[1].strip()
                    name_val = re.sub(r'[*_`#]', '', name_val) # Remove common markdown chars
                    self.world_details["name"] = name_val
                elif re.match(r"4\.\s*ATMOSPHERE/TONE:", line, re.IGNORECASE) or line.upper().startswith("ATMOSPHERE/TONE:"):
                     self.world_details["atmosphere"] = re.split(r":", line, 1)[1].strip()


            key_loc_match = re.search(r"(?:2\.|KEY LOCATIONS:)\s*((?:-\s*.*|\d\.\s*.*|\*\s*.*(?:\n|$))+)", world_details_text, re.IGNORECASE | re.DOTALL)
            if key_loc_match:
                self.world_details["key_locations"] = [loc.strip("-* ").strip().lstrip("123456789. ") for loc in key_loc_match.group(1).strip().split('\n') if loc.strip()]

            cultural_elem_match = re.search(r"(?:3\.|CULTURAL ELEMENTS:)\s*((?:-\s*.*|\d\.\s*.*|\*\s*.*(?:\n|$))+)", world_details_text, re.IGNORECASE | re.DOTALL)
            if cultural_elem_match:
                self.world_details["cultural_elements"] = [elem.strip("-* ").strip().lstrip("123456789. ") for elem in cultural_elem_match.group(1).strip().split('\n') if elem.strip()]

            rules_match = re.search(r"(?:5\.|KEY RULES/LAWS:)\s*((?:-\s*.*|\d\.\s*.*|\*\s*.*(?:\n|$))+)", world_details_text, re.IGNORECASE | re.DOTALL)
            if rules_match:
                self.world_details["rules"] = [rule.strip("-* ").strip().lstrip("123456789. ") for rule in rules_match.group(1).strip().split('\n') if rule.strip()]

            print(f"Generated World Details for '{self.world_details.get('name', 'Unnamed World')}'.")
            if not self.world_details.get("name"):
                print("Warning: World Name not parsed. Raw world details output (first 1000 chars):")
                print(world_details_text[:1000] + "..." if len(world_details_text) > 1000 else world_details_text)


        # 3. Themes and Motifs
        print("\nStep 1.3: Generating Themes and Motifs...")
        themes_system_prompt = f"You are a literary analyst identifying profound themes and recurring motifs for {self.genre} novels, in the vein of {self.author_style}."
        themes_prompt = f"""
Based on the novel's subject, genre, and initial character concepts, identify:
Novel Subject: {self.subject}
Genre: {self.genre}
Character Concepts: {json.dumps(self.characters, indent=2)}
World Atmosphere: {self.world_details.get("atmosphere", "N/A")}

IMPORTANT: Your response should ONLY contain lists under the exact headers '1. CORE THEMES:' and '2. RECURRING MOTIFS:'.
Use the specified format. Do not add any other explanations, introductory text, or unrelated information.

1.  CORE THEMES (2-4): [List abstract concepts the story will explore. Each item should be on a new line, like: "- Loss and Memory: The story explores how memories define individuals and societies, and the consequences of their loss or manipulation." Provide a brief (1-sentence) explanation for each, linking it to the context.]
2.  RECURRING MOTIFS (3-5): [List concrete symbols, objects, phrases, or imagery. Each item should be on a new line, like: "- A cracked pocket watch," "- The phrase 'shadows remember'."]
"""
        themes_motifs_text = self._ollama_generate(themes_prompt, themes_system_prompt, temperature=0.6)
        if "[OLLAMA" in themes_motifs_text:
            print(f"ERROR generating themes/motifs: {themes_motifs_text}")
        else:
            themes_match = re.search(r"(?:1\.|CORE THEMES:)\s*((?:-\s*.*|\d\.\s*.*|\*\s*.*(?:\n|$))+)", themes_motifs_text, re.IGNORECASE | re.DOTALL)
            if themes_match:
                self.themes_motifs["themes"] = [theme.strip("-* ").strip().lstrip("123456789. ") for theme in themes_match.group(1).strip().split('\n') if theme.strip()]

            motifs_match = re.search(r"(?:2\.|RECURRING MOTIFS:)\s*((?:-\s*.*|\d\.\s*.*|\*\s*.*(?:\n|$))+)", themes_motifs_text, re.IGNORECASE | re.DOTALL)
            if motifs_match:
                self.themes_motifs["motifs"] = [motif.strip("-* ").strip().lstrip("123456789. ") for motif in motifs_match.group(1).strip().split('\n') if motif.strip()]

            if not self.themes_motifs["themes"] and not self.themes_motifs["motifs"]:
                print("Warning: Themes and Motifs not parsed. Raw output (first 1000 chars):")
                print(themes_motifs_text[:1000] + "..." if len(themes_motifs_text) > 1000 else themes_motifs_text)
            print(f"Generated Themes: {self.themes_motifs['themes']}")
            print(f"Generated Motifs: {self.themes_motifs['motifs']}")

        # 4. High-Level Plot Outline & Determine Number of Chapters
        print("\nStep 1.4: Generating High-Level Plot Outline and Determining Chapter Count...")
        plot_system_prompt = f"You are a master storyteller, outlining engaging plots for {self.genre} novels in the style of {self.author_style}."
        plot_prompt = f"""
Create a high-level plot outline for a novel based on the subject.
Novel Subject: {self.subject}
Genre: {self.genre}
Characters: {json.dumps(self.characters, indent=2)}
World: {json.dumps(self.world_details, indent=2)}
Themes: {json.dumps(self.themes_motifs["themes"])}

The outline should follow a classic narrative structure (e.g., Three-Act Structure: Setup, Confrontation, Resolution).
Describe:
-   **Act I (Setup):** Introduction of protagonist, inciting incident, establishment of conflict and stakes.
-   **Act II (Confrontation):** Rising action, character development through trials, introduction of key allies/antagonists, major turning points/complications.
-   **Act III (Resolution):** Climax, falling action, resolution of main conflict, and thematic conclusion.

For each Act, provide a 2-4 sentence summary of its key developments and objectives.
IMPORTANT: Your response should ONLY contain the summaries for Act I, Act II, and Act III, followed by the SUGGESTED_CHAPTER_COUNT line.
Do NOT include character profiles, world details, or themes here as they have been generated in separate, previous steps. Your response should only be the plot act summaries and the chapter count line.

Example Structure:
**Act I (Setup):** [Summary for Act I]
**Act II (Confrontation):** [Summary for Act II]
**Act III (Resolution):** [Summary for Act III]
SUGGESTED_CHAPTER_COUNT: [Number]

Begin your response with Act I.
"""
        self.plot_outline = self._ollama_generate(plot_prompt, plot_system_prompt, temperature=0.7)

        if "[OLLAMA" in self.plot_outline:
            print(f"ERROR generating plot outline: {self.plot_outline}")
            # Defaulting num_chapters handled further down
        else:
            print("Generated High-Level Plot Outline:")
            plot_display = re.sub(r"SUGGESTED_CHAPTER_COUNT:\s*\d+", "", self.plot_outline, flags=re.IGNORECASE).strip()
            print(plot_display)
            # num_chapters parsing is handled below

        # Determine num_chapters from plot_outline or default
        suggested_chapters_match = re.search(r"SUGGESTED_CHAPTER_COUNT:\s*(\d+)", self.plot_outline, re.IGNORECASE)
        if suggested_chapters_match:
            try:
                self.num_chapters = int(suggested_chapters_match.group(1))
                if not (5 <= self.num_chapters <= 75): # Wider reasonable range
                    print(f"Warning: LLM suggested {self.num_chapters} chapters, which is outside the typical 5-75 range. Clamping to 15.")
                    self.num_chapters = 15
                else:
                    print(f"LLM suggested {self.num_chapters} chapters for the novel.")
            except ValueError:
                print("Warning: Could not parse suggested chapter count as integer. Defaulting to 15 chapters.")
                self.num_chapters = 15
        elif "[OLLAMA" in self.plot_outline: # If API error during plot gen
             print("Defaulting number of chapters to 15 due to API error in plot generation.")
             self.num_chapters = 15
        else: # If no match and no API error
            print("Warning: Could not find suggested chapter count in plot outline. Defaulting to 15 chapters.")
            self.num_chapters = 15

        if self.num_chapters < 3: # Ensure minimum for structure
            print(f"Warning: Number of chapters determined ({self.num_chapters}) is less than 3. Setting to 3.")
            self.num_chapters = 3


        # --- Validation of Foundational Elements ---
        all_steps_api_ok = not ("[OLLAMA" in character_profiles_text or \
                                "[OLLAMA" in world_details_text or \
                                "[OLLAMA" in themes_motifs_text or \
                                "[OLLAMA" in self.plot_outline)

        if not all_steps_api_ok:
            print("Halting: Foundational element generation failed due to API errors with Ollama.")
            return False

        essential_data_parsed = True
        missing_elements = []
        if not self.characters:
            essential_data_parsed = False
            missing_elements.append("Parsed character profiles")
        if not self.world_details.get("name"):
            essential_data_parsed = False
            missing_elements.append("Parsed world name/details")
        if not self.themes_motifs.get("themes") and not self.themes_motifs.get("motifs"):
            essential_data_parsed = False
            missing_elements.append("Parsed themes/motifs")
        if not self.plot_outline.strip() or self.num_chapters == 0 : # Check if plot_outline itself is empty
            essential_data_parsed = False
            missing_elements.append("Parsed plot outline or valid chapter count")
        # Also ensure plot_outline doesn't indicate an error itself, though already caught by all_steps_api_ok
        if "[OLLAMA" in self.plot_outline and "Parsed plot outline or valid chapter count" not in missing_elements:
            essential_data_parsed = False
            missing_elements.append("Plot outline contains API error")


        if not essential_data_parsed:
            print(f"Halting: Essential foundational data was not successfully parsed or generated. Missing or empty elements: {', '.join(missing_elements)}.")
            return False

        print("Foundational elements generated successfully.")
        return True


    # --- Phase 2: Detailed Planning Method ---
    def generate_detailed_chapter_plans(self):
        """
        Expands the high-level plot outline into detailed plans for each chapter.
        """
        print("\n--- Generating Detailed Chapter-by-Chapter Plans ---")
        # This check is now more robust due to changes in generate_foundational_elements
        if not self.plot_outline or not self.characters or not self.world_details.get("name") or \
           (not self.themes_motifs.get("themes") and not self.themes_motifs.get("motifs")) or \
           self.num_chapters == 0:
            print("ERROR: Cannot generate chapter plans without all foundational elements (plot, characters, world name, themes/motifs, num_chapters).")
            # For debugging, print what's missing or present
            if not self.plot_outline: print("  - Plot Outline: Missing")
            if not self.characters: print("  - Characters: Missing or Empty")
            if not self.world_details.get("name"): print("  - World Name: Missing")
            if not self.themes_motifs.get("themes") and not self.themes_motifs.get("motifs"): print("  - Themes/Motifs: Missing or Empty")
            if self.num_chapters == 0: print("  - Number of Chapters: Zero")
            return False

        system_prompt = f"You are a meticulous plot architect, detailing chapter structures for a {self.genre} novel in the style of {self.author_style}."
        character_summary_for_prompt = "\n".join([f"- {name} ({data.get('role', 'N/A')}): Motivations: {data.get('motivation', 'N/A')}. Arc: {data.get('arc_summary', 'N/A')}" for name, data in self.characters.items()])
        world_summary_for_prompt = f"World Name: {self.world_details.get('name', 'N/A')}\nKey Locations: {', '.join(self.world_details.get('key_locations',[]))}\nCultural Elements: {', '.join(self.world_details.get('cultural_elements',[]))}\nAtmosphere: {self.world_details.get('atmosphere', 'N/A')}"
        themes_for_prompt = f"Core Themes: {', '.join(self.themes_motifs.get('themes',[]))}\nRecurring Motifs: {', '.join(self.themes_motifs.get('motifs',[]))}"

        # Clean plot_outline from chapter count for this prompt
        clean_plot_outline_for_detailed_plan = re.sub(r"SUGGESTED_CHAPTER_COUNT:\s*\d+", "", self.plot_outline, flags=re.IGNORECASE).strip()


        detailed_plan_prompt = f"""
        Novel Subject: {self.subject}
        High-Level Plot Outline:
        {clean_plot_outline_for_detailed_plan}

        Character Summaries:
        {character_summary_for_prompt}

        World Summary:
        {world_summary_for_prompt}

        Themes & Motifs:
        {themes_for_prompt}

        Total Chapters to Plan: {self.num_chapters}

        Based on ALL the information above, create a VERY DETAILED plan for EACH chapter (1 through {self.num_chapters}).
        For EACH chapter, provide STRICTLY the following, in this order and clearly labeled:

        Chapter [Number] - [Evocative Title for this Chapter]:
        1.  CHAPTER GOAL: [A single paragraph (50-80 words) stating the primary narrative goal this chapter needs to achieve in the overall story and how it impacts the main character's arc or the central conflict.]
        2.  KEY SCENES (3-6 scenes): [Bulleted list. Each scene: "- Scene X: [Brief description of action/dialogue/internal monologue], Location: [Specific location from World Details or new minor one], Characters Involved: [List characters present & active], Key Revelation/Turning Point/Outcome: [What changes, is learned, or achieved? How does it advance the plot or character?]]
        3.  CHARACTER DEVELOPMENT FOCUS: [For key characters appearing: How do their motivations, relationships, knowledge, or understanding change in THIS chapter? Be specific. e.g., "Jessica: Confronts her fear of X, strengthening her resolve but creating friction with Y. Learns Z about her past."]
        4.  PLOT ADVANCEMENT: [Specific ways the main plot and any subplots move forward. What new questions are raised or old ones answered? How does this chapter build on the previous and set up the next?]
        5.  TIMELINE & PACING: [e.g., "This chapter takes place over a few hours the next day.", "Pacing: Fast, with building tension.", "Spans one week, slower reflective pace initially, then accelerates."]
        6.  EMOTIONAL TONE (End of Chapter): [e.g., "Hopeful but wary," "Tense and suspenseful," "Melancholy and reflective," "Ominous and foreboding."]
        7.  CONNECTION TO NEXT CHAPTER (Setup/Hook): [Explicitly state 1-2 elements, questions, cliffhangers, or character decisions that directly lead into the next chapter's planned events or themes.]

        This detailed plan must ensure logical progression, character consistency, integration of themes/motifs, and effective pacing.
        A character's status (location, knowledge, emotional state) at the end of one chapter MUST be the starting point for the next.
        Ensure the plans for later chapters logically follow from the resolutions and developments of earlier ones.
        The first chapter should introduce the protagonist and the inciting incident. The final chapters should bring the main plot threads to a climax and resolution.
        Do not include the "SUGGESTED_CHAPTER_COUNT" line in this response.
        Begin with "Chapter 1 - ...".
        """
        print(f"Generating detailed plan text for {self.num_chapters} chapters (this may take a while)...")
        full_detailed_plan_text = self._ollama_generate(detailed_plan_prompt, system_prompt, temperature=0.65)

        if "[OLLAMA" in full_detailed_plan_text:
            print(f"ERROR generating detailed chapter plans: {full_detailed_plan_text}")
            return False

        current_chapter_num = 0
        # Modified split pattern to handle markdown formatting (** for bold)
        chapter_blocks = re.split(r"(?=^(?:\*\*)?Chapter\s*\d+\s*[-:])", full_detailed_plan_text, flags=re.MULTILINE | re.IGNORECASE)

        for block in chapter_blocks:
            block = block.strip()
            if not block: continue

            # Modified to handle markdown formatting (** for bold) in chapter titles
            title_match = re.match(r"(?:\*\*)?Chapter\s*(\d+)\s*[-:]?\s*(.*?)(?:\*\*)?(?=\n\s*(?:1\.\s*CHAPTER GOAL:|CHAPTER GOAL:|#|$))", block, re.IGNORECASE)
            
            if title_match:
                current_chapter_num = int(title_match.group(1))
                chapter_title = title_match.group(2).strip()
                if not chapter_title: chapter_title = f"Chapter {current_chapter_num} (Untitled)"

                # Only add if chapter number is within expected range
                if current_chapter_num > self.num_chapters :
                    print(f"Warning: Parsed plan for Chapter {current_chapter_num}, which exceeds expected {self.num_chapters} chapters. Skipping.")
                    continue
                if current_chapter_num <=0:
                    print(f"Warning: Parsed plan for invalid Chapter {current_chapter_num}. Skipping.")
                    continue


                self.chapter_plans[current_chapter_num] = {"number": current_chapter_num, "title": chapter_title}
                plan_details = self.chapter_plans[current_chapter_num]

                # More robust parsing for each section, looking for the numbered headers
                goal_match = re.search(r"(?:1\.\s*CHAPTER GOAL:|CHAPTER GOAL:)\s*(.*?)(?=\n\s*(?:2\.\s*KEY SCENES:|KEY SCENES:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["goal"] = goal_match.group(1).strip() if goal_match else "N/A"

                scenes_match = re.search(r"(?:2\.\s*KEY SCENES:|KEY SCENES:)\s*(.*?)(?=\n\s*(?:3\.\s*CHARACTER DEVELOPMENT FOCUS:|CHARACTER DEVELOPMENT FOCUS:|#|$))", block, re.IGNORECASE | re.DOTALL)
                if scenes_match:
                    scenes_text = scenes_match.group(1).strip()
                    plan_details["scenes"] = [s.strip() for s in re.split(r'\n\s*(?:-\s*Scene|\*\s*Scene|\d\.\s*Scene)', scenes_text) if s.strip()]
                    plan_details["scenes"] = [re.sub(r'^(-\s*Scene\s*\d*:?\s*|-\s*|\*\s*Scene\s*\d*:?\s*|\*\s*|\d\.\s*Scene\s*\d*:?\s*|\d\.\s*)', '', s).strip() for s in plan_details["scenes"]]
                else: plan_details["scenes"] = []

                char_dev_match = re.search(r"(?:3\.\s*CHARACTER DEVELOPMENT FOCUS:|CHARACTER DEVELOPMENT FOCUS:)\s*(.*?)(?=\n\s*(?:4\.\s*PLOT ADVANCEMENT:|PLOT ADVANCEMENT:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["character_development"] = char_dev_match.group(1).strip() if char_dev_match else "N/A"

                plot_adv_match = re.search(r"(?:4\.\s*PLOT ADVANCEMENT:|PLOT ADVANCEMENT:)\s*(.*?)(?=\n\s*(?:5\.\s*TIMELINE & PACING:|TIMELINE & PACING:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["plot_advancement"] = plot_adv_match.group(1).strip() if plot_adv_match else "N/A"

                timeline_match = re.search(r"(?:5\.\s*TIMELINE & PACING:|TIMELINE & PACING:)\s*(.*?)(?=\n\s*(?:6\.\s*EMOTIONAL TONE \(End of Chapter\):|EMOTIONAL TONE \(End of Chapter\):|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["timeline_pacing"] = timeline_match.group(1).strip() if timeline_match else "N/A"

                emotion_match = re.search(r"(?:6\.\s*EMOTIONAL TONE \(End of Chapter\):|EMOTIONAL TONE \(End of Chapter\):)\s*(.*?)(?=\n\s*(?:7\.\s*CONNECTION TO NEXT CHAPTER:|CONNECTION TO NEXT CHAPTER:|#|$))", block, re.IGNORECASE | re.DOTALL)
                plan_details["emotional_tone_end"] = emotion_match.group(1).strip() if emotion_match else "N/A"

                connection_match = re.search(r"(?:7\.\s*CONNECTION TO NEXT CHAPTER:|CONNECTION TO NEXT CHAPTER:)\s*(.*?)(?=$|\n\s*#)", block, re.IGNORECASE | re.DOTALL) # Use # as a potential end marker too
                plan_details["connection_to_next"] = connection_match.group(1).strip() if connection_match else "N/A"
            else:
                 # This block might be introductory text from the LLM before "Chapter 1..."
                 if not self.chapter_plans and "Chapter 1" not in block and len(block) < 300: # Heuristic
                     print(f"Skipping potential preamble in detailed plan output: {block[:100]}...")


        if not self.chapter_plans or len(self.chapter_plans) == 0:
            print("ERROR: No chapter plans were successfully parsed. The LLM output might not conform to the expected structure.")
            print("--- LLM Output Start (first 2000 chars) ---")
            print(full_detailed_plan_text[:2000])
            print("--- LLM Output End ---")
            return False

        print(f"Successfully parsed detailed plans for {len(self.chapter_plans)} chapters out of {self.num_chapters} expected.")
        if len(self.chapter_plans) != self.num_chapters:
              print(f"WARNING: Mismatch in parsed plans ({len(self.chapter_plans)}) and expected chapters ({self.num_chapters}). Review LLM output for plan structure.")
              # This could be due to LLM not generating all requested chapters, or parsing issues.
              # For now, we proceed but this might cause issues if later chapters are expected by logic.
        return True

    # --- Phase 3: Prose Generation Loop ---
    def _get_continuity_context_for_chapter(self, chapter_num):
        """
        Gathers all relevant context from previous chapters and plans for generating the current chapter.
        """
        context = f"Overall Novel Subject: {self.subject}\n"
        context += f"Author Style: {self.author_style}, Genre: {self.genre}\n"
        # Clean plot_outline from chapter count
        cleaned_plot_outline = re.sub(r"SUGGESTED_CHAPTER_COUNT:\s*\d+", "", self.plot_outline, flags=re.IGNORECASE).strip()
        context += f"High-Level Plot Outline:\n{cleaned_plot_outline}\n\n"
        context += f"World Details: {json.dumps(self.world_details)}\n"
        context += f"Themes & Motifs: {json.dumps(self.themes_motifs)}\n\n"

        context += "Character Profiles & Current Status (as of start of this chapter):\n"
        for name, data in self.characters.items():
            context += f"- {name} ({data.get('role','N/A')}):\n"
            context += f"  Description: {data.get('description','N/A')}\n"
            context += f"  Motivations: {data.get('motivation','N/A')}. Initial Arc: {data.get('arc_summary','N/A')}\n"
            context += f"  Current Status: {data.get('current_status','unknown')}, Location: {data.get('current_location','unknown')}, Emotion: {data.get('emotional_state','unknown')}\n"
            context += f"  Known Facts: {', '.join(data.get('knowledge',[]))}\n"
            if data.get('development_log'):
                relevant_logs = [log for log in data['development_log'] if log['chapter'] < chapter_num]
                if relevant_logs:
                    last_dev = relevant_logs[-1]
                    context += f"  Last Noted Development (Ch {last_dev['chapter']}): {last_dev.get('summary', 'N/A')}\n"
        context += "\n"

        if chapter_num > 1:
            prev_chap_num = chapter_num - 1
            prev_continuity = self.chapter_continuity_data.get(prev_chap_num, {})
            prev_plan = self.chapter_plans.get(prev_chap_num, {})
            context += f"Summary of Previous Chapter ({prev_chap_num} - '{prev_plan.get('title', 'Untitled')}'):\n"
            context += f"{prev_continuity.get('summary', 'N/A')}\n"
            context += f"Ended with Emotional Tone: {prev_continuity.get('emotional_tone_end_achieved_in_summary', 'N/A')}\n"
            context += f"Timeline at end of Ch {prev_chap_num}: {prev_continuity.get('timeline_end', 'N/A')}\n"
            context += f"Hook for current chapter (from prev chapter's plan): {prev_plan.get('connection_to_next', 'N/A')}\n\n"

        return context

    def _generate_chapter_opener(self, chapter_num, current_chapter_plan):
        """Generates the opening paragraph(s) for the current chapter."""
        chapter_title_line = f"Chapter {chapter_num} - {current_chapter_plan.get('title', 'Untitled')}"

        if chapter_num == 1:
            return f"{chapter_title_line}\n\n" # For Ch1, prose gen starts right away

        prev_chap_num = chapter_num - 1
        prev_continuity = self.chapter_continuity_data.get(prev_chap_num, {})
        prev_plan = self.chapter_plans.get(prev_chap_num, {})

        system_prompt = f"You are a novelist in the style of {self.author_style}, skilled at crafting chapter openings that immediately re-orient the reader and smoothly transition from the previous chapter's ending."
        prompt = f"""
        You are writing the VERY FIRST paragraph(s) for Chapter {chapter_num}, titled "{current_chapter_plan.get('title', 'Untitled')}".
        This opening must seamlessly connect to the end of Chapter {prev_chap_num} and establish the immediate context for Chapter {chapter_num}.

        Context from End of Chapter {prev_chap_num} ('{prev_plan.get('title', 'Untitled')}'):
        - Summary of Ch {prev_chap_num}: {prev_continuity.get('summary', 'Previously...')[-1000:]}
        - Actual Ending Hook/Transition Text from Ch {prev_chap_num}: "{prev_continuity.get('ending_hook_text', 'The previous chapter ended.')}"
        - Emotional Tone at End of Ch {prev_chap_num}: {prev_continuity.get('emotional_tone_end_achieved_in_summary', 'Neutral')}
        - Timeline at End of Ch {prev_chap_num}: {prev_continuity.get('timeline_end', 'Unknown')}
        - Key Character States at End of Ch {prev_chap_num}: {json.dumps({name: {'status': data['current_status'], 'location': data['current_location'], 'emotion': data['emotional_state']} for name, data in self.characters.items()}, indent=2)}

        Plan for Current Chapter ({chapter_num} - "{current_chapter_plan.get('title', 'Untitled')}"):
        - Goal: {current_chapter_plan.get('goal', 'The story progresses.')}
        - First Planned Scene Hint: {current_chapter_plan.get('scenes', ['A new scene begins.'])[0] if current_chapter_plan.get('scenes') else 'A new scene begins.'}
        - Timeline & Pacing: {current_chapter_plan.get('timeline_pacing', 'As expected')}

        Write 1-2 compelling opening paragraphs (approx 100-200 words) for Chapter {chapter_num}. These paragraphs should:
        1. Directly acknowledge or subtly resolve the immediate hook/question left by Chapter {prev_chap_num}'s actual ending hook.
        2. Establish the setting, characters present, and the time elapsed since the previous chapter (if any significant passage, use the Timeline Indicators from this chapter's plan).
        3. Set the initial tone for Chapter {chapter_num}, which might be a continuation or a shift from the previous chapter's end.
        4. Orient the reader quickly without extensive exposition. Focus on "showing" the new situation.
        5. Make the transition feel natural and engaging, drawing the reader into the new chapter.
        6. This is the *opening text* of Chapter {chapter_num}. Do NOT repeat the chapter title.

        Opening paragraph(s) for Chapter {chapter_num}:
        """
        opener_text = self._ollama_generate(prompt, system_prompt, temperature=0.68)
        return f"{chapter_title_line}\n\n{opener_text}\n\n"


    def _generate_scene_prose(self, chapter_num, scene_index, scene_description, current_chapter_plan, continuity_context, previous_scene_prose=""):
        """Generates prose for a single scene within a chapter."""
        system_prompt = f"You are a celebrated novelist in the style of {self.author_style}, writing a {self.genre} novel. Your prose is vivid, emotionally resonant, and drives the plot forward. You excel at 'showing, not telling' and making fantastical elements relatable."

        motif_to_weave = "N/A"
        if self.themes_motifs.get("motifs") and len(self.themes_motifs["motifs"]) > 0 :
            motif_to_weave = self.themes_motifs["motifs"][(chapter_num + scene_index -1) % len(self.themes_motifs["motifs"])]

        prompt = f"""
        {continuity_context}

        Current Chapter Plan ({chapter_num} - "{current_chapter_plan.get('title', 'Untitled')}"):
        - Chapter Goal: {current_chapter_plan.get('goal', 'N/A')}
        - Full Scene Breakdown for this chapter: {json.dumps(current_chapter_plan.get('scenes',[]))}
        - Character Development Focus for this chapter: {current_chapter_plan.get('character_development', 'N/A')}
        - Plot Advancement for this chapter: {current_chapter_plan.get('plot_advancement', 'N/A')}
        - Planned Emotional Tone (End of Chapter): {current_chapter_plan.get('emotional_tone_end', 'N/A')}
        - Timeline & Pacing for this chapter: {current_chapter_plan.get('timeline_pacing', 'N/A')}

        Prose written SO FAR in THIS CHAPTER (before this scene):
        ---
        {previous_scene_prose[-2000:] if previous_scene_prose else "This is the first scene of the chapter (after the chapter title and opener)."}
        ---

        YOUR TASK: Write the narrative prose for THE FOLLOWING SPECIFIC SCENE ONLY.
        Scene {scene_index + 1} Description (from chapter plan): "{scene_description}"

        INSTRUCTIONS FOR THIS SCENE:
        1.  **Narrative Focus:** Write approximately 300-700 words (adjust based on scene importance) bringing THIS SCENE to life.
        2.  **Style & Tone:** Adhere to {self.author_style}'s style. Maintain the chapter's intended pacing and build towards its emotional goal.
        3.  **"Show, Don't Tell":** Demonstrate emotions, thoughts, and plot points through actions, dialogue, sensory details, and internal monologues. Make character reactions and decisions clear through their behavior.
        4.  **Relatability & Emotional Depth:** Even in fantastical situations, ground character experiences in relatable emotions. Describe what things *feel* like.
        5.  **Sensory Details:** Weave in vivid sensory details (sight, sound, smell, touch, taste) consistent with the world and scene location.
        6.  **Character Consistency:** Ensure characters act and speak consistently with their established profiles, motivations, current emotional state, and knowledge. Reflect their planned development for this chapter.
        7.  **Dialogue:** If dialogue is part of the scene, make it natural, character-specific, and purposeful (revealing character, advancing plot, or building tension).
        8.  **Motif Integration:** Subtly weave in the recurring motif: '{motif_to_weave}' if it fits naturally within this scene's events. Do not force it.
        9.  **Continuity:** Ensure this scene flows logically from any previous prose in this chapter. Do NOT repeat information.
        10. **Output:** Generate ONLY the newly written narrative paragraphs for THIS SCENE. Do not add scene numbers or headings.

        Begin Scene {scene_index + 1} prose now:
        """
        scene_prose = self._ollama_generate(prompt, system_prompt, temperature=0.72, top_p=0.92)
        return scene_prose

    def _analyze_inter_chapter_flow(self, previous_chapter_num, current_chapter_num, current_chapter_opening_text):
        """
        Analyzes the narrative flow between the end of the previous chapter and the start of the current chapter.
        """
        print(f"  Analyzing flow from Chapter {previous_chapter_num} to Chapter {current_chapter_num}...")
        if previous_chapter_num not in self.chapter_continuity_data or \
           previous_chapter_num not in self.chapter_plans or \
           current_chapter_num not in self.chapter_plans:
            print("    Skipping flow analysis: Missing data for previous or current chapter.")
            return "Flow analysis skipped due to missing data."

        prev_continuity = self.chapter_continuity_data[previous_chapter_num]
        prev_plan = self.chapter_plans[previous_chapter_num]
        current_plan = self.chapter_plans[current_chapter_num]

        prev_summary = prev_continuity.get("summary", "N/A")
        prev_hook = prev_continuity.get("ending_hook_text", "N/A")
        prev_title = prev_plan.get("title", "Untitled Previous Chapter")

        current_title = current_plan.get("title", "Untitled Current Chapter")
        current_goal = current_plan.get("goal", "N/A")

        opening_paragraphs_only = "\n".join(current_chapter_opening_text.split('\n\n')[1:]).strip()

        system_prompt = "You are an expert literary editor specializing in narrative coherence and flow between chapters."
        prompt = f"""
        Analyze the transition and flow from the end of Chapter {previous_chapter_num} ("{prev_title}") to the beginning of Chapter {current_chapter_num} ("{current_title}").

        CONTEXT:
        End of Chapter {previous_chapter_num} ("{prev_title}"):
        - Summary: {prev_summary[-1000:]}
        - Actual Ending Hook/Transition Text: "{prev_hook}"
        - Key Character States at End (from continuity): {json.dumps({name: {'status': data['current_status'], 'location': data['current_location'], 'emotion': data['emotional_state']} for name, data in self.characters.items() if data.get('first_appearance_chapter', 0) <= previous_chapter_num and data.get('first_appearance_chapter', 0) > 0}, indent=2)}
        - Planned Connection from Ch {previous_chapter_num} to Ch {current_chapter_num}: "{prev_plan.get('connection_to_next', 'N/A')}"


        Beginning of Chapter {current_chapter_num} ("{current_title}"):
        - Chapter Goal: {current_goal}
        - Actual Opening Paragraph(s): "{opening_paragraphs_only}"

        EVALUATION TASK:
        Provide a brief (1-2 paragraph) analysis covering these points:
        1.  **Hook Resolution:** How well does the opening of Chapter {current_chapter_num} address or follow up on the ending hook/transition of Chapter {previous_chapter_num}?
        2.  **Plot Continuity:** Is there a logical and clear progression of plot events or situation from the end of the previous chapter to the start of the current one?
        3.  **Character Consistency:** Do the characters' states (emotional, physical, location, knowledge) seem consistent and logically follow from the previous chapter's end into the current chapter's opening?
        4.  **Tone & Pacing:** Is the transition in emotional tone and pacing smooth and appropriate, or jarring?
        5.  **Overall Coherence:** How effective is the overall flow between these two chapters? Does it feel natural and engaging for the reader?

        Be concise and constructive.
        Flow Analysis:
        """
        flow_analysis_text = self._ollama_generate(prompt, system_prompt, temperature=0.5)
        print(f"    Flow Analysis Result: {flow_analysis_text[:200]}...")

        if current_chapter_num not in self.chapter_continuity_data:
            self.chapter_continuity_data[current_chapter_num] = {}
        self.chapter_continuity_data[current_chapter_num]["flow_analysis_from_previous"] = flow_analysis_text

        return flow_analysis_text


    def _update_chapter_continuity_data(self, chapter_num, full_chapter_content, is_final_pass_for_chapter=False):
        """
        Analyzes generated chapter content to update continuity data (summary, character states, timeline, emotion).
        """
        print(f"Updating continuity data for Chapter {chapter_num} ({'final pass' if is_final_pass_for_chapter else 'interim pass'})...")
        entry = self.chapter_continuity_data.get(chapter_num, {})

        summary_system_prompt = "You are a literary analyst. Your task is to summarize chapter content accurately and concisely for continuity purposes."
        summary_prompt = f"""
        Create a concise yet detailed summary of the following chapter content (Chapter {chapter_num}).
        Focus on all key plot events, character actions and significant development, setting details, revelations, emotional shifts, and how it ends.
        This summary will be crucial context for writing the NEXT chapter.
        CHAPTER CONTENT (Chapter {chapter_num} - "{self.chapter_plans.get(chapter_num, {}).get('title', 'Untitled')}"):
        ---
        {full_chapter_content}
        ---
        Detailed Summary of Chapter {chapter_num}:
        """
        entry["summary"] = self._ollama_generate(summary_prompt, summary_system_prompt, temperature=0.5)

        if is_final_pass_for_chapter:
            active_chars_in_chapter = []
            current_chapter_plan = self.chapter_plans.get(chapter_num, {})
            if current_chapter_plan:
                for scene_desc in current_chapter_plan.get("scenes", []):
                    char_involved_match = re.search(r"Characters Involved:\s*(.*?)(?:\.\s*Key Revelation|$)", scene_desc, re.IGNORECASE)
                    if char_involved_match:
                        names_str = char_involved_match.group(1)
                        potential_names = re.split(r'[,\s]+and\s+|\s*,\s*|[,\s]+with\s+', names_str)
                        for char_name_candidate in potential_names:
                            clean_name = char_name_candidate.strip().rstrip('.').strip()
                            if clean_name and clean_name in self.characters and clean_name not in active_chars_in_chapter:
                                active_chars_in_chapter.append(clean_name)
                char_dev_focus = current_chapter_plan.get("character_development", "")
                for char_name in self.characters.keys():
                    if re.search(r'\b' + re.escape(char_name) + r'\b', char_dev_focus, re.IGNORECASE) and char_name not in active_chars_in_chapter:
                            active_chars_in_chapter.append(char_name)

            if not active_chars_in_chapter: active_chars_in_chapter = list(self.characters.keys())

            char_update_system_prompt = "You are a narrative continuity expert. Update character states based on chapter events."
            char_update_prompt = f"""
            Based on the FULL content of Chapter {chapter_num} below, update the status for EACH listed character.
            Chapter {chapter_num} ("{self.chapter_plans.get(chapter_num, {}).get('title', 'Untitled')}") Content:
            ---
            {full_chapter_content}
            ---
            For EACH of these characters who appeared or were central: {', '.join(active_chars_in_chapter) if active_chars_in_chapter else "Summarize general impact if no specific characters listed."}
            Provide the following updates. If a character did not appear or had no significant change for a field, state "No change" or "Did not appear."

            CHARACTER NAME: [Character's Name]
            -   STATUS CHANGE: [e.g., "Remains alive," "Injured (twisted ankle)," "Captured," "Learned X," "Decided Y."]
            -   LOCATION AT END OF CHAPTER: [Specific location.]
            -   EMOTIONAL STATE AT END OF CHAPTER: [e.g., "Hopeful," "Grieving," "Determined," "Suspicious."]
            -   KEY DEVELOPMENT/ACTION: [Significant actions, learnings, or internal changes.]
            -   RELATIONSHIP CHANGES: [e.g., "Strained with Z over X," "New alliance with W."]
            -   NEW KNOWLEDGE/SECRETS ACQUIRED: [New critical info, clues, secrets.]

            Format clearly for each character.
            """
            character_updates_text = self._ollama_generate(char_update_prompt, char_update_system_prompt, temperature=0.55)
            entry["character_updates_text"] = character_updates_text

            current_char_name_update = None
            parsed_updates_for_log = {}
            for line in character_updates_text.split('\n'):
                line = line.strip()
                name_match = re.match(r"CHARACTER NAME:\s*(.*)", line, re.IGNORECASE)
                if name_match:
                    if current_char_name_update and parsed_updates_for_log and current_char_name_update in self.characters:
                        self.characters[current_char_name_update]["development_log"].append(
                            {"chapter": chapter_num, "summary": "Updates from chapter events", **parsed_updates_for_log}
                        )
                    current_char_name_update = name_match.group(1).strip()
                    parsed_updates_for_log = {}
                    if current_char_name_update not in self.characters:
                        current_char_name_update = None
                    elif self.characters[current_char_name_update].get("first_appearance_chapter", 0) == 0:
                            self.characters[current_char_name_update]["first_appearance_chapter"] = chapter_num
                    continue

                if current_char_name_update and current_char_name_update in self.characters:
                    char_obj = self.characters[current_char_name_update]
                    def get_value(text, key_phrase):
                        # Match key phrase at start of line, possibly after "- "
                        if re.match(r"-\s*" + re.escape(key_phrase.upper()), text.upper()) or text.upper().startswith(key_phrase.upper()):
                            val_part = re.split(":",text,1)
                            val = val_part[1].strip() if len(val_part) > 1 else ""
                            if val.lower() in ["no change", "did not appear", "n/a", "none", "no significant change."]:
                                return None
                            return val
                        return "NO_MATCH"

                    status_val = get_value(line, "STATUS CHANGE")
                    if status_val != "NO_MATCH":
                        if status_val is not None: char_obj["current_status"] = status_val
                        parsed_updates_for_log["status"] = status_val if status_val is not None else "No change"
                        continue
                    loc_val = get_value(line, "LOCATION AT END OF CHAPTER")
                    if loc_val != "NO_MATCH":
                        if loc_val is not None: char_obj["current_location"] = loc_val
                        parsed_updates_for_log["location"] = loc_val if loc_val is not None else "No change"
                        continue
                    emo_val = get_value(line, "EMOTIONAL STATE AT END OF CHAPTER")
                    if emo_val != "NO_MATCH":
                        if emo_val is not None: char_obj["emotional_state"] = emo_val
                        parsed_updates_for_log["emotion"] = emo_val if emo_val is not None else "No change"
                        continue
                    dev_val = get_value(line, "KEY DEVELOPMENT/ACTION")
                    if dev_val != "NO_MATCH":
                        parsed_updates_for_log["development"] = dev_val if dev_val is not None else "No change"
                        continue
                    rel_val = get_value(line, "RELATIONSHIP CHANGES")
                    if rel_val != "NO_MATCH":
                        parsed_updates_for_log["relationships_changed"] = rel_val if rel_val is not None else "No change"
                        continue
                    know_val = get_value(line, "NEW KNOWLEDGE/SECRETS ACQUIRED")
                    if know_val != "NO_MATCH":
                        if know_val is not None and know_val not in char_obj["knowledge"]:
                            char_obj["knowledge"].append(know_val)
                        parsed_updates_for_log["new_knowledge"] = know_val if know_val is not None else "No change"
                        continue
            if current_char_name_update and parsed_updates_for_log and current_char_name_update in self.characters:
                self.characters[current_char_name_update]["development_log"].append(
                    {"chapter": chapter_num, "summary": "Updates from chapter events", **parsed_updates_for_log}
                )

        timeline_system_prompt = "You are a temporal analyst for narratives. Extract timeline information precisely."
        timeline_prompt = f"""
        Analyze Chapter {chapter_num}'s content for timeline information:
        Chapter {chapter_num} ("{self.chapter_plans.get(chapter_num, {}).get('title', 'Untitled')}") Content:
        ---
        {full_chapter_content}
        ---
        Determine:
        1.  APPROXIMATE TIME ELAPSED DURING THIS CHAPTER: [e.g., "Several hours," "One full day," "A week."]
        2.  TIME OF DAY/DATE AT THE END OF THIS CHAPTER: [e.g., "Evening of the third day," "Midnight," "Following morning."]
        3.  ANY SPECIFIC TIME MARKERS MENTIONED: [e.g., "After dawn," "Two moons passed," "Clock struck three."]
        Reply in format:
        ELAPSED: [answer]
        END_TIME: [answer]
        MARKERS: [answer]
        """
        timeline_text = self._ollama_generate(timeline_prompt, timeline_system_prompt, temperature=0.4)
        entry["timeline_elapsed"] = re.search(r"ELAPSED:\s*(.*)", timeline_text, re.IGNORECASE).group(1).strip() if re.search(r"ELAPSED:\s*(.*)", timeline_text, re.IGNORECASE) else "N/A"
        entry["timeline_end"] = re.search(r"END_TIME:\s*(.*)", timeline_text, re.IGNORECASE).group(1).strip() if re.search(r"END_TIME:\s*(.*)", timeline_text, re.IGNORECASE) else "N/A"
        entry["timeline_markers"] = re.search(r"MARKERS:\s*(.*)", timeline_text, re.IGNORECASE).group(1).strip() if re.search(r"MARKERS:\s*(.*)", timeline_text, re.IGNORECASE) else "N/A"

        entry["emotional_tone_end_achieved_in_summary"] = entry["summary"][-300:]
        if "ending_hook_text" not in entry:
            entry["ending_hook_text"] = "N/A (Last chapter or hook not generated)"


        self.chapter_continuity_data[chapter_num] = entry

    def _generate_chapter_transition_hook(self, chapter_num, current_chapter_content, current_chapter_continuity):
        """Generates the transition hook/paragraph(s) for the end of the current chapter."""
        if chapter_num >= self.num_chapters: return "" # No hook for the very last chapter

        next_chap_num = chapter_num + 1
        next_chapter_plan = self.chapter_plans.get(next_chap_num)
        if not next_chapter_plan:
            print(f"Warning: No plan found for Chapter {next_chap_num} to generate hook from Chapter {chapter_num}.")
            return "\n\n(The story continues...)" # Generic fallback

        system_prompt = f"You are a master storyteller in the style of {self.author_style}, crafting suspenseful and engaging chapter endings that seamlessly lead into the next."
        prompt = f"""
        You are writing the VERY LAST paragraph(s) for Chapter {chapter_num} ("{self.chapter_plans.get(chapter_num,{}).get('title','Untitled')}").
        This transition must create anticipation for Chapter {next_chap_num} ("{next_chapter_plan.get('title','Untitled')}").

        End of Chapter {chapter_num} Context:
        - Current Chapter Summary (focus on ending events): {current_chapter_continuity.get('summary', 'The chapter concluded.')[-1000:]}
        - Emotional Tone at End (inferred from summary): {current_chapter_continuity.get('emotional_tone_end_achieved_in_summary', 'Neutral')}
        - Last ~500 characters of Chapter {chapter_num} (before this hook): "{current_chapter_content[-500:]}"
        - Key Character States at End of Ch {chapter_num}: {json.dumps({name: {'status': data['current_status'], 'location': data['current_location'], 'emotion': data['emotional_state']} for name, data in self.characters.items()}, indent=2)}

        Plan for NEXT Chapter ({next_chap_num} - "{next_chapter_plan.get('title', 'Untitled')}"):
        - Next Chapter Goal: {next_chapter_plan.get('goal', 'The story continues.')}
        - Next Chapter Likely Opening Scene/Focus: {next_chapter_plan.get('scenes', ['A new challenge arises.'])[0] if next_chapter_plan.get('scenes') else 'A new challenge arises.'}
        - Next Chapter Character Development Focus: {next_chapter_plan.get('character_development', 'Further growth.')}
        - Planned connection from current Ch {chapter_num} to next Ch {next_chap_num}: "{self.chapter_plans.get(chapter_num,{}).get('connection_to_next','N/A')}"

        Write 1-2 compelling paragraphs (approx 75-150 words) that will be the *final text* of Chapter {chapter_num}. This hook should:
        1.  Provide a sense of immediate closure for Chapter {chapter_num}'s main events but leave the reader wanting more.
        2.  Directly foreshadow, question, or set up the conflict, theme, or situation of Chapter {next_chap_num} based on its plan AND the planned connection.
        3.  Maintain or slightly shift the established emotional tone to build suspense, curiosity, or dread for what's next.
        4.  Avoid clichés. Be original and impactful. Ensure it feels like a natural continuation of the narrative, not an abrupt summary.
        5.  This is the *final text* of Chapter {chapter_num}.

        Final transition paragraph(s) for Chapter {chapter_num}:
        """
        hook_text = self._ollama_generate(prompt, system_prompt, temperature=0.75)

        if chapter_num in self.chapter_continuity_data:
            self.chapter_continuity_data[chapter_num]["ending_hook_text"] = hook_text.strip()
        else:
            self.chapter_continuity_data[chapter_num] = {"ending_hook_text": hook_text.strip()}

        return f"\n\n{hook_text}"


    def generate_novel_content(self):
        """
        Main loop to generate content for all chapters, applying coherence measures.
        """
        print("\n--- Generating Full Novel Content (Chapter by Chapter) ---")
        if not self.chapter_plans or self.num_chapters == 0:
            print("ERROR: Cannot generate novel content without detailed chapter plans or chapter count.")
            return False

        for i in range(1, self.num_chapters + 1):
            print(f"\n--- Generating Chapter {i} of {self.num_chapters} ---")
            current_chapter_plan = self.chapter_plans.get(i)
            if not current_chapter_plan:
                print(f"ERROR: No plan found for Chapter {i}. Skipping.")
                self.generated_chapters_content[i] = f"[ERROR: No plan found for Chapter {i}]"
                self.chapter_continuity_data[i] = {"summary": "Error: No plan.", "character_updates_text": "", "timeline_end": "Unknown", "emotional_tone_end_achieved_in_summary": "Error", "ending_hook_text": "", "flow_analysis_from_previous": "N/A"}
                continue

            continuity_context = self._get_continuity_context_for_chapter(i)

            chapter_opener_text_with_title = self._generate_chapter_opener(i, current_chapter_plan)

            chapter_prose = chapter_opener_text_with_title

            if i > 1:
                self._analyze_inter_chapter_flow(i - 1, i, chapter_opener_text_with_title)

            scenes = current_chapter_plan.get("scenes", [])
            if not scenes:
                print(f"Warning: No scenes defined in plan for Chapter {i}. Chapter might be short or only opener/hook.")
                if chapter_prose.strip() == chapter_opener_text_with_title.strip(): # If only opener was generated
                     chapter_prose += "\n\n[This chapter's plan had no specific scenes. The narrative continues based on the chapter goal.]\n\n"
            else:
                accumulated_scene_prose_for_chapter = chapter_opener_text_with_title # Start with opener
                for scene_idx, scene_desc in enumerate(scenes):
                    print(f"  Generating Scene {scene_idx + 1} of {len(scenes)} for Chapter {i}: {scene_desc[:80]}...")
                    scene_specific_prose = self._generate_scene_prose(i, scene_idx, scene_desc, current_chapter_plan, continuity_context, accumulated_scene_prose_for_chapter)
                    if "[OLLAMA" in scene_specific_prose:
                        print(f"    ERROR generating scene {scene_idx+1}: {scene_specific_prose}")
                        scene_specific_prose = f"\n\n[Error generating scene: {scene_desc[:50]}...]\n\n"

                    chapter_prose += scene_specific_prose + "\n\n" # Append scene to overall chapter prose
                    accumulated_scene_prose_for_chapter += scene_specific_prose + "\n\n" # Update context for next scene in this chapter
                    time.sleep(0.2)

            # Interim continuity update (based on content BEFORE the hook for this chapter)
            # This is useful for the hook generation itself, if it needs summary of current chapter.
            self._update_chapter_continuity_data(i, chapter_prose.strip(), is_final_pass_for_chapter=False)

            if i < self.num_chapters:
                print(f"  Generating transition hook for Chapter {i}...")
                hook_text = self._generate_chapter_transition_hook(i, chapter_prose.strip(), self.chapter_continuity_data[i])
                chapter_prose += hook_text

            self.generated_chapters_content[i] = chapter_prose.strip()
            print(f"  Chapter {i} ('{current_chapter_plan.get('title', 'Untitled')}') content generated (approx length: {len(chapter_prose)} chars).")

            # FINAL continuity update for the chapter (with opener, scenes, and hook included)
            self._update_chapter_continuity_data(i, self.generated_chapters_content[i], is_final_pass_for_chapter=True)

            if i < self.num_chapters:
                print("Pausing briefly before next chapter...")
                time.sleep(0.5)

        return True

    # --- Transition Checking Phase ---
    def _check_and_improve_transition(self, prev_chapter_num, current_chapter_num):
        """Checks transition from prev to current chapter and improves if needed."""
        print(f"\n--- Checking transition from Chapter {prev_chapter_num} to {current_chapter_num} ---")

        prev_chapter_content = self.generated_chapters_content.get(prev_chapter_num)
        current_chapter_content = self.generated_chapters_content.get(current_chapter_num)

        if not prev_chapter_content or not current_chapter_content:
            print("  Skipping transition check: Missing content for one or both chapters.")
            return

        system_prompt = """You are a professional editor specializing in narrative flow and chapter transitions."""
        prompt = f"""Analyze the transition between the end of the previous chapter and the beginning of the current chapter.

        END OF PREVIOUS CHAPTER ({prev_chapter_num}):
        ---
        {prev_chapter_content[-1000:]}
        ---

        BEGINNING OF CURRENT CHAPTER ({current_chapter_num}):
        ---
        {current_chapter_content[:1000]}
        ---

        If the transition is already smooth and logical, respond ONLY with the exact text:
        TRANSITION: SMOOTH

        Otherwise, provide an improved beginning for the current chapter (first 1-3 paragraphs, approx 100-250 words) that:
        1. Creates a smoother, more logical connection with the previous chapter's ending hook/state.
        2. Avoids repeating information already established.
        3. Maintains character and plot consistency.
        4. Progresses the timeline naturally.
        5. Matches the established author style ({self.author_style}).

        Start your response with the exact text "TRANSITION: REVISED" followed by the revised beginning paragraphs. Do NOT include the chapter title in the revised text.
        """
        transition_check_result = self._ollama_generate(prompt, system_prompt, temperature=0.6)

        if "[OLLAMA" in transition_check_result:
            print(f"  Error during transition check: {transition_check_result}")
        elif "TRANSITION: REVISED" in transition_check_result:
            try:
                revised_beginning = transition_check_result.split("TRANSITION: REVISED", 1)[1].strip()

                if revised_beginning:
                    print(f"  Transition needs improvement. Applying revised opening to Chapter {current_chapter_num}.")
                    original_lines = current_chapter_content.split('\n', 1) # Split only the first line (title)
                    original_title_line = original_lines[0]

                    # Find where the original opener text effectively ends.
                    # This is a heuristic. We assume the opener is before the first "scene" content.
                    # A simple way is to replace the first few paragraphs.
                    # Let's find the end of the original opener by looking for the second double newline after the title.

                    body_after_title = original_lines[1] if len(original_lines) > 1 else ""
                    # Split the body into paragraphs, keeping double newlines as separators
                    # The opener is usually the first paragraph block after the title line
                    # For robustness, we find the first distinct section of the original chapter body after the title.
                    # The title itself is `Chapter X - YYY\n\nActual Opener Text...`
                    # So, current_chapter_content has Title\n\nOpener\n\nScene1...

                    # Split after the first double newline (which is after the title)
                    parts_after_title_line = current_chapter_content.split('\n\n', 1)
                    if len(parts_after_title_line) > 1:
                        original_opener_and_rest = parts_after_title_line[1]
                        # Now split the original_opener_and_rest to separate the opener from scenes
                        # Assuming opener is one block of text followed by \n\n
                        original_opener_parts = original_opener_and_rest.split('\n\n', 1)
                        original_rest_of_chapter = original_opener_parts[1] if len(original_opener_parts) > 1 else ""

                        self.generated_chapters_content[current_chapter_num] = f"{original_title_line}\n\n{revised_beginning}\n\n{original_rest_of_chapter}".strip()
                        print(f"  Chapter {current_chapter_num} opening revised successfully.")
                    else: # Chapter content was just the title line or title + one block
                         self.generated_chapters_content[current_chapter_num] = f"{original_title_line}\n\n{revised_beginning}".strip()
                         print(f"  Chapter {current_chapter_num} (short) opening revised successfully.")

                else:
                    print("  Transition check indicated revision needed, but no revised text was provided by LLM.")
            except Exception as e:
                print(f"  Error applying revised transition for Chapter {current_chapter_num}: {e}")
        elif "TRANSITION: SMOOTH" in transition_check_result:
            print("  Transition is smooth. No changes needed.")
        else:
            print(f"  Transition check response was unclear: {transition_check_result[:200]}... No changes applied.")


    def _perform_final_transition_checks(self):
        """Loops through all chapters to check and improve transitions."""
        print("\n--- Performing Final Pass: Checking Chapter Transitions ---")
        if len(self.generated_chapters_content) < 2:
            print("  Skipping transition checks (less than 2 chapters generated).")
            return

        for i in range(2, self.num_chapters + 1):
            if i in self.generated_chapters_content and (i - 1) in self.generated_chapters_content:
                self._check_and_improve_transition(i - 1, i)
                time.sleep(0.5)
            else:
                print(f"  Skipping transition check for Chapter {i} (missing previous or current chapter content).")
        print("--- Finished Final Transition Checks ---")


    # --- Phase 4: Compilation & Output Methods ---
    def generate_novel_title(self):
        """Generates a compelling title for the novel."""
        print("\n--- Generating Novel Title ---")
        if not self.subject or not self.characters or not self.themes_motifs:
            self.novel_title = f"A {self.genre.replace('/', ' ')} Story"
            print(f"Warning: Insufficient data for title generation. Using placeholder: {self.novel_title}")
            return

        system_prompt = f"You are a creative book title generator, expert in {self.genre} and the style of {self.author_style}."
        prompt = f"""
        Generate ONE compelling and marketable novel title based on the following details.
        The title must be highly consistent with the genre, author's style, subject, main character(s), world, and core themes.
        It should be intriguing, memorable, and not overly generic.

        Novel Subject: {self.subject}
        Genre: {self.genre}
        Author's Style Inspiration: {self.author_style}
        Main Character(s): {', '.join([f'{name} ({data.get("role")})' for name, data in self.characters.items()])}
        World Name/Atmosphere: {self.world_details.get('name', 'N/A')} / {self.world_details.get('atmosphere', 'N/A')}
        Core Themes: {', '.join(self.themes_motifs.get('themes', []))}
        Recurring Motifs: {', '.join(self.themes_motifs.get('motifs', []))}

        Return ONLY the generated title itself, without any quotation marks, labels (like "Title:"), or explanatory text.
        Novel Title:
        """
        title_text = self._ollama_generate(prompt, system_prompt, temperature=0.8)
        if "[OLLAMA" in title_text or not title_text.strip():
            print(f"ERROR generating title: {title_text}. Using placeholder.")
            main_char_name = list(self.characters.keys())[0] if self.characters else 'Adventure'
            self.novel_title = f"A {self.genre.replace('/', ' ')} Tale of {main_char_name}"
        else:
            title_text = re.sub(r'^(title|novel title):?\s*', '', title_text, flags=re.IGNORECASE).strip()
            self.novel_title = title_text.strip('"\'')
        print(f"Generated Novel Title: {self.novel_title}")


    def compile_and_save_novel(self):
        """Compiles the generated content into a .docx file and saves it."""
        print("\n--- Compiling and Saving Novel ---")
        if not self.generated_chapters_content:
            print("ERROR: No chapter content generated. Cannot save novel.")
            return

        self.generate_novel_title()

        doc = Document()
        try:
            title_style = doc.styles['Title']
            title_style.font.name = 'Garamond'
            title_style.font.size = Pt(28)
        except KeyError:
            print("Warning: 'Title' style not found. Using default.")
            title_style = 'Title'

        try:
            # In python-docx, 'Heading 1' corresponds to level 1 heading.
            # If you want to style it, you access it via doc.styles and modify.
            # Adding a heading uses doc.add_heading('Text', level=1)
            # We will apply style to paragraphs identified as headings later.
            # For now, ensure a base normal style.
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Garamond'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing = 1.5
            normal_style.paragraph_format.space_after = Pt(0)
            normal_style.paragraph_format.first_line_indent = Pt(24)
        except KeyError:
            print("Warning: 'Normal' style not found. Using default.")
            normal_style = 'Normal'


        doc.add_paragraph(self.novel_title, style=title_style).alignment = WD_ALIGN_PARAGRAPH.CENTER

        author_para_style = doc.styles['Normal'].font # Get a copy
        author_para_style.name = 'Garamond'
        author_para_style.size = Pt(12)
        author_line = doc.add_paragraph(style='Normal')
        author_line.add_run(f"Inspired by the style of {self.author_style}").font.name = 'Garamond'
        author_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_line.paragraph_format.first_line_indent = None


        genre_line = doc.add_paragraph(style='Normal')
        genre_line.add_run(f"Genre: {self.genre}").font.name = 'Garamond'
        genre_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        genre_line.paragraph_format.first_line_indent = None
        doc.add_page_break()

        for i in sorted(self.generated_chapters_content.keys()): # Ensure chapters are in order
            chapter_content = self.generated_chapters_content.get(i, f"[ERROR: Content for Chapter {i} not found]")

            paragraphs = chapter_content.split('\n\n')

            if paragraphs:
                ch_title_line_full = paragraphs[0].strip()

                # Default heading text is the full first line
                heading_text_for_doc = ch_title_line_full

                # Try to extract a cleaner title for the heading
                # Expected format: "Chapter X - Title Text"
                title_match_for_heading = re.match(r"Chapter\s*\d+\s*[:*-]?\s*(.*)", ch_title_line_full, re.IGNORECASE)
                if title_match_for_heading and title_match_for_heading.group(1).strip():
                    heading_text_for_doc = title_match_for_heading.group(1).strip() # Use just the title part

                # If what we extracted as "title" is very short or empty, revert to full line for safety
                if not heading_text_for_doc or len(heading_text_for_doc) < 3:
                    heading_text_for_doc = ch_title_line_full


                ch_heading_para = doc.add_heading(heading_text_for_doc, level=1)
                ch_heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Accessing style for heading directly to modify if needed (python-docx limitation)
                # For more control, create custom style based on 'Heading 1'
                # For now, assume default 'Heading 1' style is acceptable or modify 'Heading 1' in styles
                # ch_heading_para.style.font.name = 'Garamond' # This would require ensuring style object is not string
                # ch_heading_para.style.font.size = Pt(18)

                # Add the "Chapter X" part if it was separated, or if the title line didn't include it for some reason
                if not ch_title_line_full.lower().startswith("chapter"):
                    sub_heading_para = doc.add_paragraph(f"Chapter {i}")
                    sub_heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sub_heading_para.paragraph_format.space_before = Pt(0)
                    sub_heading_para.paragraph_format.space_after = Pt(6)
                    if hasattr(sub_heading_para.style.font, 'name'):
                         sub_heading_para.style.font.name = 'Garamond'
                         sub_heading_para.style.font.size = Pt(14)


                for para_block in paragraphs[1:]: # Start from second block for content
                    if para_block.strip():
                        p = doc.add_paragraph(para_block.strip(), style=normal_style)
            else:
                doc.add_paragraph(chapter_content, style=normal_style)

            if i < self.num_chapters:
                doc.add_page_break()

        safe_title = re.sub(r'[^\w\s-]', '', self.novel_title).strip().replace(' ', '_')
        safe_genre = self.genre.replace('/','-').replace(' ','')
        filename = f"{safe_title[:50]}_Novel_{safe_genre}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)

        try:
            doc.save(filepath)
            print(f"Novel successfully saved to: {filepath}")
        except Exception as e:
            print(f"ERROR saving .docx file: {e}")

        metadata = {
            "title": self.novel_title,
            "subject": self.subject,
            "author_style": self.author_style,
            "genre": self.genre,
            "num_chapters_determined": self.num_chapters,
            "ollama_model_used": OLLAMA_MODEL,
            "generation_timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "characters_final": self.characters,
            "world_details": self.world_details,
            "themes_motifs": self.themes_motifs,
            "plot_outline": self.plot_outline,
            "chapter_plans": self.chapter_plans,
            "chapter_continuity_data": self.chapter_continuity_data,
        }
        meta_filename = f"{safe_title[:50]}_Novel_METADATA.json"
        meta_filepath = os.path.join(OUTPUT_DIR, meta_filename)
        try:
            with open(meta_filepath, "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
            print(f"Metadata saved to: {meta_filepath}")
        except Exception as e:
            print(f"ERROR saving metadata JSON: {e}")


    def orchestrate_generation(self):
        """
        Main public method to run the entire novel generation pipeline.
        """
        start_time = time.time()
        print("--- Starting Novel Generation Pipeline ---")

        if not self.generate_foundational_elements():
            # Error message already printed in generate_foundational_elements
            return

        if not self.generate_detailed_chapter_plans():
            # Error message already printed in generate_detailed_chapter_plans
            return

        if not self.generate_novel_content():
            print("Halting: Novel content generation failed.")
            return

        self._perform_final_transition_checks()

        self.compile_and_save_novel()

        end_time = time.time()
        total_time_minutes = (end_time - start_time) / 60
        print(f"--- Novel Generation Pipeline Finished ---")
        print(f"Total time taken: {total_time_minutes:.2f} minutes.")


def get_user_input_multiline(prompt_message):
    print(prompt_message + " (Type 'ENDINPUT' on a new line when done, or just press Enter if input is short):")
    lines = []
    first_line = input()
    if not first_line.strip() and not lines: # Handle immediate Enter press for short input
        return first_line

    if first_line.strip().upper() == 'ENDINPUT':
        return "" # Empty if ENDINPUT is the first thing

    lines.append(first_line)

    # If first line is short and no explicit ENDINPUT, assume it's a single line input
    if len(first_line) < 70 and "ENDINPUT" not in first_line.strip().upper() :
        is_multiline_intent = False # Heuristic: assume single line if short
        # Check if user might still want multiline
        if any(kw in first_line for kw in ["\n", "\\n"]): # A bit of a guess
            is_multiline_intent = True
        if not is_multiline_intent: # If it looks like a single line, return it.
             # Check if there's an accidental ENDINPUT at the end of a short first line.
            if first_line.strip().upper().endswith("ENDINPUT"):
                return first_line.strip()[:-(len("ENDINPUT"))].strip()
            return first_line # Return as is

    # Proceed to read more lines if it seems like multiline input was intended
    while True:
        try:
            line = input()
            if line.strip().upper() == 'ENDINPUT':
                break
            lines.append(line)
        except EOFError: # Handle Ctrl+D or unexpected end of input stream
            print("INFO: EOF reached while reading multiline input.")
            break
    return "\n".join(lines)

def load_resume_text(file_path):
    """Loads text from a file, attempting PDF extraction if it's a .pdf file."""
    if not file_path:
        return ""

    resume_text = ""
    try:
        if file_path.lower().endswith(".pdf"):
            try:
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    # Check if PDF is encrypted - implementation depends on pypdf version
                    is_encrypted = getattr(reader, "is_encrypted", None)
                    if is_encrypted is not None:  # New pypdf versions
                        if reader.is_encrypted:
                            try:
                                # Attempt to decrypt with empty password
                                if reader.decrypt('') == 0:  # 0 means failed
                                    print(f"Warning: PDF file '{file_path}' is encrypted and could not be decrypted with empty password.")
                                    return ""
                            except Exception as de:
                                print(f"Warning: PDF file '{file_path}' is encrypted. Attempt to decrypt failed: {de}. Cannot extract text.")
                                return ""
                    else:  # Older pypdf versions
                        try:
                            if hasattr(reader, "decrypt"):  # Try decryption if method exists
                                reader.decrypt('')
                        except:
                            print(f"Warning: PDF file '{file_path}' might be encrypted and couldn't be processed.")
                            return ""

                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            resume_text += page_text + "\n"
                if resume_text.strip():
                    print(f"Successfully extracted text from PDF: {file_path}")
                else:
                    print(f"Warning: No text could be extracted from PDF: {file_path}. It might be an image-based PDF, scanned, or corrupted.")
            except Exception as e:
                print(f"Warning: Error processing PDF file '{file_path}': {e}. Proceeding without resume content from this file.")
                resume_text = ""
        else:
            encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
            loaded_successfully = False
            for enc in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        resume_text = f.read()
                    print(f"Resume loaded from text file: {file_path} (using {enc})")
                    loaded_successfully = True
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e_text:
                    print(f"Warning: Error loading text file '{file_path}' with {enc}: {e_text}.")
                    break
            if not loaded_successfully and not resume_text:
                print(f"Warning: Could not load resume from text file '{file_path}' after trying multiple encodings. Proceeding without it.")

    except FileNotFoundError:
        print(f"Warning: Resume file not found at '{file_path}'. Proceeding without resume.")
    except Exception as e_general:
        print(f"Warning: An unexpected error occurred while trying to load '{file_path}': {e_general}. Proceeding without resume.")

    return resume_text.strip()


if __name__ == "__main__":
    print("Welcome to the AI Novel Generator!")
    print("Please ensure your Ollama server is running and the model is available.")
    # OLLAMA_MODEL can be overridden by user input if desired, or set here.
    user_ollama_model = input(f"Enter Ollama model name (default: {OLLAMA_MODEL}): ").strip()
    if user_ollama_model:
        OLLAMA_MODEL = user_ollama_model
    print(f"Using Ollama Model: {OLLAMA_MODEL} at {OLLAMA_BASE_URL}")
    print("----------------------------------------------------")

    resume_file_path_input = input("Enter path to resume file (text or PDF) (or press Enter to skip): ").strip()
    resume_text_content = load_resume_text(resume_file_path_input)

    novel_subject_input = get_user_input_multiline("Enter the novel's subject/premise")
    author_style_input_str = input("Enter the desired author style (e.g., 'Stephen King', 'Jane Austen'): ").strip()
    genre_input_str = input("Enter the genre(s) (e.g., 'Sci-Fi/Thriller', 'Historical Romance'): ").strip()

    # Basic input validation
    if not novel_subject_input:
        print("Novel subject/premise cannot be empty. Exiting.")
        exit()
    if not author_style_input_str:
        print("Author style cannot be empty. Using 'Generic'.")
        author_style_input_str = "Generic"
    if not genre_input_str:
        print("Genre cannot be empty. Using 'Fiction'.")
        genre_input_str = "Fiction"


    generator = NovelGenerator(
        resume_content=resume_text_content,
        subject=novel_subject_input,
        author_style=author_style_input_str,
        genre=genre_input_str
    )
    generator.orchestrate_generation()
    print("----------------------------------------------------")



# Welcome to the AI Novel Generator!
# Please ensure your Ollama server is running and the model is available.
# Enter Ollama model name (default: gemma3:12b): gemma3:12b
# Using Ollama Model: gemma3:12b at http://localhost:11434/api/generate
# ----------------------------------------------------
# Enter path to resume file (text or PDF) (or press Enter to skip): /content/Jessica_Holdens_resume.pdf
# Successfully extracted text from PDF: /content/Jessica_Holdens_resume.pdf
# Enter the novel's subject/premise (Type 'ENDINPUT' on a new line when done, or just press Enter if input is short):
# The salt-laced air of Los Angeles had always felt like a distant cousin to Jessica's soul, a mere whisper of the vast, blue heart she felt beating within her. Eighteen years had etched a quiet yearning onto her spirit, a pull towards the ocean that went deeper than simple admiration. It was a recognition, a silent conversation held across the sandy divide.    Then came the dreams. At first, they were fleeting glimpses – the shimmer of unseen depths, the murmur of voices she couldn't quite grasp, a profound sense of belonging to a place she'd never seen. But since her eighteenth birthday, the whispers had grown insistent, the glimmers brighter, each nocturnal voyage tugging her further into a reality that felt both alien and intimately familiar. A world veiled in a sapphire haze, humming with an ancient magic, was calling her home.    One ordinary Sunday morning, the call shattered the fragile membrane of her waking life. An overwhelming vision flooded her senses – a swirling vortex of turquoise, the mournful cry of unseen creatures – followed by a blinding flash, and then… nothing.    She awoke to a world reborn. Gone was the familiar clutter of her bedroom, replaced by an endless expanse of ocean, its surface alive with a shimmering, sentient energy. The Blueward Deep. A name that echoed in the deepest chambers of her heart, a forgotten world whispered about only in the faintest of myths, had claimed her entirely.    The revelation that she was more than a visitor struck her with the force of a tidal wave. Heir. The word resonated with the ocean's hum, a legacy woven into the very currents. Tideborn. Endowed with a rare and potent gift called Soulcurrent, she could feel the ocean's pulse, a symphony of ebb and flow she instinctively understood. The waters responded to her, currents bending at her unspoken will, and within their depths, she could almost hear the echoes of ancient memories.    But this homecoming was shadowed by peril. Selmyra was not at peace. A malevolent force, the Drowned Court, clawed at the ocean's magic, their hunger for power twisting the once-harmonious currents into weapons of destruction. Monstrous sea beasts, their eyes gleaming with unnatural malice, patrolled the depths. Dark rifts scarred the ocean's surface, spewing forth chaotic energies. And a chilling dread hung in the water – the slow, insidious awakening of The Abyss, a slumbering, god-like entity whose stirring threatened to swallow Selmyra whole.    Amidst this breathtaking yet terrifying reality, Jessica’s innate empathy for the ocean became her unexpected strength. The silent suffering of the corrupted waters resonated within her, fueling a fierce protectiveness. And then there was Kaelen.    He moved with the fluid grace of the tides, his eyes the deep, calming blue of a twilight sea. A Tideborn himself, his connection to Selmyra was profound, yet it was the unexpected resonance with Jessica’s Soulcurrent that sparked a different kind of current between them perhaps romance. In his steady gaze, she found not just an ally, but a silent understanding that mirrored her own deep connection to the ocean. His presence was a comforting anchor in this bewildering world, a warmth that spread through her like sunlight on the water’s surface.    Guided by fragmented visions – a girl lost to the churning waves, cryptic flashbacks that felt both alien and achingly familiar – Jessica knew her arrival was no mere chance. The mysterious rescue on a beach years ago, a memory that flickered at the edges of her awareness, seemed to hold a key. And the truth of the Tideborn lineage, the ancient magic that pulsed within her veins, was a mystery she desperately needed to unravel.    Could she, a dreamer pulled from a world she barely knew she'd left, become the anchor of Selmyra against the rising tide of madness? And could the tentative warmth and feelings she felt for Kaelen blossom amidst the looming war, a fragile bloom in a world teetering on the brink?    Her journey would plunge her into the heart of Selmyra’s deepest secrets – where trust was as treacherous as a hidden reef, alliances shifted with the currents of hidden agendas, and the ocean itself seemed to hold its breath, waiting. As a Tideborn wielding the rare Soulcurrent, Jessica would have to awaken the full extent of her power, to face the shadows that lurked beneath the waves, and to choose: would she rise with the tide, or be swallowed by the depths? The fate of Selmyra, and perhaps her own heart, hung in the balance.
# ENDINPUT
# Enter the desired author style (e.g., 'Stephen King', 'Jane Austen'): Victoria Aveyard- \n Genre: ROMANCE,FANTASY,DYSTOPIAN,YOUNG ADULT,HOT,SEXY
# Enter the genre(s) (e.g., 'Sci-Fi/Thriller', 'Historical Romance'):  ROMANCE,FANTASY,DYSTOPIAN,YOUNG ADULT,HOT,SEXY
# NovelGenerator initialized.
#   Subject: The salt-laced air of Los Angeles had always felt like a distant cousin to Jessica's soul, a mere wh...
#   Author Style: Victoria Aveyard- \n Genre: ROMANCE,FANTASY,DYSTOPIAN,YOUNG ADULT,HOT,SEXY
#   Genre: ROMANCE,FANTASY,DYSTOPIAN,YOUNG ADULT,HOT,SEXY
#   Resume provided: Yes
#   Number of chapters will be determined automatically.
# --- Starting Novel Generation Pipeline ---

# --- Generating Foundational Elements ---
# Step 1.1: Generating Character Profiles...
# Generated 3 character profiles: Jessica Holdens, Kaelen Vesper, Lady Morwen Ashworth

# Step 1.2: Generating World Details...
# Generated World Details for 'Selmyra'.

# Step 1.3: Generating Themes and Motifs...
# Generated Themes: ["Destiny vs. Free Will: Jessica’s arrival and “heir” status challenge the notion of predetermined fate, while her choices shape Selmyra's future.", 'Environmental Corruption & Responsibility: The Drowned Court’s actions and The Abyss’s stirring highlight the consequences of exploiting natural resources and the importance of ecological stewardship.', 'Trauma and Healing: Kaelen’s past and Lady Morwen’s descent explore the long-term effects of trauma and the possibility of finding solace and redemption.', 'Belonging & Identity: Jessica’s journey is about finding her place in a world she never knew, grappling with her dual heritage and defining her own identity.', 'RECURRING MOTIFS:', 'Saltwater: Represents both life and decay, purity and corruption.', 'Bioluminescence: Symbolizes hidden knowledge, magic, and the beauty that can exist even in darkness.', 'Echoes/Resonance: Represents fragmented memories, the weight of the past, and the interconnectedness of Selmyra’s inhabitants.', 'Shells: Symbolize protection, secrets, and the fragility of life.', "Turquoise/Sapphire Hues: Represent the alluring and deceptive nature of Selmyra's magic."]
# Generated Motifs: ['Saltwater: Represents both life and decay, purity and corruption.', 'Bioluminescence: Symbolizes hidden knowledge, magic, and the beauty that can exist even in darkness.', 'Echoes/Resonance: Represents fragmented memories, the weight of the past, and the interconnectedness of Selmyra’s inhabitants.', 'Shells: Symbolize protection, secrets, and the fragility of life.', "Turquoise/Sapphire Hues: Represent the alluring and deceptive nature of Selmyra's magic."]

# Step 1.4: Generating High-Level Plot Outline and Determining Chapter Count...
# Generated High-Level Plot Outline:
# **Act I (Setup):** Jessica’s ordinary life is shattered when she's pulled from her bedroom and awakens in the underwater realm of Selmyra, discovering she is a Tideborn heir with a rare Soulcurrent ability. She meets Kaelen, who becomes her guide, while the looming threat of the Drowned Court and the stirring of The Abyss are introduced, establishing the immediate danger to Selmyra. The inciting incident is her arrival and the revelation of her heritage, setting her on a path to uncover her purpose and the secrets of her past.

# **Act II (Confrontation):** Jessica undergoes rigorous training to harness her Soulcurrent, facing trials that test her abilities and reveal fragmented memories of a past life connected to Selmyra. She uncovers a conspiracy within the Drowned Court, led by Lady Morwen, who seeks to awaken The Abyss for power, forcing Jessica and Kaelen to confront escalating conflicts and navigate treacherous alliances. A pivotal moment occurs when Jessica discovers a hidden connection to Lady Morwen, complicating her understanding of the conflict and forcing her to question everything she thought she knew.

# **Act III (Resolution):** Jessica confronts Lady Morwen in a final, desperate battle to prevent the awakening of The Abyss, utilizing her fully realized Soulcurrent abilities and confronting the truth about her past and Lady Morwen’s tragic history. The conflict culminates in a choice: sacrifice a piece of herself to stabilize the ocean’s magic or risk Selmyra’s destruction, ultimately leading to a bittersweet victory and a new era of balance. Jessica embraces her role as a protector of Selmyra, forging a path towards a future where her human and Tideborn identities intertwine.
# LLM suggested 35 chapters for the novel.
# Foundational elements generated successfully.

# --- Generating Detailed Chapter-by-Chapter Plans ---
# Generating detailed plan text for 35 chapters (this may take a while)...
# ERROR: No chapter plans were successfully parsed. The LLM output might not conform to the expected structure.
# --- LLM Output Start (first 2000 chars) ---
# Okay, here's a detailed chapter-by-chapter plan for your romance, fantasy, dystopian YA novel, structured as requested. It's a substantial amount of detail, designed to be a working guide. I'm aiming for a Victoria Aveyard-esque feel – morally grey characters, a strong sense of atmosphere, and a focus on internal struggles alongside external conflicts.

# **Chapter 1 - The Salt-Kissed Echo:**
# 1.  CHAPTER GOAL: Introduce Jessica, establish her connection to the ocean, and initiate the inciting incident - her abrupt transportation to Selmyra. This chapter establishes the mystery and disorientation that will drive the initial plot.
# 2.  KEY SCENES:
#     - Scene 1: Jessica's morning routine, feeling a vague restlessness. Location: Jessica's bedroom. Characters Involved: Jessica. Key Revelation/Turning Point/Outcome: Jessica experiences a particularly vivid dream, more intense than usual.
#     - Scene 2: The vision floods Jessica’s senses – turquoise vortex, mournful cries, blinding flash. Location: Jessica's bedroom. Characters Involved: Jessica. Key Revelation/Turning Point/Outcome: Jessica is pulled from her world.
#     - Scene 3: Jessica awakens in the Blueward Deep, disoriented and surrounded by strange, bioluminescent flora. Location: The Blueward Deep. Characters Involved: Jessica. Key Revelation/Turning Point/Outcome: Jessica realizes she's no longer in her bedroom.
# 3.  CHARACTER DEVELOPMENT FOCUS: Jessica: Initially bewildered and frightened, demonstrating a quiet strength as she begins to assess her situation.
# 4.  PLOT ADVANCEMENT: Introduces the central conflict – Jessica’s displacement and the mystery of Selmyra.
# 5.  TIMELINE & PACING: Takes place within a few hours. Pacing: Rapid, disorienting.
# 6.  EMOTIONAL TONE (End of Chapter): Disoriented, frightened, but with a burgeoning sense of wonder.
# 7.  CONNECTION TO NEXT CHAPTER (Setup/Hook): Jessica notices a faint, familiar resonance within the ocean’s hum – a feeling she can't explain.

# **Chapter 2 - Currents of Rec
# --- LLM Output End ---
# ----------------------------------------------------
