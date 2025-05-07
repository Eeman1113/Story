# app_v2_refined.py
from dotenv import load_dotenv
import os
import docx
import traceback
import time
import re
import logging # MOD: Added logging

# --- Langchain Imports ---
from langchain_ollama import OllamaLLM
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import PyPDFLoader
# --- End Imports ---

# Load environment variables
load_dotenv()

# --- MOD: Setup Structured Logging ---
# For a production system, you might configure this more extensively (e.g., file output, log rotation)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(name)s - %(message)s')
logger = logging.getLogger(__name__)
# --- END MOD ---

# --- Constants ---
DEFAULT_MODEL = os.getenv("OLLAMA_MODEL", "gemma3:4b") # Example, ensure this model is pulled
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OUTPUT_FOLDER = './docs'
LLM_CALL_DELAY_SECONDS = 0.5
DEFAULT_RESUME_FILENAME = 'kenji_gamer_resume.pdf'

# --- MOD: Custom Exceptions for Better Error Handling ---
class GenerationError(Exception):
    """Base class for errors during content generation."""
    pass

class CriticalGenerationError(GenerationError):
    """An error that should likely halt the entire process."""
    pass

class ProfileGenerationError(CriticalGenerationError): pass
class PlotGenerationError(CriticalGenerationError): pass
class ChapterGenerationError(CriticalGenerationError): pass
class EventGenerationError(GenerationError): pass # May not be critical for all chapters
class SettingGenerationError(GenerationError): pass
class ThemeGenerationError(GenerationError): pass
class TitleGenerationError(GenerationError): pass
class WriterError(GenerationError): pass
class RefinementError(GenerationError): pass
class ParsingError(GenerationError): pass
# --- END MOD ---

# --- LLM Initialization Function ---
def create_llm(temperature=0.7, top_p=0.9, top_k=40):
    logger.info(f"Connecting to Ollama at: {OLLAMA_BASE_URL} with Model: {DEFAULT_MODEL}")
    try:
        llm = OllamaLLM(
            model=DEFAULT_MODEL,
            temperature=temperature,
            top_p=top_p,
            top_k=top_k,
            base_url=OLLAMA_BASE_URL,
            request_timeout=180.0,
        )
        logger.info(f"LLM Instance Created (Temp: {temperature}, Top P: {top_p}, Top K: {top_k})")
        return llm
    except Exception as e:
        logger.error(f"FATAL ERROR: Could not create OllamaLLM instance. Error: {e}")
        logger.error(f"Attempted Base URL: {OLLAMA_BASE_URL}, Model: {DEFAULT_MODEL}")
        logger.exception("Details:") # MOD: Use logger.exception for tracebacks
        raise # Reraise to stop execution if LLM can't be created

# --- Chain Classes (Prompts remain largely the same as previous 'relatable' version) ---
# MOD: Added more specific error raising and logging within chains

class MainCharacterChain:
    PROMPT = """
    Analyze the following resume text to create a rich, multi-faceted character profile suitable for a literary narrative. Go beyond surface-level skills.
    Infer and extrapolate:
    1.  **Name:** (If mentioned, otherwise suggest one based on context or leave blank)
    2.  **Core Identity & Demeanor:** A brief summary (2-3 sentences) of their apparent personality, job/role, and how they present themselves.
    3.  **Key Skills/Strengths:** List 3-5 notable skills or positive attributes relevant to potential plot points.
    4.  **Potential Motivations:** What might drive this character? What deep-seated desires or goals could be hinted at? (Infer 2-3)
    5.  **Potential Flaws/Weaknesses:** What vulnerabilities, biases, or negative traits might they possess? (Infer 2-3)
    6.  **Internal Conflicts:** What inner struggles or contradictions could define their character arc? (Infer 1-2)
    7.  **Secrets/Hidden Depths:** What might this character be hiding from others, or even themselves? (Infer 1-2)
    8.  **Narrative Potential:** Briefly suggest (1-2 sentences) how these traits could fuel a compelling story in the {genre} genre.
    Resume Text:
    {text}
    Detailed Character Profile:"""
    def __init__(self):
        self.llm = create_llm(temperature=0.6, top_p=0.85)
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False) # MOD: verbose to False for cleaner logs

    def load_resume(self, file_name):
        file_path = os.path.join(OUTPUT_FOLDER, file_name)
        if not os.path.exists(file_path):
             raise FileNotFoundError(f"Resume file not found at: {file_path}")
        try:
            logger.info(f"Loading PDF from: {file_path}")
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            if not docs:
                logger.warning(f"PyPDFLoader didn't extract any documents/pages from {file_name}.")
                return None
            logger.info(f"Successfully loaded {len(docs)} page(s) from PDF.")
            full_text = '\n\n'.join([doc.page_content for doc in docs if doc.page_content])
            return re.sub(r'\s+', ' ', full_text).strip()
        except Exception as e:
             logger.error(f"Error loading or processing PDF {file_path}: {e}")
             logger.exception("PDF Loading Traceback:")
             raise

    def run(self, file_name, genre):
        try:
            resume_text = self.load_resume(file_name)
            if not resume_text or not resume_text.strip():
                 logger.error("Could not load or resume content is empty.")
                 raise ProfileGenerationError("Missing or empty resume content.")

            logger.info("Invoking MainCharacterChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({"text": resume_text, "genre": genre})
            profile = result.get('text', "").strip()

            if not profile or len(profile) < 100:
                logger.warning(f"Generated profile seems invalid or too short: '{profile[:100]}...'")
                raise ProfileGenerationError(f"Failed to generate a meaningful profile. LLM Output: {profile[:100]}...")
            return profile
        except Exception as e:
            logger.error(f"An error occurred in MainCharacterChain.run: {e}")
            if not isinstance(e, ProfileGenerationError): # Don't re-wrap if already specific
                raise ProfileGenerationError(f"Profile generation failed: {e}") from e
            raise


class SettingChain:
    PROMPT = """
    You are a world-building assistant specializing in atmospheric settings for novels.
    Based on the novel's subject, genre, and main character profile, generate a concise description of the primary setting(s).
    Focus on:
    1.  **Key Locations:** Identify 2-4 significant places where the story unfolds.
    2.  **Time Period/Atmosphere:** Describe the general era, mood, and sensory feeling (e.g., oppressive, nostalgic, futuristic, decaying, magical).
    3.  **Sensory Details:** Suggest 3-5 recurring sensory elements (specific sights, sounds, smells, textures) that define the world's feel. **Make these details tangible and relatable.**
    4.  **Relation to Character/Plot:** Briefly explain (1-2 sentences) how the setting reflects or influences the main character and potential plot events.
    Novel Subject: {subject}
    Genre(s): {genre}
    Main Character Profile: {profile}
    Setting Description:"""
    def __init__(self):
        self.llm = create_llm(temperature=0.7, top_p=0.9)
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)

    def run(self, subject, genre, profile):
        try:
            logger.info("Invoking SettingChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({"subject": subject, "genre": genre, "profile": profile})
            setting = result.get('text', "").strip()
            if not setting or len(setting) < 50:
                logger.warning(f"Generated setting seems invalid or too short: '{setting[:100]}...'")
                raise SettingGenerationError(f"Failed to generate a meaningful setting. LLM Output: {setting[:100]}...")
            return setting
        except Exception as e:
            logger.error(f"An error occurred in SettingChain.run: {e}")
            if not isinstance(e, SettingGenerationError):
                raise SettingGenerationError(f"Setting generation failed: {e}") from e
            raise

class ThemeChain:
    PROMPT = """
    You are a literary analyst identifying core themes.
    Based on the novel's subject, genre, character profile, and setting, identify 2-4 central themes that the story explores.
    Themes should be abstract concepts (e.g., "Loss and Memory," "Identity vs. Society," "The Nature of Reality," "Redemption," "Man vs. Nature").
    Provide a brief (1-sentence) explanation for each theme, linking it to the provided context. **Keep explanations concise and clear.**
    Novel Subject: {subject}
    Genre(s): {genre}
    Main Character Profile: {profile}
    Setting Description: {setting}
    Identified Themes:
    1.  [Theme 1]: [1-sentence explanation]
    2.  [Theme 2]: [1-sentence explanation]
    3.  [Theme 3 (Optional)]:[1-sentence explanation]
    4.  [Theme 4 (Optional)]:[1-sentence explanation]"""
    def __init__(self):
        self.llm = create_llm(temperature=0.5, top_p=0.8)
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)

    def parse_themes(self, response_text):
        themes = {}
        # Regex to capture "X. [Theme Name]: [Explanation]"
        pattern = re.compile(r"^\s*\d+\.\s*\[?([^:\]]+)\]?:\s*(.*)", re.MULTILINE)
        matches = pattern.findall(response_text)
        if not matches:
             # Fallback: Try splitting by lines if regex fails
             lines = [line.strip() for line in response_text.strip().split('\n') if ':' in line]
             for line in lines:
                 parts = line.split(':', 1)
                 theme_name = re.sub(r"^\s*\d+\.\s*", "", parts[0]).strip() # Remove leading numbers and dots
                 theme_name = theme_name.replace('[', '').replace(']', '') # Remove brackets if any
                 explanation = parts[1].strip()
                 if theme_name and explanation:
                     themes[theme_name] = explanation
        else:
             for match in matches:
                 theme_name = match[0].strip().replace('[', '').replace(']', '') # Remove brackets
                 explanation = match[1].strip()
                 if theme_name and explanation:
                     themes[theme_name] = explanation
        if not themes:
            logger.warning(f"Could not parse themes from response:\n---\n{response_text}\n---")
            # MOD: Raising a specific parsing error.
            raise ParsingError(f"Theme parsing failed. Raw response: {response_text[:200]}...")
        return themes

    def run(self, subject, genre, profile, setting):
        try:
            logger.info("Invoking ThemeChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({"subject": subject, "genre": genre, "profile": profile, "setting": setting})
            raw_themes_text = result.get('text', "").strip()
            if not raw_themes_text:
                raise ThemeGenerationError("LLM returned empty response for themes.")
            return self.parse_themes(raw_themes_text)
        except ParsingError: # Let parsing errors propagate
            raise
        except Exception as e:
            logger.error(f"An error occurred in ThemeChain.run: {e}")
            if not isinstance(e, ThemeGenerationError):
                raise ThemeGenerationError(f"Theme generation failed: {e}") from e
            raise

class TitleChain:
    PROMPT = """
    You are a creative book title generator known for evocative and genre-appropriate titles.
    Generate ONE compelling novel title based on the provided details.
    The title must be highly consistent with the genre(s), author's style, subject, main character, setting, and core themes.
    It should be intriguing and memorable. **Aim for clarity and impact.**
    Subject: {subject}
    Genre(s): {genre}
    Author's Style Inspiration: {author}
    Main Character Profile: {profile}
    Setting Description: {setting}
    Core Themes: {themes}
    Return ONLY the generated title itself, without any quotation marks, labels (like "Title:"), or explanatory text.
    Novel Title:"""
    def __init__(self):
        self.llm = create_llm(temperature=0.85, top_p=0.95) # Higher temp for creativity
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)

    def run(self, subject, genre, author, profile, setting, themes_str):
        try:
            logger.info("Invoking TitleChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({"subject": subject, "genre": genre, "author": author, "profile": profile, "setting": setting, "themes": themes_str })
            title = result.get('text', "Untitled Novel").strip()
            title = re.sub(r'^(Title:|Novel Title:)\s*', '', title, flags=re.IGNORECASE)
            title = title.strip('"\'')

            if not title or len(title) < 3 or title == "Untitled Novel":
                 logger.warning(f"Generated title seems invalid: '{title}'. Using placeholder.")
                 genre_tag = genre.split('/')[0].split(',')[0].strip()
                 # Not raising an error here, as a placeholder might be acceptable
                 return f"Placeholder - {genre_tag} Story about {profile.split('.')[0]}"
            return title
        except Exception as e: # Catch any unexpected error
            logger.error(f"An error occurred in TitleChain.run: {e}")
            # Raise a specific error, but title failure might not be critical to halt all
            raise TitleGenerationError(f"Title generation failed: {e}") from e


class PlotChain:
    PROMPT = """
    You are a master storyteller and plot architect, crafting narratives in the style of {author}.
    Develop a detailed, multi-act plot outline (e.g., Act I, Act II, Act III or Beginning, Middle, End) for a novel.
    The plot must be engaging, coherent, thematically resonant, and build towards a satisfying climax and resolution appropriate for the {genre} genre.
    Integrate the following elements seamlessly:
    - Main Character Arc: Show how the character (profile below) changes or is challenged, driven by their motivations/flaws.
    - Core Themes: Weave the themes ({themes}) into the events and character journey.
    - Setting Influence: Show how the setting ({setting}) impacts the mood, obstacles, or events.
    - Compelling Attributes: Incorporate story elements like: {features}.
    - Supporting Characters: Introduce necessary allies, antagonists, or foils.
    - Conflict: Include both internal (character-based) and external (plot-based) conflicts.
    - Turning Points: Define key moments that shift the narrative direction.
    **Ensure the plot points are clear and logically connected.**
    Novel Subject: {subject}
    Genre(s): {genre}
    Novel Title: "{title}"
    Main Character Profile: {profile}
    Setting Description: {setting}
    Core Themes: {themes}
    Detailed Plot Outline:"""
    HELPER_PROMPT = """
    Generate a comma-separated list of 5-7 diverse and compelling story attributes or narrative devices suitable for the specified genre(s) and author style. Avoid generic terms. Be specific and evocative.
    **Focus on attributes that lead to engaging and understandable plot developments.**
    Genre(s): {genre}
    Author's Style Inspiration: {author}
    List of Attributes (comma-separated):"""
    def __init__(self):
        self.llm = create_llm(temperature=0.75, top_p=0.9)
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)
        self.helper_chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.HELPER_PROMPT), verbose=False)

    def run(self, subject, genre, author, profile, title, setting, themes_str):
        try:
            logger.info(f"Generating plot features for genre '{genre}' and author style '{author}'...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            features_result = self.helper_chain.invoke({"genre": genre, "author": author})
            features = features_result.get('text', "Compelling conflict, Character depth, Unexpected twists").strip()
            logger.info(f"Generated plot features: {features}")

            logger.info(f"Generating main plot outline for title: {title}")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            plot_result = self.chain.invoke({
                "features": features, "subject": subject, "genre": genre, "author": author,
                "profile": profile, "title": title, "setting": setting, "themes": themes_str
            })
            plot = plot_result.get('text', "").strip()

            if not plot or len(plot) < 200:
                 logger.warning(f"Generated plot seems invalid or too short: '{plot[:200]}...'")
                 raise PlotGenerationError(f"Failed to generate a detailed plot. LLM Output: {plot[:200]}...")
            return plot
        except Exception as e:
            logger.error(f"An error occurred in PlotChain.run: {e}")
            if not isinstance(e, PlotGenerationError):
                raise PlotGenerationError(f"Plot generation failed: {e}") from e
            raise

class ChaptersChain:
    PROMPT = """
    You are a meticulous book editor outlining chapter structure in the style of {author}.
    Based on the detailed plot outline, generate a list of chapter titles AND brief, one-sentence descriptions capturing the core focus or turning point of each chapter.
    Aim for a realistic number of chapters (e.g., 15-35) that logically progress through the plot.
    Include a Prologue and/or Epilogue ONLY if appropriate for the story structure, genre, and author's style.
    Ensure strict consistency with: Novel Title: "{title}", Genre(s): {genre}, Author's Style: {author}, Main Character Arc: {profile}, Setting: {setting}, Core Themes: {themes}, Detailed Plot Outline (provided below)
    **Chapter descriptions should be concise, clear, and hint at concrete actions or revelations.**
    Use this EXACT format for each entry, with NO blank lines between entries:
    [Chapter Type] [Number (if not Prologue/Epilogue)]: [Concise, evocative one-sentence Description]
    Example Format:
    Prologue: Whispers of the past in the decaying manor.
    Chapter 1: A mundane routine shattered by a cryptic message.
    Detailed Plot Outline:
    <PLOT_START>
    {plot}
    <PLOT_END>
    Chapters List (Strict Format Adherence Required):"""
    def __init__(self):
        self.llm = create_llm(temperature=0.65, top_p=0.9, top_k=50)
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)

    def parse_chapters(self, response_text):
        chapters = {}
        lines = [line.strip() for line in response_text.strip().split('\n') if line.strip()]
        parsed_count = 0
        # Regex to capture "Prologue/Epilogue/Chapter [Num]: Description"
        # It's more flexible with spacing around the number and colon.
        pattern = re.compile(r"^(Prologue|Epilogue|Chapter\s*(\d+))\s*:\s*(.+)$", re.IGNORECASE | re.MULTILINE)
        processed_lines = set() # To avoid double-parsing with fallback

        matches = pattern.findall(response_text)
        for match in matches:
            # match[0] is the full "Chapter X" or "Prologue"
            # match[1] is the number if it's "Chapter X", otherwise empty
            # match[2] is the description
            full_title_prefix = match[0].strip() # e.g., "Chapter  1", "Prologue"
            chapter_number_str = match[1] # e.g., "1" or ""
            description = match[2].strip()

            # Normalize title
            if "prologue" in full_title_prefix.lower():
                full_title = "Prologue"
            elif "epilogue" in full_title_prefix.lower():
                full_title = "Epilogue"
            elif chapter_number_str:
                full_title = f"Chapter {chapter_number_str.strip()}"
            else: # Should not happen if regex is correct, but as a safeguard
                full_title = full_title_prefix

            if description:
                chapters[full_title] = description
                parsed_count += 1
                # Mark line as processed (approximate by checking if description is in the original line)
                # This is a bit heuristic; a more robust way would be to track line indices if exact match is needed.
                for i, line_content in enumerate(lines):
                    if description in line_content and full_title_prefix in line_content:
                        processed_lines.add(i)
                        break
            else:
                logger.warning(f"Skipping chapter match due to empty description: '{full_title_prefix}'")

        # Fallback parsing for lines not caught by precise regex
        for i, line in enumerate(lines):
            if i in processed_lines: continue # Skip if already processed by regex
            if ':' in line:
                parts = line.split(':', 1)
                potential_title = parts[0].strip()
                potential_desc = parts[1].strip()
                # Basic check to see if it looks like a chapter title
                is_pro_epi_chap = any(kw in potential_title.lower() for kw in ["prologue", "epilogue", "chapter"])

                if potential_title not in chapters and is_pro_epi_chap and potential_desc:
                    # Further normalize potential_title from fallback
                    if "prologue" in potential_title.lower():
                        normalized_title = "Prologue"
                    elif "epilogue" in potential_title.lower():
                        normalized_title = "Epilogue"
                    else:
                        num_match = re.search(r"chapter\s*(\d+)", potential_title, re.IGNORECASE)
                        if num_match:
                            normalized_title = f"Chapter {num_match.group(1).strip()}"
                        else:
                            normalized_title = potential_title # Use as is if no number found

                    if normalized_title not in chapters: # Check again after normalization
                        logger.debug(f"Fallback Parsed Chapter Line {i}: '{potential_title}' (Normalized: '{normalized_title}') -> '{potential_desc}'")
                        chapters[normalized_title] = potential_desc
                        parsed_count += 1

        if not chapters:
             logger.error(f"CRITICAL WARNING: Could not parse *any* chapters from response:\n---\n{response_text}\n---")
             raise ParsingError(f"Chapter parsing failed completely. Raw response: {response_text[:200]}...")
        logger.info(f"Successfully parsed {parsed_count} chapters.")

        # --- THE FIX IS HERE ---
        # sort_chapters expects an iterable of (key, value) pairs, which chapters.items() provides.
        # It returns a sorted list of (key, value) tuples, which dict() can then correctly process.
        return dict(sort_chapters(chapters.items()))

    def run(self, subject, genre, author, profile, title, plot, setting, themes_str):
        try:
            logger.info("Invoking ChaptersChain...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            response_result = self.chain.invoke({
                "subject": subject, "genre": genre, "author": author, "profile": profile,
                "title": title, "plot": plot, "setting": setting, "themes": themes_str
            })
            raw_response = response_result.get('text', "")
            if not raw_response.strip():
                raise ChapterGenerationError("LLM returned empty response for chapters.")
            return self.parse_chapters(raw_response)
        except ParsingError: # Propagate parsing errors
            raise
        except Exception as e:
            logger.error(f"An error occurred in ChaptersChain.run: {e}")
            if not isinstance(e, (ChapterGenerationError, ParsingError)): # Don't re-wrap if already specific
                raise ChapterGenerationError(f"Chapter list generation failed: {e}") from e
            raise # Re-raise original or wrapped error


class EventChain:
    PROMPT = """
    You are a narrative strategist breaking down a chapter into key scenes or events, following the style of {author}.
    Based on the overall plot, the chapter's details, character profile, and themes, generate an ordered list of 3-7 key plot events or scenes that MUST occur within this chapter.
    Focus on actions, decisions, revelations, significant interactions, or internal character moments that advance the plot AND reflect the character's journey or the novel's themes.
    Ensure events flow logically from the chapter summary and overall plot.
    **Event descriptions should be concrete and focus on what happens or is revealed, rather than abstract feelings alone.**
    Overall Plot Summary: <PLOT_SUMMARY>{plot}</PLOT_SUMMARY>
    Main Character Profile: <CHARACTER>{profile}</CHARACTER>
    Core Themes: {themes}
    Current Chapter Title: "{chapter_title}"
    Current Chapter Summary: {chapter_summary}
    Return ONLY the numbered list of events, one event per line. Start numbering from 1. Be concise but evocative.
    Numbered Event List for Chapter "{chapter_title}": """
    def __init__(self):
        self.llm = create_llm(temperature=0.7, top_p=0.9)
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)

    def parse_events(self, response_text, chapter_title): # MOD: Added chapter_title for logging
        events = []
        # Regex to find lines starting with a number, followed by a dot or parenthesis, then whitespace, then the event text.
        pattern = re.compile(r"^\s*\d+[\.\)\s]+(.*)", re.MULTILINE)
        matches = pattern.findall(response_text)

        if matches:
            for event_text in matches:
                if event_text.strip(): # Ensure the captured group is not just whitespace
                    events.append(event_text.strip())
        else: # Fallback: if no numbered list is found, try to split by lines and clean them
            lines = [line.strip() for line in response_text.strip().split('\n') if line.strip()]
            for line in lines:
                 # Attempt to remove any leading "numbering-like" characters if the LLM didn't follow format
                 cleaned_line = re.sub(r"^\s*\d*[\.\)\s-]*", "", line).strip()
                 if cleaned_line: # Only add if there's content after cleaning
                     events.append(cleaned_line)

        if not events:
             logger.warning(f"Could not parse any numbered events for chapter '{chapter_title}' from response:\n---\n{response_text}\n---")
             # Not raising ParsingError here, as an empty event list might be handled by write_book
             # Or could return a placeholder: return ["Placeholder event due to parsing failure."]
        return events

    def run(self, plot, profile, themes_str, chapter_title, chapter_summary, author):
        try:
            logger.info(f"Invoking EventChain for: {chapter_title}")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "plot": plot, "profile": profile, "themes": themes_str,
                "chapter_title": chapter_title, "chapter_summary": chapter_summary, "author": author
            })
            raw_events = result.get('text', "").strip()
            if not raw_events: # If LLM returns nothing
                logger.warning(f"EventChain returned empty response for '{chapter_title}'.")
                return [] # Return empty list, not an error for now
            return self.parse_events(raw_events, chapter_title)
        except Exception as e:
            logger.error(f"Error generating events for '{chapter_title}': {e}")
            # This error might not be critical enough to stop the whole book
            raise EventGenerationError(f"Event generation for '{chapter_title}' failed: {e}") from e


class WriterChain:
    PROMPT = """
    You are a master novelist embodying the distinct writing style of {author}. Your task is to write a segment of the novel "{title}", a {genre} work.
    Your primary goal is to make the scene **clear, engaging, and relatable** to the reader, even while reflecting {author}'s style.
    **ESTABLISHED CONTEXT (DO NOT REPEAT):**
    <MAIN_CHARACTER_PROFILE>{profile}</MAIN_CHARACTER_PROFILE>
    <SETTING_DESCRIPTION>{setting}</SETTING_DESCRIPTION>
    <CORE_THEMES>{themes}</CORE_THEMES>
    <NOVEL_PLOT_SUMMARY>{plot}</NOVEL_PLOT_SUMMARY>
    <CURRENT_CHAPTER_INFO>Chapter Title: {chapter_name}, Chapter Summary: {summary}</CURRENT_CHAPTER_INFO>
    <NARRATIVE_HISTORY>Key events already occurred (prior chapters/events):{previous_events}</NARRATIVE_HISTORY>
    <CHAPTER_PROGRESS_SO_FAR>Narrative written previously within THIS chapter ({chapter_name}):{previous_paragraphs}</CHAPTER_PROGRESS_SO_FAR>
    **YOUR CURRENT WRITING TASK:**
    Bring the following specific event to life. Focus *exclusively* on this moment: <CURRENT_EVENT_TO_WRITE>{current_event}</CURRENT_EVENT_TO_WRITE>
    **WRITING INSTRUCTIONS (CRITICAL):**
    1.  **Style & Tone:** Write 1-3 detailed paragraphs (approx. 150-400 words total for this event, adjust based on significance) in the voice and style of {author}. Match their typical sentence structure, vocabulary, pacing, and tone, **BUT prioritize clarity and readability. Avoid overly complex sentences or obscure vocabulary unless it's truly essential for character voice or a specific, impactful stylistic effect. If a simpler way to say something exists without losing the core style, prefer it.**
    2.  **Show, Don't Tell (Concretely):** Demonstrate emotions, character thoughts, and plot developments through **specific actions, concrete sensory details, relatable internal monologues, and clear dialogue (if appropriate)**, rather than stating them abstractly. For example, instead of "he felt profound desolation," describe "a hollowness spread through his chest, and he stared blankly at his hands, noticing for the first time how they trembled."
    3.  **Sensory Integration (Make it Tangible):** Weave in specific sensory details (sight, sound, smell, touch, taste) consistent with the established <SETTING_DESCRIPTION>. **Focus on details that are easy for a reader to imagine and connect with. Ground abstract sensations with concrete comparisons where possible.** For instance, if describing a confusing mental state, compare it to something familiar like "his thoughts felt like a radio dial spinning between stations, never quite landing on a clear signal," or "it was like trying to see through a fogged-up window."
    4.  **Character Depth (Accessible Emotions):** Reflect the <MAIN_CHARACTER_PROFILE> (motivations, flaws, internal conflicts) in their reactions, thoughts, and actions during this event. **Ensure the character's internal experience is conveyed in a way that allows the reader to empathize or understand their state of mind easily.**
    5.  **Thematic Resonance (Subtly):** Subtly connect the events or the character's internal experience to the <CORE_THEMES> where appropriate, without being overt or resorting to dense, philosophical language.
    6.  **Continuity:** Ensure a seamless transition from the <CHAPTER_PROGRESS_SO_FAR>. Do NOT repeat information already covered. Do NOT summarize the event itself.
    7.  **Output:** Generate ONLY the newly written narrative paragraphs for the <CURRENT_EVENT_TO_WRITE>. No explanations, labels, or summaries.
    8.  **Relatability Focus:** When describing complex or fantastical events, try to include details or reactions that a reader can connect to from ordinary life. What would it *feel* like physically? What common emotion would it evoke? For example, sudden disorientation could be "like standing up too fast, the world tilting on its axis," or a strange magic could feel "like a faint electrical current humming under the skin."
    New Narrative Paragraphs (Style: {author}, Prioritizing Clarity & Relatability):"""
    def __init__(self):
        self.llm = create_llm(temperature=0.7, top_p=0.85, top_k=40) # Values from user's last request
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)

    def run(self, genre, author, title, profile, plot, setting, themes_str, chapter_name,
            previous_events_history_str, chapter_summary, previous_paragraphs_str, current_event):
        if not previous_events_history_str: previous_events_history_str = "None (This is the beginning of the story)."
        if not previous_paragraphs_str: previous_paragraphs_str = "None (This is the beginning of the chapter)."
        try:
            logger.info(f"WriterChain: Event: {current_event[:80]}... in Ch: {chapter_name}")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "genre": genre, "author": author, "title": title, "profile": profile, "plot": plot,
                "setting": setting, "themes": themes_str, "chapter_name": chapter_name,
                "previous_events": previous_events_history_str, "summary": chapter_summary,
                "previous_paragraphs": previous_paragraphs_str, "current_event": current_event
            })
            generated_text = result.get('text', "").strip()
            if not generated_text or len(generated_text) < 50:
                 logger.warning(f"WriterChain generated short/empty output for event '{current_event[:50]}...'. Output: '{generated_text[:100]}'")
                 # Return a placeholder rather than raising an error that stops the whole chapter
                 return f"[Writer Error - Event: '{current_event[:50]}...'. LLM Output was too short or empty.]"
            return generated_text
        except Exception as e:
            logger.error(f"Error in WriterChain.run for event '{current_event[:50]}...': {e}")
            # Again, return placeholder to allow chapter to complete if possible
            raise WriterError(f"WriterChain failed for event '{current_event[:50]}...': {e}") from e


class RefinementChain:
    PROMPT = """
    You are an expert literary editor refining a draft chapter written in the style of {author}.
    The chapter is part of the novel "{title}" ({genre}).
    **CONTEXT:** <MAIN_CHARACTER_PROFILE>{profile}</MAIN_CHARACTER_PROFILE> <SETTING_DESCRIPTION>{setting}</SETTING_DESCRIPTION> <CORE_THEMES>{themes}</CORE_THEMES> <NOVEL_PLOT_SUMMARY>{plot}</NOVEL_PLOT_SUMMARY> <CHAPTER_INFO>Title: {chapter_name}, Summary: {summary}</CHAPTER_INFO>
    **DRAFT TEXT FOR REFINEMENT:** <DRAFT_START>{draft_text}</DRAFT_END>
    **YOUR TASK:** Review and refine the provided draft text. Focus on ONE OR TWO of the following aspects:
    1.  **Stylistic Consistency:** Enhance adherence to {author}'s unique voice, **while maintaining clarity.**
    2.  **Prose Flow & Pacing:** Improve readability, sentence variety, and transitions.
    3.  **Sensory Detail Enhancement:** Weave in more vivid and **relatable** sensory details.
    4.  **Dialogue Polish:** Make dialogue more natural and character-specific.
    5.  **Thematic Resonance:** Subtly strengthen connections to themes ({themes}) **without becoming overly abstract.**
    6.  **Show, Don't Tell:** Convert telling into **concrete showing**.
    7.  **Conciseness:** Trim unnecessary words.
    8.  **Enhance Clarity & Relatability:** If dense or complex, simplify sentences and ground abstract descriptions.
    **INSTRUCTIONS:** Make subtle but impactful changes. Preserve original plot points. Output ONLY the refined text.
    Refined Chapter Text (Style: {author}):"""
    def __init__(self):
        self.llm = create_llm(temperature=0.6, top_p=0.95, top_k=50)
        self.chain = LLMChain(llm=self.llm, prompt=PromptTemplate.from_template(self.PROMPT), verbose=False)

    def run(self, author, title, genre, profile, setting, themes_str, plot, chapter_name, summary, draft_text):
        if not draft_text or draft_text.startswith("[Content generation skipped") or "[Writer Error" in draft_text or "[FATAL WRITER ERROR" in draft_text :
            logger.info(f"Skipping refinement for '{chapter_name}' due to missing or errored draft content.")
            return draft_text
        try:
            logger.info(f"Invoking RefinementChain for chapter: {chapter_name}...")
            time.sleep(LLM_CALL_DELAY_SECONDS)
            result = self.chain.invoke({
                "author": author, "title": title, "genre": genre, "profile": profile, "setting": setting,
                "themes": themes_str, "plot": plot, "chapter_name": chapter_name,
                "summary": summary, "draft_text": draft_text
            })
            refined_text = result.get('text', "").strip()
            if not refined_text or len(refined_text) < len(draft_text) * 0.7: # Allow some shrinkage
                 logger.warning(f"RefinementChain generated potentially invalid output for '{chapter_name}'. Original length: {len(draft_text)}, Refined length: {len(refined_text)}. Output: '{refined_text[:100]}...' Using original.")
                 return draft_text # Fallback to original
            return refined_text
        except Exception as e:
            logger.error(f"Error in RefinementChain.run for chapter '{chapter_name}': {e}")
            raise RefinementError(f"Refinement for '{chapter_name}' failed: {e}") from e


# --- Book Writing Orchestration ---

def format_themes_string(themes_dict):
    if not themes_dict or isinstance(themes_dict.get("Error"), str) : # Check if "Error" key exists and its value is a string
        return "N/A (Themes not available or errored)"
    return "; ".join([f"{name}: {desc}" for name, desc in themes_dict.items()])

def sort_chapters(chapter_dict_items):
    """Sorts chapter items (list of (title, summary) tuples)."""
    def chapter_sort_key(item): # item is a (title, summary) tuple
        chapter_title = item[0] # Get the title string for sorting
        title_lower = chapter_title.lower().strip()
        if title_lower == "prologue": return (-1, 0)
        if title_lower == "epilogue": return (9999, 0) # Large number to ensure epilogue is last
        match = re.match(r"chapter\s*(\d+)", title_lower)
        if match:
            try: return (1, int(match.group(1))) # Sort by chapter number
            except (IndexError, ValueError): return (2, chapter_title) # Fallback for malformed chapter numbers
        return (3, chapter_title) # Fallback for non-standard chapter titles
    return sorted(chapter_dict_items, key=chapter_sort_key)


def generate_events_for_all_chapters(plot, profile, themes_str, sorted_chapters_list_of_tuples, author):
    logger.info("Generating Events for Each Chapter")
    event_generator = EventChain()
    event_dict = {} # chapter_title -> list_of_events
    for chapter_title, summary in sorted_chapters_list_of_tuples:
         logger.info(f"Generating events for: {chapter_title}")
         try:
            events = event_generator.run(plot, profile, themes_str, chapter_title, summary, author)
            # events is already a list from event_generator.run (or [] on error/empty)
            event_dict[chapter_title] = events
            if events:
                logger.info(f"Generated {len(events)} events for '{chapter_title}'.")
            else:
                logger.warning(f"No events generated or parsed for '{chapter_title}'.")
         except EventGenerationError as e:
            logger.warning(f"Could not generate events for '{chapter_title}': {e}. Skipping events for this chapter.")
            event_dict[chapter_title] = [] # Assign empty list on error
    return event_dict

def _write_single_chapter_content(writer_chain, chapter_title, chapter_summary, chapter_events, book_context):
    """Helper function to write content for a single chapter."""
    logger.info(f"Writing Chapter: {chapter_title} (Summary: {chapter_summary[:50]}...)")
    if not chapter_events:
        logger.warning(f"No events found for '{chapter_title}'. Skipping content generation for this chapter.")
        return "[Content generation skipped: No events defined for this chapter.]"

    chapter_paragraphs_accumulator = ""
    total_events = len(chapter_events)

    for event_idx, event_description in enumerate(chapter_events):
        MAX_HISTORY_EVENTS = 20 # Reduced history for token limits
        limited_history = book_context['previous_events_history'][-MAX_HISTORY_EVENTS:]
        history_str = '\n'.join(f"- {evt}" for evt in limited_history) if limited_history else "None (Start of Story)"
        progress_str = chapter_paragraphs_accumulator if chapter_paragraphs_accumulator else "None (Start of Chapter)"

        logger.info(f"  Writing Event {event_idx+1}/{total_events}: {event_description[:70]}...")

        try:
            new_paragraphs = writer_chain.run(
                genre=book_context['genre'], author=book_context['author_style'], title=book_context['title'],
                profile=book_context['profile'], plot=book_context['plot'], setting=book_context['setting'],
                themes_str=book_context['themes_str'], chapter_name=chapter_title,
                previous_events_history_str=history_str, chapter_summary=chapter_summary,
                previous_paragraphs_str=progress_str, current_event=event_description
            )
        except WriterError as e: # Catch specific WriterError
            logger.error(f"WriterError encountered for event {event_idx+1} in '{chapter_title}': {e}. Adding error message to content.")
            new_paragraphs = f"[WRITER ERROR for event: '{event_description[:50]}...'. Details: {e}]"
        except Exception as e: # Catch any other unexpected error from WriterChain
            logger.error(f"Unexpected error writing event {event_idx+1} for '{chapter_title}': {e}. Adding error message.")
            logger.exception("WriterChain Unexpected Error Details:")
            new_paragraphs = f"[UNEXPECTED WRITER ERROR for event: '{event_description[:50]}...'. Check logs.]"


        if "[FATAL WRITER ERROR" in new_paragraphs or "[Writer Error" in new_paragraphs or "[WRITER ERROR" in new_paragraphs or "[UNEXPECTED WRITER ERROR" in new_paragraphs:
            logger.error(f"Error detected writing event {event_idx+1} for '{chapter_title}'. Error message added to content.")
        
        if chapter_paragraphs_accumulator and not chapter_paragraphs_accumulator.endswith(('\n\n', '\n')):
            chapter_paragraphs_accumulator += "\n\n"
        elif chapter_paragraphs_accumulator and not chapter_paragraphs_accumulator.endswith('\n\n'):
            chapter_paragraphs_accumulator += "\n" # Ensure at least one newline separation

        chapter_paragraphs_accumulator += new_paragraphs
        # Only add successful event descriptions to history, or a placeholder if it failed
        if "ERROR for event" not in new_paragraphs:
            book_context['previous_events_history'].append(f"[{chapter_title}] {event_description}")
        else:
            book_context['previous_events_history'].append(f"[{chapter_title}] (Error writing event: {event_description[:30]}...)")

    return chapter_paragraphs_accumulator

def write_book(genre, author_style, title, profile, plot, setting, themes_str,
               sorted_chapters_list_of_tuples, event_dict, refine_chapters=False):
    logger.info("Starting Detailed Book Writing Process")
    writer_chain = WriterChain()
    refiner_chain = RefinementChain() if refine_chapters else None

    book_content_map = {} # Stores final text: {chapter_title: "Full chapter text..."}
    
    book_writing_context = {
        'genre': genre, 'author_style': author_style, 'title': title,
        'profile': profile, 'plot': plot, 'setting': setting, 'themes_str': themes_str,
        'previous_events_history': [] # Running list of event descriptions written so far
    }
    
    total_chapters = len(sorted_chapters_list_of_tuples)

    for chap_idx, (chapter_title, chapter_summary) in enumerate(sorted_chapters_list_of_tuples):
        logger.info(f"--- Processing Chapter {chap_idx+1}/{total_chapters}: {chapter_title} ---")
        
        chapter_events = event_dict.get(chapter_title, [])
        raw_chapter_text = _write_single_chapter_content(
            writer_chain, chapter_title, chapter_summary, chapter_events, book_writing_context
        )

        final_chapter_text = raw_chapter_text
        if refiner_chain:
            logger.info(f"Refining chapter: {chapter_title}...")
            try:
                refined_text = refiner_chain.run(
                    author=author_style, title=title, genre=genre, profile=profile, setting=setting,
                    themes_str=themes_str, plot=plot, chapter_name=chapter_title,
                    summary=chapter_summary, draft_text=raw_chapter_text
                )
                # Check if refinement actually changed something and didn't just return an error placeholder
                if refined_text != raw_chapter_text and not refined_text.startswith(("[Writer Error", "[FATAL WRITER ERROR", "[Content generation skipped")):
                     logger.info(f"Refinement applied for '{chapter_title}'.")
                     final_chapter_text = refined_text
                elif refined_text.startswith(("[Writer Error", "[FATAL WRITER ERROR", "[Content generation skipped")):
                    logger.warning(f"Refinement resulted in an error or skipped content for '{chapter_title}', using unrefined text.")
                else: # Refinement didn't change or error, so log it
                     logger.info(f"Refinement resulted in no significant changes for '{chapter_title}' or returned original due to issues.")
            except RefinementError as e:
                logger.warning(f"Refinement process error for '{chapter_title}', using unrefined text. Error: {e}")
            except Exception as e: # Catch-all for unexpected refinement errors
                logger.error(f"Unexpected error during refinement of '{chapter_title}': {e}. Using unrefined text.")
                logger.exception("Refinement Unexpected Error Details:")

        book_content_map[chapter_title] = final_chapter_text
        logger.info(f"--- Finished Chapter: {chapter_title} ---")

    logger.info("Book Writing Process Complete")
    return book_content_map


class DocWriter:
    def __init__(self, output_folder=OUTPUT_FOLDER):
        self.output_folder = output_folder
        os.makedirs(self.output_folder, exist_ok=True)
        logger.info(f"Ensured output directory exists: {os.path.abspath(self.output_folder)}")

    def _sanitize_filename(self, name):
        if not isinstance(name, str): name = str(name)
        name = re.sub(r'[\\/*?:"<>|]', "", name) # Remove illegal characters
        name = re.sub(r'\s+', '_', name)         # Replace whitespace with underscores
        name = re.sub(r'[_.-]{2,}', '_', name)   # Replace multiple separators with one
        name = name.strip('._- ')               # Strip leading/trailing separators
        max_len = 180 # Max filename length (conservative)
        if len(name) > max_len:
            # Truncate in the middle to preserve beginning and end
            name = name[:max_len//2 - 3] + "..." + name[-(max_len//2):]
        return name if name else "Untitled_Novel"

    def write_doc(self, book_content_map, sorted_chapters_list_of_tuples, title, genre, author, themes_dict, setting_desc):
        logger.info("Assembling and Writing Document")
        doc = docx.Document()
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Genre: {genre}")
        doc.add_paragraph(f"Inspired by the style of: {author}")
        doc.add_paragraph("\n")

        doc.add_heading("Core Themes", level=2)
        current_themes_str = format_themes_string(themes_dict)
        if "N/A" in current_themes_str or (isinstance(themes_dict, dict) and isinstance(themes_dict.get("Error"), str)):
             doc.add_paragraph(current_themes_str) # Will show N/A or error message
        elif isinstance(themes_dict, dict):
            for theme_name, theme_desc in themes_dict.items():
                # Use a run for bold to allow special characters in theme_name
                p = doc.add_paragraph(style='ListBullet')
                p.add_run(f"{theme_name}:").bold = True
                p.add_run(f" {theme_desc}")
        else:
            doc.add_paragraph("Themes data is unavailable or in an unexpected format.")
        doc.add_paragraph("\n")

        doc.add_heading("Setting Summary", level=2)
        if isinstance(setting_desc, str) and setting_desc.startswith("Error:"):
             doc.add_paragraph("Setting description could not be generated or errored.")
        elif isinstance(setting_desc, str):
             for para in [p.strip() for p in setting_desc.split('\n') if p.strip()]: doc.add_paragraph(para)
        else:
            doc.add_paragraph("Setting data is unavailable or in an unexpected format.")
        doc.add_paragraph("\n")

        doc.add_heading("Chapters Overview", level=1) # Changed from "Chapters" to avoid confusion with content
        for chapter_title, description in sorted_chapters_list_of_tuples:
            p = doc.add_paragraph(style='ListBullet')
            p.add_run(f"{chapter_title}:").bold = True
            p.add_run(f" {description}")
        doc.add_page_break()

        logger.info("Adding chapter content to document...")
        total_chapters = len(sorted_chapters_list_of_tuples)
        for i, (chapter_title, _) in enumerate(sorted_chapters_list_of_tuples):
            doc.add_heading(chapter_title.strip(), level=1)
            logger.info(f"  Adding Chapter {i+1}/{total_chapters}: {chapter_title}")
            chapter_text = book_content_map.get(chapter_title, "[Error: Chapter content not found in map]")
            # Split by one or more newlines to handle different paragraph separations from LLM
            para_blocks = re.split(r'\n+', chapter_text)
            for para_block in para_blocks:
                if para_block.strip(): doc.add_paragraph(para_block.strip())
            if i < total_chapters - 1: doc.add_page_break()
            logger.info(f"    Added content for '{chapter_title}'")

        safe_basename = self._sanitize_filename(title)
        safe_author = self._sanitize_filename(author)
        safe_genre = self._sanitize_filename(genre.split('/')[0].split(',')[0]) # First part of genre
        filename = f"{safe_basename}_by_{safe_author}_{safe_genre}.docx"
        output_path = os.path.join(self.output_folder, filename)

        try:
            logger.info(f"Attempting to save document to: {output_path}")
            doc.save(output_path)
            logger.info("Document saved successfully!")
            return output_path
        except PermissionError:
            logger.error(f"Permission denied trying to save '{output_path}'. Check permissions or if file is open.")
            raise WriterError(f"Permission denied saving document: {output_path}") from None # Reraise as WriterError
        except Exception as e:
            logger.error(f"Error saving document to {output_path}: {e}")
            logger.exception("Doc Saving Traceback:")
            raise WriterError(f"Failed to save document: {e}") from e # Reraise as WriterError
        # return None # This line is unreachable if exceptions are raised

# --- Main Execution ---

def initialize_components():
    """Initializes all chain and utility components."""
    logger.info("Initializing Core Components...")
    components = {
        "main_character_chain": MainCharacterChain(),
        "setting_chain": SettingChain(),
        "theme_chain": ThemeChain(),
        "title_chain": TitleChain(),
        "plot_chain": PlotChain(),
        "chapters_chain": ChaptersChain(),
        # EventChain, WriterChain, RefinementChain are instantiated dynamically or within write_book
        "doc_writer": DocWriter(output_folder=OUTPUT_FOLDER)
    }
    logger.info("Core Components Initialized Successfully.")
    return components

def run_generation_pipeline(components, resume_filename, subject, author_style, genre, enable_refinement):
    """Runs the main novel generation pipeline."""
    data_store = {} # To hold generated artifacts like profile, plot, etc.

    try:
        logger.info("--- Step 1: Generating Main Character Profile ---")
        start_time = time.time()
        data_store['profile'] = components["main_character_chain"].run(resume_filename, genre)
        logger.info(f"Profile Generation Time: {time.time() - start_time:.2f}s")
        logger.info(f"Profile Generated (first 300 chars):\n---\n{data_store['profile'][:300]}...\n---")
    except ProfileGenerationError as e:
        logger.error(f"CRITICAL FAILURE in Profile Generation: {e}")
        raise # This is critical, so re-raise to stop
    except FileNotFoundError as e: # Catch FileNotFoundError from load_resume specifically
        logger.error(f"CRITICAL FILE ERROR for resume: {e}")
        raise CriticalGenerationError(f"Resume file not found: {e}") from e


    try:
        logger.info("--- Step 2: Generating Setting Description ---")
        start_time = time.time()
        data_store['setting'] = components["setting_chain"].run(subject, genre, data_store['profile'])
        logger.info(f"Setting Generation Time: {time.time() - start_time:.2f}s")
        logger.info(f"Setting Generated (first 300 chars):\n---\n{data_store['setting'][:300]}...\n---")
    except SettingGenerationError as e:
        logger.warning(f"Setting Generation Failed: {e}. Proceeding with placeholder/error state.")
        data_store['setting'] = f"Error: Setting generation failed. Details: {e}"


    try:
        logger.info("--- Step 3: Generating Core Themes ---")
        start_time = time.time()
        data_store['themes_dict'] = components["theme_chain"].run(subject, genre, data_store['profile'], data_store['setting'])
        logger.info(f"Theme Generation Time: {time.time() - start_time:.2f}s")
        data_store['themes_str'] = format_themes_string(data_store['themes_dict'])
        logger.info(f"Themes Generated:\n---\n{data_store['themes_str']}\n---")
    except (ThemeGenerationError, ParsingError) as e: # Catch both generation and parsing issues
        logger.warning(f"Theme Generation/Parsing Failed: {e}. Proceeding with N/A themes.")
        data_store['themes_dict'] = {"Error": f"Theme generation/parsing failed: {e}"}
        data_store['themes_str'] = format_themes_string(data_store['themes_dict'])


    try:
        logger.info("--- Step 4: Generating Novel Title ---")
        start_time = time.time()
        data_store['title'] = components["title_chain"].run(subject, genre, author_style, data_store['profile'], data_store['setting'], data_store['themes_str'])
        logger.info(f"Title Generation Time: {time.time() - start_time:.2f}s")
        if "Placeholder" in data_store['title'] or data_store['title'].startswith("Error") or data_store['title'] == "Untitled Novel (Generation Failed)":
             logger.warning(f"TITLE GENERATION USING PLACEHOLDER/ERROR: {data_store['title']}")
        else:
             logger.info(f"Title Generated: '{data_store['title']}'")
    except TitleGenerationError as e:
        logger.warning(f"Title Generation Failed: {e}. Using default placeholder.")
        data_store['title'] = "Untitled Novel (Generation Failed)"


    try:
        logger.info("--- Step 5: Generating Detailed Plot Outline ---")
        start_time = time.time()
        data_store['plot'] = components["plot_chain"].run(subject, genre, author_style, data_store['profile'], data_store['title'], data_store['setting'], data_store['themes_str'])
        logger.info(f"Plot Generation Time: {time.time() - start_time:.2f}s")
        logger.info(f"Plot Generated (first 500 chars):\n---\n{data_store['plot'][:500]}...\n---")
    except PlotGenerationError as e:
        logger.error(f"CRITICAL FAILURE in Plot Generation: {e}")
        raise # Plot is critical

    try:
        logger.info("--- Step 6: Generating Chapter List & Summaries ---")
        start_time = time.time()
        # chapter_dict is already sorted by parse_chapters -> sort_chapters and returned as a dict
        chapter_dict = components["chapters_chain"].run(subject, genre, author_style, data_store['profile'], data_store['title'], data_store['plot'], data_store['setting'], data_store['themes_str'])
        # Convert to list of tuples for consistent processing order later, though it's already sorted if from dict(sorted_items)
        data_store['sorted_chapters_list_of_tuples'] = list(chapter_dict.items())
        logger.info(f"Chapter List Generation Time: {time.time() - start_time:.2f}s")
        logger.info(f"Chapters Generated & Parsed ({len(data_store['sorted_chapters_list_of_tuples'])} total):")
        for i, (chap_title, chap_desc) in enumerate(data_store['sorted_chapters_list_of_tuples']):
            if i < 5: logger.info(f"  - {chap_title}: {chap_desc[:60]}...") # Log first 5
        if len(data_store['sorted_chapters_list_of_tuples']) > 5: logger.info(f"  ... and {len(data_store['sorted_chapters_list_of_tuples'])-5} more chapters.")
    except (ChapterGenerationError, ParsingError) as e:
        logger.error(f"CRITICAL FAILURE in Chapter Generation/Parsing: {e}")
        raise # Chapters are critical

    logger.info("--- Step 7: Generating Events for All Chapters ---")
    start_time = time.time()
    data_store['event_dict'] = generate_events_for_all_chapters(data_store['plot'], data_store['profile'], data_store['themes_str'], data_store['sorted_chapters_list_of_tuples'], author_style)
    logger.info(f"Event Generation Time: {time.time() - start_time:.2f}s")
    if not any(data_store['event_dict'].values()): # Check if all chapters have empty event lists
        logger.warning("No events were generated for *any* chapter. Book content might be minimal or contain placeholders.")
    else:
        # Log a sample of events for the first chapter that has them
        for chap_title_events, events_list in data_store['event_dict'].items():
            if events_list:
                logger.info(f"Sample events for '{chap_title_events}' (first 3):")
                for i, event_desc_log in enumerate(events_list):
                    if i < 3: logger.info(f"    - {event_desc_log[:70]}...")
                break # Only log for the first chapter with events

    logger.info("--- Step 8: Writing Full Book Content ---")
    start_time = time.time()
    data_store['book_content_map'] = write_book(
        genre, author_style, data_store['title'], data_store['profile'], data_store['plot'], data_store['setting'], data_store['themes_str'],
        data_store['sorted_chapters_list_of_tuples'], data_store['event_dict'],
        refine_chapters=enable_refinement
    )
    logger.info(f"Book Writing Time: {time.time() - start_time:.2f}s")

    if data_store.get('book_content_map') and any(data_store['book_content_map'].values()):
        logger.info("--- Step 9: Saving Document ---")
        start_time = time.time()
        try:
            saved_path = components["doc_writer"].write_doc(
                data_store['book_content_map'], data_store['sorted_chapters_list_of_tuples'], data_store['title'],
                genre, author_style, data_store['themes_dict'], data_store['setting']
            )
            logger.info(f"Document Saving Time: {time.time() - start_time:.2f}s")
            if saved_path:
                logger.info(f"Success! Novel saved to: {saved_path}")
            else:
                # This case should ideally be covered by exceptions from write_doc
                logger.error("Failed to save the document (write_doc returned None without exception).")
        except WriterError as e: # Catch errors from DocWriter
            logger.error(f"DOCUMENT SAVING FAILED: {e}")
        except Exception as e: # Catch any other unexpected error during doc saving
            logger.error(f"UNEXPECTED ERROR DURING DOCUMENT SAVING: {e}")
            logger.exception("Doc Saving Unexpected Error Details:")
    else:
        logger.warning("Skipping Document Saving as no substantial book content was generated or book_content_map is empty.")
    
    return data_store


def main():
    process_start_time = time.time()
    logger.info("=============================================")
    logger.info("=== ENHANCED NOVEL GENERATION SYSTEM V3.1 ===") # Incremented version for fix
    logger.info(f"=== Timestamp: {time.strftime('%Y-%m-%d %H:%M:%S')} ===")
    logger.info(f"=== Using Model: {DEFAULT_MODEL} ===")
    logger.info(f"=== Ollama Base URL: {OLLAMA_BASE_URL} ===")
    logger.info("=============================================")

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    logger.info(f"Output Folder: {os.path.abspath(OUTPUT_FOLDER)}")

#__________________________________________________________________________________________
    # --- USER INPUTS ---
    # Ensure the resume PDF is in the OUTPUT_FOLDER (./docs/)
    # or update DEFAULT_RESUME_FILENAME at the top, or 'resume_filename' variable below.

    resume_filename = DEFAULT_RESUME_FILENAME # e.g., "my_resume.pdf"

    subject = ("""
Kenji Tanaka, an unassuming yet brilliant strategy gamer and hobbyist survivalist, is abruptly summoned mid-match to the vibrant, chaotic fantasy world of Aethel. Due to a cosmic fluke in the summoning ritual, he's accidentally bestowed with a unique, god-like ability: an 'Administrative Interface' to the world's fundamental laws of magic, physics, and skills. Initially bewildered, Kenji, with his gamer mindset, quickly grasps that he can manipulate reality, learn and master skills instantaneously, and essentially 'debug' or 'mod' his own powers and even aspects of the world around him, bypassing all traditional limitations.
The Kingdom of Eldoria, his summoners, expect a prophesied hero to combat the encroaching 'Demon Lord Malakor' and his monstrous legions. While Kenji has the raw power to be that hero, his methods are anything but conventional. He approaches epic quests like game levels, min-maxes his absurd abilities, exploits magical 'bugs,' and crafts hilariously overpowered 'builds' to resolve conflicts with shocking efficiency, often leaving both allies and enemies dumbfounded.
As he navigates the politics of Eldoria, the expectations of being a 'Hero,' and the true nature of the 'Demon Lord' (who might be more than just a final boss), Kenji must also decide his ultimate goal: Is it to find a way back to his old life, or to fully embrace his newfound omnipotence in a world that feels like the ultimate immersive RPG? His journey is one of overwhelming power, clever outsmarting of ancient rules, and perhaps, accidentally reshaping Aethel in his own, uniquely optimized image.""")

    author_style = 'Will Wight (emulating his engaging progression and action, but with an emphasis on clear and relatable prose for broader appeal)'

    genre = 'Isekai Fantasy / Overpowered Main Character / LitRPG Adventure Comedy'

    ENABLE_REFINEMENT_PASS = False # Set to True to enable the experimental refinement pass
#__________________________________________________________________________________________

    logger.info("--- Input Parameters ---")
    logger.info(f"Resume File: {resume_filename} (Expected in: {OUTPUT_FOLDER})")
    logger.info(f"Subject (first 100 chars): {subject[:100]}...")
    logger.info(f"Author Style Inspiration: {author_style}")
    logger.info(f"Genre(s): {genre}")
    logger.info(f"Refinement Pass Enabled: {ENABLE_REFINEMENT_PASS}")
    logger.info("------------------------")

    try:
        # Check for resume file existence early
        if not os.path.exists(os.path.join(OUTPUT_FOLDER, resume_filename)):
            logger.critical(f"CRITICAL FILE ERROR: Resume file '{resume_filename}' not found in '{os.path.abspath(OUTPUT_FOLDER)}'. Please ensure it exists.")
            raise FileNotFoundError(f"Resume file '{resume_filename}' not found in '{os.path.abspath(OUTPUT_FOLDER)}'")

        components = initialize_components()
        run_generation_pipeline(components, resume_filename, subject, author_style, genre, ENABLE_REFINEMENT_PASS)
    except FileNotFoundError as e: # Should be caught above, but as a safeguard
         logger.critical(f"CRITICAL FILE ERROR: {e}")
    except CriticalGenerationError as e: # Catch critical errors that should stop the process
        logger.critical(f"A critical error occurred, halting generation: {e}")
        logger.exception("Critical Error Details:")
    except Exception as e: # Catch any other unexpected errors
        logger.critical("AN UNEXPECTED CATASTROPHIC ERROR OCCURRED DURING NOVEL GENERATION.")
        logger.exception(f"Error Type: {type(e).__name__}, Details: {e}")
    finally:
        end_time = time.time()
        total_time = end_time - process_start_time
        logger.info("============================================")
        logger.info("=== NOVEL GENERATION PROCESS FINISHED ===")
        logger.info(f"=== Total Execution Time: {total_time:.2f} seconds ({total_time/60:.2f} minutes) ===")
        logger.info("============================================")

if __name__ == "__main__":
    main()
